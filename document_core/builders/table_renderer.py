from __future__ import annotations

try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    OxmlElement = None
    qn = None

from ..models.blocks import resolve_table_column_count
from .docx_utils import (
    _DEFAULT_FONT_NAME,
    clear_paragraph,
    format_run,
    resolve_alignment,
    rgb,
)

DOCX_TABLE_STYLES = {
    "report_grid": "Table Grid",
    "metrics_compact": "Light List Accent 1",
    "minimal": "Table Grid",
}


class TableRenderer:
    def render(
        self,
        doc,
        block,
        theme: dict,
        *,
        default_table_style: str,
    ) -> None:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        column_count = resolve_table_column_count(block.headers, block.rows)
        if column_count <= 0:
            return

        title_row_offset = 1 if block.caption else 0
        group_header_row_offset = 1 if block.header_groups else 0
        header_row_offset = 1 if block.headers else 0
        row_count = (
            len(block.rows)
            + title_row_offset
            + group_header_row_offset
            + header_row_offset
        )
        table = doc.add_table(rows=row_count, cols=column_count)
        style_name = (
            getattr(block.style, "table_grid", None)
            or getattr(block, "table_style", "")
            or default_table_style
            or theme["table_style"]
        )
        table.alignment = self._table_alignment(
            block.table_align or theme.get("table_align"),
            WD_TABLE_ALIGNMENT,
        )
        resolved_style = self.resolve_docx_table_style(style_name)
        if resolved_style:
            try:
                table.style = resolved_style
            except (KeyError, ValueError):
                if style_name == "metrics_compact":
                    try:
                        table.style = DOCX_TABLE_STYLES["report_grid"]
                    except (KeyError, ValueError):
                        pass
        self._apply_table_borders(
            table,
            border_style=block.border_style or theme.get("table_border_style"),
            accent_color=theme["accent"],
        )

        # Table styling precedence stays consistent across helper methods:
        # block-level fields > document_style.table_defaults in theme > preset/theme fallbacks.
        body_alignment = getattr(block.style, "cell_align", None) or theme.get(
            "table_cell_align"
        )
        numeric_columns = set(getattr(block, "numeric_columns", []) or [])

        if getattr(block, "column_widths", None):
            table.autofit = False
            for col_index, width_cm in enumerate(block.column_widths[:column_count]):
                if width_cm <= 0:
                    continue
                target_width = Cm(width_cm)
                for row in table.rows:
                    row.cells[col_index].width = target_width

        row_index = 0
        if block.caption:
            title_cell = table.rows[0].cells[0]
            for col_index in range(1, column_count):
                title_cell = title_cell.merge(table.rows[0].cells[col_index])
            self._set_cell_text(
                title_cell,
                block.caption,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                font_size=Pt(
                    self._caption_font_size(
                        block.caption_emphasis or theme.get("table_caption_emphasis"),
                        theme,
                    )
                ),
                font_name=theme["font_name"],
                color=self._caption_color(block, theme),
                background=self._caption_fill(block, theme),
            )
            row_index = 1

        if block.header_groups:
            self._render_group_header_row(
                table.rows[row_index],
                block.header_groups,
                block=block,
                style_name=style_name,
                theme=theme,
            )
            row_index += 1

        if block.headers:
            header_row = table.rows[row_index]
            for col_index, value in enumerate(block.headers):
                self._set_cell_text(
                    header_row.cells[col_index],
                    value,
                    bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    font_size=Pt(self._table_font_size(style_name, theme, header=True)),
                    font_name=theme["font_name"],
                    color=self._table_header_color(block, style_name, theme),
                    background=self._header_fill(block, style_name, theme),
                )
            row_index += 1

        for data_row_index, values in enumerate(block.rows):
            current_row = row_index + data_row_index
            for col_index in range(column_count):
                default_alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if style_name == "metrics_compact" and col_index > 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                if body_alignment is None and col_index in numeric_columns:
                    default_alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._set_cell_text(
                    table.rows[current_row].cells[col_index],
                    values[col_index] if col_index < len(values) else "",
                    bold=bool(self._first_column_bold(block, theme) and col_index == 0),
                    alignment=resolve_alignment(
                        body_alignment,
                        default=default_alignment,
                    ),
                    font_size=Pt(
                        self._table_font_size(style_name, theme, header=False)
                    ),
                    font_name=theme["font_name"],
                    background=self._table_row_fill(
                        block,
                        style_name,
                        data_row_index + 1,
                        theme,
                    ),
                )

    def _render_group_header_row(
        self,
        row,
        header_groups,
        *,
        block,
        style_name: str,
        theme: dict,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        current_col = 0
        for group in header_groups:
            merged_cell = row.cells[current_col]
            span_end = current_col + group.span
            for merge_col in range(current_col + 1, span_end):
                merged_cell = merged_cell.merge(row.cells[merge_col])
            self._set_cell_text(
                merged_cell,
                group.title,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                font_size=Pt(self._table_font_size(style_name, theme, header=True)),
                font_name=theme["font_name"],
                color=self._table_header_color(block, style_name, theme),
                background=self._header_fill(block, style_name, theme),
            )
            current_col = span_end

    @staticmethod
    def resolve_docx_table_style(style_name: str) -> str | None:
        candidate = (style_name or "").strip()
        if not candidate:
            return DOCX_TABLE_STYLES["report_grid"]
        return DOCX_TABLE_STYLES.get(candidate, candidate)

    @staticmethod
    def _set_cell_text(
        cell,
        value: str,
        *,
        bold: bool,
        alignment,
        font_size,
        font_name: str = _DEFAULT_FONT_NAME,
        color=None,
        background: str | None = None,
    ) -> None:
        text = value or ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        clear_paragraph(paragraph)
        run = paragraph.add_run(text)
        format_run(
            run,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            color=color,
        )
        paragraph.paragraph_format.space_after = 0
        if background:
            TableRenderer._set_cell_background(cell, background)

    @staticmethod
    def _header_fill(block, style_name: str, theme: dict) -> str | None:
        if block.header_fill:
            return block.header_fill
        if theme.get("table_header_fill"):
            return theme["table_header_fill"]
        if style_name == "minimal":
            return theme["accent_soft"]
        return theme["accent"]

    @staticmethod
    def _table_header_color(block, style_name: str, theme: dict):
        if block.header_text_color:
            return rgb(block.header_text_color)
        if theme.get("table_header_text_color"):
            return rgb(theme["table_header_text_color"])
        if style_name == "minimal":
            return rgb(theme["accent"])
        return rgb("FFFFFF")

    @staticmethod
    def _table_row_fill(
        block,
        style_name: str,
        row_index: int,
        theme: dict,
    ) -> str | None:
        if block.banded_rows is False:
            return None
        if block.banded_rows is True and row_index % 2 == 1:
            return block.banded_row_fill or "F7FBFF"
        if theme.get("table_banded_rows") is True:
            if row_index % 2 == 1:
                return theme.get("table_banded_row_fill") or "F7FBFF"
            return None
        if style_name == "report_grid" and row_index % 2 == 1:
            return "F7FBFF"
        return None

    @staticmethod
    def _first_column_bold(block, theme: dict) -> bool:
        if block.first_column_bold is not None:
            return block.first_column_bold
        return bool(theme.get("table_first_column_bold"))

    @staticmethod
    def _table_font_size(style_name: str, theme: dict, *, header: bool) -> float:
        base_size = theme["table_font_size"]
        if style_name == "metrics_compact":
            return max(base_size - 0.5, 9)
        if style_name == "minimal" and header:
            return max(base_size, theme["body_size"])
        return base_size

    @staticmethod
    def _set_cell_background(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _table_alignment(table_align: str | None, alignment_enum):
        if table_align == "left":
            return alignment_enum.LEFT
        return alignment_enum.CENTER

    @staticmethod
    def _caption_fill(block, theme: dict) -> str:
        resolved_emphasis = block.caption_emphasis or theme.get(
            "table_caption_emphasis"
        )
        if resolved_emphasis == "strong":
            return (
                block.header_fill or theme.get("table_header_fill") or theme["accent"]
            )
        return theme["accent_soft"]

    @staticmethod
    def _caption_color(block, theme: dict):
        resolved_emphasis = block.caption_emphasis or theme.get(
            "table_caption_emphasis"
        )
        if resolved_emphasis == "strong":
            return rgb(
                block.header_text_color
                or theme.get("table_header_text_color")
                or "FFFFFF"
            )
        return rgb(theme["accent"])

    @staticmethod
    def _caption_font_size(caption_emphasis: str | None, theme: dict) -> float:
        base_size = max(theme["body_size"], 11)
        if caption_emphasis == "strong":
            return base_size + 1
        return base_size

    @staticmethod
    def _apply_table_borders(
        table, *, border_style: str | None, accent_color: str
    ) -> None:
        if not border_style:
            return

        border_map = {
            "minimal": {"size": "4", "color": "D0D7DE"},
            "standard": {"size": "8", "color": "7A7A7A"},
            "strong": {"size": "16", "color": accent_color},
        }
        border_spec = border_map[border_style]

        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = tbl_borders.find(qn(f"w:{edge_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{edge_name}")
                tbl_borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), border_spec["size"])
            edge.set(qn("w:color"), border_spec["color"])
