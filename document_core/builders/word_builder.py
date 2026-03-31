from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from ..macros import (
    build_summary_card_group,
    expand_summary_card_block,
    summary_card_defaults_from_theme,
)
from ..models import blocks as block_models
from ..models.document import DocumentModel
from .docx_utils import (
    _DEFAULT_CODE_FONT_NAME,
    _DEFAULT_FONT_NAME,
    format_run,
    resolve_alignment,
    rgb,
)
from .table_renderer import TableRenderer

if TYPE_CHECKING:
    from docx.document import Document as WordDocument

ColumnsBlock = block_models.ColumnsBlock
GroupBlock = block_models.GroupBlock
HeaderFooterConfig = block_models.HeaderFooterConfig
HeadingBlock = block_models.HeadingBlock
ImageBlock = block_models.ImageBlock
ListBlock = block_models.ListBlock
PageBreakBlock = block_models.PageBreakBlock
ParagraphBlock = block_models.ParagraphBlock
SectionBreakBlock = block_models.SectionBreakBlock
TableBlock = block_models.TableBlock
SummaryCardBlock = getattr(block_models, "SummaryCardBlock", None)
TocBlock = block_models.TocBlock

THEMES = {
    "business_report": {
        "accent": "1F4E79",
        "accent_soft": "DCE6F1",
        "title_size": 20,
        "heading_size": 14,
        "body_size": 11,
        "table_style": "report_grid",
        "summary_fill": "EEF4FA",
        "font_name": _DEFAULT_FONT_NAME,
    },
    "project_review": {
        "accent": "0F766E",
        "accent_soft": "D9F3EE",
        "title_size": 19,
        "heading_size": 13,
        "body_size": 10.5,
        "table_style": "metrics_compact",
        "summary_fill": "E8F6F3",
        "font_name": _DEFAULT_FONT_NAME,
    },
    "executive_brief": {
        "accent": "B45309",
        "accent_soft": "FDEBD8",
        "title_size": 18.5,
        "heading_size": 12.5,
        "body_size": 10.5,
        "table_style": "minimal",
        "summary_fill": "FFF7ED",
        "font_name": _DEFAULT_FONT_NAME,
    },
}

_DOCUMENT_STYLE_THEME_FIELDS = {
    "heading_color": "heading_color",
    "title_align": "title_align",
    "body_font_size": "body_size",
    "body_line_spacing": "body_line_spacing",
    "paragraph_space_after": "body_space_after",
    "list_space_after": "list_space_after",
}

_SUMMARY_CARD_THEME_FIELDS = {
    "title_align": "summary_card_title_align",
    "title_emphasis": "summary_card_title_emphasis",
    "title_font_scale": "summary_card_title_font_scale",
    "title_space_before": "summary_card_title_space_before",
    "title_space_after": "summary_card_title_space_after",
    "list_space_after": "summary_card_list_space_after",
}

_TABLE_DEFAULT_THEME_FIELDS = {
    "preset": "table_style",
    "header_fill": "table_header_fill",
    "header_text_color": "table_header_text_color",
    "banded_rows": "table_banded_rows",
    "banded_row_fill": "table_banded_row_fill",
    "first_column_bold": "table_first_column_bold",
    "table_align": "table_align",
    "border_style": "table_border_style",
    "caption_emphasis": "table_caption_emphasis",
    "cell_align": "table_cell_align",
}


def _merge_style_attrs(source, theme: dict, attr_map: dict[str, str]) -> None:
    for source_attr, theme_key in attr_map.items():
        value = getattr(source, source_attr, None)
        if value is not None:
            theme[theme_key] = value


class WordDocumentBuilder:
    def __init__(self, table_renderer: TableRenderer | None = None) -> None:
        self._table_renderer = table_renderer or TableRenderer()

    def build(self, document_model: DocumentModel, output_path: Path) -> Path:
        try:
            from docx import Document
            from docx.enum.section import WD_SECTION
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required to export Word documents."
            ) from exc

        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        section = doc.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        theme = self._resolve_theme(document_model)
        self._enable_update_fields_on_open(doc)
        self._configure_document_header_footer_modes(doc, document_model)
        self._apply_section_layout(section, margins=theme["margins"])
        current_header_footer = getattr(
            document_model.metadata, "header_footer", HeaderFooterConfig()
        )
        self._apply_section_header_footer(
            section,
            current_header_footer,
            theme,
        )

        if document_model.metadata.title:
            self._add_title(doc, document_model.metadata.title, theme)

        workspace_dir = output_path.parent.resolve()
        for block in document_model.blocks:
            current_header_footer = self._append_block(
                doc,
                block,
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )

        doc.save(str(output_path))
        return output_path

    def _append_block(
        self,
        doc: WordDocument,
        block,
        theme: dict,
        document_model: DocumentModel,
        workspace_dir: Path,
        current_header_footer: HeaderFooterConfig,
    ) -> HeaderFooterConfig:
        if isinstance(block, HeadingBlock):
            self._add_heading(doc, block, theme)
        elif isinstance(block, ParagraphBlock):
            self._add_paragraph(
                doc,
                block,
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )
        elif isinstance(block, ListBlock):
            self._add_list(doc, block, theme)
        elif isinstance(block, TableBlock):
            self._table_renderer.render(
                doc,
                block,
                theme,
                default_table_style=getattr(
                    document_model.metadata, "table_template", ""
                ),
            )
        elif isinstance(block, ImageBlock):
            self._add_image(doc, block, workspace_dir)
        elif isinstance(block, GroupBlock):
            current_header_footer = self._add_group(
                doc,
                block,
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )
        elif isinstance(block, ColumnsBlock):
            current_header_footer = self._add_columns(
                doc,
                block,
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )
        elif isinstance(block, PageBreakBlock):
            self._add_page_break(doc)
        elif isinstance(block, SectionBreakBlock):
            current_header_footer = self._add_section_break(
                doc, block, theme, current_header_footer
            )
        elif isinstance(block, TocBlock):
            self._add_toc(doc, block, theme)
        elif SummaryCardBlock is not None and isinstance(block, SummaryCardBlock):
            current_header_footer = self._add_group(
                doc,
                expand_summary_card_block(
                    block,
                    **summary_card_defaults_from_theme(theme),
                ),
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )
        return current_header_footer

    def _resolve_theme(self, document_model: DocumentModel) -> dict:
        theme_name = getattr(document_model.metadata, "theme_name", "business_report")
        theme = dict(THEMES.get(theme_name, THEMES["business_report"]))
        density = getattr(document_model.metadata, "density", "comfortable")
        if density == "compact":
            theme.update(
                {
                    "margins": {"top": 2.2, "bottom": 2.1, "left": 2.4, "right": 2.3},
                    "title_align": "center",
                    "title_spacing_after": 14,
                    "heading_space_before": 10,
                    "heading_space_after": 5,
                    "body_indent": 18,
                    "body_space_after": 6,
                    "body_line_spacing": 1.2,
                    "list_space_after": 4,
                    "table_font_size": 9.5,
                }
            )
        else:
            theme.update(
                {
                    "margins": {"top": 2.8, "bottom": 2.6, "left": 2.8, "right": 2.6},
                    "title_align": "center",
                    "title_spacing_after": 18,
                    "heading_space_before": 14,
                    "heading_space_after": 8,
                    "body_indent": 24,
                    "body_space_after": 10,
                    "body_line_spacing": 1.5,
                    "list_space_after": 6,
                    "table_font_size": 10.5,
                }
            )

        accent_color = getattr(document_model.metadata, "accent_color", "")
        if accent_color:
            theme["accent"] = accent_color
            theme["accent_soft"] = self._blend_hex(accent_color, "FFFFFF", 0.84)
            theme["summary_fill"] = self._blend_hex(accent_color, "FFFFFF", 0.92)

        document_style = getattr(document_model.metadata, "document_style", None)
        if document_style is not None:
            _merge_style_attrs(document_style, theme, _DOCUMENT_STYLE_THEME_FIELDS)
            _merge_style_attrs(
                document_style.summary_card_defaults,
                theme,
                _SUMMARY_CARD_THEME_FIELDS,
            )
            _merge_style_attrs(
                document_style.table_defaults,
                theme,
                _TABLE_DEFAULT_THEME_FIELDS,
            )

        return theme

    def _add_title(self, doc: WordDocument, text: str, theme: dict) -> None:
        from docx.shared import Pt

        paragraph = doc.add_paragraph()
        paragraph.alignment = resolve_alignment(theme.get("title_align"), default=None)
        paragraph.paragraph_format.space_after = Pt(theme["title_spacing_after"])
        run = paragraph.add_run(text)
        format_run(
            run,
            font_name=theme["font_name"],
            font_size=Pt(theme["title_size"]),
            bold=True,
            color=rgb(theme["accent"]),
        )

    def _add_heading(self, doc: WordDocument, block: HeadingBlock, theme: dict) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        paragraph = doc.add_paragraph()
        paragraph.alignment = resolve_alignment(
            getattr(block.style, "align", None),
            default=WD_ALIGN_PARAGRAPH.LEFT,
        )
        paragraph.paragraph_format.space_before = Pt(
            self._resolved_spacing(
                getattr(block.layout, "spacing_before", None),
                theme["heading_space_before"],
            )
        )
        paragraph.paragraph_format.space_after = Pt(
            self._resolved_spacing(
                getattr(block.layout, "spacing_after", None),
                theme["heading_space_after"],
            )
        )
        run = paragraph.add_run(block.text)
        format_run(
            run,
            font_name=theme["font_name"],
            font_size=Pt(
                self._scaled_size(
                    theme["heading_size"]
                    if block.level <= 1
                    else max(theme["body_size"] + 1, 11.5),
                    getattr(block.style, "font_scale", None),
                )
            ),
            bold=True,
            color=self._resolve_text_color(
                theme=theme,
                emphasis=getattr(block.style, "emphasis", None),
                default_color=rgb(theme.get("heading_color", theme["accent"])),
            ),
        )

    def _add_paragraph(
        self,
        doc: WordDocument,
        block: ParagraphBlock,
        theme: dict,
        document_model: DocumentModel,
        workspace_dir: Path,
        current_header_footer: HeaderFooterConfig,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Pt

        variant = getattr(block, "variant", "body")
        paragraph_text = self._paragraph_text(block)
        if variant in {"summary_box", "key_takeaway"}:
            card_title = getattr(block, "title", "") or (
                "Summary" if variant == "summary_box" else "Key Takeaway"
            )
            self._add_group(
                doc,
                build_summary_card_group(
                    title=card_title,
                    items=[paragraph_text],
                    variant="summary",
                    style=block.style,
                    layout=block.layout,
                    **summary_card_defaults_from_theme(theme),
                ),
                theme,
                document_model,
                workspace_dir=workspace_dir,
                current_header_footer=current_header_footer,
            )
            return

        paragraph = doc.add_paragraph()
        paragraph.alignment = resolve_alignment(
            getattr(block.style, "align", None),
            default=WD_ALIGN_PARAGRAPH.LEFT,
        )
        paragraph.paragraph_format.first_line_indent = Pt(theme["body_indent"])
        paragraph.paragraph_format.space_before = Pt(
            self._resolved_spacing(getattr(block.layout, "spacing_before", None), 0)
        )
        paragraph.paragraph_format.space_after = Pt(
            self._resolved_spacing(
                getattr(block.layout, "spacing_after", None),
                theme["body_space_after"],
            )
        )
        paragraph.paragraph_format.line_spacing = theme["body_line_spacing"]
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        self._append_paragraph_runs(paragraph, block, theme)

    def _add_list(self, doc: WordDocument, block: ListBlock, theme: dict) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        for index, item in enumerate(block.items, start=1):
            paragraph = doc.add_paragraph()
            paragraph.alignment = resolve_alignment(
                getattr(block.style, "align", None),
                default=WD_ALIGN_PARAGRAPH.LEFT,
            )
            paragraph.paragraph_format.left_indent = Pt(
                max(theme["body_indent"] - 6, 12)
            )
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(
                self._resolved_spacing(
                    getattr(block.layout, "spacing_before", None),
                    0 if index > 1 else 0,
                )
            )
            paragraph.paragraph_format.space_after = Pt(
                self._resolved_spacing(
                    getattr(block.layout, "spacing_after", None),
                    theme["list_space_after"],
                )
            )
            marker = f"{index}. " if block.ordered else "• "
            run = paragraph.add_run(f"{marker}{item}")
            format_run(
                run,
                font_name=theme["font_name"],
                font_size=Pt(
                    self._scaled_size(
                        theme["body_size"], getattr(block.style, "font_scale", None)
                    )
                ),
                bold=self._resolved_bold(False, getattr(block.style, "emphasis", None)),
                color=self._resolve_text_color(
                    theme=theme,
                    emphasis=getattr(block.style, "emphasis", None),
                ),
            )

    def _add_group(
        self,
        doc: WordDocument,
        block: GroupBlock,
        theme: dict,
        document_model: DocumentModel,
        workspace_dir: Path,
        current_header_footer: HeaderFooterConfig,
    ) -> HeaderFooterConfig:
        for child in block.blocks:
            current_header_footer = self._append_block(
                doc,
                child,
                theme,
                document_model,
                workspace_dir,
                current_header_footer,
            )
        return current_header_footer

    def _add_columns(
        self,
        doc: WordDocument,
        block: ColumnsBlock,
        theme: dict,
        document_model: DocumentModel,
        workspace_dir: Path,
        current_header_footer: HeaderFooterConfig,
    ) -> HeaderFooterConfig:
        # Word multi-column layout requires section-level changes. For now we
        # keep the primitive in the model and render it sequentially as a safe
        # fallback instead of introducing layout-specific template branches.
        for column_index, column in enumerate(block.columns):
            if column_index > 0:
                doc.add_paragraph("")
            for child in column.blocks:
                current_header_footer = self._append_block(
                    doc,
                    child,
                    theme,
                    document_model,
                    workspace_dir,
                    current_header_footer,
                )
        return current_header_footer

    @staticmethod
    def _add_page_break(doc: WordDocument) -> None:
        doc.add_page_break()

    def _add_section_break(
        self,
        doc: WordDocument,
        block: SectionBreakBlock,
        theme: dict,
        current_header_footer: HeaderFooterConfig,
    ) -> HeaderFooterConfig:
        from docx.enum.section import WD_SECTION

        start_type_map = {
            "continuous": WD_SECTION.CONTINUOUS,
            "odd_page": WD_SECTION.ODD_PAGE,
            "even_page": WD_SECTION.EVEN_PAGE,
            "new_column": WD_SECTION.NEW_COLUMN,
            "new_page": WD_SECTION.NEW_PAGE,
        }
        start_type = start_type_map.get(block.start_type, WD_SECTION.NEW_PAGE)
        section = doc.add_section(start_type)
        self._apply_section_layout(
            section,
            page_orientation=block.page_orientation,
            margins=self._resolve_section_margins(block, theme),
        )
        self._apply_section_page_numbering(section, block)
        if block.inherit_header_footer and not self._has_header_footer_override(
            block.header_footer
        ):
            return current_header_footer
        effective_header_footer = (
            self._merge_header_footer_config(
                current_header_footer,
                block.header_footer,
            )
            if block.inherit_header_footer
            else block.header_footer.model_copy(deep=True)
        )
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        self._apply_section_header_footer(section, effective_header_footer, theme)
        return effective_header_footer

    def _add_toc(self, doc: WordDocument, block: TocBlock, theme: dict) -> None:
        from docx.shared import Pt

        if block.start_on_new_page:
            doc.add_page_break()

        if block.title.strip():
            title_paragraph = doc.add_paragraph()
            title_paragraph.paragraph_format.space_after = Pt(
                self._resolved_spacing(
                    getattr(block.layout, "spacing_after", None),
                    theme["heading_space_after"],
                )
            )
            title_run = title_paragraph.add_run(block.title)
            format_run(
                title_run,
                font_name=theme["font_name"],
                font_size=Pt(theme["heading_size"]),
                bold=True,
                color=rgb(theme.get("heading_color", theme["accent"])),
            )

        toc_paragraph = doc.add_paragraph()
        toc_paragraph.paragraph_format.space_after = Pt(theme["body_space_after"])
        self._append_field_code(
            toc_paragraph,
            f'TOC \\o "1-{block.levels}" \\h \\z \\u',
        )

    @staticmethod
    def _blend_hex(source: str, target: str, ratio: float) -> str:
        ratio = min(max(ratio, 0.0), 1.0)
        source_channels = [int(source[index : index + 2], 16) for index in (0, 2, 4)]
        target_channels = [int(target[index : index + 2], 16) for index in (0, 2, 4)]
        blended = [
            round(src * (1 - ratio) + dst * ratio)
            for src, dst in zip(source_channels, target_channels, strict=False)
        ]
        return "".join(f"{channel:02X}" for channel in blended)

    def _add_image(
        self, doc: WordDocument, block: ImageBlock, workspace_dir: Path
    ) -> None:
        from docx.shared import Inches

        image_path = self._resolve_workspace_image_path(block.path, workspace_dir)
        if image_path is None:
            return
        section = doc.sections[0]
        max_width = section.page_width - section.left_margin - section.right_margin
        picture = doc.add_picture(str(image_path))

        if block.width_px is not None:
            picture.width = min(Inches(block.width_px / 96.0), max_width)
        elif picture.width > max_width:
            picture.width = max_width

        if block.caption:
            doc.add_paragraph(block.caption)

    @staticmethod
    def _resolve_workspace_image_path(
        path_value: str, workspace_dir: Path
    ) -> Path | None:
        workspace_root = workspace_dir.resolve()
        candidate = Path(path_value)
        if not candidate.is_absolute():
            candidate = workspace_root / candidate

        resolved_candidate = candidate.resolve(strict=False)
        try:
            resolved_candidate.relative_to(workspace_root)
        except ValueError:
            return None

        return resolved_candidate if resolved_candidate.is_file() else None

    def _configure_document_header_footer_modes(
        self, doc: WordDocument, document_model: DocumentModel
    ) -> None:
        doc.settings.odd_and_even_pages_header_footer = any(
            self._uses_even_page_variants(config)
            for config in self._iter_header_footer_configs(document_model)
        )

    def _iter_header_footer_configs(self, document_model: DocumentModel):
        yield getattr(document_model.metadata, "header_footer", HeaderFooterConfig())
        yield from self._iter_nested_header_footer_configs(document_model.blocks)

    def _iter_nested_header_footer_configs(self, blocks):
        for block in blocks:
            if isinstance(block, SectionBreakBlock):
                yield block.header_footer
            elif isinstance(block, GroupBlock):
                yield from self._iter_nested_header_footer_configs(block.blocks)
            elif isinstance(block, ColumnsBlock):
                for column in block.columns:
                    yield from self._iter_nested_header_footer_configs(
                        column.blocks
                    )

    def _apply_section_layout(
        self,
        section,
        *,
        page_orientation: str | None = None,
        margins: dict[str, float] | None = None,
    ) -> None:
        from docx.enum.section import WD_ORIENT
        from docx.shared import Cm

        if page_orientation == "landscape":
            if section.orientation != WD_ORIENT.LANDSCAPE:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )
        elif page_orientation == "portrait":
            if section.orientation != WD_ORIENT.PORTRAIT:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )

        if margins is None:
            return
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])

    @staticmethod
    def _resolve_section_margins(block: SectionBreakBlock, theme: dict) -> dict[str, float] | None:
        overrides = block.margins
        if not any(
            value is not None
            for value in (
                overrides.top_cm,
                overrides.bottom_cm,
                overrides.left_cm,
                overrides.right_cm,
            )
        ):
            return None
        base_margins = theme["margins"]
        return {
            "top": overrides.top_cm if overrides.top_cm is not None else base_margins["top"],
            "bottom": (
                overrides.bottom_cm
                if overrides.bottom_cm is not None
                else base_margins["bottom"]
            ),
            "left": overrides.left_cm if overrides.left_cm is not None else base_margins["left"],
            "right": (
                overrides.right_cm
                if overrides.right_cm is not None
                else base_margins["right"]
            ),
        }

    def _apply_section_page_numbering(
        self, section, block: SectionBreakBlock
    ) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        section_properties = section._sectPr
        page_number = section_properties.find(qn("w:pgNumType"))
        if not block.restart_page_numbering:
            if page_number is not None:
                section_properties.remove(page_number)
            return
        if page_number is None:
            page_number = OxmlElement("w:pgNumType")
            section_properties.append(page_number)
        page_number.set(qn("w:start"), str(block.page_number_start or 1))

    def _apply_section_header_footer(
        self, section, config: HeaderFooterConfig, theme: dict
    ) -> None:
        uses_first_page = self._uses_first_page_variants(config)
        uses_even_page = self._uses_even_page_variants(config)

        section.different_first_page_header_footer = uses_first_page
        self._set_story_text(section.header, config.header_text, theme)
        self._set_footer_content(
            section.footer,
            text=config.footer_text,
            show_page_number=bool(config.show_page_number),
            page_number_align=config.page_number_align,
            theme=theme,
        )

        if uses_first_page:
            section.first_page_header.is_linked_to_previous = False
            section.first_page_footer.is_linked_to_previous = False
            self._set_story_text(
                section.first_page_header,
                config.first_page_header_text,
                theme,
            )
            self._set_footer_content(
                section.first_page_footer,
                text=config.first_page_footer_text,
                show_page_number=(
                    config.first_page_show_page_number
                    if config.first_page_show_page_number is not None
                    else bool(config.show_page_number)
                ),
                page_number_align=config.page_number_align,
                theme=theme,
            )
        else:
            self._set_story_text(section.first_page_header, "", theme)
            self._set_footer_content(
                section.first_page_footer,
                text="",
                show_page_number=False,
                page_number_align=config.page_number_align,
                theme=theme,
            )

        if uses_even_page:
            section.even_page_header.is_linked_to_previous = False
            section.even_page_footer.is_linked_to_previous = False
            self._set_story_text(
                section.even_page_header,
                config.even_page_header_text or config.header_text,
                theme,
            )
            self._set_footer_content(
                section.even_page_footer,
                text=config.even_page_footer_text or config.footer_text,
                show_page_number=(
                    config.even_page_show_page_number
                    if config.even_page_show_page_number is not None
                    else bool(config.show_page_number)
                ),
                page_number_align=config.page_number_align,
                theme=theme,
            )
        else:
            self._set_story_text(section.even_page_header, "", theme)
            self._set_footer_content(
                section.even_page_footer,
                text="",
                show_page_number=False,
                page_number_align=config.page_number_align,
                theme=theme,
            )

    def _set_footer_content(
        self,
        story,
        *,
        text: str,
        show_page_number: bool,
        page_number_align: str,
        theme: dict,
    ) -> None:
        paragraph_count = 2 if text.strip() and show_page_number else 1
        paragraphs = self._reset_story(story, paragraph_count=paragraph_count)
        if text.strip():
            text_paragraph = paragraphs[0]
            text_paragraph.alignment = resolve_alignment("left", default=None)
            run = text_paragraph.add_run(text)
            format_run(
                run,
                font_name=theme["font_name"],
                font_size=None,
                bold=False,
            )
        if show_page_number:
            page_paragraph = paragraphs[-1]
            page_paragraph.alignment = resolve_alignment(
                page_number_align,
                default=None,
            )
            self._append_page_number_field(page_paragraph)

    def _set_story_text(self, story, text: str, theme: dict) -> None:
        paragraphs = self._reset_story(story)
        if not text.strip():
            return
        run = paragraphs[0].add_run(text)
        format_run(
            run,
            font_name=theme["font_name"],
            font_size=None,
            bold=False,
        )

    def _reset_story(self, story, *, paragraph_count: int = 1):
        paragraphs = list(story.paragraphs)
        if not paragraphs:
            paragraphs = [story.add_paragraph()]
        for paragraph in paragraphs[1:]:
            paragraph._element.getparent().remove(paragraph._element)
        base_paragraph = story.paragraphs[0]
        self._clear_paragraph(base_paragraph)
        paragraphs = [base_paragraph]
        while len(paragraphs) < paragraph_count:
            new_paragraph = story.add_paragraph()
            self._clear_paragraph(new_paragraph)
            paragraphs.append(new_paragraph)
        return paragraphs

    def _append_page_number_field(self, paragraph) -> None:
        self._append_field_code(paragraph, "PAGE")

    def _append_field_code(self, paragraph, instruction: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        paragraph._p.append(begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        paragraph._p.append(instr)

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        paragraph._p.append(separate)

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        paragraph._p.append(end)

    def _enable_update_fields_on_open(self, doc: WordDocument) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    @staticmethod
    def _clear_paragraph(paragraph) -> None:
        from docx.oxml.ns import qn

        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

    @staticmethod
    def _has_header_footer_override(config: HeaderFooterConfig) -> bool:
        return any(
            WordDocumentBuilder._config_field_is_set(config, field_name)
            for field_name in HeaderFooterConfig.model_fields
        )

    @staticmethod
    def _uses_first_page_variants(config: HeaderFooterConfig) -> bool:
        return any(
            [
                config.different_first_page,
                config.first_page_header_text.strip(),
                config.first_page_footer_text.strip(),
                config.first_page_show_page_number is not None,
            ]
        )

    @staticmethod
    def _uses_even_page_variants(config: HeaderFooterConfig) -> bool:
        return any(
            [
                config.different_odd_even,
                config.even_page_header_text.strip(),
                config.even_page_footer_text.strip(),
                config.even_page_show_page_number is not None,
            ]
        )

    @staticmethod
    def _config_field_is_set(
        config: HeaderFooterConfig, field_name: str
    ) -> bool:
        return field_name in getattr(config, "model_fields_set", set())

    def _merge_header_footer_config(
        self,
        base_config: HeaderFooterConfig,
        override_config: HeaderFooterConfig,
    ) -> HeaderFooterConfig:
        merged_config = base_config.model_copy(deep=True)
        for field_name in HeaderFooterConfig.model_fields:
            if self._config_field_is_set(override_config, field_name):
                setattr(
                    merged_config,
                    field_name,
                    getattr(override_config, field_name),
                )
        return merged_config

    @staticmethod
    def _resolved_spacing(value: float | None, default: float) -> float:
        return default if value is None else value

    @staticmethod
    def _scaled_size(base_size: float, scale: float | None) -> float:
        return base_size if scale is None else base_size * scale

    @staticmethod
    def _resolved_bold(default_bold: bool, emphasis: str | None) -> bool:
        if emphasis == "strong":
            return True
        return default_bold

    @staticmethod
    def _paragraph_text(block: ParagraphBlock) -> str:
        if block.runs:
            return "".join(run.text for run in block.runs)
        return block.text

    def _append_paragraph_runs(
        self, paragraph, block: ParagraphBlock, theme: dict
    ) -> None:
        from docx.shared import Pt

        font_size = Pt(
            self._scaled_size(
                theme["body_size"], getattr(block.style, "font_scale", None)
            )
        )
        default_color = self._resolve_text_color(
            theme=theme,
            emphasis=getattr(block.style, "emphasis", None),
            default_color=None,
        )

        if block.runs:
            for run_block in block.runs:
                font_name = (
                    _DEFAULT_CODE_FONT_NAME if run_block.code else theme["font_name"]
                )
                run = paragraph.add_run(run_block.text)
                format_run(
                    run,
                    font_name=font_name,
                    font_size=font_size,
                    bold=self._resolved_bold(
                        run_block.bold, getattr(block.style, "emphasis", None)
                    ),
                    italic=run_block.italic,
                    underline=run_block.underline,
                    color=default_color,
                )
        else:
            run = paragraph.add_run(self._paragraph_text(block))
            format_run(
                run,
                font_name=theme["font_name"],
                font_size=font_size,
                bold=self._resolved_bold(False, getattr(block.style, "emphasis", None)),
                color=default_color,
            )

    def _resolve_text_color(
        self,
        *,
        theme: dict,
        emphasis: str | None,
        default_color=None,
    ):
        if emphasis == "subtle":
            return rgb(theme["accent"])
        return default_color
