from __future__ import annotations


def clear_paragraph(paragraph) -> None:
    from docx.oxml.ns import qn

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def reset_story(story, *, paragraph_count: int = 1):
    paragraphs = list(story.paragraphs)
    if not paragraphs:
        paragraphs = [story.add_paragraph()]
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    base_paragraph = story.paragraphs[0]
    clear_paragraph(base_paragraph)
    paragraphs = [base_paragraph]
    while len(paragraphs) < paragraph_count:
        new_paragraph = story.add_paragraph()
        clear_paragraph(new_paragraph)
        paragraphs.append(new_paragraph)
    return paragraphs


def append_field_code(paragraph, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run.append(instr)
    paragraph._p.append(instr_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def append_page_number_field(paragraph) -> None:
    append_field_code(paragraph, "PAGE")


def enable_update_fields_on_open(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
