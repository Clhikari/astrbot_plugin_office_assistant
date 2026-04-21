import base64
import contextlib
import json
import io
import shutil
import traceback
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import astrbot_plugin_office_assistant.utils as office_utils
import mcp
import pytest
from astrbot_plugin_office_assistant.constants import (
    DEFAULT_MAX_INLINE_DOCX_IMAGE_COUNT,
    DEFAULT_MAX_INLINE_DOCX_IMAGE_MB,
    DOC_COMMAND_TRIGGER_EVENT_KEY,
    EXPLICIT_FILE_TOOL_EVENT_KEY,
    OfficeType,
)
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendConfig,
)
from astrbot_plugin_office_assistant.domain.document.session_store import (
    get_document_style_defaults,
)
from astrbot_plugin_office_assistant.office_generator import OfficeGenerator
from astrbot_plugin_office_assistant.internal_hooks import NoticeBuildContext
from astrbot_plugin_office_assistant.message_buffer import BufferedMessage
from astrbot_plugin_office_assistant.services import (
    AccessPolicyService,
    CommandService,
    DeliveryService,
    ErrorHookService,
    ExportHookService,
    FileReadService,
    FileToolService,
    GeneratedFileDeliveryService,
    IncomingMessageService,
    LLMRequestPolicy,
    PostExportHookService,
    RequestHookService,
    UploadPromptService,
    UploadSessionService,
    WordReadService,
    WorkspaceService,
    build_plugin_runtime,
)
from astrbot_plugin_office_assistant.services.runtime_builder import (
    _build_document_summary_lookup,
    _build_workbook_toolset,
    _build_workbook_summary_lookup,
)
from astrbot_plugin_office_assistant.services.excel_intent_router import (
    ExcelIntentRouter,
)
from astrbot_plugin_office_assistant.services.prompt_context_service import (
    SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP,
    SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP,
    SECTION_SCENE_UPLOADED_CONTEXT,
    SECTION_STATIC_DOCUMENT_TOOLS,
    SECTION_STATIC_DOCUMENT_TOOLS_DETAIL,
    SECTION_STATIC_EXCEL_READ,
    SECTION_STATIC_EXCEL_ROUTING,
    SECTION_STATIC_EXCEL_SCRIPT,
    SECTION_STATIC_WORKBOOK_TOOLS,
    SECTION_STATIC_WORKBOOK_TOOLS_DETAIL,
    PromptContextService,
)
from astrbot_plugin_office_assistant.utils import (
    ExtractedWordContent,
    ExtractedWordItem,
    extract_word_text,
    format_extracted_word_content,
)

import astrbot.api.message_components as Comp
from astrbot.core.agent.tool import FunctionTool, ToolSet
from astrbot.core.platform.message_type import MessageType
from astrbot.core.provider.entities import ProviderRequest
from conftest import build_notice_once_callback as _build_notice_once_callback


def _build_event(
    *,
    sender_id: str = "user-1",
    message_type=MessageType.FRIEND_MESSAGE,
):
    event = MagicMock()
    extras: dict[str, object] = {}
    event.message_obj = SimpleNamespace(type=message_type, message=[], self_id="bot-1")
    event.get_sender_id.return_value = sender_id
    event.get_platform_id.return_value = "platform-1"
    event.unified_msg_origin = "session-1"
    event.message_str = ""
    event.is_admin.return_value = False
    event._buffer_reentry_count = 0
    event._buffered = False
    event.set_extra.side_effect = lambda key, value: extras.__setitem__(key, value)
    event.get_extra.side_effect = lambda key=None, default=None: (
        dict(extras) if key is None else extras.get(key, default)
    )
    return event


def _tool(name: str) -> FunctionTool:
    return FunctionTool(
        name=name,
        description=f"tool {name}",
        parameters={"type": "object", "properties": {}},
        handler=None,
    )


def _build_upload_infos(
    count: int,
    *,
    original_name_template: str = "file-{idx}.txt",
    stored_name_template: str = "file_{idx}.txt",
    source_path_template: str = "",
    file_suffix: str = ".txt",
    type_desc: str = "文本/代码文件",
    is_supported: bool = True,
) -> list[dict[str, object]]:
    return [
        {
            "original_name": original_name_template.format(idx=idx),
            "file_suffix": file_suffix,
            "type_desc": type_desc,
            "is_supported": is_supported,
            "stored_name": stored_name_template.format(idx=idx),
            "source_path": source_path_template.format(idx=idx)
            if source_path_template
            else "",
        }
        for idx in range(count)
    ]


def _make_workspace(name: str) -> Path:
    workspace_base = Path(__file__).resolve().parent / ".tmp_services"
    workspace_base.mkdir(parents=True, exist_ok=True)
    workspace_dir = workspace_base / f"{name}-{uuid4().hex}"
    workspace_dir.mkdir(parents=True, exist_ok=True)
    return workspace_dir


def _build_astrbot_context(*, runtime: str = "sandbox"):
    context = MagicMock()
    context.get_config.side_effect = lambda *args, **kwargs: {
        "provider_settings": {"computer_use_runtime": runtime}
    }
    return context


class _FakeSandboxPython:
    def __init__(self, workspace_root: Path):
        self._workspace_root = workspace_root

    async def exec(
        self,
        code: str,
        kernel_id: str | None = None,
        timeout: int = 30,
        silent: bool = False,
    ) -> dict[str, object]:
        _ = kernel_id
        _ = timeout
        globals_dict = {"__name__": "__main__"}
        stdout_buffer = io.StringIO()
        stderr_buffer = io.StringIO()
        try:
            with contextlib.redirect_stdout(stdout_buffer), contextlib.redirect_stderr(
                stderr_buffer
            ):
                exec(compile(code, "<fake-sandbox>", "exec"), globals_dict, globals_dict)
        except Exception:
            if not stderr_buffer.getvalue():
                stderr_buffer.write(traceback.format_exc())
            error_text = stderr_buffer.getvalue()
            return {
                "success": False,
                "output": "" if silent else stdout_buffer.getvalue(),
                "error": error_text,
                "data": {
                    "output": {"text": "" if silent else stdout_buffer.getvalue()},
                    "error": error_text,
                },
            }

        output_text = "" if silent else stdout_buffer.getvalue()
        return {
            "success": True,
            "output": output_text,
            "error": stderr_buffer.getvalue(),
            "data": {
                "output": {"text": output_text},
                "error": stderr_buffer.getvalue(),
            },
        }


class _FakeSandboxBooter:
    def __init__(
        self,
        workspace_root: Path,
        *,
        capabilities: tuple[str, ...] = ("python", "filesystem"),
    ):
        self.workspace_root = str(workspace_root)
        self.capabilities = capabilities
        self.python = _FakeSandboxPython(workspace_root)

    def _resolve_remote_path(self, path: str) -> Path:
        path_obj = Path(path)
        if path_obj.is_absolute():
            return path_obj
        return Path(self.workspace_root) / path_obj

    async def upload_file(self, path: str, file_name: str) -> dict[str, object]:
        source = Path(path)
        destination = Path(self.workspace_root) / Path(file_name)
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
        return {
            "success": True,
            "message": "uploaded",
            "file_path": str(destination),
        }

    async def download_file(self, remote_path: str, local_path: str) -> None:
        source = self._resolve_remote_path(remote_path)
        if not source.exists():
            raise FileNotFoundError(remote_path)
        destination = Path(local_path)
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)


def _patch_excel_sandbox(booter):
    return patch(
        "astrbot_plugin_office_assistant.services.excel_script_service._get_computer_booter",
        new=AsyncMock(return_value=booter),
    )


def _write_png(path: Path) -> None:
    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aF9kAAAAASUVORK5CYII="
    )
    path.write_bytes(png_bytes)


def _import_docx():
    return pytest.importorskip("docx")


def _find_paragraph(doc, text: str):
    return next(paragraph for paragraph in doc.paragraphs if paragraph.text == text)


def _node_render_backend_config_for_tests() -> DocumentRenderBackendConfig:
    renderer_entry = (
        Path(__file__).resolve().parents[1] / "word_renderer_js" / "dist" / "cli.js"
    )
    if shutil.which("node") is None or not renderer_entry.exists():
        pytest.skip("node renderer build is not available")
    return DocumentRenderBackendConfig(
        preferred_backend="node",
        fallback_enabled=False,
        node_renderer_entry=str(renderer_entry),
    )


def _build_file_tool_service(
    *,
    astrbot_context=None,
    workspace_service,
    office_generator=None,
    pdf_converter=None,
    delivery_service=None,
    office_libs=None,
    allow_external_input_files: bool = False,
    enable_docx_image_review: bool = True,
    max_inline_docx_image_bytes: int = DEFAULT_MAX_INLINE_DOCX_IMAGE_MB * 1024 * 1024,
    max_inline_docx_image_count: int = DEFAULT_MAX_INLINE_DOCX_IMAGE_COUNT,
    is_group_feature_enabled=None,
    check_permission=None,
    group_feature_disabled_error=None,
):
    delivery_service = delivery_service or MagicMock()
    astrbot_context = astrbot_context or _build_astrbot_context()
    return FileToolService(
        astrbot_context=astrbot_context,
        workspace_service=workspace_service,
        office_generator=office_generator or MagicMock(),
        pdf_converter=pdf_converter or MagicMock(),
        delivery_service=delivery_service,
        generated_file_delivery_service=GeneratedFileDeliveryService(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
        ),
        word_read_service=WordReadService(
            workspace_service=workspace_service,
            enable_docx_image_review=enable_docx_image_review,
            max_inline_docx_image_bytes=max_inline_docx_image_bytes,
            max_inline_docx_image_count=max_inline_docx_image_count,
        ),
        office_libs=office_libs or {},
        allow_external_input_files=allow_external_input_files,
        enable_docx_image_review=enable_docx_image_review,
        max_inline_docx_image_bytes=max_inline_docx_image_bytes,
        max_inline_docx_image_count=max_inline_docx_image_count,
        is_group_feature_enabled=is_group_feature_enabled or (lambda _event: True),
        check_permission=check_permission or (lambda _event: True),
        group_feature_disabled_error=group_feature_disabled_error
        or (lambda: "group disabled"),
    )


def test_file_tool_service_builds_default_word_and_delivery_services():
    workspace_service = MagicMock()
    delivery_service = MagicMock()

    service = FileToolService(
        workspace_service=workspace_service,
        office_generator=MagicMock(),
        pdf_converter=MagicMock(),
        delivery_service=delivery_service,
        generated_file_delivery_service=None,
        word_read_service=None,
        office_libs={},
        allow_external_input_files=False,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        group_feature_disabled_error=lambda: "group disabled",
    )

    assert isinstance(service._file_read_service._word_read_service, WordReadService)
    assert (
        service._file_read_service._word_read_service._max_inline_docx_image_bytes
        == DEFAULT_MAX_INLINE_DOCX_IMAGE_MB * 1024 * 1024
    )
    assert (
        service._file_read_service._word_read_service._max_inline_docx_image_count
        == DEFAULT_MAX_INLINE_DOCX_IMAGE_COUNT
    )
    assert (
        service._office_generate_service._file_delivery_service._generated_file_delivery_service
        is not None
    )


def test_file_tool_service_requires_workspace_for_default_file_read_service():
    with pytest.raises(
        ValueError,
        match="file_read_service requires injected service or dependencies: workspace_service",
    ):
        FileToolService(
            workspace_service=None,
            office_generator=MagicMock(),
            pdf_converter=MagicMock(),
            delivery_service=MagicMock(),
            generated_file_delivery_service=MagicMock(),
            word_read_service=None,
            office_libs={},
            allow_external_input_files=False,
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )


def test_file_tool_service_requires_permission_callbacks_for_default_services():
    with pytest.raises(
        ValueError,
        match=(
            "permission callbacks requires injected service or dependencies: "
            "is_group_feature_enabled, check_permission, group_feature_disabled_error"
        ),
    ):
        FileToolService(
            workspace_service=MagicMock(),
            office_generator=MagicMock(),
            pdf_converter=MagicMock(),
            delivery_service=MagicMock(),
            generated_file_delivery_service=MagicMock(),
            word_read_service=MagicMock(),
            office_libs={},
            allow_external_input_files=False,
        )


def test_file_tool_service_allows_missing_callbacks_when_all_services_injected():
    service = FileToolService(
        workspace_service=None,
        office_generator=None,
        pdf_converter=None,
        delivery_service=None,
        generated_file_delivery_service=None,
        word_read_service=None,
        office_libs={},
        allow_external_input_files=False,
        file_read_service=MagicMock(),
        office_generate_service=MagicMock(),
        pdf_convert_service=MagicMock(),
    )

    assert service._file_read_service is not None
    assert service._office_generate_service is not None
    assert service._pdf_convert_service is not None


def _rewrite_docx_document_xml(path: Path, transform) -> None:
    with zipfile.ZipFile(path, "r") as source_zip:
        file_map = {
            info.filename: source_zip.read(info.filename)
            for info in source_zip.infolist()
        }

    document_xml = file_map["word/document.xml"].decode("utf-8")
    file_map["word/document.xml"] = transform(document_xml).encode("utf-8")

    with zipfile.ZipFile(path, "w") as target_zip:
        for filename, content in file_map.items():
            target_zip.writestr(filename, content)


def test_access_policy_service_handles_whitelist_and_group_flags():
    service = AccessPolicyService(
        whitelist_users=["user-1"],
        admin_users=[],
        enable_features_in_group=False,
    )
    friend_event = _build_event(message_type=MessageType.FRIEND_MESSAGE)
    group_event = _build_event(message_type=MessageType.GROUP_MESSAGE)

    assert service.check_permission(friend_event) is True
    assert service.is_group_feature_enabled(friend_event) is True
    assert service.is_group_feature_enabled(group_event) is False


def test_access_policy_service_detects_bot_mention():
    service = AccessPolicyService(
        whitelist_users=["user-1"],
        admin_users=[],
        enable_features_in_group=True,
    )
    event = _build_event()
    event.message_obj.message = [Comp.At(qq="bot-1")]

    assert service.is_bot_mentioned(event) is True


def test_access_policy_service_detects_platform_level_bot_mention():
    service = AccessPolicyService(
        whitelist_users=["user-1"],
        admin_users=[],
        enable_features_in_group=True,
    )
    event = _build_event()
    event.message_obj.message = []
    event.is_mentioned.return_value = True

    assert service.is_bot_mentioned(event) is True


def test_access_policy_service_allows_framework_admin_sender_id():
    service = AccessPolicyService(
        whitelist_users=[],
        admin_users=["1474436119298048127"],
        enable_features_in_group=True,
    )
    event = _build_event(sender_id="1474436119298048127")
    event.is_admin.return_value = False

    assert service.check_permission(event) is True


def test_access_policy_service_reads_framework_admins_dynamically():
    admin_state = {"admins_id": set()}
    service = AccessPolicyService(
        whitelist_users=[],
        admin_users=[],
        get_admin_users=lambda: admin_state["admins_id"],
        enable_features_in_group=True,
    )
    event = _build_event(sender_id="1474436119298048127")
    event.is_admin.return_value = False

    assert service.check_permission(event) is False

    admin_state["admins_id"] = {"1474436119298048127"}

    assert service.check_permission(event) is True


@pytest.mark.asyncio
async def test_llm_request_policy_logs_missing_permission():
    event = _build_event(
        sender_id="1474436119298048127",
        message_type=MessageType.GROUP_MESSAGE,
    )
    request = ProviderRequest(
        prompt="请读取 report.docx",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    policy = LLMRequestPolicy(
        document_toolset=ToolSet([_tool("read_file")]),
        require_at_in_group=True,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: False,
        is_bot_mentioned=lambda _event: True,
        notice_hooks=[],
        tool_exposure_hooks=[],
    )

    with patch(
        "astrbot_plugin_office_assistant.services.llm_request_policy.logger.debug"
    ) as logger_debug:
        await policy.apply(event, request)

    assert "read_file" not in set(request.func_tool.names())
    logger_debug.assert_any_call(
        "[文件管理] 用户 1474436119298048127 无文件权限，已隐藏文件工具"
    )


@pytest.mark.asyncio
async def test_llm_request_policy_logs_missing_group_trigger():
    event = _build_event(
        sender_id="1474436119298048127",
        message_type=MessageType.GROUP_MESSAGE,
    )
    request = ProviderRequest(
        prompt="请读取 report.docx",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    policy = LLMRequestPolicy(
        document_toolset=ToolSet([_tool("read_file")]),
        require_at_in_group=True,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        is_bot_mentioned=lambda _event: False,
        notice_hooks=[],
        tool_exposure_hooks=[],
    )

    with patch(
        "astrbot_plugin_office_assistant.services.llm_request_policy.logger.debug"
    ) as logger_debug:
        await policy.apply(event, request)

    assert "read_file" not in set(request.func_tool.names())
    logger_debug.assert_any_call(
        "[文件管理] 用户 1474436119298048127 未满足群聊触发条件，已隐藏文件工具"
    )


@pytest.mark.asyncio
async def test_llm_request_policy_handles_events_without_get_extra():
    event = SimpleNamespace(
        message_obj=SimpleNamespace(type=MessageType.GROUP_MESSAGE),
        get_sender_id=lambda: "1474436119298048127",
        is_admin=lambda: False,
    )
    request = ProviderRequest(
        prompt="请读取 report.docx",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    policy = LLMRequestPolicy(
        document_toolset=ToolSet([_tool("read_file")]),
        require_at_in_group=True,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        is_bot_mentioned=lambda _event: False,
        notice_hooks=[],
        tool_exposure_hooks=[],
    )

    await policy.apply(event, request)

    assert "read_file" not in set(request.func_tool.names())
    assert "当前聊天不可使用文件/Office/PDF 相关功能" in request.system_prompt


@pytest.mark.asyncio
async def test_llm_request_policy_exposes_document_and_workbook_toolsets_together():
    event = _build_event(
        sender_id="1474436119298048127",
        message_type=MessageType.FRIEND_MESSAGE,
    )
    request = ProviderRequest(
        prompt="请生成 Excel 报表",
        system_prompt="base",
        func_tool=ToolSet([_tool("existing_tool")]),
    )
    policy = LLMRequestPolicy(
        document_toolset=ToolSet([_tool("create_document")]),
        workbook_toolset=ToolSet([_tool("create_workbook"), _tool("write_rows")]),
        require_at_in_group=True,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        is_bot_mentioned=lambda _event: True,
        notice_hooks=[],
        tool_exposure_hooks=[],
    )

    await policy.apply(event, request)

    tool_names = set(request.func_tool.names())
    assert "create_document" in tool_names
    assert "create_workbook" in tool_names
    assert "write_rows" in tool_names
    assert "existing_tool" in tool_names


def test_build_plugin_runtime_returns_temp_workspace_and_services():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["admin-1"]}
    config = {
        "file_settings": {
            "auto_delete_files": True,
            "max_file_size_mb": 8,
            "enable_docx_image_review": False,
            "max_inline_docx_image_mb": 3,
            "max_inline_docx_image_count": 4,
            "message_buffer_seconds": 4,
            "upload_session_ttl_seconds": 600,
        },
        "trigger_settings": {
            "reply_to_user": True,
            "require_at_in_group": True,
            "enable_features_in_group": False,
            "auto_block_execution_tools": True,
        },
        "preview_settings": {
            "enable": False,
            "dpi": 150,
        },
        "path_settings": {
            "allow_external_input_files": False,
        },
        "permission_settings": {
            "whitelist_users": ["user-1"],
        },
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        assert runtime.settings.auto_delete is True
        assert runtime.plugin_data_path.exists()
        assert runtime.temp_dir is not None
        assert runtime.settings.enable_docx_image_review is False
        assert runtime.settings.max_inline_docx_image_bytes == 3 * 1024 * 1024
        assert runtime.settings.max_inline_docx_image_count == 4
        assert runtime.settings.recent_text_ttl_seconds == 20
        assert runtime.settings.upload_session_ttl_seconds == 600
        assert runtime.settings.recent_text_cleanup_interval_seconds == 20
        assert runtime.settings.upload_session_cleanup_interval_seconds == 300
        assert runtime.settings.ppt_render_backend == "node"
        assert runtime.settings.excel_render_backend == "python"
        assert runtime.settings.js_renderer_entry == ""
        export_tool = next(
            tool
            for tool in runtime.document_toolset.tools
            if getattr(tool, "name", "") == "export_document"
        )
        assert [backend.name for backend in export_tool.render_backends] == ["node"]
        assert runtime.workspace_service.plugin_data_path == runtime.plugin_data_path
        assert runtime.post_export_hook_service is not None
        assert runtime.message_buffer is not None
        assert runtime.incoming_message_service is not None
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_ignores_legacy_word_render_settings():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["admin-1"]}
    config = {
        "render_settings": {
            "word_render_backend": "python",
            "word_render_fallback_enabled": False,
            "ppt_render_backend": "node",
            "excel_render_backend": "python",
            "js_renderer_entry": "D:/custom/js-renderer.js",
        }
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        export_tool = next(
            tool
            for tool in runtime.document_toolset.tools
            if getattr(tool, "name", "") == "export_document"
        )
        assert runtime.settings.ppt_render_backend == "node"
        assert runtime.settings.excel_render_backend == "python"
        assert runtime.settings.js_renderer_entry == "D:/custom/js-renderer.js"
        assert [backend.name for backend in export_tool.render_backends] == ["node"]
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_accepts_legacy_node_renderer_entry_alias():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["admin-1"]}
    config = {
        "render_settings": {
            "node_renderer_entry": "D:/custom/legacy-renderer.js",
        }
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        assert runtime.settings.js_renderer_entry == "D:/custom/legacy-renderer.js"
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_applies_word_style_defaults():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["admin-1"]}
    config = {
        "word_style_settings": {
            "default_font_name": "Arial",
            "default_heading_font_name": "Arial",
            "default_table_font_name": "Arial",
            "default_code_font_name": "JetBrains Mono",
        }
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        assert runtime.settings.default_word_font_name == "Arial"
        assert runtime.settings.default_word_heading_font_name == "Arial"
        assert runtime.settings.default_word_table_font_name == "Arial"
        assert runtime.settings.default_word_code_font_name == "JetBrains Mono"
        assert runtime.office_gen._default_document_style == {
            "font_name": "Arial",
            "heading_font_name": "Arial",
            "table_font_name": "Arial",
            "code_font_name": "JetBrains Mono",
        }
        assert get_document_style_defaults(runtime.document_toolset.document_store) == {
            "font_name": "Arial",
            "heading_font_name": "Arial",
            "table_font_name": "Arial",
            "code_font_name": "JetBrains Mono",
        }
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_keeps_zero_admin_id():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": 0}

    runtime = build_plugin_runtime(
        context=context,
        config={},
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id="0")
        event.is_admin.return_value = False

        assert runtime.access_policy_service.check_permission(event) is True
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


@pytest.mark.parametrize(
    ("admins_id", "sender_id", "expected"),
    [
        ([], "admin-1", False),
        ((), "admin-1", False),
        ("admin-1", "admin-1", True),
        (["admin-1", "admin-2"], "admin-2", True),
        (("admin-1", "admin-2"), "admin-2", True),
    ],
)
def test_build_plugin_runtime_handles_multiple_admin_id_shapes(
    admins_id,
    sender_id: str,
    expected: bool,
):
    context = MagicMock()
    context.get_config.return_value = {"admins_id": admins_id}

    runtime = build_plugin_runtime(
        context=context,
        config={},
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id=sender_id)
        event.is_admin.return_value = False

        assert runtime.access_policy_service.check_permission(event) is expected
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_handles_iterable_admin_ids_without_len():
    class _AdminIterable:
        def __iter__(self):
            yield "admin-iterable"

    context = MagicMock()
    context.get_config.return_value = {"admins_id": _AdminIterable()}

    runtime = build_plugin_runtime(
        context=context,
        config={},
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id="admin-iterable")
        event.is_admin.return_value = False

        assert runtime.access_policy_service.check_permission(event) is True
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_uses_persistent_workspace_when_auto_delete_disabled(
    monkeypatch: pytest.MonkeyPatch,
):
    data_root = _make_workspace("runtime-builder-data-root")
    called: dict[str, str | None] = {}
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["admin-2"]}
    config = {
        "file_settings": {
            "auto_delete_files": False,
            "max_file_size_mb": 16,
            "max_inline_docx_image_mb": 5,
            "max_inline_docx_image_count": 6,
            "message_buffer_seconds": 7,
            "recent_text_ttl_seconds": 45,
            "upload_session_ttl_seconds": 900,
        },
        "trigger_settings": {
            "reply_to_user": False,
            "require_at_in_group": False,
            "enable_features_in_group": True,
            "auto_block_execution_tools": False,
        },
        "preview_settings": {
            "enable": True,
            "dpi": 180,
        },
        "path_settings": {
            "allow_external_input_files": True,
        },
        "permission_settings": {
            "whitelist_users": ["user-2"],
        },
        "feature_settings": {
            "enable_office_files": True,
        },
    }

    def fake_get_data_dir(plugin_name=None):
        called["plugin_name"] = plugin_name
        return data_root

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.services.runtime_builder.StarTools.get_data_dir",
        fake_get_data_dir,
    )

    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        assert runtime.temp_dir is None
        assert runtime.plugin_data_path == data_root / "files"
        assert runtime.plugin_data_path.exists()
        assert runtime.settings.auto_delete is False
        assert runtime.settings.max_file_size == 16 * 1024 * 1024
        assert runtime.settings.max_inline_docx_image_bytes == 5 * 1024 * 1024
        assert runtime.settings.max_inline_docx_image_count == 6
        assert runtime.settings.reply_to_user is False
        assert runtime.settings.require_at_in_group is False
        assert runtime.settings.enable_features_in_group is True
        assert runtime.settings.auto_block_execution_tools is False
        assert runtime.settings.enable_preview is True
        assert runtime.settings.preview_dpi == 180
        assert runtime.settings.allow_external_input_files is True
        assert runtime.settings.recent_text_ttl_seconds == 45
        assert runtime.settings.upload_session_ttl_seconds == 900
        assert runtime.settings.recent_text_cleanup_interval_seconds == 45
        assert runtime.settings.upload_session_cleanup_interval_seconds == 300
        assert runtime.command_service._plugin_data_path == data_root / "files"
        assert runtime.workspace_service.plugin_data_path == data_root / "files"
        assert runtime.post_export_hook_service is not None
        assert called["plugin_name"] == "astrbot_plugin_office_assistant"
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        shutil.rmtree(data_root, ignore_errors=True)


def test_build_plugin_runtime_reads_admin_ids_from_context_get_config():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": ["1474436119298048127"]}
    config = {
        "file_settings": {
            "auto_delete_files": True,
        },
        "trigger_settings": {},
        "permission_settings": {},
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id="1474436119298048127")
        event.is_admin.return_value = False
        assert runtime.access_policy_service.check_permission(event) is True

        context.get_config.return_value = {"admins_id": ["new-admin-id"]}
        runtime.access_policy_service._get_admin_users.refresh()

        previous_event = _build_event(sender_id="1474436119298048127")
        previous_event.is_admin.return_value = False
        refreshed_event = _build_event(sender_id="new-admin-id")
        refreshed_event.is_admin.return_value = False
        assert runtime.access_policy_service.check_permission(previous_event) is False
        assert runtime.access_policy_service.check_permission(refreshed_event) is True
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_reads_admin_ids_from_legacy_astrbot_config():
    context = MagicMock()
    context.get_config.return_value = None
    context.astrbot_config = {"admins_id": ["admin-x"]}
    config = {
        "file_settings": {
            "auto_delete_files": True,
        },
        "trigger_settings": {},
        "permission_settings": {},
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id="admin-x")
        event.is_admin.return_value = False
        assert runtime.access_policy_service.check_permission(event) is True
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


def test_build_plugin_runtime_handles_scalar_admin_id_config():
    context = MagicMock()
    context.get_config.return_value = {"admins_id": 1474436119298048127}
    config = {
        "file_settings": {
            "auto_delete_files": True,
        },
        "trigger_settings": {},
        "permission_settings": {},
    }
    runtime = build_plugin_runtime(
        context=context,
        config=config,
        plugin_name="astrbot_plugin_office_assistant",
        handle_exported_document_tool=AsyncMock(),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
    )

    try:
        event = _build_event(sender_id="1474436119298048127")
        event.is_admin.return_value = False
        assert runtime.access_policy_service.check_permission(event) is True
    finally:
        runtime.executor.shutdown(wait=False)
        runtime.office_gen.cleanup()
        runtime.pdf_converter.cleanup()
        if runtime.temp_dir is not None:
            try:
                runtime.temp_dir.cleanup()
            except PermissionError:
                pass


@pytest.mark.asyncio
async def test_post_export_hook_service_handles_exported_document_tool():
    event = _build_event()
    event.send = AsyncMock()
    context = SimpleNamespace(context=SimpleNamespace(event=event))
    service = PostExportHookService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=False,
        reply_to_user=False,
        exported_message="✅ 文档已导出",
    )
    file_path = Path(__file__).resolve()

    try:
        result = await service.handle_exported_document_tool(
            context,
            str(file_path),
        )
    finally:
        service._executor.shutdown(wait=False)

    assert result == f"文档已导出并发送给用户：{file_path.name}"
    assert event.send.await_count == 2


@pytest.mark.asyncio
async def test_post_export_hook_service_returns_missing_message_without_sending():
    event = _build_event()
    event.send = AsyncMock()
    service = PostExportHookService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=False,
        reply_to_user=False,
        exported_message="✅ 文档已导出",
    )
    missing_path = Path(__file__).resolve().parent / "missing-export.docx"

    try:
        result = await service.send_exported_document(event, missing_path)
    finally:
        service._executor.shutdown(wait=False)

    assert "不存在" in result
    assert missing_path.name in result
    assert str(missing_path) not in result
    event.send.assert_not_awaited()


@pytest.mark.asyncio
async def test_post_export_hook_service_sends_preview_reply_and_deletes_files():
    workspace_dir = _make_workspace("post-export-preview")
    event = _build_event()
    event.send = AsyncMock()
    file_path = workspace_dir / "report.docx"
    preview_path = workspace_dir / "report-preview.png"
    file_path.write_text("docx", encoding="utf-8")
    _write_png(preview_path)
    preview_generator = MagicMock()
    preview_generator.generate_preview.return_value = preview_path
    service = PostExportHookService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=preview_generator,
        enable_preview=True,
        auto_delete=True,
        reply_to_user=True,
        exported_message="✅ 文档已导出",
    )

    try:
        result = await service.send_exported_document(event, file_path)
    finally:
        service._executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result == f"文档已导出并发送给用户：{file_path.name}"
    assert event.send.await_count == 3

    success_chain = event.send.await_args_list[0].args[0]
    assert "✅ 文档已导出" in success_chain.chain[0].text
    assert any(isinstance(component, Comp.At) for component in success_chain.chain)

    preview_chain = event.send.await_args_list[1].args[0]
    assert isinstance(preview_chain.chain[0], Comp.Image)
    assert preview_chain.chain[0].file == str(preview_path.resolve())

    file_chain = event.send.await_args_list[2].args[0]
    assert isinstance(file_chain.chain[0], Comp.File)
    assert file_chain.chain[0].name == file_path.name

    assert not preview_path.exists()
    assert not file_path.exists()


@pytest.mark.asyncio
async def test_post_export_hook_service_skips_preview_when_generation_fails():
    workspace_dir = _make_workspace("post-export-preview-fail")
    event = _build_event()
    event.send = AsyncMock()
    file_path = workspace_dir / "report.docx"
    file_path.write_text("docx", encoding="utf-8")
    preview_generator = MagicMock()
    preview_generator.generate_preview.side_effect = RuntimeError("preview boom")
    service = PostExportHookService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=preview_generator,
        enable_preview=True,
        auto_delete=True,
        reply_to_user=False,
        exported_message="✅ 文档已导出",
    )

    try:
        with patch(
            "astrbot_plugin_office_assistant.services.post_export_hook_service.logger.warning"
        ) as logger_warning:
            result = await service.send_exported_document(event, file_path)
    finally:
        service._executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result == f"文档已导出并发送给用户：{file_path.name}"
    assert event.send.await_count == 2
    logger_warning.assert_called()
    assert not file_path.exists()


@pytest.mark.asyncio
async def test_post_export_hook_service_logs_main_file_delete_failure_without_failing():
    workspace_dir = _make_workspace("post-export-delete-fail")
    event = _build_event()
    event.send = AsyncMock()
    file_path = workspace_dir / "report.docx"
    file_path.write_text("docx", encoding="utf-8")
    service = PostExportHookService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=True,
        reply_to_user=False,
        exported_message="✅ 文档已导出",
    )
    original_unlink = Path.unlink

    def _fake_unlink(path: Path, *args, **kwargs):
        if path == file_path:
            raise OSError("unlink boom")
        return original_unlink(path, *args, **kwargs)

    try:
        with patch.object(Path, "unlink", _fake_unlink):
            with patch(
                "astrbot_plugin_office_assistant.services.post_export_hook_service.logger.warning"
            ) as logger_warning:
                result = await service.send_exported_document(event, file_path)
    finally:
        service._executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result == f"文档已导出并发送给用户：{file_path.name}"
    assert event.send.await_count == 2
    logger_warning.assert_called_once()


@pytest.mark.asyncio
async def test_export_hook_service_aliases_post_export_hook_service():
    assert ExportHookService is PostExportHookService


@pytest.mark.asyncio
async def test_request_hook_service_builds_default_tool_hooks():
    request = ProviderRequest(
        prompt="请用 read_file 看一下 report.docx",
        system_prompt="base",
        func_tool=ToolSet(
            [
                _tool("read_file"),
                _tool("create_document"),
                _tool("astrbot_execute_shell"),
                _tool("astrbot_execute_python"),
            ]
        ),
    )
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )
    context = SimpleNamespace(
        event=_build_event(),
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name="read_file",
    )

    for hook in service.build_tool_exposure_hooks():
        context = await hook(context)

    tool_names = set(request.func_tool.names())
    assert "read_file" in tool_names
    assert "create_document" not in tool_names
    assert "astrbot_execute_shell" not in tool_names
    assert "astrbot_execute_python" not in tool_names


@pytest.mark.asyncio
async def test_request_hook_service_skips_explicit_restriction_for_unavailable_workbook_tool():
    request = ProviderRequest(
        prompt="请调用 create_workbook",
        system_prompt="base",
        func_tool=ToolSet(
            [
                _tool("read_file"),
                _tool("create_document"),
                _tool("astrbot_execute_shell"),
            ]
        ),
    )
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )
    context = SimpleNamespace(
        event=_build_event(),
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name="create_workbook",
    )

    for hook in service.build_tool_exposure_hooks():
        context = await hook(context)

    tool_names = set(request.func_tool.names())
    assert "read_file" in tool_names
    assert "create_document" in tool_names
    assert "astrbot_execute_shell" not in tool_names
    assert "create_workbook" not in tool_names


@pytest.mark.asyncio
async def test_request_hook_service_merges_multiple_uploaded_files_into_one_notice():
    request = ProviderRequest(
        prompt="根据上传文件整理内容",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event.message_obj.message = [
        Comp.File(name="report.docx", file="report.docx"),
        Comp.File(name="notes.txt", file="notes.txt"),
    ]
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [
            {
                "original_name": "report.docx",
                "file_suffix": ".docx",
                "type_desc": "Office文档 (Word/Excel/PPT)",
                "is_supported": True,
                "stored_name": "report_1.docx",
                "source_path": "/AstrBot/data/temp/report.docx",
            },
            {
                "original_name": "notes.txt",
                "file_suffix": ".txt",
                "type_desc": "文本/代码文件",
                "is_supported": True,
                "stored_name": "notes_1.txt",
                "source_path": "/AstrBot/data/temp/notes.txt",
            },
        ],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=True,
    )
    context = NoticeBuildContext(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
        system_notices=[],
        system_section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    assert len(context.notices) == 1
    notice = context.notices[0]
    assert notice.count("[System Notice] [ACTION REQUIRED] 已收到上传文件") == 1
    assert "文件数量：2" in notice
    assert "report.docx" in notice
    assert "notes.txt" in notice
    assert "report_1.docx" in notice
    assert "notes_1.txt" in notice
    assert "先调用 `read_file` 依次读取这些文件" in notice
    assert "不要猜文件名，不要列目录，不要调用 shell" in notice
    assert "读取前不要创建新文档" in notice
    assert "source_path" not in notice
    assert context.section_names == [SECTION_SCENE_UPLOADED_CONTEXT]
    assert context.system_notices == []
    assert context.system_section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_limits_multi_file_notice_details():
    request = ProviderRequest(
        prompt="根据上传文件整理内容",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event.message_obj.message = [
        Comp.File(name=f"file-{idx}.txt", file=f"file-{idx}.txt") for idx in range(5)
    ]
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: _build_upload_infos(5),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )
    context = SimpleNamespace(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    notice = context.notices[0]
    assert "文件数量：5" in notice
    assert "file-0.txt" in notice
    assert "file-2.txt" in notice
    assert "file_3.txt" in notice
    assert "file_4.txt" in notice
    assert "其余 2 个文件：" in notice
    assert "未展开详细信息" in notice


@pytest.mark.asyncio
async def test_request_hook_service_skips_scene_notice_for_file_only_buffered_prompt():
    request = ProviderRequest(
        prompt=(
            "\n[System Notice] 用户上传了 2 个文件\n\n"
            "[文件信息]\n"
            "- 原始文件名: file-0.txt\n"
            "  工作区文件名: file_0.txt\n"
        ),
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event._buffered = True
    event.message_obj.message = [
        Comp.File(name=f"file-{idx}.txt", file=f"file-{idx}.txt") for idx in range(2)
    ]
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: _build_upload_infos(2),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )
    context = SimpleNamespace(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    assert context.notices == []
    assert context.section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_skips_additional_notice_for_buffered_prompt_with_instruction():
    request = ProviderRequest(
        prompt=(
            "\n[System Notice] 用户上传了 1 个文件\n\n"
            "[文件信息]\n"
            "- 原始文件名: report.docx\n"
            "  工作区文件名: report_1.docx\n\n"
            "[用户指令]\n"
            "根据文件整理成正式汇报\n"
        ),
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event._buffered = True
    event.message_obj.message = [Comp.File(name="report.docx", file="report.docx")]
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [
            {
                "original_name": "report.docx",
                "file_suffix": ".docx",
                "type_desc": "Office文档 (Word/Excel/PPT)",
                "is_supported": True,
                "stored_name": "report_1.docx",
                "source_path": "",
            }
        ],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )
    context = SimpleNamespace(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    assert context.notices == []
    assert context.section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_omitted_files_use_names_without_external_paths():
    request = ProviderRequest(
        prompt="根据上传文件整理内容",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event.message_obj.message = [
        Comp.File(name=f"file-{idx}.txt", file=f"file-{idx}.txt") for idx in range(5)
    ]
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: _build_upload_infos(
            5,
            source_path_template="/tmp/file_{idx}.txt",
        ),
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=True,
    )
    context = SimpleNamespace(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    notice = context.notices[0]
    assert "file_3.txt" in notice
    assert "file_4.txt" in notice
    assert "/tmp/file_3.txt" not in notice
    assert "/tmp/file_4.txt" not in notice
    assert "其余 2 个文件：" in notice
    assert "未展开详细信息" in notice


@pytest.mark.asyncio
async def test_request_hook_service_keeps_omitted_count_aligned_with_mixed_path_items():
    request = ProviderRequest(
        prompt="根据上传文件整理内容",
        system_prompt="base",
        func_tool=ToolSet([_tool("read_file")]),
    )
    event = _build_event()
    event.message_obj.message = [
        Comp.File(name=f"file-{idx}.txt", file=f"file-{idx}.txt") for idx in range(5)
    ]
    upload_infos = _build_upload_infos(5)
    upload_infos[3]["source_path"] = "/tmp/file_3.txt"
    upload_infos[4]["source_path"] = ""
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: upload_infos,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=True,
    )
    context = SimpleNamespace(
        event=event,
        request=request,
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
        section_names=[],
    )

    context = await service.append_uploaded_file_notices(context)

    notice = context.notices[0]
    assert "其余 2 个文件：" in notice
    assert "file_4.txt" in notice
    assert "/tmp/file_3.txt" not in notice
    assert "其余 1 个文件：" not in notice


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_draft():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 3,
        },
    )
    context = NoticeBuildContext(
        event=_build_event(),
        request=ProviderRequest(
            prompt='继续完善 document_id="doc-1" 的内容',
            system_prompt="base",
            func_tool=ToolSet([_tool("create_document")]),
        ),
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
    )

    context = await service.append_document_tool_guide_notice(context)

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-1` 仍是 draft" in context.notices[0]
    assert "继续调用 `add_blocks`" in context.notices[0]
    assert context.system_notices == []
    assert context.system_section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_injects_workbook_follow_up_notice_for_draft():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=lambda workbook_id: {
            "workbook_id": workbook_id,
            "status": "draft",
            "sheet_names": ["总表", "华东"],
            "sheet_count": 2,
            "latest_written_sheets": ["华东"],
            "next_allowed_actions": ["write_rows", "export_workbook"],
        },
    )
    context = NoticeBuildContext(
        event=_build_event(),
        request=ProviderRequest(
            prompt='继续补充 workbook_id="wb-1" 的数据',
            system_prompt="base",
            func_tool=ToolSet([_tool("create_workbook")]),
        ),
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
    )

    context = await service.append_document_tool_guide_notice(context)

    assert context.section_names == [SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP]
    assert "当前 `workbook_id=wb-1` 仍是 draft" in context.notices[0]
    assert "继续调用 `write_rows`" in context.notices[0]
    assert context.system_notices == []
    assert context.system_section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_filters_none_values_from_workbook_follow_up_notice():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=lambda workbook_id: {
            "workbook_id": workbook_id,
            "status": "draft",
            "sheet_names": [None, "总表", "   "],
            "sheet_count": 1,
            "latest_written_sheets": [None, "华东"],
            "next_allowed_actions": [None, "write_rows", " "],
        },
    )
    context = NoticeBuildContext(
        event=_build_event(),
        request=ProviderRequest(
            prompt='继续补充 workbook_id="wb-2" 的数据',
            system_prompt="base",
            func_tool=ToolSet([_tool("write_rows")]),
        ),
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
    )

    context = await service.append_document_tool_guide_notice(context)

    assert "None" not in context.notices[0]
    assert "总表" in context.notices[0]
    assert "华东" in context.notices[0]
    assert "write_rows" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_workbook_follow_up_notice_for_missing_summary():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=lambda _workbook_id: None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='继续处理 workbook_id="wb-missing"',
                system_prompt="base",
                func_tool=ToolSet([_tool("write_rows")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP]
    assert "没有找到 `workbook_id=wb-missing` 对应的工作簿会话" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_prioritizes_workbook_follow_up_when_workbook_id_present():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 1,
        },
        lookup_workbook_summary=lambda workbook_id: {
            "workbook_id": workbook_id,
            "status": "draft",
            "sheet_names": ["Sheet1"],
            "sheet_count": 1,
            "latest_written_sheets": ["Sheet1"],
            "next_allowed_actions": ["write_rows"],
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='document_id="doc-7"，同时继续 workbook_id="wb-7"',
                system_prompt="base",
                func_tool=ToolSet([_tool("write_rows")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP]
    assert "workbook_id=wb-7" in context.notices[0]
    assert "document_id=doc-7" not in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_space_separated_document_id():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 6,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="继续完善 document_id doc-3 的内容",
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-3` 仍是 draft" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_english_is_document_id():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 4,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="document_id is doc-7, please continue",
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-7` 仍是 draft" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_backtick_quoted_document_id():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 5,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="继续处理 document_id=`doc-8`",
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-8` 仍是 draft" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_fullwidth_colon_document_id():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 6,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="继续处理 document_id：doc-9",
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-9` 仍是 draft" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_compact_chinese_separator():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "draft",
            "block_count": 7,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="继续处理 document_id为doc-10",
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "当前 `document_id=doc-10` 仍是 draft" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_follow_up_notice_for_finalized():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda document_id: {
            "document_id": document_id,
            "status": "finalized",
            "block_count": 15,
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='请导出 document_id="doc-2"',
                system_prompt="base",
                func_tool=ToolSet([_tool("export_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "下一步只能调用 `export_document`" in context.notices[0]
    assert "不要再调用 `add_blocks`、`finalize_document` 或 `create_document`" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_injects_missing_notice_when_document_id_is_unknown():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda _document_id: None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='继续处理 document_id="missing-doc"',
                system_prompt="base",
                func_tool=ToolSet([_tool("add_blocks")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert context.section_names == [SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP]
    assert "没有找到 `document_id=missing-doc` 对应的文档会话" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_falls_back_to_document_guide_when_document_id_is_not_parseable():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda _document_id: None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请返回 document_id 并导出成 Word 报告",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert "create_document" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_falls_back_to_document_guide_for_english_non_id_token():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda _document_id: None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="please return document_id and export this as Word",
                system_prompt="base",
                func_tool=ToolSet([_tool("export_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert "export_document" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_falls_back_to_document_guide_for_non_id_bare_token():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda _document_id: None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请根据 document_id token-1 导出成 Word",
                system_prompt="base",
                func_tool=ToolSet([_tool("export_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert "export_document" in context.notices[0]


@pytest.mark.asyncio
async def test_request_hook_service_falls_back_to_document_guide_when_summary_lookup_is_disabled():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='继续导出 document_id="doc-11" 的 Word 报告',
                system_prompt="base",
                func_tool=ToolSet([_tool("export_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert "export_document" in context.notices[0]
    assert "没有找到 `document_id=" not in context.notices[0]


def test_runtime_builder_returns_none_when_document_summary_lookup_is_unavailable():
    document_toolset = SimpleNamespace(document_store=SimpleNamespace())

    with patch(
        "astrbot_plugin_office_assistant.services.runtime_builder.logger.warning"
    ) as mock_warning:
        lookup = _build_document_summary_lookup(document_toolset)

    assert lookup is None
    mock_warning.assert_called_once()


def test_runtime_builder_uses_document_summary_lookup_when_available():
    build_prompt_summary = MagicMock(return_value={"status": "draft"})
    document_toolset = SimpleNamespace(
        document_store=SimpleNamespace(build_prompt_summary=build_prompt_summary)
    )

    lookup = _build_document_summary_lookup(document_toolset)

    assert lookup is build_prompt_summary
    assert lookup("doc-1") == {"status": "draft"}


def test_runtime_builder_returns_none_when_workbook_summary_lookup_is_unavailable():
    workbook_toolset = SimpleNamespace(workbook_store=SimpleNamespace())

    with patch(
        "astrbot_plugin_office_assistant.services.runtime_builder.logger.warning"
    ) as mock_warning:
        lookup = _build_workbook_summary_lookup(workbook_toolset)

    assert lookup is None
    mock_warning.assert_called_once()


def test_runtime_builder_uses_workbook_summary_lookup_when_available():
    build_prompt_summary = MagicMock(return_value={"status": "draft"})
    workbook_toolset = SimpleNamespace(
        workbook_store=SimpleNamespace(build_prompt_summary=build_prompt_summary)
    )

    lookup = _build_workbook_summary_lookup(workbook_toolset)

    assert lookup is build_prompt_summary
    assert lookup("wb-1") == {"status": "draft"}


def test_runtime_builder_skips_workbook_toolset_when_builder_raises_import_error(
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.services.runtime_builder.build_workbook_toolset",
        MagicMock(side_effect=ModuleNotFoundError("missing workbook dependency")),
    )

    with patch(
        "astrbot_plugin_office_assistant.services.runtime_builder.logger.warning"
    ) as mock_warning:
        toolset = _build_workbook_toolset(
            workspace_dir=Path.cwd(),
            after_export=AsyncMock(),
        )

    assert toolset is None
    mock_warning.assert_called_once()
    assert "workbook 原语工具依赖不可用" in mock_warning.call_args.args[0]


@pytest.mark.asyncio
async def test_request_hook_service_reraises_unexpected_summary_lookup_errors():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_document_summary=lambda _document_id: (_ for _ in ()).throw(
            RuntimeError("boom")
        ),
    )

    context = NoticeBuildContext(
        event=_build_event(),
        request=ProviderRequest(
            prompt='继续处理 document_id="doc-9"',
            system_prompt="base",
            func_tool=ToolSet([_tool("add_blocks")]),
        ),
        should_expose=True,
        can_process_upload=True,
        explicit_tool_name=None,
        notices=[],
    )

    with patch(
        "astrbot_plugin_office_assistant.services.request_hook_service.logger.exception"
    ) as mock_exception:
        with pytest.raises(RuntimeError, match="boom"):
            await service.append_document_tool_guide_notice(context)

    mock_exception.assert_called_once()


def test_prompt_context_service_orders_dynamic_document_notice_after_scene_notice():
    service = PromptContextService(allow_external_input_files=False)

    ordered_names, ordered_notices = service.order_notice_sections(
        section_names=[
            SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP,
            SECTION_SCENE_UPLOADED_CONTEXT,
            SECTION_STATIC_DOCUMENT_TOOLS,
        ],
        notices=[
            "dynamic",
            "scene",
            "static",
        ],
    )

    assert ordered_names == [
        SECTION_STATIC_DOCUMENT_TOOLS,
        SECTION_SCENE_UPLOADED_CONTEXT,
        SECTION_DYNAMIC_DOCUMENT_FOLLOW_UP,
    ]
    assert ordered_notices == ["static", "scene", "dynamic"]


def test_prompt_context_service_orders_dynamic_workbook_notice_after_scene_notice():
    service = PromptContextService(allow_external_input_files=False)

    ordered_names, ordered_notices = service.order_notice_sections(
        section_names=[
            SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP,
            SECTION_SCENE_UPLOADED_CONTEXT,
            SECTION_STATIC_WORKBOOK_TOOLS,
        ],
        notices=[
            "dynamic",
            "scene",
            "static",
        ],
    )

    assert ordered_names == [
        SECTION_STATIC_WORKBOOK_TOOLS,
        SECTION_SCENE_UPLOADED_CONTEXT,
        SECTION_DYNAMIC_WORKBOOK_FOLLOW_UP,
    ]
    assert ordered_notices == ["static", "scene", "dynamic"]


@pytest.mark.asyncio
async def test_request_hook_service_injects_document_core_notice_only_once_per_session():
    consume_once = _build_notice_once_callback()
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=consume_once,
        allow_external_input_files=False,
    )
    event = _build_event()

    first = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=event,
            request=ProviderRequest(
                prompt="请生成一份 Word 报告，并导出给我",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )
    second = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=event,
            request=ProviderRequest(
                prompt="再生成一份 Word 报告",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert first.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert second.section_names == []


@pytest.mark.asyncio
async def test_request_hook_service_injects_detail_notice_later_without_repeating_core():
    consume_once = _build_notice_once_callback()
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=consume_once,
        allow_external_input_files=False,
    )
    event = _build_event()

    first = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=event,
            request=ProviderRequest(
                prompt="请生成一份 Word 报告",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )
    second = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=event,
            request=ProviderRequest(
                prompt="请用 executive_brief 主题和 document_style 生成 Word 报告",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert first.section_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert second.section_names == [SECTION_STATIC_DOCUMENT_TOOLS_DETAIL]


@pytest.mark.asyncio
async def test_request_hook_service_allows_detail_notice_without_core_when_only_style_signal_exists():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请使用 executive_brief 和 accent_color=112233",
                system_prompt="base",
                func_tool=ToolSet([_tool("create_document")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert context.section_names == [SECTION_STATIC_DOCUMENT_TOOLS_DETAIL]


@pytest.mark.asyncio
async def test_request_hook_service_injects_workbook_core_and_detail_notices():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请用 create_workbook 和 write_rows 生成 xlsx 报表，多 sheet，start_row=2",
                system_prompt="base",
                func_tool=ToolSet(
                    [
                        _tool("create_workbook"),
                        _tool("write_rows"),
                        _tool("export_workbook"),
                    ]
                ),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_WORKBOOK_TOOLS,
        SECTION_STATIC_WORKBOOK_TOOLS_DETAIL,
    ]
    assert "read_workbook" in context.notices[0]
    assert "execute_excel_script" in context.notices[0]
    assert "create_workbook" in context.notices[1]
    assert "write_rows" in context.notices[1]
    assert "execute_excel_script" in context.notices[2]


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_guide_when_workbook_tools_are_unavailable():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请用 create_workbook 和 write_rows 生成 xlsx 报表，多 sheet，start_row=2",
                system_prompt="base",
                func_tool=ToolSet([_tool("existing_tool")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_guide_for_reading_xlsx_requests():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请读取这个 xlsx 文件并告诉我内容",
                system_prompt="base",
                func_tool=ToolSet(
                    [
                        _tool("create_workbook"),
                        _tool("write_rows"),
                        _tool("export_workbook"),
                    ]
                ),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


@pytest.mark.asyncio
async def test_request_hook_service_injects_excel_read_notice_for_reading_xlsx_requests():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [
            {
                "original_name": "sales.xlsx",
                "file_suffix": ".xlsx",
                "stored_name": "sales_1.xlsx",
                "source_path": "",
                "is_supported": True,
            }
        ],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请读取这个 xlsx 文件并告诉我内容",
                system_prompt="base",
                func_tool=ToolSet([_tool("read_workbook")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_EXCEL_READ,
    ]
    assert "read_workbook" in context.notices[0]
    assert "read_workbook" in context.notices[1]
    assert SECTION_STATIC_WORKBOOK_TOOLS not in context.section_names


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_guide_for_xlsx_to_pdf_conversion_requests():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请把这个 xlsx 导出成 PDF",
                system_prompt="base",
                func_tool=ToolSet(
                    [
                        _tool("create_workbook"),
                        _tool("write_rows"),
                        _tool("export_workbook"),
                    ]
                ),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


@pytest.mark.asyncio
async def test_request_hook_service_injects_excel_script_notice_for_complex_excel_generation():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请生成一个带公式和条件格式的 Excel 报表",
                system_prompt="base",
                func_tool=ToolSet([_tool("execute_excel_script")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_EXCEL_SCRIPT,
    ]
    assert "execute_excel_script" in context.notices[0]
    assert "execute_excel_script" in context.notices[1]
    assert "create_workbook" not in context.notices[1]


@pytest.mark.asyncio
async def test_request_hook_service_injects_excel_script_notice_for_existing_xlsx_modification():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [
            {
                "original_name": "sales.xlsx",
                "file_suffix": ".xlsx",
                "stored_name": "sales_1.xlsx",
                "source_path": "",
                "is_supported": True,
            }
        ],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请给这个 xlsx 增加一列公式并导出新版本",
                system_prompt="base",
                func_tool=ToolSet([_tool("execute_excel_script")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_EXCEL_SCRIPT,
    ]
    assert "read_workbook" in context.notices[0]
    assert "execute_excel_script" in context.notices[1]


@pytest.mark.asyncio
async def test_request_hook_service_injects_excel_script_notice_for_mixed_read_modify_xlsx_requests():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [
            {
                "original_name": "sales.xlsx",
                "file_suffix": ".xlsx",
                "stored_name": "sales_1.xlsx",
                "source_path": "",
                "is_supported": True,
            }
        ],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请读取这个 xlsx 后新增一列公式并导出新版本",
                system_prompt="base",
                func_tool=ToolSet([_tool("execute_excel_script")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_EXCEL_SCRIPT,
    ]
    assert SECTION_STATIC_EXCEL_READ not in context.section_names
    assert "execute_excel_script" in context.notices[1]


def test_excel_intent_router_keeps_export_worded_existing_filename():
    decision = ExcelIntentRouter.decide(
        request_text="请导出 report.xlsx 的内容并总结",
        upload_infos=[],
        explicit_tool_name=None,
        exposed_tool_names={"read_workbook"},
    )

    assert decision is not None
    assert decision.route == "read_existing"
    assert decision.matched_files == ("report.xlsx",)


def test_excel_intent_router_ignores_output_filename_when_upload_present():
    decision = ExcelIntentRouter.decide(
        request_text="请读取 input.xlsx 并生成 result.xlsx 作为输出文件",
        upload_infos=[
            {
                "original_name": "input.xlsx",
                "file_suffix": ".xlsx",
                "stored_name": "input.xlsx",
                "source_path": "",
                "is_supported": True,
            }
        ],
        explicit_tool_name=None,
        exposed_tool_names={"read_workbook"},
    )

    assert decision is not None
    assert decision.route == "read_existing"
    assert decision.matched_files == ("input.xlsx",)


@pytest.mark.asyncio
async def test_request_hook_service_treats_output_xlsx_name_as_new_workbook_request():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt="请生成 sales.xlsx 并汇总 Q1 销售数据",
                system_prompt="base",
                func_tool=ToolSet(
                    [
                        _tool("read_workbook"),
                        _tool("create_workbook"),
                        _tool("write_rows"),
                        _tool("export_workbook"),
                    ]
                ),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_WORKBOOK_TOOLS,
    ]
    assert SECTION_STATIC_EXCEL_READ not in context.section_names
    assert "create_workbook" in context.notices[1]


@pytest.mark.asyncio
async def test_request_hook_service_falls_back_to_workbook_guide_when_workbook_lookup_disabled():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='请返回 workbook_id 并导出 workbook_id="wb-11" 的 xlsx',
                system_prompt="base",
                func_tool=ToolSet(
                    [
                        _tool("create_workbook"),
                        _tool("write_rows"),
                        _tool("export_workbook"),
                    ]
                ),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == [
        SECTION_STATIC_EXCEL_ROUTING,
        SECTION_STATIC_WORKBOOK_TOOLS,
    ]
    assert "read_workbook" in context.notices[0]
    assert "create_workbook" in context.notices[1]
    assert "export_workbook" in context.notices[1]


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_guide_for_partial_workbook_toolset():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=None,
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='请返回 workbook_id 并导出 workbook_id="wb-11" 的 xlsx',
                system_prompt="base",
                func_tool=ToolSet([_tool("export_workbook")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_follow_up_when_workbook_tools_are_unavailable():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=lambda workbook_id: {
            "workbook_id": workbook_id,
            "status": "draft",
            "sheet_names": ["Sheet1"],
            "sheet_count": 1,
            "latest_written_sheets": ["Sheet1"],
            "next_allowed_actions": ["write_rows"],
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='继续补充 workbook_id="wb-11" 的数据',
                system_prompt="base",
                func_tool=ToolSet([_tool("existing_tool")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name=None,
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


@pytest.mark.asyncio
async def test_request_hook_service_skips_workbook_follow_up_for_partial_workbook_toolset():
    service = RequestHookService(
        auto_block_execution_tools=True,
        get_cached_upload_infos=lambda _event: [],
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        consume_session_notice_once=_build_notice_once_callback(),
        allow_external_input_files=False,
        lookup_workbook_summary=lambda workbook_id: {
            "workbook_id": workbook_id,
            "status": "draft",
            "sheet_names": ["Sheet1"],
            "sheet_count": 1,
            "latest_written_sheets": ["Sheet1"],
            "next_allowed_actions": ["write_rows"],
        },
    )

    context = await service.append_document_tool_guide_notice(
        NoticeBuildContext(
            event=_build_event(),
            request=ProviderRequest(
                prompt='继续补充 workbook_id="wb-11" 的数据',
                system_prompt="base",
                func_tool=ToolSet([_tool("export_workbook")]),
            ),
            should_expose=True,
            can_process_upload=True,
            explicit_tool_name="export_workbook",
            notices=[],
        )
    )

    assert context.section_names == []
    assert context.notices == []


def test_prompt_context_service_orders_notice_sections_by_stability():
    service = PromptContextService(allow_external_input_files=False)

    ordered_names, ordered_notices = service.order_notice_sections(
        section_names=[
            SECTION_SCENE_UPLOADED_CONTEXT,
            SECTION_STATIC_DOCUMENT_TOOLS,
        ],
        notices=[
            "scene",
            "static",
        ],
    )

    assert ordered_names == [
        SECTION_STATIC_DOCUMENT_TOOLS,
        SECTION_SCENE_UPLOADED_CONTEXT,
    ]
    assert ordered_notices == ["static", "scene"]
    trace = service.build_section_trace(
        section_names=ordered_names,
        notices=ordered_notices,
    )
    assert trace.startswith(
        f"{SECTION_STATIC_DOCUMENT_TOOLS}, {SECTION_SCENE_UPLOADED_CONTEXT}"
    )
    assert (
        f"[len={SECTION_STATIC_DOCUMENT_TOOLS}:6, "
        f"{SECTION_SCENE_UPLOADED_CONTEXT}:5]" in trace
    )
    assert "[groups=static:6, scene:5]" in trace
    assert "[total=11]" in trace


def test_prompt_context_service_logs_section_length_mismatch():
    service = PromptContextService(allow_external_input_files=False)

    with patch(
        "astrbot_plugin_office_assistant.services.prompt_context_service.logger.debug"
    ) as logger_debug:
        ordered_names, ordered_notices = service.order_notice_sections(
            section_names=[SECTION_STATIC_DOCUMENT_TOOLS],
            notices=["static", "dynamic"],
        )

    assert ordered_names == [SECTION_STATIC_DOCUMENT_TOOLS]
    assert ordered_notices == ["static", "dynamic"]
    logger_debug.assert_called_once_with(
        "[文件管理] Prompt section mismatch: sections=%s notices=%s",
        1,
        2,
    )


def test_prompt_context_service_build_section_trace_tolerates_length_mismatch():
    service = PromptContextService(allow_external_input_files=False)

    trace = service.build_section_trace(
        section_names=[
            SECTION_STATIC_DOCUMENT_TOOLS,
            SECTION_SCENE_UPLOADED_CONTEXT,
        ],
        notices=["static only"],
    )

    assert trace.startswith(
        f"{SECTION_STATIC_DOCUMENT_TOOLS}, {SECTION_SCENE_UPLOADED_CONTEXT}"
    )
    assert f"[len={SECTION_STATIC_DOCUMENT_TOOLS}:11]" in trace
    assert "[groups=static:11]" in trace
    assert "[total=11]" in trace


def test_upload_prompt_service_builds_instructional_notice_for_readable_files():
    service = UploadPromptService(allow_external_input_files=True)

    prompt_text = service.build_prompt(
        upload_infos=[
            {
                "original_name": "report.docx",
                "file_suffix": ".docx",
                "stored_name": "report_1.docx",
                "source_path": "/AstrBot/data/temp/report.docx",
                "is_supported": True,
            }
        ],
        user_instruction="看看里面的内容",
    )

    assert "[用户指令]" in prompt_text
    assert "看看里面的内容" in prompt_text
    assert "工作区文件名: report_1.docx" in prompt_text
    assert "外部绝对路径" not in prompt_text
    assert "先调用 `read_file` 读取文件" in prompt_text
    assert "读取后按用户指令继续调用工具，不要只回复过渡说明" in prompt_text


def test_upload_prompt_service_prefers_read_workbook_for_excel_uploads():
    service = UploadPromptService(allow_external_input_files=True)

    prompt_text = service.build_prompt(
        upload_infos=[
            {
                "original_name": "sales.xlsx",
                "file_suffix": ".xlsx",
                "stored_name": "sales_1.xlsx",
                "source_path": "/AstrBot/data/temp/sales.xlsx",
                "is_supported": True,
            }
        ],
        user_instruction="读取内容并汇总华东数据",
    )

    assert "先调用 `read_workbook` 读取文件" in prompt_text
    assert "先调用 `read_file` 读取文件" not in prompt_text


def test_upload_prompt_service_builds_notice_for_readable_files_without_instruction():
    service = UploadPromptService(allow_external_input_files=False)

    prompt_text = service.build_prompt(
        upload_infos=[
            {
                "original_name": "report.docx",
                "file_suffix": ".docx",
                "stored_name": "report_1.docx",
                "source_path": "/AstrBot/data/temp/report.docx",
                "is_supported": True,
            }
        ],
        user_instruction="",
    )

    assert "用户上传了可读取文件，后续应优先围绕这些文件处理" in prompt_text
    assert "[用户指令]" not in prompt_text
    assert "工作区文件名: report_1.docx" in prompt_text
    assert "外部绝对路径" not in prompt_text


def test_upload_prompt_service_handles_empty_upload_infos():
    service = UploadPromptService(allow_external_input_files=True)

    prompt_text = service.build_prompt(
        upload_infos=[],
        user_instruction="随便看看",
    )

    assert isinstance(prompt_text, str)
    assert "[文件信息]" in prompt_text
    assert "[操作要求]" in prompt_text
    assert "[用户指令]" not in prompt_text
    assert "工作区文件名" not in prompt_text
    assert "外部绝对路径" not in prompt_text


def test_upload_prompt_service_handles_mixed_readable_and_unreadable_files():
    service = UploadPromptService(allow_external_input_files=True)

    prompt_text = service.build_prompt(
        upload_infos=[
            {
                "original_name": "readable.txt",
                "file_suffix": ".txt",
                "stored_name": "readable_1.txt",
                "source_path": "/AstrBot/data/temp/readable.txt",
                "is_supported": True,
            },
            {
                "original_name": "unreadable.bin",
                "file_suffix": ".bin",
                "stored_name": "unreadable_1.bin",
                "source_path": "/AstrBot/data/temp/unreadable.bin",
                "is_supported": False,
            },
        ],
        user_instruction="阅读可用文件",
    )

    assert "[文件信息]" in prompt_text
    assert "工作区文件名: readable_1.txt" in prompt_text
    assert "工作区文件名: unreadable_1.bin" in prompt_text
    assert "外部绝对路径" not in prompt_text
    assert "先调用 `read_file` 读取文件" in prompt_text


def test_upload_prompt_service_limits_file_details_for_many_uploads():
    service = UploadPromptService(allow_external_input_files=False)

    prompt_text = service.build_prompt(
        upload_infos=_build_upload_infos(5),
        user_instruction="整理成报告",
    )

    assert "用户上传了 5 个文件" in prompt_text
    assert "file-0.txt" in prompt_text
    assert "file-2.txt" in prompt_text
    assert "file_3.txt" in prompt_text
    assert "file_4.txt" in prompt_text
    assert "其余 2 个文件：" in prompt_text
    assert "未展开详细信息" in prompt_text


def test_upload_prompt_service_omitted_files_only_show_names():
    service = UploadPromptService(allow_external_input_files=True)

    prompt_text = service.build_prompt(
        upload_infos=_build_upload_infos(
            5,
            source_path_template="/tmp/file_{idx}.txt",
        ),
        user_instruction="整理成报告",
    )

    assert "file_3.txt" in prompt_text
    assert "file_4.txt" in prompt_text
    assert "/tmp/file_3.txt" not in prompt_text
    assert "/tmp/file_4.txt" not in prompt_text
    assert "其余 2 个文件：" in prompt_text
    assert "未展开详细信息" in prompt_text


def test_upload_prompt_service_keeps_omitted_count_aligned_with_mixed_path_items():
    service = UploadPromptService(allow_external_input_files=True)
    upload_infos = _build_upload_infos(5)
    upload_infos[3]["source_path"] = "/tmp/file_3.txt"
    upload_infos[4]["source_path"] = ""

    prompt_text = service.build_prompt(
        upload_infos=upload_infos,
        user_instruction="整理成报告",
    )

    assert "其余 2 个文件：" in prompt_text
    assert "file_4.txt" in prompt_text
    assert "/tmp/file_3.txt" not in prompt_text
    assert "其余 1 个文件：" not in prompt_text


def test_upload_prompt_service_builds_generic_notice_for_unreadable_files():
    service = UploadPromptService(allow_external_input_files=False)

    prompt_text = service.build_prompt(
        upload_infos=[
            {
                "original_name": "archive.bin",
                "file_suffix": ".bin",
                "stored_name": "",
                "source_path": "",
                "is_supported": False,
            }
        ],
        user_instruction="",
    )

    assert "[操作要求]" in prompt_text
    assert "请根据用户要求处理这些文件，使用中文与用户沟通。" in prompt_text
    assert "[用户指令]" not in prompt_text
    assert "工作区文件名" not in prompt_text


@pytest.mark.asyncio
async def test_word_read_service_returns_disabled_image_notice_for_image_only_docx():
    workspace_dir = _make_workspace("word-read-service-image-only")
    executor = ThreadPoolExecutor(max_workers=1)
    docx_path = workspace_dir / "image-only.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        docx = _import_docx()

        _write_png(image_path)
        document = docx.Document()
        document.add_picture(str(image_path), width=docx.shared.Inches(1))
        document.save(docx_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = WordReadService(
            workspace_service=workspace_service,
            enable_docx_image_review=False,
        )

        results = [
            result
            async for result in service.iter_word_results(
                docx_path,
                docx_path.name,
                docx_path.suffix.lower(),
                docx_path.stat().st_size,
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert len(results) == 1
    assert isinstance(results[0], str)
    assert "该 Word 文档仅包含图片内容，当前未启用图片理解。" in results[0]


def test_workspace_service_pre_check_rejects_outside_workspace():
    workspace_dir = Path(__file__).resolve().parent / "workspace-root"
    outside_file = Path(__file__).resolve()

    executor = ThreadPoolExecutor(max_workers=1)
    try:
        service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024,
            feature_settings={},
        )

        ok, resolved_path, err = service.pre_check(
            _build_event(),
            str(outside_file),
            require_exists=True,
            is_group_feature_enabled=lambda _event: True,
            check_permission_fn=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )

        assert not ok
        assert resolved_path is None
        assert err == "错误：非法路径：禁止访问工作区外的文件"
    finally:
        executor.shutdown(wait=False)


def test_upload_session_service_preserves_input_upload_infos_when_caching():
    service = UploadSessionService(
        context=MagicMock(),
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        allow_external_input_files=False,
    )
    event = _build_event()
    upload_infos = [
        {
            "original_name": "report.docx",
            "file_suffix": ".docx",
            "type_desc": "Office文档 (Word/Excel/PPT)",
            "is_supported": True,
            "stored_name": "report_1.docx",
            "source_path": "D:/tmp/report.docx",
        }
    ]

    assigned_infos = service._cache_session_upload_infos(event, upload_infos)

    assert "file_id" not in upload_infos[0]
    assert assigned_infos[0]["file_id"] == "f1"
    assert service.list_session_upload_infos(event)[0]["file_id"] == "f1"


def test_upload_session_service_consumes_notice_once_per_session():
    service = UploadSessionService(
        context=MagicMock(),
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        allow_external_input_files=False,
    )
    event = _build_event()

    assert service.consume_session_notice_once(event, "document_core_guide") is True
    assert service.consume_session_notice_once(event, "document_core_guide") is False
    assert service.consume_session_notice_once(event, "document_detail_guide") is True
    assert (
        service.consume_session_notice_once(
            _build_event(sender_id="user-2"),
            "document_core_guide",
        )
        is True
    )


@pytest.mark.asyncio
async def test_upload_session_service_requeues_file_only_buffered_upload_in_friend_chat():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=True,
    )
    event = _build_event()
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=event, files=[upload], texts=[])

    await service.on_buffer_complete(buf)

    queued_event = event_queue.put.await_args.args[0]
    prompt_text = queued_event.message_obj.message[0].text
    upload_infos = service.list_session_upload_infos(event)
    assert len(upload_infos) == 1
    assert upload_infos[0]["original_name"] == "report.docx"
    assert upload_infos[0]["stored_name"] == "report_1.docx"
    assert upload_infos[0]["file_id"] == "f1"
    assert "[用户指令]" not in prompt_text
    assert "用户意图尚不明确时，再用中文询问用户想要如何处理" in prompt_text


@pytest.mark.asyncio
async def test_upload_session_service_caches_file_only_buffered_upload_in_group_chat():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=True,
    )
    event = _build_event(message_type=MessageType.GROUP_MESSAGE)
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=event, files=[upload], texts=[])

    await service.on_buffer_complete(buf)

    event_queue.put.assert_not_awaited()
    upload_infos = service.list_session_upload_infos(event)
    assert len(upload_infos) == 1
    assert upload_infos[0]["original_name"] == "report.docx"
    assert upload_infos[0]["stored_name"] == "report_1.docx"
    assert upload_infos[0]["file_id"] == "f1"


@pytest.mark.asyncio
async def test_upload_session_service_uses_recent_text_for_file_only_buffered_upload():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=True,
    )
    text_event = _build_event()
    text_event.message_obj.message = [Comp.Plain("根据这份文件整理成正式汇报")]
    file_event = _build_event()
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=file_event, files=[upload], texts=[])

    service.remember_recent_text(text_event)
    await service.on_buffer_complete(buf)

    queued_event = event_queue.put.await_args.args[0]
    prompt_text = queued_event.message_obj.message[0].text
    assert "[用户指令]" in prompt_text
    assert "根据这份文件整理成正式汇报" in prompt_text
    assert "工作区文件名: report_1.docx" in prompt_text


@pytest.mark.asyncio
async def test_upload_session_service_ignores_command_like_recent_text():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=True,
    )
    text_event = _build_event()
    text_event.message_obj.message = [Comp.Plain("/doc list")]
    file_event = _build_event(message_type=MessageType.GROUP_MESSAGE)
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=file_event, files=[upload], texts=[])

    service.remember_recent_text(text_event)
    await service.on_buffer_complete(buf)

    event_queue.put.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_session_service_preserves_raw_message_for_file_only_cache():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=False,
    )
    event = _build_event(message_type=MessageType.GROUP_MESSAGE)
    raw_message = SimpleNamespace(mentions=[SimpleNamespace(id="bot-1")])
    event.message_obj.raw_message = raw_message
    event.is_mentioned.side_effect = lambda: hasattr(
        event.message_obj.raw_message, "mentions"
    ) and any(
        str(mention.id) == str(event.message_obj.self_id)
        for mention in event.message_obj.raw_message.mentions
    )
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=event, files=[upload], texts=[])

    await service.on_buffer_complete(buf)

    assert event.message_obj.raw_message is raw_message
    assert event.is_mentioned() is True
    assert event.message_obj.raw_message is raw_message
    assert event.is_mentioned() is True
    event_queue.put.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_session_service_omits_external_path_when_disabled():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=False,
    )
    event = _build_event()
    upload = Comp.File(name="report.docx", file="report.docx")
    buf = BufferedMessage(event=event, files=[upload], texts=[])

    await service.on_buffer_complete(buf)

    queued_event = event_queue.put.await_args.args[0]
    prompt_text = queued_event.message_obj.message[0].text
    upload_infos = service.list_session_upload_infos(event)
    assert len(upload_infos) == 1
    assert upload_infos[0]["stored_name"] == "report_1.docx"
    assert upload_infos[0]["source_path"]
    assert "外部绝对路径:" not in prompt_text


@pytest.mark.asyncio
async def test_upload_session_service_uses_extracted_filename_for_type_detection():
    context = MagicMock()
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "report.docx")),
        store_uploaded_file=MagicMock(return_value=Path("report_1.docx")),
        allow_external_input_files=False,
    )
    event = _build_event()
    upload = Comp.File(name="upload-token", file="upload-token")

    infos = await service._ensure_upload_infos(event, [upload])

    assert len(infos) == 1
    info = infos[0]
    assert info["original_name"] == "report.docx"
    assert info["file_suffix"] == ".docx"
    assert info["type_desc"] == "Office文档 (Word/Excel/PPT)"
    assert info["is_supported"] is True
    assert info["stored_name"] == "report_1.docx"
    assert info["source_path"] == str(source_path.resolve())


@pytest.mark.asyncio
async def test_upload_session_service_assigns_file_ids_per_session_user():
    context = MagicMock()
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(
            side_effect=[
                (source_path, "A.docx"),
                (source_path, "B.xlsx"),
                (source_path, "C.docx"),
            ]
        ),
        store_uploaded_file=MagicMock(
            side_effect=[Path("A_1.docx"), Path("B_1.xlsx"), Path("C_1.docx")]
        ),
        allow_external_input_files=False,
    )
    user_a = _build_event(sender_id="user-a")
    user_b = _build_event(sender_id="user-b")

    infos_a1 = await service._ensure_upload_infos(
        user_a,
        [Comp.File(name="A.docx", file="A.docx")],
    )
    infos_a2 = await service._ensure_upload_infos(
        _build_event(sender_id="user-a"),
        [Comp.File(name="B.xlsx", file="B.xlsx")],
    )
    infos_b1 = await service._ensure_upload_infos(
        user_b,
        [Comp.File(name="C.docx", file="C.docx")],
    )

    assert infos_a1[0]["file_id"] == "f1"
    assert infos_a2[0]["file_id"] == "f2"
    assert infos_b1[0]["file_id"] == "f1"

    listed_a = service.list_session_upload_infos(_build_event(sender_id="user-a"))
    listed_b = service.list_session_upload_infos(user_b)
    assert [info["file_id"] for info in listed_a] == ["f1", "f2"]
    assert [info["file_id"] for info in listed_b] == ["f1"]


@pytest.mark.asyncio
async def test_upload_session_service_uses_independent_upload_ttl():
    context = MagicMock()
    source_path = Path(__file__).resolve()
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=20,
        upload_session_ttl_seconds=120,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(return_value=(source_path, "A.docx")),
        store_uploaded_file=MagicMock(return_value=Path("A_1.docx")),
        allow_external_input_files=False,
    )
    event = _build_event(sender_id="user-a")

    with patch(
        "astrbot_plugin_office_assistant.services.upload_session_service.time.time"
    ) as mocked_time:
        mocked_time.return_value = 1000.0
        await service._ensure_upload_infos(
            event,
            [Comp.File(name="A.docx", file="A.docx")],
        )

        mocked_time.return_value = 1050.0
        assert [
            info["file_id"] for info in service.list_session_upload_infos(event)
        ] == ["f1"]

        mocked_time.return_value = 1121.0
        assert service.list_session_upload_infos(event) == []


@pytest.mark.asyncio
async def test_command_service_doc_lists_selects_and_clears_uploaded_files():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    context.get_config.return_value = {"wake_prefix": ["/"]}
    source_path = Path(__file__).resolve()
    upload_session_service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(
            side_effect=[
                (source_path, "A.docx"),
                (source_path, "B.xlsx"),
            ]
        ),
        store_uploaded_file=MagicMock(side_effect=[Path("A_1.docx"), Path("B_1.xlsx")]),
        allow_external_input_files=False,
    )
    workspace_dir = _make_workspace("command-doc")
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.capabilities = {
            "office_to_pdf": False,
            "pdf_to_word": False,
            "pdf_to_excel": False,
        }
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=workspace_dir,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=upload_session_service,
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )
        upload_event_a = _build_event(sender_id="user-a")
        upload_event_b = _build_event(sender_id="user-a")
        await upload_session_service._ensure_upload_infos(
            upload_event_a,
            [Comp.File(name="A.docx", file="A.docx")],
        )
        await upload_session_service._ensure_upload_infos(
            upload_event_b,
            [Comp.File(name="B.xlsx", file="B.xlsx")],
        )

        command_event = _build_event(sender_id="user-a")
        listed = service.doc_list(command_event)
        assert "[f1] A.docx" in listed
        assert "[f2] B.xlsx" in listed

        selected = await service.doc_use(
            command_event,
            "f2 根据这份文件整理成正式汇报",
        )
        assert selected is None
        assert command_event.get_extra(DOC_COMMAND_TRIGGER_EVENT_KEY) is True
        queued_event = event_queue.put.await_args.args[0]
        assert queued_event is not command_event
        cached_infos = upload_session_service.get_cached_upload_infos(queued_event)
        assert len(cached_infos) == 1
        assert cached_infos[0]["file_id"] == "f2"
        assert cached_infos[0]["original_name"] == "B.xlsx"
        assert queued_event.message_str.startswith("/")
        assert "[用户指令]" in queued_event.message_str
        assert "根据这份文件整理成正式汇报" in queued_event.message_str
        event_queue.put.assert_awaited_once()

        cleared = service.doc_clear(_build_event(sender_id="user-a"), "f1")
        assert cleared == "✅ 已清除文件 f1。"

        remaining = service.doc_list(_build_event(sender_id="user-a"))
        assert "[f1] A.docx" not in remaining
        assert "[f2] B.xlsx" in remaining
    finally:
        shutil.rmtree(workspace_dir, ignore_errors=True)
        executor.shutdown(wait=False)


@pytest.mark.asyncio
async def test_command_service_doc_use_supports_multiple_file_ids():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    context.get_config.return_value = {"wake_prefix": ["/"]}
    source_path = Path(__file__).resolve()
    upload_session_service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(
            side_effect=[
                (source_path, "A.docx"),
                (source_path, "B.xlsx"),
            ]
        ),
        store_uploaded_file=MagicMock(side_effect=[Path("A_1.docx"), Path("B_1.xlsx")]),
        allow_external_input_files=False,
    )
    workspace_dir = _make_workspace("command-doc-multi")
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.capabilities = {
            "office_to_pdf": False,
            "pdf_to_word": False,
            "pdf_to_excel": False,
        }
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=workspace_dir,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=upload_session_service,
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )
        await upload_session_service._ensure_upload_infos(
            _build_event(sender_id="user-a"),
            [Comp.File(name="A.docx", file="A.docx")],
        )
        await upload_session_service._ensure_upload_infos(
            _build_event(sender_id="user-a"),
            [Comp.File(name="B.xlsx", file="B.xlsx")],
        )

        command_event = _build_event(sender_id="user-a")
        selected = await service.doc_use(
            command_event,
            "f1 f2 根据这些文件整理成正式汇报",
        )

        assert selected is None
        queued_event = event_queue.put.await_args.args[0]
        cached_infos = upload_session_service.get_cached_upload_infos(queued_event)
        assert [info["file_id"] for info in cached_infos] == ["f1", "f2"]
        assert "根据这些文件整理成正式汇报" in queued_event.message_str
    finally:
        shutil.rmtree(workspace_dir, ignore_errors=True)
        executor.shutdown(wait=False)


@pytest.mark.asyncio
async def test_upload_session_service_requeue_uses_configured_wake_prefix():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    context.get_config.return_value = {"wake_prefix": ["!"]}
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        allow_external_input_files=False,
    )
    event = _build_event(sender_id="user-a")
    upload_info = {
        "file_id": "f1",
        "original_name": "A.docx",
        "stored_name": "A_1.docx",
        "source_path": "",
        "file_suffix": ".docx",
        "type_desc": "Office文档 (Word/Excel/PPT)",
        "is_supported": True,
    }

    await service.requeue_upload_request(
        event,
        upload_infos=[upload_info],
        user_instruction="整理成正式汇报",
    )

    queued_event = event_queue.put.await_args.args[0]
    assert queued_event is not event
    assert queued_event.message_str.startswith("!")
    assert "[用户指令]" in queued_event.message_str
    event_queue.put.assert_awaited_once()


@pytest.mark.asyncio
async def test_upload_session_service_requeues_buffered_upload_without_command_prefix():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    context.get_config.return_value = {"wake_prefix": ["!"]}
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        allow_external_input_files=False,
    )
    event = _build_event(sender_id="user-a")
    upload_info = {
        "file_id": "f1",
        "original_name": "A.docx",
        "stored_name": "A_1.docx",
        "source_path": "",
        "file_suffix": ".docx",
        "type_desc": "Office文档 (Word/Excel/PPT)",
        "is_supported": True,
    }

    with patch(
        "astrbot_plugin_office_assistant.services.upload_session_service.active_event_registry.request_agent_stop_all"
    ) as request_agent_stop_all:
        await service.requeue_buffered_upload_request(
            event,
            upload_infos=[upload_info],
            user_instruction="整理成正式汇报",
        )

    request_agent_stop_all.assert_called_once_with(
        event.unified_msg_origin,
        exclude=event,
    )

    queued_event = event_queue.put.await_args.args[0]
    assert queued_event is not event
    assert not queued_event.message_str.startswith("!")
    assert not queued_event.message_str.startswith("/")
    assert "[用户指令]" in queued_event.message_str
    event_queue.put.assert_awaited_once()


@pytest.mark.asyncio
async def test_upload_session_service_requests_active_agent_stop_before_command_requeue():
    context = MagicMock()
    event_queue = AsyncMock()
    context.get_event_queue.return_value = event_queue
    context.get_config.return_value = {"wake_prefix": ["/"]}
    service = UploadSessionService(
        context=context,
        recent_text_ttl_seconds=30,
        upload_session_ttl_seconds=300,
        recent_text_max_entries=32,
        recent_text_cleanup_interval_seconds=10,
        upload_session_cleanup_interval_seconds=30,
        extract_upload_source=AsyncMock(),
        store_uploaded_file=MagicMock(),
        allow_external_input_files=False,
    )
    event = _build_event(sender_id="user-a")
    upload_info = {
        "file_id": "f1",
        "original_name": "A.docx",
        "stored_name": "A_1.docx",
        "source_path": "",
        "file_suffix": ".docx",
        "type_desc": "Office文档 (Word/Excel/PPT)",
        "is_supported": True,
    }

    with patch(
        "astrbot_plugin_office_assistant.services.upload_session_service.active_event_registry.request_agent_stop_all"
    ) as request_agent_stop_all:
        await service.requeue_upload_request(
            event,
            upload_infos=[upload_info],
            user_instruction="整理成正式汇报",
        )

    request_agent_stop_all.assert_called_once_with(
        event.unified_msg_origin,
        exclude=event,
    )
    event_queue.put.assert_awaited_once()


@pytest.mark.asyncio
async def test_delivery_service_sends_message_and_file_for_existing_path():
    event = _build_event()
    event.send = AsyncMock()
    service = DeliveryService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=False,
        reply_to_user=True,
    )
    file_path = Path(__file__).resolve()
    try:
        await service.send_file_with_preview(event, file_path, "✅ 已发送")
    finally:
        service._executor.shutdown(wait=False)

    assert event.send.await_count == 2


@pytest.mark.asyncio
async def test_delivery_service_returns_missing_message_for_absent_export():
    event = _build_event()
    service = DeliveryService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=False,
        reply_to_user=False,
    )
    missing_path = Path(__file__).resolve().parent / "missing-output.docx"
    try:
        result = await service.send_exported_document(
            event,
            missing_path,
            "✅ exported",
        )
    finally:
        service._executor.shutdown(wait=False)

    assert "but the file does not exist" in result


@pytest.mark.asyncio
async def test_delivery_service_handles_exported_document_tool_via_context_event():
    event = _build_event()
    event.send = AsyncMock()
    context = SimpleNamespace(context=SimpleNamespace(event=event))
    service = DeliveryService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=MagicMock(),
        enable_preview=False,
        auto_delete=False,
        reply_to_user=False,
    )
    file_path = Path(__file__).resolve()
    try:
        result = await service.handle_exported_document_tool(
            context,
            str(file_path),
            "✅ exported",
        )
    finally:
        service._executor.shutdown(wait=False)

    assert result == f"Document exported and sent to the user: {file_path.name}"
    assert event.send.await_count == 2


@pytest.mark.asyncio
async def test_generated_file_delivery_service_rejects_oversized_output():
    workspace_dir = _make_workspace("generated-file-delivery")
    event = _build_event()
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()
    executor = ThreadPoolExecutor(max_workers=1)
    output_path = workspace_dir / "oversized.pdf"
    output_path.write_bytes(b"x" * 32)

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=8,
            feature_settings={},
        )
        service = GeneratedFileDeliveryService(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
        )

        result = await service.deliver_generated_file(event, output_path)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result.status == "oversized"
    assert result.file_size == 32
    assert result.max_size == 8
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_generated_file_delivery_service_sends_existing_output_with_expected_args():
    workspace_dir = _make_workspace("generated-file-delivery-sent")
    event = _build_event()
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()
    executor = ThreadPoolExecutor(max_workers=1)
    output_path = workspace_dir / "report.pdf"
    output_path.write_bytes(b"small-pdf")
    file_size = output_path.stat().st_size

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=64,
            feature_settings={},
        )
        service = GeneratedFileDeliveryService(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
        )

        result_without_message = await service.deliver_generated_file(
            event, output_path
        )
        result_with_message = await service.deliver_generated_file(
            event,
            output_path,
            success_message="✅ 已发送",
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result_without_message.status == "sent"
    assert result_with_message.status == "sent"
    assert result_without_message.file_size == file_size
    assert result_with_message.file_size == file_size
    assert result_without_message.max_size == 64
    assert result_with_message.max_size == 64
    assert delivery_service.send_file_with_preview.await_count == 2
    assert delivery_service.send_file_with_preview.await_args_list[0].args == (
        event,
        output_path,
    )
    assert delivery_service.send_file_with_preview.await_args_list[1].args == (
        event,
        output_path,
        "✅ 已发送",
    )


@pytest.mark.asyncio
async def test_generated_file_delivery_service_logs_missing_output_path():
    workspace_dir = _make_workspace("generated-file-delivery-missing")
    event = _build_event()
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()
    executor = ThreadPoolExecutor(max_workers=1)

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=8,
            feature_settings={},
        )
        service = GeneratedFileDeliveryService(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
        )

        with patch(
            "astrbot_plugin_office_assistant.services.generated_file_delivery_service.logger.info"
        ) as logger_info:
            result = await service.deliver_generated_file(event, None)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result.status == "missing"
    delivery_service.send_file_with_preview.assert_not_called()
    logger_info.assert_called_once()


@pytest.mark.asyncio
async def test_delivery_service_logs_preview_generation_failure():
    event = _build_event()
    event.send = AsyncMock()
    preview_generator = MagicMock()
    preview_generator.generate_preview.side_effect = RuntimeError("preview boom")
    service = DeliveryService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=preview_generator,
        enable_preview=True,
        auto_delete=False,
        reply_to_user=False,
    )
    file_path = Path(__file__).resolve()
    try:
        with patch(
            "astrbot_plugin_office_assistant.services.delivery_service.logger.warning"
        ) as warning_mock:
            await service.send_file_with_preview(event, file_path, "✅ 已发送")
    finally:
        service._executor.shutdown(wait=False)

    warning_mock.assert_called_once()
    assert "生成预览图失败" in warning_mock.call_args.args[0]
    assert event.send.await_count == 2


@pytest.mark.asyncio
async def test_delivery_service_logs_file_cleanup_failure():
    event = _build_event()
    event.send = AsyncMock()
    workspace_dir = _make_workspace("delivery-cleanup")
    preview_path = workspace_dir / "preview.png"
    preview_path.write_text("preview", encoding="utf-8")
    file_path = workspace_dir / "result.docx"
    file_path.write_text("result", encoding="utf-8")

    preview_generator = MagicMock()
    preview_generator.generate_preview.return_value = preview_path

    service = DeliveryService(
        executor=ThreadPoolExecutor(max_workers=1),
        preview_generator=preview_generator,
        enable_preview=True,
        auto_delete=True,
        reply_to_user=False,
    )

    with patch.object(Path, "unlink", side_effect=PermissionError("locked")):
        with patch(
            "astrbot_plugin_office_assistant.services.delivery_service.logger.warning"
        ) as warning_mock:
            try:
                await service.send_file_with_preview(event, file_path, "✅ 已发送")
            finally:
                service._executor.shutdown(wait=False)
                shutil.rmtree(workspace_dir, ignore_errors=True)

    assert warning_mock.call_count == 2
    logged_messages = [call.args[0] for call in warning_mock.call_args_list]
    assert any("自动删除预览文件失败" in message for message in logged_messages)
    assert any("自动删除文件失败" in message for message in logged_messages)


@pytest.mark.asyncio
async def test_incoming_message_service_buffers_supported_file_and_stops_event():
    message_buffer = MagicMock()
    message_buffer.add_message = AsyncMock(return_value=True)
    message_buffer.is_buffering.return_value = False
    remember_recent_text = MagicMock()
    service = IncomingMessageService(
        message_buffer=message_buffer,
        remember_recent_text=remember_recent_text,
        is_group_feature_enabled=lambda _event: True,
    )
    event = _build_event()
    event.stop_event = MagicMock()
    event.message_obj.message = [Comp.File(name="report.docx", file="report.docx")]

    await service.handle_file_message(event)

    remember_recent_text.assert_not_called()
    message_buffer.add_message.assert_awaited_once_with(event)
    event.stop_event.assert_called_once()


@pytest.mark.asyncio
async def test_incoming_message_service_ignores_non_file_messages_during_buffer():
    message_buffer = MagicMock()
    message_buffer.add_message = AsyncMock(return_value=True)
    message_buffer.is_buffering.return_value = True
    remember_recent_text = MagicMock()
    service = IncomingMessageService(
        message_buffer=message_buffer,
        remember_recent_text=remember_recent_text,
        is_group_feature_enabled=lambda _event: True,
    )
    event = _build_event()
    event.stop_event = MagicMock()
    event.message_obj.message = [Comp.File(name="avatar.png", file="avatar.png")]

    await service.handle_file_message(event)

    remember_recent_text.assert_not_called()
    message_buffer.is_buffering.assert_called_once_with(event)
    message_buffer.add_message.assert_not_awaited()
    event.stop_event.assert_not_called()


@pytest.mark.asyncio
async def test_incoming_message_service_buffers_follow_up_plain_text_while_file_waits():
    message_buffer = MagicMock()
    message_buffer.add_message = AsyncMock(return_value=True)
    message_buffer.is_buffering.return_value = True
    remember_recent_text = MagicMock()
    service = IncomingMessageService(
        message_buffer=message_buffer,
        remember_recent_text=remember_recent_text,
        is_group_feature_enabled=lambda _event: True,
    )
    event = _build_event()
    event.stop_event = MagicMock()
    event.message_obj.message = [Comp.Plain("reset")]

    await service.handle_file_message(event)

    remember_recent_text.assert_not_called()
    message_buffer.is_buffering.assert_called_once_with(event)
    message_buffer.add_message.assert_awaited_once_with(event)
    event.stop_event.assert_called_once()


@pytest.mark.asyncio
async def test_incoming_message_service_ignores_command_like_text_while_file_waits():
    message_buffer = MagicMock()
    message_buffer.add_message = AsyncMock(return_value=True)
    message_buffer.is_buffering.return_value = True
    remember_recent_text = MagicMock()
    service = IncomingMessageService(
        message_buffer=message_buffer,
        remember_recent_text=remember_recent_text,
        is_group_feature_enabled=lambda _event: True,
    )
    event = _build_event()
    event.stop_event = MagicMock()
    event.message_obj.message = [Comp.Plain("/doc list")]

    await service.handle_file_message(event)

    remember_recent_text.assert_not_called()
    message_buffer.is_buffering.assert_called_once_with(event)
    message_buffer.add_message.assert_not_awaited()
    event.stop_event.assert_not_called()


@pytest.mark.asyncio
async def test_incoming_message_service_remembers_plain_text_when_no_buffer_active():
    message_buffer = MagicMock()
    message_buffer.add_message = AsyncMock(return_value=True)
    message_buffer.is_buffering.return_value = False
    remember_recent_text = MagicMock()
    service = IncomingMessageService(
        message_buffer=message_buffer,
        remember_recent_text=remember_recent_text,
        is_group_feature_enabled=lambda _event: True,
    )
    event = _build_event()
    event.stop_event = MagicMock()
    event.message_obj.message = [Comp.Plain("根据这份文件整理成正式汇报")]

    await service.handle_file_message(event)

    remember_recent_text.assert_called_once_with(event)
    message_buffer.is_buffering.assert_called_once_with(event)
    message_buffer.add_message.assert_not_awaited()
    event.stop_event.assert_not_called()


@pytest.mark.asyncio
async def test_error_hook_service_uses_event_session_when_target_not_configured():
    context = MagicMock()
    context.send_message = AsyncMock(return_value=True)
    service = ErrorHookService(
        context=context,
        config={},
        plugin_name="astrbot_plugin_office_assistant",
    )
    event = _build_event()
    event.stop_event = MagicMock()

    await service.handle_plugin_error(
        event,
        "astrbot_plugin_office_assistant",
        "handler_name",
        RuntimeError("boom"),
        "traceback line 1\ntraceback line 2",
    )

    context.send_message.assert_awaited_once()
    args, _kwargs = context.send_message.await_args
    assert args[0] == "session-1"
    assert "handler=handler_name" in args[1].chain[0].text
    event.stop_event.assert_called_once()


@pytest.mark.asyncio
async def test_file_tool_service_reads_text_from_workspace_file():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={},
        )

        result = await service.read_file(event, Path(__file__).resolve().name)
    finally:
        executor.shutdown(wait=False)

    assert result is not None
    assert "[文件:" in result
    assert "test_file_tool_service_reads_text_from_workspace_file" in result


@pytest.mark.asyncio
async def test_file_tool_service_reads_docx_and_extracts_embedded_images():
    workspace_dir = _make_workspace("read-docx-images")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-report.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        docx = _import_docx()

        _write_png(image_path)
        document = docx.Document()
        document.add_paragraph("文档正文")
        document.add_picture(str(image_path), width=docx.shared.Inches(1))
        document.add_paragraph("图片后的说明")
        document.save(docx_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
        )

        result = await service.read_file(event, docx_path.name)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result is not None
    assert "文档正文" in result
    assert "[插图1]" in result
    assert "图片后的说明" in result


def test_extract_word_text_ignores_deleted_and_field_code_runs():
    workspace_dir = _make_workspace("extract-docx-hidden-runs")
    docx_path = workspace_dir / "tracked.docx"

    try:
        docx = _import_docx()

        document = docx.Document()
        document.add_paragraph("保留文本")
        document.save(docx_path)

        _rewrite_docx_document_xml(
            docx_path,
            lambda xml: xml.replace(
                "<w:t>保留文本</w:t>",
                (
                    "<w:t>保留文本</w:t>"
                    "<w:delText>删除内容</w:delText>"
                    "<w:instrText> MERGEFIELD secret </w:instrText>"
                ),
            ),
        )

        extracted = extract_word_text(docx_path, workspace_dir)
    finally:
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert extracted is not None
    assert "保留文本" in extracted
    assert "删除内容" not in extracted
    assert "MERGEFIELD secret" not in extracted


def test_format_extracted_word_content_uses_basename_outside_workspace():
    workspace_dir = _make_workspace("format-word-content-paths")
    outside_image = Path(__file__).resolve()

    try:
        content = ExtractedWordContent(
            items=[
                ExtractedWordItem(type="text", text="正文"),
                ExtractedWordItem(
                    type="image",
                    image_path=outside_image,
                    image_index=1,
                ),
            ]
        )

        formatted = format_extracted_word_content(
            content,
            workspace_root=workspace_dir,
            include_image_paths=True,
        )
    finally:
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert formatted is not None
    assert "[插图1]" in formatted
    assert outside_image.name in formatted
    assert outside_image.as_posix() not in formatted


def test_extract_word_content_returns_structured_result_for_legacy_doc(
    monkeypatch: pytest.MonkeyPatch,
):
    workspace_dir = _make_workspace("extract-legacy-doc")
    doc_path = workspace_dir / "legacy.doc"
    doc_path.write_bytes(b"legacy")

    monkeypatch.setattr(office_utils, "_ANTIWORD_AVAILABLE", True)
    monkeypatch.setattr(office_utils, "_WIN32COM_AVAILABLE", False)
    monkeypatch.setattr(
        office_utils,
        "_extract_doc_text_antiword",
        lambda _path: "Legacy Word content",
    )

    try:
        extracted = office_utils.extract_word_content(doc_path, workspace_dir)
    finally:
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert extracted == ExtractedWordContent(
        text="Legacy Word content",
        image_paths=[],
        items=[],
    )


def test_workspace_service_extract_office_text_reads_legacy_doc(
    monkeypatch: pytest.MonkeyPatch,
):
    workspace_dir = _make_workspace("workspace-legacy-doc")
    executor = ThreadPoolExecutor(max_workers=1)
    doc_path = workspace_dir / "legacy.doc"
    doc_path.write_bytes(b"legacy")

    monkeypatch.setattr(office_utils, "_ANTIWORD_AVAILABLE", True)
    monkeypatch.setattr(office_utils, "_WIN32COM_AVAILABLE", False)
    monkeypatch.setattr(
        office_utils,
        "_extract_doc_text_antiword",
        lambda _path: "Legacy document text",
    )

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        extracted = workspace_service.extract_office_text(doc_path, OfficeType.WORD)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert extracted == "Legacy document text"


def test_workspace_service_extract_word_content_skips_docx_image_materialization_when_disabled():
    workspace_dir = _make_workspace("workspace-docx-no-image-materialize")
    executor = ThreadPoolExecutor(max_workers=1)
    docx_path = workspace_dir / "image-report.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        docx = _import_docx()

        _write_png(image_path)
        document = docx.Document()
        document.add_paragraph("图前说明")
        document.add_picture(str(image_path), width=docx.shared.Inches(1))
        document.add_paragraph("图后说明")
        document.save(docx_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )

        extracted = workspace_service.extract_word_content(
            docx_path,
            include_images=False,
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert extracted is not None
    assert extracted.image_count == 1
    assert extracted.image_paths == []
    assert all(item.type == "text" for item in extracted.items)


def test_workspace_service_extract_word_content_formats_relative_image_paths_when_enabled():
    workspace_dir = _make_workspace("workspace-docx-image-materialize")
    executor = ThreadPoolExecutor(max_workers=1)
    docx_path = workspace_dir / "image-report.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        docx = _import_docx()

        _write_png(image_path)
        document = docx.Document()
        document.add_paragraph("图前说明")
        document.add_picture(str(image_path), width=docx.shared.Inches(1))
        document.add_paragraph("图后说明")
        document.save(docx_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )

        extracted = workspace_service.extract_word_content(
            docx_path,
            include_images=True,
        )
        formatted = workspace_service.format_word_content(
            extracted,
            include_image_paths=True,
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert extracted is not None
    assert extracted.image_count == 1
    assert len(extracted.image_paths) == 1
    assert formatted is not None
    assert "图前说明" in formatted
    assert "图后说明" in formatted
    assert "[插图1]" in formatted
    assert ".read_assets/" in formatted
    assert str(workspace_dir) not in formatted
    assert str(workspace_dir.resolve()) not in formatted


@pytest.mark.asyncio
async def test_file_tool_service_streams_docx_images_as_tool_results():
    workspace_dir = _make_workspace("stream-docx-images")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-report.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        docx = _import_docx()

        _write_png(image_path)
        document = docx.Document()
        document.add_paragraph("文档正文")
        document.add_picture(str(image_path), width=docx.shared.Inches(1))
        document.add_paragraph("图片后的说明")
        document.save(docx_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert isinstance(results[0], str)
    assert "文档正文" in results[0]
    assert "[插图1]" in results[0]
    assert len(results) == 2
    assert isinstance(results[1], mcp.types.CallToolResult)
    assert isinstance(results[1].content[0], mcp.types.ImageContent)
    assert results[1].content[0].mimeType == "image/png"
    assert results[1].content[0].data
    assert "图片后的说明" in results[0]


@pytest.mark.asyncio
async def test_file_tool_service_returns_error_when_docx_library_missing():
    workspace_dir = _make_workspace("missing-docx-lib")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "missing-lib.docx"
    docx_path.write_bytes(b"not-a-real-docx")

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={},
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert results == ["错误：文件 'missing-lib.docx' 无法读取，可能未安装对应解析库"]


@pytest.mark.asyncio
async def test_file_tool_service_skips_unreadable_docx_image_bytes(
    monkeypatch: pytest.MonkeyPatch,
):
    workspace_dir = _make_workspace("broken-docx-image")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "broken-image.docx"
    docx_path.write_bytes(b"placeholder")
    broken_image_path = workspace_dir / "broken.png"
    healthy_image_path = workspace_dir / "healthy.png"
    broken_image_path.write_bytes(b"broken")
    _write_png(healthy_image_path)

    extracted = ExtractedWordContent(
        items=[
            ExtractedWordItem(type="text", text="文档正文"),
            ExtractedWordItem(
                type="image",
                image_path=broken_image_path,
                image_index=1,
            ),
            ExtractedWordItem(type="text", text="收尾说明"),
            ExtractedWordItem(
                type="image",
                image_path=healthy_image_path,
                image_index=2,
            ),
        ],
        image_paths=[broken_image_path, healthy_image_path],
    )

    original_read_bytes = Path.read_bytes

    def fake_read_bytes(path: Path) -> bytes:
        if path == broken_image_path:
            raise OSError("broken image")
        return original_read_bytes(path)

    monkeypatch.setattr(Path, "read_bytes", fake_read_bytes)

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        monkeypatch.setattr(
            workspace_service,
            "extract_word_content",
            lambda _path, include_images=True: extracted,
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert len(results) == 2
    assert isinstance(results[0], str)
    assert "文档正文" in results[0]
    assert "[插图1]" in results[0]
    assert "[插图2]" in results[0]
    assert "收尾说明" in results[0]
    assert isinstance(results[1], mcp.types.CallToolResult)
    assert len(results[1].content) == 1
    assert isinstance(results[1].content[0], mcp.types.ImageContent)


@pytest.mark.asyncio
async def test_file_tool_service_uses_item_image_index_for_skip_reasons():
    workspace_dir = _make_workspace("stream-docx-image-index")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-report.docx"
    large_image = workspace_dir / "embedded-large.png"
    small_image = workspace_dir / "embedded-small.png"

    try:
        docx_path.write_bytes(b"docx")
        large_image.write_bytes(b"a" * 40)
        small_image.write_bytes(b"b" * 10)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        workspace_service.extract_word_content = MagicMock(
            return_value=SimpleNamespace(
                text="文档正文",
                image_paths=[large_image, small_image],
                image_count=2,
                items=[
                    SimpleNamespace(type="text", text="文档正文"),
                    SimpleNamespace(
                        type="image",
                        image_path=small_image,
                        image_index=2,
                    ),
                    SimpleNamespace(
                        type="image",
                        image_path=large_image,
                        image_index=1,
                    ),
                    SimpleNamespace(type="text", text="收尾说明"),
                ],
            )
        )

        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
            max_inline_docx_image_bytes=20,
            max_inline_docx_image_count=2,
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert isinstance(results[0], str)
    assert "[插图2]" in results[0]
    assert "[插图1]（未注入模型上下文" in results[0]
    assert "超过 20.00 B 限制" in results[0]
    assert isinstance(results[1], mcp.types.CallToolResult)
    assert len(results[1].content) == 1


@pytest.mark.asyncio
async def test_file_tool_service_limits_inline_docx_images():
    workspace_dir = _make_workspace("stream-docx-image-limits")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-report.docx"
    small_image_1 = workspace_dir / "embedded-1.png"
    large_image = workspace_dir / "embedded-large.png"
    small_image_2 = workspace_dir / "embedded-2.png"
    small_image_3 = workspace_dir / "embedded-3.png"

    try:
        docx_path.write_bytes(b"docx")
        small_image_1.write_bytes(b"a" * 10)
        large_image.write_bytes(b"b" * 40)
        small_image_2.write_bytes(b"c" * 10)
        small_image_3.write_bytes(b"d" * 10)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        workspace_service.extract_word_content = MagicMock(
            return_value=SimpleNamespace(
                text="文档正文",
                image_paths=[
                    small_image_1,
                    large_image,
                    small_image_2,
                    small_image_3,
                ],
                image_count=4,
                items=[
                    SimpleNamespace(type="text", text="文档正文"),
                    SimpleNamespace(type="image", image_path=small_image_1),
                    SimpleNamespace(type="image", image_path=large_image),
                    SimpleNamespace(type="image", image_path=small_image_2),
                    SimpleNamespace(type="image", image_path=small_image_3),
                    SimpleNamespace(type="text", text="收尾说明"),
                ],
            )
        )

        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
            max_inline_docx_image_bytes=20,
            max_inline_docx_image_count=2,
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert isinstance(results[0], str)
    assert "文档正文" in results[0]
    assert "[插图1]" in results[0]
    assert "插图2" in results[0]
    assert "超过 20.00 B 限制" in results[0]
    assert "插图4" in results[0]
    assert "超过单文档最多 2 张限制" in results[0]
    assert "收尾说明" in results[0]
    assert isinstance(results[1], mcp.types.CallToolResult)
    assert len(results[1].content) == 2
    assert all(isinstance(item, mcp.types.ImageContent) for item in results[1].content)
    assert len(results) == 2


@pytest.mark.asyncio
async def test_file_tool_service_skips_docx_image_review_when_disabled():
    workspace_dir = _make_workspace("stream-docx-image-review-disabled")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-report.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        _write_png(image_path)
        docx_path.write_bytes(b"docx")

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        workspace_service.extract_word_content = MagicMock(
            return_value=SimpleNamespace(
                text="文档正文",
                image_paths=[image_path],
                image_count=1,
                items=[
                    SimpleNamespace(type="text", text="图前说明"),
                    SimpleNamespace(
                        type="image",
                        image_path=image_path,
                        image_index=1,
                    ),
                    SimpleNamespace(type="text", text="图后说明"),
                ],
            )
        )

        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
            enable_docx_image_review=False,
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert len(results) == 1
    assert isinstance(results[0], str)
    assert "图前说明" in results[0]
    assert "图后说明" in results[0]
    assert "[插图1]" not in results[0]
    workspace_service.extract_word_content.assert_called_once_with(
        docx_path,
        include_images=False,
    )


@pytest.mark.asyncio
async def test_file_tool_service_returns_message_for_image_only_docx_when_review_disabled():
    workspace_dir = _make_workspace("stream-docx-image-only-disabled")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    docx_path = workspace_dir / "image-only.docx"
    image_path = workspace_dir / "embedded.png"

    try:
        _write_png(image_path)
        docx_path.write_bytes(b"docx")

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        workspace_service.extract_word_content = MagicMock(
            return_value=SimpleNamespace(
                text=None,
                image_paths=[image_path],
                image_count=1,
                items=[
                    SimpleNamespace(
                        type="image",
                        image_path=image_path,
                        image_index=1,
                    )
                ],
            )
        )

        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"docx": object()},
            enable_docx_image_review=False,
        )

        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, docx_path.name
            )
        ]
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert len(results) == 1
    assert isinstance(results[0], str)
    assert "仅包含图片内容" in results[0]
    assert "[插图1]" not in results[0]
    workspace_service.extract_word_content.assert_called_once_with(
        docx_path,
        include_images=False,
    )


@pytest.mark.asyncio
async def test_file_tool_service_returns_local_guidance_for_missing_file():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={},
        )

        result = await service.read_file(event, "CLAUDE.md")
    finally:
        executor.shutdown(wait=False)

    assert result is not None
    assert "错误：文件 'CLAUDE.md' 不存在。" in result
    assert "不要联网搜索" in result
    assert "重新上传文件或提供正确的本地路径" in result


@pytest.mark.asyncio
async def test_file_tool_service_returns_error_when_precheck_lacks_resolved_path():
    event = _build_event()
    workspace_service = MagicMock()
    workspace_service.pre_check.return_value = (True, None, None)
    workspace_service.get_max_file_size.return_value = 1024 * 1024
    service = FileReadService(
        workspace_service=workspace_service,
        word_read_service=MagicMock(),
        allow_external_input_files=False,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        group_feature_disabled_error=lambda: "group disabled",
    )

    results = [
        result async for result in service.iter_read_file_tool_results(event, "a.txt")
    ]

    assert results == ["错误：文件路径解析失败"]


@pytest.mark.asyncio
async def test_file_tool_service_returns_error_when_stat_fails():
    event = _build_event()
    workspace_service = MagicMock()
    broken_path = Path("missing.txt")
    workspace_service.pre_check.return_value = (True, broken_path, None)
    workspace_service.display_name.return_value = "missing.txt"
    service = FileReadService(
        workspace_service=workspace_service,
        word_read_service=MagicMock(),
        allow_external_input_files=False,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        group_feature_disabled_error=lambda: "group disabled",
    )

    with patch(
        "astrbot_plugin_office_assistant.services.file_read_service.asyncio.to_thread",
        side_effect=FileNotFoundError("gone"),
    ):
        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, "missing.txt"
            )
        ]

    assert results == ["错误：无法读取文件信息 (missing.txt)"]


@pytest.mark.asyncio
async def test_file_tool_service_offloads_office_text_extraction_to_thread():
    event = _build_event()
    workspace_service = MagicMock()
    resolved_path = Path("report.xlsx")
    workspace_service.pre_check.return_value = (True, resolved_path, None)
    workspace_service.display_name.return_value = "report.xlsx"
    workspace_service.get_max_file_size.return_value = 1024 * 1024
    workspace_service.format_file_result.return_value = "formatted"
    service = FileReadService(
        workspace_service=workspace_service,
        word_read_service=MagicMock(),
        allow_external_input_files=False,
        is_group_feature_enabled=lambda _event: True,
        check_permission=lambda _event: True,
        group_feature_disabled_error=lambda: "group disabled",
    )

    async def _fake_to_thread(func, *args, **kwargs):
        if (
            getattr(func, "__self__", None) == resolved_path
            and getattr(func, "__name__", "") == "stat"
        ):
            return SimpleNamespace(st_size=16)
        if func is workspace_service.extract_office_text:
            return "sheet text"
        raise AssertionError(f"unexpected function: {func}")

    with patch(
        "astrbot_plugin_office_assistant.services.file_read_service.asyncio.to_thread",
        side_effect=_fake_to_thread,
    ) as to_thread:
        results = [
            result
            async for result in service.iter_read_file_tool_results(
                event, "report.xlsx"
            )
        ]

    assert results == ["formatted"]
    assert to_thread.await_count == 2


@pytest.mark.asyncio
async def test_file_tool_service_reads_workbook_with_sheet_sections():
    workspace_dir = _make_workspace("read-workbook")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        openpyxl = pytest.importorskip("openpyxl")
        workbook_path = workspace_dir / "sales.xlsx"
        workbook = openpyxl.Workbook()
        summary_sheet = workbook.active
        summary_sheet.title = "总表"
        summary_sheet.append(["月份", "销售额"])
        summary_sheet.append(["2026-01", 100])
        east_sheet = workbook.create_sheet("华东")
        east_sheet.append(["城市", "销售额"])
        east_sheet.append(["上海", 80])
        workbook.save(workbook_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )

        result = await service.read_workbook(event, workbook_path.name)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result is not None
    assert "[文件信息] 文件名: sales.xlsx" in result
    assert "[Sheet 列表] 总表, 华东" in result
    assert "[Sheet: 总表]" in result
    assert "[Sheet: 华东]" in result
    assert "2026-01\t100" in result
    assert "上海\t80" in result


@pytest.mark.asyncio
async def test_file_tool_service_read_workbook_rejects_non_excel_suffix():
    workspace_dir = _make_workspace("read-workbook-invalid")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        text_path = workspace_dir / "notes.txt"
        text_path.write_text("demo", encoding="utf-8")
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={},
        )

        result = await service.read_workbook(event, text_path.name)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert result is not None
    assert "错误：不支持的文件格式 '.txt'" in result
    assert ".xlsx" in result
    assert ".xls" in result


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_returns_text_result():
    workspace_dir = _make_workspace("execute-excel-script-text")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        delivery_service = MagicMock()
        delivery_service.send_file_with_preview = AsyncMock()
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script=(
                    "workbook = Workbook()\n"
                    "worksheet = workbook.active\n"
                    "worksheet.title = '总表'\n"
                    "worksheet['A1'] = '值'\n"
                    "result_text = worksheet['A1'].value\n"
                ),
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "text"
    assert payload["result_text"] == "值"
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_returns_generated_file():
    workspace_dir = _make_workspace("execute-excel-script-file")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        delivery_service = MagicMock()
        delivery_service.send_file_with_preview = AsyncMock()
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            delivery_service=delivery_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script=(
                    "workbook = Workbook()\n"
                    "worksheet = workbook.active\n"
                    "worksheet['A1'] = '导出'\n"
                    "workbook.save(output_path)\n"
                ),
                output_name="script-output.xlsx",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "file"
    assert payload["output_name"] == "script-output.xlsx"
    delivery_service.send_file_with_preview.assert_awaited_once()


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("is_group_feature_enabled", "check_permission", "expected_error"),
    [
        (lambda _event: False, lambda _event: True, "group disabled"),
        (lambda _event: True, lambda _event: False, "错误：权限不足"),
    ],
)
async def test_file_tool_service_execute_excel_script_validation_errors_do_not_consume_retry_budget(
    is_group_feature_enabled,
    check_permission,
    expected_error,
):
    workspace_dir = _make_workspace("execute-excel-script-validation-error")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
            is_group_feature_enabled=is_group_feature_enabled,
            check_permission=check_permission,
        )
        runner = MagicMock()
        service._excel_script_service._run_script_process = runner

        first = json.loads(
            await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
        )
        second = json.loads(
            await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert first["success"] is False
    assert first["error"] == expected_error
    assert first["retry_count"] == 0
    assert first["retry_exhausted"] is False
    assert second["success"] is False
    assert second["error"] == expected_error
    assert second["retry_count"] == 0
    assert second["retry_exhausted"] is False
    assert runner.call_count == 0


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_rejects_non_sandbox_runtime():
    workspace_dir = _make_workspace("execute-excel-script-non-sandbox")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            astrbot_context=_build_astrbot_context(runtime="local"),
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )

        with patch(
            "astrbot_plugin_office_assistant.services.excel_script_service._get_computer_booter",
            new=AsyncMock(),
        ) as mocked_get_booter:
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert "仅支持 sandbox runtime" in payload["error"]
    assert payload["retry_count"] == 0
    assert payload["retry_exhausted"] is False
    mocked_get_booter.assert_not_awaited()


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_accepts_positional_session_config():
    class _PositionalSessionContext:
        @staticmethod
        def get_config(session_id):
            assert session_id == "session-1"
            return {"provider_settings": {"computer_use_runtime": "sandbox"}}

    workspace_dir = _make_workspace("execute-excel-script-positional-config")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            astrbot_context=_PositionalSessionContext(),
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        service._excel_script_service._run_script_process = AsyncMock(
            return_value=SimpleNamespace(
                success=True,
                mode="text",
                result_text="ok",
                output_path=None,
            )
        )

        with patch(
            "astrbot_plugin_office_assistant.services.excel_script_service._get_computer_booter",
            new=AsyncMock(return_value=MagicMock(capabilities=("python", "filesystem"))),
        ) as mocked_get_booter:
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "text"
    assert payload["result_text"] == "ok"
    mocked_get_booter.assert_awaited_once()


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_accepts_noarg_session_config_when_signature_unavailable():
    class _NoArgSessionContext:
        @staticmethod
        def get_config():
            return {"provider_settings": {"computer_use_runtime": "sandbox"}}

    workspace_dir = _make_workspace("execute-excel-script-noarg-config")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            astrbot_context=_NoArgSessionContext(),
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        service._excel_script_service._run_script_process = AsyncMock(
            return_value=SimpleNamespace(
                success=True,
                mode="text",
                result_text="ok",
                output_path=None,
            )
        )

        with patch(
            "astrbot_plugin_office_assistant.services.excel_script_service.inspect.signature",
            side_effect=ValueError("signature unavailable"),
        ), patch(
            "astrbot_plugin_office_assistant.services.excel_script_service._get_computer_booter",
            new=AsyncMock(return_value=MagicMock(capabilities=("python", "filesystem"))),
        ) as mocked_get_booter:
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "text"
    assert payload["result_text"] == "ok"
    mocked_get_booter.assert_awaited_once()


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_rejects_sandbox_missing_capabilities():
    workspace_dir = _make_workspace("execute-excel-script-missing-capabilities")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            astrbot_context=_build_astrbot_context(runtime="sandbox"),
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        runner = AsyncMock()
        service._excel_script_service._run_script_process = runner

        with patch(
            "astrbot_plugin_office_assistant.services.excel_script_service._get_computer_booter",
            new=AsyncMock(return_value=MagicMock(capabilities=("python",))),
        ):
            first = json.loads(
                await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                )
            )
            second = json.loads(
                await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                )
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert first["success"] is False
    assert "当前 sandbox profile 缺少必要能力：filesystem" in first["error"]
    assert first["retry_count"] == 0
    assert first["retry_exhausted"] is False
    assert second["success"] is False
    assert second["retry_count"] == 0
    assert second["retry_exhausted"] is False
    runner.assert_not_awaited()


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_handles_sandbox_exec_exception():
    class _ExecFailureSandboxPython(_FakeSandboxPython):
        async def exec(
            self,
            code: str,
            kernel_id: str | None = None,
            timeout: int = 30,
            silent: bool = False,
        ) -> dict[str, object]:
            _ = code
            _ = kernel_id
            _ = timeout
            _ = silent
            raise RuntimeError("sandbox rpc failed")

    class _ExecFailureSandboxBooter(_FakeSandboxBooter):
        def __init__(self, workspace_root: Path):
            super().__init__(workspace_root)
            self.python = _ExecFailureSandboxPython(workspace_root)

    workspace_dir = _make_workspace("execute-excel-script-exec-exception")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_ExecFailureSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "Excel sandbox 初始化失败"
    assert "sandbox rpc failed" in payload["traceback"]
    assert payload["retry_count"] == 0
    assert payload["retry_exhausted"] is False


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_upload_failure_does_not_consume_retry_budget():
    class _UploadFailureSandboxBooter(_FakeSandboxBooter):
        async def upload_file(self, path: str, file_name: str) -> dict[str, object]:
            _ = path
            _ = file_name
            return {"success": False, "message": "upload failed"}

    workspace_dir = _make_workspace("execute-excel-script-upload-failure")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        openpyxl = pytest.importorskip("openpyxl")
        source_path = workspace_dir / "source.xlsx"
        workbook = openpyxl.Workbook()
        workbook.active["A1"] = "原始"
        workbook.save(source_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_UploadFailureSandboxBooter(sandbox_dir)):
            first = json.loads(
                await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                    input_files=[source_path.name],
                )
            )
            second = json.loads(
                await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                    input_files=[source_path.name],
                )
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert first["success"] is False
    assert first["error"] == "Excel 输入文件上传失败"
    assert first["retry_count"] == 0
    assert first["retry_exhausted"] is False
    assert second["success"] is False
    assert second["error"] == "Excel 输入文件上传失败"
    assert second["retry_count"] == 0
    assert second["retry_exhausted"] is False


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_handles_upload_exception():
    class _UploadExceptionSandboxBooter(_FakeSandboxBooter):
        async def upload_file(self, path: str, file_name: str) -> dict[str, object]:
            _ = path
            _ = file_name
            raise RuntimeError("upload rpc failed")

    workspace_dir = _make_workspace("execute-excel-script-upload-exception")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        openpyxl = pytest.importorskip("openpyxl")
        source_path = workspace_dir / "source.xlsx"
        workbook = openpyxl.Workbook()
        workbook.active["A1"] = "原始"
        workbook.save(source_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_UploadExceptionSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
                input_files=[source_path.name],
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "Excel 输入文件上传失败"
    assert "upload rpc failed" in payload["traceback"]
    assert payload["retry_count"] == 0
    assert payload["retry_exhausted"] is False


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_uses_input_file_copies():
    workspace_dir = _make_workspace("execute-excel-script-input-copy")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        openpyxl = pytest.importorskip("openpyxl")
        source_path = workspace_dir / "source.xlsx"
        workbook = openpyxl.Workbook()
        workbook.active["A1"] = "原始"
        workbook.save(source_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script=(
                    "workbook = load_workbook(input_files[0])\n"
                    "worksheet = workbook.active\n"
                    "worksheet['A1'] = '已修改'\n"
                    "workbook.save(input_files[0])\n"
                    "result_text = worksheet['A1'].value\n"
                ),
                input_files=[source_path.name],
            )
        reloaded = openpyxl.load_workbook(source_path)
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "text"
    assert payload["result_text"] == "已修改"
    assert reloaded.active["A1"].value == "原始"


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_keeps_same_named_inputs_separate():
    workspace_dir = _make_workspace("execute-excel-script-same-name-inputs")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        openpyxl = pytest.importorskip("openpyxl")
        first_dir = workspace_dir / "north"
        second_dir = workspace_dir / "south"
        first_dir.mkdir()
        second_dir.mkdir()
        first_path = first_dir / "sales.xlsx"
        second_path = second_dir / "sales.xlsx"

        first_workbook = openpyxl.Workbook()
        first_workbook.active["A1"] = "north"
        first_workbook.save(first_path)

        second_workbook = openpyxl.Workbook()
        second_workbook.active["A1"] = "south"
        second_workbook.save(second_path)

        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script=(
                    "values = []\n"
                    "for path in input_files:\n"
                    "    workbook = load_workbook(path)\n"
                    "    values.append(workbook.active['A1'].value)\n"
                    "result_text = '|'.join(values)\n"
                ),
                input_files=["north/sales.xlsx", "south/sales.xlsx"],
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is True
    assert payload["mode"] == "text"
    assert payload["result_text"] == "north|south"


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_rejects_text_and_file_together():
    workspace_dir = _make_workspace("execute-excel-script-conflict")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script=(
                    "workbook = Workbook()\n"
                    "workbook.save(output_path)\n"
                    "result_text = 'done'\n"
                ),
                output_name="conflict.xlsx",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert "不能同时设置 result_text" in payload["error"]
    assert "workbook.save(output_path)" in payload["script"]


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_returns_traceback_on_failure():
    workspace_dir = _make_workspace("execute-excel-script-error")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "boom"
    assert "ValueError: boom" in payload["traceback"]
    assert payload["script"] == "raise ValueError('boom')"


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_handles_invalid_result_json():
    workspace_dir = _make_workspace("execute-excel-script-invalid-result-json")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )

        def _invalid_result_runner(*, result_path: str, **_kwargs) -> str:
            result_file = Path(result_path)
            return (
                "from pathlib import Path\n"
                f"Path({str(result_file)!r}).write_text('not json', encoding='utf-8')\n"
            )

        with patch.object(
            service._excel_script_service,
            "_build_runner_script",
            side_effect=_invalid_result_runner,
        ):
            with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
                result = await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "Excel 脚本结果解析失败"
    assert "not json" not in payload["traceback"]
    assert payload["script"] == "result_text = 'ok'"


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("result_content", "traceback_fragment"),
    [
        ("[]", "list"),
        ("null", "NoneType"),
    ],
)
async def test_file_tool_service_execute_excel_script_handles_non_object_result_json(
    result_content,
    traceback_fragment,
):
    workspace_dir = _make_workspace("execute-excel-script-non-object-result-json")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )

        def _non_object_result_runner(*, result_path: str, **_kwargs) -> str:
            result_file = Path(result_path)
            return (
                "from pathlib import Path\n"
                f"Path({str(result_file)!r}).write_text({result_content!r}, encoding='utf-8')\n"
            )

        with patch.object(
            service._excel_script_service,
            "_build_runner_script",
            side_effect=_non_object_result_runner,
        ):
            with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
                result = await service.execute_excel_script(
                    event,
                    script="result_text = 'ok'",
                )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "Excel 脚本结果解析失败"
    assert traceback_fragment in payload["traceback"]
    assert payload["script"] == "result_text = 'ok'"


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("payload_text", "expected_error", "traceback_fragment"),
    [
        (
            '{"success": true, "mode": "file"}',
            "脚本返回了 file 模式，但缺少 output_path",
            "",
        ),
        (
            '{"success": true, "mode": "file", "output_path": "unexpected.xlsx"}',
            "脚本返回了无效的 output_path",
            "unexpected.xlsx",
        ),
    ],
)
async def test_file_tool_service_execute_excel_script_validates_file_mode_output_path(
    payload_text,
    expected_error,
    traceback_fragment,
):
    workspace_dir = _make_workspace("execute-excel-script-invalid-file-output-path")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )

        def _file_payload_runner(*, result_path: str, **_kwargs) -> str:
            result_file = Path(result_path)
            return (
                "from pathlib import Path\n"
                f"Path({str(result_file)!r}).write_text({payload_text!r}, encoding='utf-8')\n"
            )

        with patch.object(
            service._excel_script_service,
            "_build_runner_script",
            side_effect=_file_payload_runner,
        ):
            with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
                result = await service.execute_excel_script(
                    event,
                    script="workbook = Workbook()",
                    output_name="script-output.xlsx",
                )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == expected_error
    if traceback_fragment:
        assert traceback_fragment in payload["traceback"]
    else:
        assert payload["traceback"] == ""


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_handles_result_read_error(
    monkeypatch,
):
    workspace_dir = _make_workspace("execute-excel-script-result-read-error")
    sandbox_dir = workspace_dir / "fake-sandbox"
    sandbox_dir.mkdir()
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        original_read_text = Path.read_text

        def _broken_read_text(self, *args, **kwargs):
            if self.name == "result.json":
                raise OSError("read failed")
            return original_read_text(self, *args, **kwargs)

        monkeypatch.setattr(Path, "read_text", _broken_read_text)
        with _patch_excel_sandbox(_FakeSandboxBooter(sandbox_dir)):
            result = await service.execute_excel_script(
                event,
                script="result_text = 'ok'",
            )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    payload = json.loads(result)
    assert payload["success"] is False
    assert payload["error"] == "Excel 脚本结果解析失败"
    assert "read failed" in payload["traceback"]
    assert payload["script"] == "result_text = 'ok'"


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_marks_retry_exhausted_after_third_failure():
    workspace_dir = _make_workspace("execute-excel-script-retry-exhausted")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        runner = MagicMock(
            return_value=SimpleNamespace(
                success=False,
                mode="error",
                error="boom",
                traceback="traceback",
                script="raise ValueError('boom')",
            )
        )
        service._excel_script_service._run_script_process = AsyncMock(side_effect=runner)
        service._excel_script_service._acquire_sandbox_booter = AsyncMock(
            return_value=(MagicMock(capabilities=("python", "filesystem")), None)
        )

        first = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        )
        second = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        )
        third = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert first["success"] is False
    assert first["retry_count"] == 0
    assert first["retry_exhausted"] is False
    assert second["success"] is False
    assert second["retry_count"] == 1
    assert second["retry_exhausted"] is False
    assert third["success"] is False
    assert third["retry_count"] == 2
    assert third["retry_exhausted"] is True
    assert "已超过最多 2 次脚本重试" in third["error"]
    assert runner.call_count == 3


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_stops_running_after_retry_limit():
    workspace_dir = _make_workspace("execute-excel-script-retry-stop")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        runner = MagicMock(
            return_value=SimpleNamespace(
                success=False,
                mode="error",
                error="boom",
                traceback="traceback",
                script="raise ValueError('boom')",
            )
        )
        service._excel_script_service._run_script_process = AsyncMock(side_effect=runner)
        service._excel_script_service._acquire_sandbox_booter = AsyncMock(
            return_value=(MagicMock(capabilities=("python", "filesystem")), None)
        )

        for _ in range(3):
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        fourth = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert fourth["success"] is False
    assert fourth["retry_count"] == 2
    assert fourth["retry_exhausted"] is True
    assert "脚本重试次数已用尽" in fourth["error"]
    assert runner.call_count == 3


@pytest.mark.asyncio
async def test_file_tool_service_execute_excel_script_resets_retry_count_after_success():
    workspace_dir = _make_workspace("execute-excel-script-retry-reset")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_libs={"openpyxl": object()},
        )
        runner = MagicMock(
            side_effect=[
                SimpleNamespace(
                    success=False,
                    mode="error",
                    error="boom",
                    traceback="traceback",
                    script="raise ValueError('boom')",
                ),
                SimpleNamespace(
                    success=True,
                    mode="text",
                    result_text="done",
                    output_path=None,
                ),
                SimpleNamespace(
                    success=False,
                    mode="error",
                    error="boom-again",
                    traceback="traceback",
                    script="raise ValueError('boom-again')",
                ),
            ]
        )
        service._excel_script_service._run_script_process = AsyncMock(side_effect=runner)
        service._excel_script_service._acquire_sandbox_booter = AsyncMock(
            return_value=(MagicMock(capabilities=("python", "filesystem")), None)
        )

        first = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom')",
            )
        )
        second = json.loads(
            await service.execute_excel_script(
                event,
                script="result_text = 'done'",
            )
        )
        third = json.loads(
            await service.execute_excel_script(
                event,
                script="raise ValueError('boom-again')",
            )
        )
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)

    assert first["success"] is False
    assert first["retry_count"] == 0
    assert second["success"] is True
    assert second["mode"] == "text"
    assert second["result_text"] == "done"
    assert third["success"] is False
    assert third["retry_count"] == 0
    assert third["retry_exhausted"] is False
    assert runner.call_count == 3


@pytest.mark.asyncio
async def test_file_tool_service_creates_office_file_via_generator_and_delivery():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    event.send = AsyncMock()
    office_generator = MagicMock()
    office_generator.generate = AsyncMock(return_value=Path(__file__).resolve())
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report.docx",
            content="hello world",
            file_type="word",
        )
    finally:
        executor.shutdown(wait=False)

    office_generator.generate.assert_awaited_once()
    delivery_service.send_file_with_preview.assert_awaited_once()
    assert result is None


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_exports_word_via_node_backend():
    docx = _import_docx()
    workspace_dir = _make_workspace("create-office-word-node")
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    event.send = AsyncMock()
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        office_generator = OfficeGenerator(
            data_path=workspace_dir,
            render_backend_config=_node_render_backend_config_for_tests(),
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report.docx",
            content={
                "metadata": {
                    "title": "Node Legacy Entry",
                    "document_style": {
                        "summary_card_defaults": {
                            "title_align": "center",
                            "title_emphasis": "strong",
                            "title_font_scale": 1.2,
                            "title_space_before": 12,
                            "title_space_after": 4,
                            "list_space_after": 8,
                        }
                    },
                },
                "blocks": [
                    {"type": "heading", "text": "一、经营总览", "level": 1},
                    {
                        "type": "table",
                        "headers": ["日期", "时间", "内容"],
                        "rows": [
                            ["第一天", "09:00", "课程 A"],
                            ["第二天", "13:00", "课程 B"],
                        ],
                    },
                    {
                        "type": "summary_card",
                        "title": "Highlights",
                        "items": ["Stable revenue", "Lower churn"],
                        "variant": "conclusion",
                    },
                ],
            },
            file_type="word",
        )
        assert result is None
        delivery_service.send_file_with_preview.assert_awaited_once()
        delivered_path = Path(delivery_service.send_file_with_preview.await_args.args[1])
        loaded_doc = docx.Document(delivered_path)
        table = loaded_doc.tables[0]

        assert any(
            paragraph.text == "Node Legacy Entry" for paragraph in loaded_doc.paragraphs
        )
        assert any(
            paragraph.text == "一、经营总览" for paragraph in loaded_doc.paragraphs
        )
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        summary_title = _find_paragraph(loaded_doc, "Highlights")
        summary_item = _find_paragraph(loaded_doc, "• Stable revenue")
        assert summary_title.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert summary_title.runs[0].bold is True
        assert summary_title.paragraph_format.space_before.pt == pytest.approx(12, abs=0.5)
        assert summary_title.paragraph_format.space_after.pt == pytest.approx(4, abs=0.5)
        assert summary_item.paragraph_format.space_after.pt == pytest.approx(8, abs=0.5)
        assert table.rows[0].cells[0].text == "日期"
        assert len(table.rows) >= 3
    finally:
        executor.shutdown(wait=False)
        shutil.rmtree(workspace_dir, ignore_errors=True)


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_returns_error_without_sending():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    event.send = AsyncMock()
    office_generator = MagicMock()
    delivery_service = MagicMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={},
        )

        result = await service.create_office_file(
            event,
            filename="report.unsupported",
            content="hello world",
            file_type="unknown",
        )
    finally:
        executor.shutdown(wait=False)

    assert result == "错误：不支持的文件类型 'unknown'。允许值：excel/powerpoint"
    event.send.assert_not_called()
    office_generator.generate.assert_not_called()
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_returns_error_when_generated_file_missing():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    office_generator = MagicMock()
    office_generator.generate = AsyncMock(return_value=None)
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report.docx",
            content="hello world",
            file_type="word",
        )
    finally:
        executor.shutdown(wait=False)

    assert result == "错误：文件生成失败，未找到输出文件"
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_requires_explicit_type_without_suffix():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    office_generator = MagicMock()
    delivery_service = MagicMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"openpyxl": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"openpyxl": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report",
            content="hello world",
            file_type="",
        )
    finally:
        executor.shutdown(wait=False)

    assert (
        result
        == "错误：未指定文件类型。请提供带后缀的文件名，或显式传入 file_type（excel/powerpoint）。"
    )
    office_generator.generate.assert_not_called()
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_rejects_word_fallback_without_suffix():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    office_generator = MagicMock()
    delivery_service = MagicMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report",
            content="hello world",
            file_type="word",
        )
    finally:
        executor.shutdown(wait=False)

    assert (
        result
        == "错误：Word 文档请直接提供 .docx/.doc 文件名，或改用 create_document → "
        "add_blocks → finalize_document → export_document。"
    )
    office_generator.generate.assert_not_called()
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_create_office_file_returns_direct_result_for_explicit_tool_error():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    event.get_extra.side_effect = lambda key, default=None: {
        EXPLICIT_FILE_TOOL_EVENT_KEY: "create_office_file"
    }.get(key, default)
    event.plain_result.side_effect = lambda text: f"DIRECT::{text}"
    office_generator = MagicMock()
    delivery_service = MagicMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_office_files": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=office_generator,
            pdf_converter=MagicMock(),
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.create_office_file(
            event,
            filename="report",
            content="hello world",
            file_type="word",
        )
    finally:
        executor.shutdown(wait=False)

    assert (
        result
        == "DIRECT::错误：Word 文档请直接提供 .docx/.doc 文件名，或改用 create_document → "
        "add_blocks → finalize_document → export_document。"
    )
    event.plain_result.assert_called_once()
    office_generator.generate.assert_not_called()
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_convert_to_pdf_returns_none_after_delivery():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    source_path = workspace_dir / "convert-source.docx"
    source_path.write_text("demo", encoding="utf-8")
    output_path = workspace_dir / "convert-source.pdf"
    pdf_converter = MagicMock()
    pdf_converter.is_available.return_value = True
    pdf_converter.office_to_pdf = AsyncMock(return_value=output_path)
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_pdf_conversion": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=MagicMock(),
            pdf_converter=pdf_converter,
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )
        output_path.write_text("pdf", encoding="utf-8")

        result = await service.convert_to_pdf(
            event,
            filename=source_path.name,
        )
    finally:
        source_path.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)
        executor.shutdown(wait=False)

    pdf_converter.office_to_pdf.assert_awaited_once_with(source_path)
    delivery_service.send_file_with_preview.assert_awaited_once()
    assert result is None


@pytest.mark.asyncio
async def test_file_tool_service_convert_to_pdf_returns_error_when_generated_file_missing():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    source_path = workspace_dir / "convert-source-missing.docx"
    source_path.write_text("demo", encoding="utf-8")
    pdf_converter = MagicMock()
    pdf_converter.is_available.return_value = True
    pdf_converter.office_to_pdf = AsyncMock(return_value=None)
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={"docx": object()},
            max_file_size=1024 * 1024,
            feature_settings={"enable_pdf_conversion": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=MagicMock(),
            pdf_converter=pdf_converter,
            delivery_service=delivery_service,
            office_libs={"docx": object()},
        )

        result = await service.convert_to_pdf(
            event,
            filename=source_path.name,
        )
    finally:
        source_path.unlink(missing_ok=True)
        executor.shutdown(wait=False)

    assert result == "错误：PDF 转换失败，未找到生成的 PDF 文件"
    delivery_service.send_file_with_preview.assert_not_called()


@pytest.mark.asyncio
async def test_file_tool_service_convert_from_pdf_returns_error_when_generated_file_missing():
    workspace_dir = Path(__file__).resolve().parent
    executor = ThreadPoolExecutor(max_workers=1)
    event = _build_event()
    source_path = workspace_dir / "convert-back-missing.pdf"
    source_path.write_text("pdf", encoding="utf-8")
    pdf_converter = MagicMock()
    pdf_converter.is_available.return_value = True
    pdf_converter.get_missing_dependencies.return_value = []
    pdf_converter.pdf_to_word = AsyncMock(return_value=None)
    delivery_service = MagicMock()
    delivery_service.send_file_with_preview = AsyncMock()

    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={"enable_pdf_conversion": True},
        )
        service = _build_file_tool_service(
            workspace_service=workspace_service,
            office_generator=MagicMock(),
            pdf_converter=pdf_converter,
            delivery_service=delivery_service,
            office_libs={},
        )

        result = await service.convert_from_pdf(
            event,
            filename=source_path.name,
            target_format="word",
        )
    finally:
        source_path.unlink(missing_ok=True)
        executor.shutdown(wait=False)

    assert result == "错误：PDF→Word 文档 转换失败，未找到生成的文件"
    delivery_service.send_file_with_preview.assert_not_called()


def test_command_service_lists_office_files_in_workspace():
    workspace_dir = Path(__file__).resolve().parent / "workspace-command-list"
    workspace_dir.mkdir(exist_ok=True)
    sample_file = workspace_dir / "report.docx"
    sample_file.write_text("demo", encoding="utf-8")

    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=workspace_dir,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.capabilities = {
            "office_to_pdf": False,
            "pdf_to_word": False,
            "pdf_to_excel": False,
        }
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=workspace_dir,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=MagicMock(),
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )

        result = service.list_files(_build_event())
    finally:
        sample_file.unlink(missing_ok=True)
        workspace_dir.rmdir()
        executor.shutdown(wait=False)

    assert "机器人工作区 Office 文件列表" in result
    assert "report.docx" in result


def test_command_service_builds_pdf_status_summary():
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=Path(__file__).resolve().parent,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.get_detailed_status.return_value = {
            "capabilities": {
                "office_to_pdf": True,
                "pdf_to_word": False,
                "pdf_to_excel": True,
            },
            "office_to_pdf_backend": "libreoffice",
            "word_backend": None,
            "excel_backend": "tabula",
            "is_windows": False,
            "java_available": True,
            "libreoffice_path": "/usr/bin/libreoffice",
            "libs": {"pdf2docx": False, "tabula-py": True},
        }
        pdf_converter.get_missing_dependencies.return_value = ["pdf2docx"]
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=Path(__file__).resolve().parent,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=MagicMock(),
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
        )

        result = service.pdf_status(_build_event())
    finally:
        executor.shutdown(wait=False)

    assert "Office→PDF: ✅ 可用 (libreoffice)" in result
    assert "PDF→Excel:  ✅ 可用 (tabula)" in result
    assert "缺失依赖" in result
    assert "pdf2docx" in result


def test_command_service_fileinfo_reports_word_toolchain_available():
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=Path(__file__).resolve().parent,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.capabilities = {
            "office_to_pdf": False,
            "pdf_to_word": False,
            "pdf_to_excel": False,
        }
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=Path(__file__).resolve().parent,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=MagicMock(),
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
            node_renderer_entry="D:/custom/renderer.js",
        )

        with patch(
            "astrbot_plugin_office_assistant.services.command_service.NodeDocumentRenderBackend.is_available",
            return_value=True,
        ):
            result = service.fileinfo(_build_event())
    finally:
        executor.shutdown(wait=False)

    assert "Word工具链: ✅ Node 渲染可用" in result
    assert "renderer.js" in result


def test_command_service_fileinfo_reports_word_toolchain_unavailable():
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        workspace_service = WorkspaceService(
            plugin_data_path=Path(__file__).resolve().parent,
            executor=executor,
            office_libs={},
            max_file_size=1024 * 1024,
            feature_settings={},
        )
        pdf_converter = MagicMock()
        pdf_converter.capabilities = {
            "office_to_pdf": False,
            "pdf_to_word": False,
            "pdf_to_excel": False,
        }
        service = CommandService(
            workspace_service=workspace_service,
            pdf_converter=pdf_converter,
            plugin_data_path=Path(__file__).resolve().parent,
            auto_delete=False,
            allow_external_input_files=False,
            enable_features_in_group=True,
            auto_block_execution_tools=True,
            reply_to_user=True,
            upload_session_service=MagicMock(),
            is_group_feature_enabled=lambda _event: True,
            check_permission=lambda _event: True,
            group_feature_disabled_error=lambda: "group disabled",
            node_renderer_entry="D:/custom/missing-renderer.js",
        )

        with patch(
            "astrbot_plugin_office_assistant.services.command_service.NodeDocumentRenderBackend.is_available",
            return_value=False,
        ):
            result = service.fileinfo(_build_event())
    finally:
        executor.shutdown(wait=False)

    assert "Word工具链: ❌ Node 渲染不可用" in result
    assert "missing-renderer.js" in result
