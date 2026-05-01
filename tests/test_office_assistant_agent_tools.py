import builtins
import importlib
import json
import shutil
import subprocess
import struct
import sys
import zipfile
import zlib
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch
from uuid import uuid4

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from astrbot_plugin_office_assistant.agent_tools import (
    build_document_toolset,
    build_workbook_toolset,
)
from astrbot_plugin_office_assistant.agent_tools.workbook_tools import (
    CreateWorkbookTool,
    ExportWorkbookTool,
    WriteRowsTool,
)
from astrbot_plugin_office_assistant.agent_tools.document_tools import (
    CreateDocumentTool,
)
from astrbot_plugin_office_assistant.constants import (
    EXCEL_SCRIPT_RETRY_EXHAUSTED_EVENT_KEY,
)
from astrbot_plugin_office_assistant.document_core.builders.table_renderer import (
    DOCX_TABLE_STYLES,
    TableRenderer,
)
from astrbot_plugin_office_assistant.document_core.macros.summary_card import (
    build_summary_card_group,
)
from astrbot_plugin_office_assistant.domain.document.session_store import (
    DocumentSessionStore,
)
from astrbot_plugin_office_assistant.domain.workbook.contracts import (
    CreateWorkbookRequest,
    MAX_WORKBOOK_ROW_INDEX,
)
from astrbot_plugin_office_assistant.domain.workbook.session_store import (
    WorkbookSessionStore,
)
from astrbot_plugin_office_assistant.domain.document.export_pipeline import (
    export_document_via_pipeline,
)
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendConfig,
    DocumentRenderBackendError,
    NodeDocumentRenderBackend,
    build_document_render_backends,
    RenderResult,
    build_document_render_payload,
)
from astrbot_plugin_office_assistant.document_core.models.blocks import (
    BlockStyle,
    BusinessReviewCoverData,
    ColumnBlock,
    ColumnsBlock,
    GroupBlock,
    HeaderFooterConfig,
    HeadingBlock,
    HeroBannerBlock,
    PageTemplateBlock,
    ParagraphBlock,
    ParagraphRun,
    SectionBreakBlock,
    SectionMarginsConfig,
    SummaryCardBlock,
    TableBlock,
    TocBlock,
)
from astrbot_plugin_office_assistant.document_core.models.document import (
    DocumentMetadata,
    DocumentModel,
    DocumentStyleConfig,
    DocumentSummaryCardDefaults,
)
from astrbot_plugin_office_assistant.domain.document.contracts import (
    AddBlocksRequest,
    AddHeadingRequest,
    AddListRequest,
    AddParagraphRequest,
    AddTableRequest,
    BlockHeadingInput,
    CreateDocumentRequest,
    ExportDocumentRequest,
    FinalizeDocumentRequest,
    SectionListInput,
    SectionParagraphInput,
    SectionTableInput,
    normalize_create_document_kwargs,
    normalize_raw_block_payloads,
)
import astrbot_plugin_office_assistant.mcp_server.server as mcp_server_module
from astrbot_plugin_office_assistant.mcp_server.server import (
    create_server,
)
from astrbot_plugin_office_assistant.tools.mcp_adapter import (
    register_document_tools_from_registry,
)
from astrbot_plugin_office_assistant.tools.astrbot_adapter import (
    build_document_toolset_from_registry,
)
from astrbot_plugin_office_assistant.tools.registry import (
    DocumentToolSpec,
    get_document_tool_specs,
)
from pydantic import ValidationError

from astrbot.core.utils.astrbot_path import get_astrbot_plugin_data_path


from tests._docx_test_helpers import *  # noqa: F401,F403
from tests._schema_test_helpers import _schema_contains_key, _schema_contains_type_list


def _build_agent_tool_context(*, excel_script_retry_exhausted: bool = False):
    event = MagicMock()
    event.get_extra.side_effect = lambda key, default=None: (
        excel_script_retry_exhausted
        if key == EXCEL_SCRIPT_RETRY_EXHAUSTED_EVENT_KEY
        else default
    )
    return SimpleNamespace(context=SimpleNamespace(event=event))


def test_build_document_toolset_uses_shared_store_and_default_workspace():
    toolset = build_document_toolset()
    tool_names = [tool.name for tool in toolset.tools]

    assert tool_names == [
        "create_document",
        "add_blocks",
        "finalize_document",
        "export_document",
    ]

    stores = [tool.store for tool in toolset.tools if hasattr(tool, "store")]
    assert len(stores) == len(tool_names)
    assert len({id(store) for store in stores}) == 1

    expected_workspace = (
        Path(get_astrbot_plugin_data_path())
        / "astrbot_plugin_office_assistant"
        / "documents"
    )
    assert stores[0].workspace_dir == expected_workspace


def test_build_workbook_toolset_uses_shared_store_and_default_workspace():
    toolset = build_workbook_toolset()
    tool_names = [tool.name for tool in toolset.tools]

    assert tool_names == [
        "create_workbook",
        "write_rows",
        "export_workbook",
    ]

    stores = [tool.store for tool in toolset.tools if hasattr(tool, "store")]
    assert len(stores) == len(tool_names)
    assert len({id(store) for store in stores}) == 1

    expected_workspace = (
        Path(get_astrbot_plugin_data_path())
        / "astrbot_plugin_office_assistant"
        / "workbooks"
    )
    assert stores[0].workspace_dir == expected_workspace


@pytest.mark.asyncio
async def test_write_rows_tool_returns_targeted_retry_message_for_validation_errors(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = WriteRowsTool(store=store)

    result = json.loads(
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["=SUM(A1:A2)"]],
        )
    )

    assert result["success"] is False
    assert "only fix invalid fields" in result["message"]
    assert "Original error" in result["message"]


@pytest.mark.asyncio
async def test_create_workbook_tool_wraps_expected_store_errors(
    workspace_root: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    tool = CreateWorkbookTool(store=store)
    monkeypatch.setattr(
        store, "create_workbook", MagicMock(side_effect=OSError("disk full"))
    )

    result = json.loads(await tool.call(None, filename="rows.xlsx"))

    assert result["success"] is False
    assert result["message"] == "create_workbook failed: disk full"


@pytest.mark.asyncio
async def test_create_workbook_tool_reraises_unexpected_store_errors(
    workspace_root: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    tool = CreateWorkbookTool(store=store)
    monkeypatch.setattr(
        store, "create_workbook", MagicMock(side_effect=RuntimeError("boom"))
    )

    with pytest.raises(RuntimeError, match="boom"):
        await tool.call(None, filename="rows.xlsx")


@pytest.mark.asyncio
async def test_workbook_tools_refuse_after_excel_script_retry_exhaustion(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    context = _build_agent_tool_context(excel_script_retry_exhausted=True)

    create_result = json.loads(
        await CreateWorkbookTool(store=store).call(context, filename="new.xlsx")
    )
    write_result = json.loads(
        await WriteRowsTool(store=store).call(
            context,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["value"]],
        )
    )
    export_result = json.loads(
        await ExportWorkbookTool(store=store).call(
            context,
            workbook_id=workbook.workbook_id,
        )
    )

    for result in (create_result, write_result, export_result):
        assert result["success"] is False
        assert "Excel 脚本重试次数已用尽" in result["message"]
        assert "请停止调用工具" in result["message"]


@pytest.mark.asyncio
async def test_write_rows_tool_preserves_invalid_zero_start_row(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = WriteRowsTool(store=store)

    result = json.loads(
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["value"]],
            start_row=0,
        )
    )

    assert result["success"] is False
    assert "only fix invalid fields" in result["message"]
    assert "greater than or equal to 1" in result["message"]


@pytest.mark.asyncio
async def test_write_rows_tool_rejects_oversized_start_row(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = WriteRowsTool(store=store)

    result = json.loads(
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["value"]],
            start_row=MAX_WORKBOOK_ROW_INDEX + 1,
        )
    )

    assert result["success"] is False
    assert "only fix invalid fields" in result["message"]
    assert "less than or equal to" in result["message"]


@pytest.mark.asyncio
async def test_write_rows_tool_returns_neutral_message_for_state_errors(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    workbook.status = "exported"
    tool = WriteRowsTool(store=store)

    result = json.loads(
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["value"]],
        )
    )

    assert result["success"] is False
    assert result["message"].startswith("write_rows failed:")
    assert "only fix invalid fields" not in result["message"]


@pytest.mark.asyncio
async def test_write_rows_tool_reraises_unexpected_store_errors(
    workspace_root: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = WriteRowsTool(store=store)
    monkeypatch.setattr(
        store, "write_rows", MagicMock(side_effect=RuntimeError("boom"))
    )

    with pytest.raises(RuntimeError, match="boom"):
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            sheet="Data",
            rows=[["value"]],
        )


@pytest.mark.asyncio
async def test_export_workbook_tool_wraps_expected_validation_errors(
    workspace_root: Path,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = ExportWorkbookTool(store=store)

    result = json.loads(
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
            output_name="C:/temp/final.xlsx",
        )
    )

    assert result["success"] is False
    assert result["message"].startswith("export_workbook failed:")
    assert "must not be an absolute path" in result["message"]


@pytest.mark.asyncio
async def test_export_workbook_tool_reraises_unexpected_store_errors(
    workspace_root: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    store = WorkbookSessionStore(workspace_dir=workspace_root)
    workbook = store.create_workbook(CreateWorkbookRequest(filename="rows.xlsx"))
    tool = ExportWorkbookTool(store=store)
    monkeypatch.setattr(
        store, "export_workbook", MagicMock(side_effect=RuntimeError("boom"))
    )

    with pytest.raises(RuntimeError, match="boom"):
        await tool.call(
            None,
            workbook_id=workbook.workbook_id,
        )


@pytest.mark.asyncio
async def test_document_tool_messages_enforce_workflow_after_create_and_finalize(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-messages")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Workflow Message",
            output_name="workflow-message.docx",
        )
    )
    document_id = created["document"]["document_id"]

    assert "下一步只能调用 add_blocks 添加内容" in created["message"]
    assert "不要提前调用 finalize_document 或 export_document" in created["message"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[{"type": "paragraph", "text": "正文"}],
    )
    finalized = json.loads(
        await tool_by_name["finalize_document"].call(
            None,
            document_id=document_id,
        )
    )

    assert "下一步只能调用 export_document 导出文件" in finalized["message"]
    assert (
        "不要再调用 add_blocks、create_document 或 finalize_document"
        in finalized["message"]
    )


@pytest.mark.asyncio
async def test_add_blocks_failure_message_keeps_model_on_same_tool(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-tools-add-blocks-failure"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Add Blocks Failure",
            output_name="add-blocks-failure.docx",
        )
    )

    failed = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks=[{"type": "table", "style": {"table_grid": "none"}, "rows": []}],
        )
    )

    assert failed["success"] is False
    assert "继续使用同一个 document_id 再次调用 add_blocks" in failed["message"]
    assert "不要改调 finalize_document 或 export_document" in failed["message"]


@pytest.mark.asyncio
async def test_add_blocks_tool_accepts_json_string_blocks(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-json-blocks")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            output_name="json-blocks.docx",
        )
    )

    added = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks='[{"type": "paragraph", "text": "正文"}]',
        )
    )

    assert added["success"] is True
    assert added["document"]["block_count"] == 1


@pytest.mark.asyncio
async def test_create_document_tool_accepts_json_string_document_style(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-json-style")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            output_name="json-style.docx",
            document_style=(
                '{"font_name": "Microsoft YaHei", "heading_color": "1F4E79"}'
            ),
        )
    )

    assert created["success"] is True
    assert created["document"]["document_style"]["heading_color"] == "1F4E79"


@pytest.mark.asyncio
async def test_create_document_tool_returns_incrementing_short_document_ids():
    tool = CreateDocumentTool()

    first = json.loads(await tool.call(None, title="第一份"))
    second = json.loads(await tool.call(None, title="第二份"))

    assert first["document"]["document_id"] == "doc-1"
    assert second["document"]["document_id"] == "doc-2"


def test_document_tool_registry_keeps_document_tool_order():
    assert [spec.name for spec in get_document_tool_specs()] == [
        "create_document",
        "add_blocks",
        "finalize_document",
        "export_document",
    ]


def test_astrbot_toolset_preserves_document_tool_registry_order():
    toolset = build_document_toolset_from_registry()

    assert [tool.name for tool in toolset.tools] == [
        spec.name for spec in get_document_tool_specs()
    ]


def test_build_document_toolset_defaults_to_node_only_for_word():
    toolset = build_document_toolset()
    export_tool = next(tool for tool in toolset.tools if tool.name == "export_document")

    assert [backend.name for backend in export_tool.render_backends] == ["node"]


@pytest.mark.parametrize(
    ("export_name", "expected_name"),
    [
        ("DocumentRenderBackendConfig", "DocumentRenderBackendConfig"),
        ("DocumentRenderBackend", "DocumentRenderBackend"),
        ("NodeDocumentRenderBackend", "NodeDocumentRenderBackend"),
    ],
)
def test_domain_document_exports_document_render_interfaces(
    export_name: str,
    expected_name: str,
):
    import astrbot_plugin_office_assistant.domain.document as domain_document

    exported = getattr(domain_document, export_name)
    assert getattr(exported, "__name__", None) == expected_name
    assert export_name in getattr(domain_document, "__all__", [])


@pytest.mark.parametrize(
    ("export_name", "expected_name"),
    [
        ("DocumentRenderBackendConfig", "DocumentRenderBackendConfig"),
        ("DocumentRenderBackend", "DocumentRenderBackend"),
        ("DocumentRenderBackendError", "DocumentRenderBackendError"),
        ("NodeDocumentRenderBackend", "NodeDocumentRenderBackend"),
    ],
)
def test_render_backends_exports_document_render_interfaces(
    export_name: str,
    expected_name: str,
):
    import astrbot_plugin_office_assistant.domain.document.render_backends as render_backends

    exported = getattr(render_backends, export_name)
    assert getattr(exported, "__name__", None) == expected_name
    assert export_name in getattr(render_backends, "__all__", [])


def test_registry_import_does_not_require_fastmcp_at_runtime():
    module_name = "astrbot_plugin_office_assistant.tools.registry"
    cached_module = sys.modules.pop(module_name, None)
    original_import = builtins.__import__

    def _guard_fastmcp_import(name, globals=None, locals=None, fromlist=(), level=0):
        if name == "mcp.server.fastmcp":
            raise ModuleNotFoundError("fastmcp unavailable")
        return original_import(name, globals, locals, fromlist, level)

    try:
        with patch("builtins.__import__", side_effect=_guard_fastmcp_import):
            imported = importlib.import_module(module_name)
        assert [spec.name for spec in imported.get_document_tool_specs()] == [
            "create_document",
            "add_blocks",
            "finalize_document",
            "export_document",
        ]
    finally:
        sys.modules.pop(module_name, None)
        if cached_module is not None:
            sys.modules[module_name] = cached_module
        else:
            importlib.import_module(module_name)


def test_astrbot_toolset_passes_export_hooks_and_callback(
    monkeypatch: pytest.MonkeyPatch,
):
    captured_kwargs: dict[str, object] = {}
    before_hooks = [MagicMock()]
    after_hooks = [MagicMock()]

    async def _after_export(_context, _path):
        return None

    base_specs = get_document_tool_specs()

    def _record_export_factory(
        store, before_export_hooks, after_export_hooks, after_export
    ):
        captured_kwargs["store"] = store
        captured_kwargs["before_export_hooks"] = before_export_hooks
        captured_kwargs["after_export_hooks"] = after_export_hooks
        captured_kwargs["after_export"] = after_export
        return CreateDocumentTool(store=store, name="export_document")

    patched_specs = base_specs[:-1] + (
        DocumentToolSpec(
            name="export_document",
            astrbot_factory=_record_export_factory,
            mcp_registrar=base_specs[-1].mcp_registrar,
        ),
    )

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.tools.astrbot_adapter.get_document_tool_specs",
        lambda: patched_specs,
    )

    toolset = build_document_toolset_from_registry(
        before_export_hooks=before_hooks,
        after_export_hooks=after_hooks,
        after_export=_after_export,
    )

    assert captured_kwargs["store"] is toolset.document_store
    assert captured_kwargs["before_export_hooks"] is before_hooks
    assert captured_kwargs["after_export_hooks"] is after_hooks
    assert captured_kwargs["after_export"] is _after_export


def test_mcp_document_tool_registration_matches_registry_order(
    monkeypatch: pytest.MonkeyPatch,
):
    registered_names: list[str] = []

    def _make_registrar(name: str):
        def _record(*_args, **_kwargs):
            registered_names.append(name)

        return _record

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.create_document.register_create_document_tool",
        _make_registrar("create_document"),
    )
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.add_blocks.register_add_blocks_tool",
        _make_registrar("add_blocks"),
    )
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.finalize_document.register_finalize_document_tool",
        _make_registrar("finalize_document"),
    )
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.export_document.register_export_document_tool",
        _make_registrar("export_document"),
    )

    register_document_tools_from_registry(
        server=MagicMock(),
        store=DocumentSessionStore(),
    )

    assert registered_names == [spec.name for spec in get_document_tool_specs()]


def test_mcp_document_tool_registration_passes_export_hooks(
    monkeypatch: pytest.MonkeyPatch,
):
    export_call_kwargs: dict[str, object] = {}
    before_hooks = [MagicMock()]
    after_hooks = [MagicMock()]

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.create_document.register_create_document_tool",
        lambda *_args, **_kwargs: None,
    )
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.add_blocks.register_add_blocks_tool",
        lambda *_args, **_kwargs: None,
    )
    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.finalize_document.register_finalize_document_tool",
        lambda *_args, **_kwargs: None,
    )

    def _record_export(*_args, **kwargs):
        export_call_kwargs.update(kwargs)

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.tools.export_document.register_export_document_tool",
        _record_export,
    )

    register_document_tools_from_registry(
        server=MagicMock(),
        store=DocumentSessionStore(),
        before_export_hooks=before_hooks,
        after_export_hooks=after_hooks,
    )

    assert export_call_kwargs["before_export_hooks"] is before_hooks
    assert export_call_kwargs["after_export_hooks"] is after_hooks


@pytest.mark.asyncio
async def test_create_document_tool_does_not_stringify_missing_session():
    tool = CreateDocumentTool()

    created = json.loads(await tool.call(None, title="No Session"))

    assert created["success"] is True
    document = tool.store.require_document(created["document"]["document_id"])
    assert document.session_id == ""


@pytest.mark.asyncio
async def test_document_toolset_smoke_export(workspace_root: Path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Pytest Smoke",
            output_name="pytest-smoke.docx",
            theme_name="executive_brief",
            table_template="minimal",
            density="compact",
            accent_color="#AA5500",
        )
    )
    document_id = created["document"]["document_id"]
    assert created["document"]["theme_name"] == "executive_brief"
    assert created["document"]["table_template"] == "minimal"
    assert created["document"]["density"] == "compact"
    assert created["document"]["accent_color"] == "AA5500"

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "Section 1", "level": 1},
            {
                "type": "paragraph",
                "text": "Hello from pytest.",
                "style": {
                    "align": "center",
                    "emphasis": "strong",
                    "font_scale": 1.1,
                },
                "layout": {"spacing_after": 9},
            },
            {
                "type": "list",
                "items": ["Point A", "Point B"],
                "ordered": True,
                "style": {"emphasis": "subtle"},
            },
            {
                "type": "table",
                "headers": ["Metric", "Jan", "Feb"],
                "rows": [["Users", "120", "140"]],
                "table_style": "minimal",
            },
            _summary_card_block(items=["The new layout should look more intentional."]),
            {"type": "page_break"},
            {"type": "heading", "text": "Appendix", "level": 1},
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    assert exported["success"] is True
    assert Path(exported["file_path"]).exists()
    assert Path(exported["file_path"]).parent == workspace_dir
    loaded_doc = docx.Document(exported["file_path"])
    assert loaded_doc.paragraphs[0].text == "Pytest Smoke"
    assert loaded_doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert loaded_doc.paragraphs[0].runs[0].bold is True
    assert loaded_doc.paragraphs[0].runs[0].font.color.rgb == RGBColor.from_string(
        "000000"
    )
    assert loaded_doc.paragraphs[1].text == "Section 1"
    assert loaded_doc.paragraphs[1].runs[0].bold is True
    assert loaded_doc.paragraphs[2].text == "Hello from pytest."
    assert loaded_doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert loaded_doc.paragraphs[2].runs[0].bold is True
    assert loaded_doc.paragraphs[
        2
    ].paragraph_format.first_line_indent.pt == pytest.approx(18, abs=0.5)
    assert loaded_doc.paragraphs[2].paragraph_format.space_after.pt == pytest.approx(
        9, abs=0.5
    )


@pytest.mark.asyncio
async def test_create_document_tool_uses_configured_default_fonts_when_missing():
    toolset = build_document_toolset(
        default_document_style={
            "font_name": "Arial",
            "heading_font_name": "Arial",
            "table_font_name": "Arial",
            "code_font_name": "JetBrains Mono",
        }
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Font Defaults",
            output_name="font-defaults.docx",
            theme_name="business_report",
            document_style={"heading_color": "1F4E79"},
        )
    )

    assert created["document"]["document_style"]["font_name"] == "Arial"
    assert created["document"]["document_style"]["heading_font_name"] == "Arial"
    assert created["document"]["document_style"]["table_font_name"] == "Arial"
    assert created["document"]["document_style"]["code_font_name"] == "JetBrains Mono"
    assert created["document"]["document_style"]["heading_color"] == "1F4E79"


@pytest.mark.asyncio
async def test_create_document_tool_applies_document_style_defaults(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    workspace_dir = _make_workspace(workspace_root, "pytest-document-style")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Styled Report",
            output_name="styled-report.docx",
            theme_name="business_report",
            document_style={
                "brief": "deep blue business report",
                "heading_color": "0F4C81",
                "title_align": "left",
                "body_font_size": 12,
                "body_line_spacing": 1.25,
                "paragraph_space_after": 14,
                "list_space_after": 11,
                "summary_card_defaults": _summary_card_defaults(),
                "table_defaults": {
                    "preset": "minimal",
                    "header_fill": "DCE6F1",
                    "header_text_color": "123456",
                    "banded_rows": True,
                    "banded_row_fill": "EEF4FA",
                    "first_column_bold": True,
                    "table_align": "left",
                    "border_style": "standard",
                    "caption_emphasis": "strong",
                    "cell_align": "center",
                },
            },
        )
    )
    document_id = created["document"]["document_id"]
    assert created["document"]["document_style"]["brief"] == "deep blue business report"
    assert created["document"]["document_style"]["heading_color"] == "0F4C81"
    assert created["document"]["document_style"]["title_align"] == "left"

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "Overview", "level": 1},
            {"type": "paragraph", "text": "Styled body paragraph."},
            {"type": "list", "items": ["Alpha", "Beta"]},
            _summary_card_block(
                title="Highlights",
                items=["Stable revenue", "Lower churn"],
            ),
            {
                "type": "table",
                "caption": "Quarterly Summary",
                "headers": ["Region", "Score"],
                "rows": [["East", "92"], ["West", "88"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    heading_paragraph = _find_paragraph(loaded_doc, "Overview")
    title_paragraph = _find_paragraph(loaded_doc, "Styled Report")
    body_paragraph = _find_paragraph(loaded_doc, "Styled body paragraph.")
    list_paragraph = _find_paragraph(loaded_doc, "• Alpha")
    summary_title_paragraph = _find_paragraph(loaded_doc, "Highlights")
    summary_item_paragraph = _find_paragraph(loaded_doc, "• Stable revenue")
    table = loaded_doc.tables[0]

    assert title_paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert _paragraph_run_rgb(title_paragraph) == "0F4C81"
    assert _paragraph_run_rgb(heading_paragraph) == "0F4C81"
    assert _paragraph_run_size(body_paragraph) == 12
    assert float(body_paragraph.paragraph_format.line_spacing) == pytest.approx(1.25)
    assert body_paragraph.paragraph_format.space_after.pt == pytest.approx(14, abs=0.5)
    assert list_paragraph.paragraph_format.space_after.pt == pytest.approx(11, abs=0.5)
    assert summary_title_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert summary_title_paragraph.runs[0].bold is True
    assert _paragraph_run_size(summary_title_paragraph) == pytest.approx(14.0, abs=0.5)
    assert summary_title_paragraph.paragraph_format.space_before.pt == pytest.approx(
        12, abs=0.5
    )
    assert summary_title_paragraph.paragraph_format.space_after.pt == pytest.approx(
        4, abs=0.5
    )
    assert summary_item_paragraph.paragraph_format.space_after.pt == pytest.approx(
        8, abs=0.5
    )
    assert _cell_fill(table.rows[0].cells[0]) == "DCE6F1"
    assert _run_rgb(table.rows[0].cells[0]) == "123456"
    assert _cell_fill(table.rows[2].cells[0]) == "EEF4FA"
    assert _run_bold(table.rows[2].cells[0]) is True
    assert _run_bold(table.rows[2].cells[1]) is False
    assert table.rows[2].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _table_border_size(table, "top") == "8"
    assert _table_border_color(table, "top") == "7A7A7A"


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_unfilled_nonbold_table_headers(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-header-style")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="表头样式",
            output_name="table-header-style.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "headers": ["Date", "Time", "Session Title"],
                "rows": [["Sep 20", "09:00", "Kickoff"]],
                "header_fill_enabled": False,
                "header_text_color": "888888",
                "header_bold": False,
                "border_style": "minimal",
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    header_cell = loaded_doc.tables[0].rows[0].cells[0]

    assert _cell_fill(header_cell) is None
    assert _run_rgb(header_cell) == "888888"
    assert _run_bold(header_cell) is False


@pytest.mark.asyncio
async def test_create_document_tool_prefers_table_block_over_document_style_defaults(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    workspace_dir = _make_workspace(workspace_root, "pytest-document-style-precedence")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Table Precedence",
            output_name="table-precedence.docx",
            document_style={
                "table_defaults": {
                    "header_fill": "DCE6F1",
                    "header_text_color": "123456",
                    "banded_rows": True,
                    "banded_row_fill": "EEF4FA",
                    "first_column_bold": True,
                    "table_align": "left",
                    "border_style": "standard",
                    "caption_emphasis": "normal",
                    "cell_align": "center",
                },
            },
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "caption": "Override Table",
                "headers": ["Metric", "Value"],
                "rows": [["North", "100"], ["South", "200"]],
                "header_fill": "1F4E79",
                "header_text_color": "FFFFFF",
                "banded_rows": False,
                "first_column_bold": False,
                "table_align": "center",
                "border_style": "strong",
                "caption_emphasis": "strong",
                "style": {"cell_align": "right"},
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
    assert _cell_fill(table.rows[0].cells[0]) == "1F4E79"
    assert _run_rgb(table.rows[0].cells[0]) == "FFFFFF"
    assert _cell_fill(table.rows[2].cells[0]) is None
    assert _run_bold(table.rows[2].cells[0]) is False
    assert table.rows[2].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert _table_border_size(table, "top") == "16"
    assert _table_border_color(table, "top") == "1F4E79"


@pytest.mark.asyncio
async def test_create_document_tool_uses_theme_banded_fill_when_block_enables_banding(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-document-style-banded-fill-fallback"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Banded Fill Fallback",
            output_name="banded-fill-fallback.docx",
            document_style={
                "table_defaults": {
                    "banded_rows": True,
                    "banded_row_fill": "EEF4FA",
                },
            },
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "headers": ["Metric", "Value"],
                "rows": [["North", "100"], ["South", "200"]],
                "banded_rows": True,
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert _cell_fill(table.rows[1].cells[0]) == "EEF4FA"
    assert _cell_fill(table.rows[2].cells[0]) is None


def test_build_summary_card_group_prefers_block_style_over_defaults():
    group = build_summary_card_group(
        title="Summary Override",
        items=["Item A"],
        style=BlockStyle(
            align="right",
            emphasis="normal",
            font_scale=1.3,
        ),
        title_align="center",
        title_emphasis="strong",
        title_font_scale=1.05,
    )

    title_block = group.blocks[0]
    list_block = group.blocks[1]

    assert title_block.style.align == "right"
    assert title_block.style.emphasis == "normal"
    assert title_block.style.font_scale == pytest.approx(1.3)
    assert list_block.style.emphasis == "normal"


@pytest.mark.asyncio
async def test_create_document_tool_applies_border_style_color_mapping(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-document-style-borders")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Border Mapping",
            output_name="border-mapping.docx",
            accent_color="#AA5500",
        )
    )

    await tool_by_name["add_blocks"].call(
        None,
        document_id=created["document"]["document_id"],
        blocks=[
            {
                "type": "table",
                "caption": "Minimal Table",
                "headers": ["Metric", "Value"],
                "rows": [["North", "100"]],
                "border_style": "minimal",
            },
            {
                "type": "table",
                "caption": "Standard Table",
                "headers": ["Metric", "Value"],
                "rows": [["East", "120"]],
                "border_style": "standard",
            },
            {
                "type": "table",
                "caption": "Strong Table",
                "headers": ["Metric", "Value"],
                "rows": [["West", "140"]],
                "border_style": "strong",
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    minimal_table, standard_table, strong_table = loaded_doc.tables

    assert _table_border_size(minimal_table, "top") == "4"
    assert _table_border_color(minimal_table, "top") == "D0D7DE"
    assert _table_border_size(standard_table, "top") == "8"
    assert _table_border_color(standard_table, "top") == "7A7A7A"
    assert _table_border_size(strong_table, "top") == "16"
    assert _table_border_color(strong_table, "top") == "AA5500"


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("table_template", "expected_border_size", "expected_left_margin"),
    [
        ("report_grid", "3", "108"),
        ("minimal", "2", "72"),
    ],
)
async def test_create_document_tool_applies_light_default_table_borders_and_margins(
    workspace_root: Path,
    table_template: str,
    expected_border_size: str,
    expected_left_margin: str,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, f"pytest-document-default-light-table-{table_template}"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Default Light Table",
            output_name=f"default-light-table-{table_template}.docx",
            table_template=table_template,
        )
    )

    await tool_by_name["add_blocks"].call(
        None,
        document_id=created["document"]["document_id"],
        blocks=[
            {
                "type": "table",
                "headers": ["Metric", "Value"],
                "rows": [["North", "100"], ["South", "200"]],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert _table_border_size(table, "top") == expected_border_size
    assert _table_border_color(table, "top") == "D9E1E8"
    assert _table_cell_margin(table, "left") == expected_left_margin
    assert _table_cell_margin(table, "right") == expected_left_margin


@pytest.mark.asyncio
async def test_create_document_tool_prefers_summary_block_over_document_style_defaults(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    workspace_dir = _make_workspace(
        workspace_root, "pytest-document-style-summary-precedence"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Summary Precedence",
            output_name="summary-precedence.docx",
            document_style={
                "summary_card_defaults": _summary_card_defaults(),
            },
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                **_summary_card_block(
                    title="Summary Override",
                    items=["Item A", "Item B"],
                ),
                "style": {
                    "align": "right",
                    "emphasis": "normal",
                },
                "layout": {
                    "spacing_before": 20,
                },
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    summary_title_paragraph = _find_paragraph(loaded_doc, "Summary Override")
    summary_item_paragraph = _find_paragraph(loaded_doc, "• Item A")

    assert summary_title_paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert summary_title_paragraph.runs[0].bold is False
    assert summary_title_paragraph.paragraph_format.space_before.pt == pytest.approx(
        20, abs=0.5
    )
    assert summary_title_paragraph.paragraph_format.space_after.pt == pytest.approx(
        4, abs=0.5
    )
    assert summary_item_paragraph.paragraph_format.space_after.pt == pytest.approx(
        8, abs=0.5
    )


@pytest.mark.asyncio
async def test_create_document_tool_applies_compact_density_and_paragraph_variants(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-document-compact-summary-variants"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Compact Summary",
            output_name="compact-summary.docx",
            theme_name="business_report",
            density="compact",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "paragraph",
                "variant": "summary_box",
                "text": "Compact summary body.",
            },
            {
                "type": "paragraph",
                "variant": "key_takeaway",
                "title": "Action Note",
                "text": "Execute the next step.",
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    summary_title = _find_paragraph(loaded_doc, "Summary")
    summary_item = _find_paragraph(loaded_doc, "• Compact summary body.")
    takeaway_title = _find_paragraph(loaded_doc, "Action Note")
    takeaway_item = _find_paragraph(loaded_doc, "• Execute the next step.")
    section = loaded_doc.sections[0]

    assert section.top_margin.cm == pytest.approx(2.2, abs=0.01)
    assert section.bottom_margin.cm == pytest.approx(2.1, abs=0.01)
    assert section.left_margin.cm == pytest.approx(2.4, abs=0.01)
    assert section.right_margin.cm == pytest.approx(2.3, abs=0.01)
    assert summary_title.runs[0].bold is True
    assert summary_item.paragraph_format.space_after.pt == pytest.approx(4, abs=0.5)
    assert takeaway_title.runs[0].bold is True
    assert takeaway_item.text == "• Execute the next step."


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_rich_text_paragraph_runs(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-rich-paragraph")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="段落富文本",
            output_name="rich-paragraph.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "paragraph",
                "runs": [
                    {"text": "粗体", "bold": True},
                    {"text": " / "},
                    {"text": "斜体", "italic": True},
                    {"text": " / "},
                    {"text": "下划线", "underline": True},
                    {"text": " / "},
                    {"text": "代码", "code": True},
                ],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    rich_paragraph = loaded_doc.paragraphs[1]

    assert rich_paragraph.runs[0].bold is True
    assert rich_paragraph.runs[2].italic is True
    assert rich_paragraph.runs[4].underline is True
    assert rich_paragraph.runs[6].font.name == "Consolas"
    assert rich_paragraph.runs[0].text == "粗体"
    assert rich_paragraph.runs[6].text == "代码"


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_paragraph_run_colors(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-rich-paragraph-color")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="段落颜色",
            output_name="rich-paragraph-color.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "paragraph",
                "runs": [
                    {"text": "标题", "bold": True, "color": "666666"},
                    {"text": ": "},
                    {"text": "Advanced Project Management"},
                ],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    rich_paragraph = loaded_doc.paragraphs[1]

    assert rich_paragraph.runs[0].font.color.rgb == docx.shared.RGBColor.from_string(
        "666666"
    )
    assert rich_paragraph.runs[2].font.color.rgb is None


@pytest.mark.asyncio
async def test_add_blocks_tool_normalizes_legacy_paragraph_bottom_border_aliases(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-paragraph-border-alias"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="段落边框别名",
            output_name="paragraph-border-alias.docx",
        )
    )
    document_id = created["document"]["document_id"]

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=[
                {
                    "type": "paragraph",
                    "text": "带底边框正文",
                    "bottom_border": True,
                    "bottom_border_style": "single",
                    "bottom_border_color": "1F4E79",
                    "bottom_border_size_pt": 1.0,
                }
            ],
        )
    )

    assert add_blocks_result["success"] is True

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph = _find_paragraph(loaded_doc, "带底边框正文")

    assert _paragraph_bottom_border_color(paragraph) == "1F4E79"
    assert _paragraph_bottom_border_size(paragraph) == "8"


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_rich_text_list_items(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-rich-list")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="列表富文本",
            output_name="rich-list.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "list",
                "ordered": True,
                "items": [
                    {
                        "runs": [
                            {"text": "Enhance project management skills", "bold": True},
                            {
                                "text": " by exploring advanced methodologies such as Agile and Scrum.",
                            },
                        ]
                    },
                    "Standard retrospective and action tracking.",
                ],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    first_item = _find_paragraph(
        loaded_doc,
        "1. Enhance project management skills by exploring advanced methodologies such as Agile and Scrum.",
    )
    second_item = _find_paragraph(
        loaded_doc, "2. Standard retrospective and action tracking."
    )

    assert first_item.runs[0].text == "1. "
    assert first_item.runs[1].bold is True
    assert first_item.runs[1].text == "Enhance project management skills"
    assert first_item.runs[2].bold in {False, None}
    assert second_item.text.startswith("2. ")


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_nested_primitives(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-add-blocks")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Composite Blocks",
            output_name="composite-blocks.docx",
        )
    )
    document_id = created["document"]["document_id"]

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=[
                {"type": "heading", "text": "Overview", "level": 1},
                {
                    "type": "group",
                    "blocks": [
                        {"type": "paragraph", "text": "Nested intro."},
                        {
                            "type": "list",
                            "items": ["Left detail", "Right detail"],
                            "ordered": False,
                        },
                    ],
                },
                {
                    "type": "columns",
                    "columns": [
                        {
                            "blocks": [
                                {"type": "paragraph", "text": "Column A body."},
                            ]
                        },
                        {
                            "blocks": [
                                {"type": "paragraph", "text": "Column B body."},
                            ]
                        },
                    ],
                },
                {"type": "page_break"},
                {"type": "heading", "text": "Appendix", "level": 1},
            ],
        )
    )

    assert add_blocks_result["success"] is True
    assert add_blocks_result["document"]["block_count"] == 5

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph_texts = [paragraph.text for paragraph in loaded_doc.paragraphs]

    assert paragraph_texts[1] == "Overview"
    assert "Nested intro." in paragraph_texts
    assert "• Left detail" in paragraph_texts
    assert "• Right detail" in paragraph_texts
    assert "Column A body." in paragraph_texts
    assert "Column B body." in paragraph_texts
    assert "Appendix" in paragraph_texts
    assert 'w:type="page"' in loaded_doc.element.body.xml


@pytest.mark.asyncio
async def test_add_blocks_tool_supports_enhanced_tables(workspace_root: Path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-enhanced-table")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="增强表格",
            output_name="enhanced-table.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "caption": "季度经营指标",
                "headers": ["区域", "目标", "完成率"],
                "rows": [["华东", "120", "98%"], ["华南", "88", "103%"]],
                "column_widths": [4.2, 3.0, 3.0],
                "numeric_columns": [1, 2],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    assert len(loaded_doc.tables) == 1

    table = loaded_doc.tables[0]
    assert table.rows[0].cells[0].text == "季度经营指标"
    assert table.rows[1].cells[0].text == "区域"
    assert table.rows[2].cells[0].text == "华东"
    assert _row_is_repeated_header(table.rows[0]) is True
    assert _row_has_cant_split(table.rows[0]) is True
    assert _paragraph_has_keep_next(table.rows[0].cells[0].paragraphs[0]) is True
    assert _row_is_repeated_header(table.rows[1]) is True
    assert _row_has_cant_split(table.rows[1]) is True
    assert table.rows[2].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.rows[3].cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert abs(table.rows[1].cells[0].width - Cm(4.2)) < 20000
    assert abs(table.rows[1].cells[1].width - Cm(3.0)) < 20000
    assert _cell_fill(table.rows[2].cells[0]) == "F7FBFF"
    assert _cell_fill(table.rows[3].cells[0]) is None


@pytest.mark.asyncio
@pytest.mark.parametrize("table_style", ["report_grid", "metrics_compact", "minimal"])
async def test_add_blocks_tool_supports_grouped_table_headers(
    workspace_root: Path,
    table_style: str,
):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    workspace_dir = _make_workspace(
        workspace_root, f"pytest-agent-grouped-table-{table_style}"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="分组表头",
            output_name=f"grouped-table-{table_style}.docx",
            table_template=table_style,
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "caption": "季度经营指标",
                "header_groups": [
                    {"title": "经营数据", "span": 2},
                    {"title": "结果", "span": 2},
                ],
                "headers": ["区域", "目标", "完成值", "完成率"],
                "rows": [
                    ["华东", "120", "118", "98%"],
                    ["华南", "88", "91", "103%"],
                ],
                "column_widths": [3.2, 2.4, 2.4, 2.4],
                "numeric_columns": [1, 2, 3],
                "table_style": table_style,
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert len(table.rows) == 5
    assert table.rows[0].cells[0].text == "季度经营指标"
    assert _row_is_repeated_header(table.rows[0]) is True
    assert _row_has_cant_split(table.rows[0]) is True
    assert _paragraph_has_keep_next(table.rows[0].cells[0].paragraphs[0]) is True
    assert table.rows[1].cells[0].text == "经营数据"
    assert _grid_span(table.rows[1].cells[0]) == 2
    assert table.rows[1].cells[2].text == "结果"
    assert _grid_span(table.rows[1].cells[2]) == 2
    assert _row_is_repeated_header(table.rows[1]) is True
    assert _row_has_cant_split(table.rows[1]) is True
    assert table.rows[2].cells[0].text == "区域"
    assert _row_is_repeated_header(table.rows[2]) is True
    assert _row_has_cant_split(table.rows[2]) is True
    assert _row_is_repeated_header(table.rows[3]) is False
    assert _row_has_cant_split(table.rows[3]) is True
    assert _row_is_repeated_header(table.rows[4]) is False
    assert _row_has_cant_split(table.rows[4]) is True
    assert table.rows[3].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.rows[4].cells[3].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert abs(table.rows[2].cells[0].width - Cm(3.2)) < 20000


@pytest.mark.asyncio
async def test_add_blocks_tool_marks_standard_header_row_as_repeated_and_non_split(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-header-repeat")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="表头重复",
            output_name="table-header-repeat.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"], ["华南", "980"]],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    # Data rows should not be marked as repeated headers
    assert _row_is_repeated_header(table.rows[1]) is False
    assert _row_is_repeated_header(table.rows[2]) is False

    assert _row_is_repeated_header(table.rows[0]) is True
    assert _row_has_cant_split(table.rows[0]) is True
    assert _row_has_cant_split(table.rows[1]) is True
    assert _row_has_cant_split(table.rows[2]) is True


def test_add_blocks_tool_schema_hides_table_cell_spans():
    toolset = build_document_toolset()
    add_blocks_tool = next(tool for tool in toolset.tools if tool.name == "add_blocks")
    block_schema = add_blocks_tool.parameters["properties"]["blocks"]["items"]
    properties = block_schema["properties"]
    list_item_schema = properties["items"]["items"]
    row_cell_schema = properties["rows"]["items"]["items"]
    resume_sections = properties["data"]["properties"]["sections"]["items"][
        "properties"
    ]
    resume_section_schema = properties["data"]["properties"]["sections"]["items"]
    resume_detail_schema = resume_sections["entries"]["items"]["properties"]["details"][
        "items"
    ]
    resume_line_schema = resume_sections["lines"]["items"]

    assert row_cell_schema == {"type": "string"}
    assert list_item_schema == {"type": "string"}
    assert resume_section_schema["required"] == ["title"]
    assert resume_sections["title"]["description"].endswith(
        "Use title here; do not use heading."
    )
    assert "type" not in resume_detail_schema
    assert (
        resume_detail_schema["properties"]["runs"]["items"]["properties"]["bold"][
            "type"
        ]
        == "boolean"
    )
    assert "type" not in resume_line_schema
    assert (
        resume_line_schema["properties"]["runs"]["items"]["properties"]["bold"]["type"]
        == "boolean"
    )
    assert not _schema_contains_key(add_blocks_tool.parameters, "anyOf")
    assert not _schema_contains_key(add_blocks_tool.parameters, "oneOf")
    assert not _schema_contains_type_list(add_blocks_tool.parameters)
    assert not _schema_contains_key(row_cell_schema, "row_span")
    assert not _schema_contains_key(row_cell_schema, "col_span")


@pytest.mark.asyncio
async def test_add_blocks_tool_rejects_table_body_row_span_vertical_merge(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-row-span")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="课程日程",
            output_name="table-row-span.docx",
        )
    )
    document_id = created["document"]["document_id"]

    result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=[
                {
                    "type": "table",
                    "headers": ["Date", "Time", "Session Title"],
                    "rows": [
                        [
                            {"text": "September 20", "row_span": 2},
                            "09:00 AM - 12:00 PM",
                            "Advanced Project Management",
                        ],
                        ["01:00 PM - 04:00 PM", "Stakeholder Alignment Workshop"],
                    ],
                }
            ],
        )
    )

    assert result["success"] is False
    assert "row_span" in result["message"]


@pytest.mark.asyncio
async def test_add_blocks_tool_rejects_table_body_col_span_horizontal_merge(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-table-row-span-placeholder"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="课程日程占位",
            output_name="table-row-span-placeholder.docx",
        )
    )
    document_id = created["document"]["document_id"]

    result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=[
                {
                    "type": "table",
                    "headers": ["Quarter", "Revenue", "Profit", "Status"],
                    "rows": [
                        [
                            {"text": "Q3 Summary", "col_span": 2},
                            "18%",
                            "On Track",
                        ],
                    ],
                }
            ],
        )
    )

    assert result["success"] is False
    assert "col_span" in result["message"]


def test_table_renderer_sets_cant_split_value_to_true_even_when_row_had_false():
    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document()
    row = document.add_table(rows=1, cols=1).rows[0]
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "false")
    tr_pr.append(cant_split)

    TableRenderer._set_row_cant_split(row)

    assert _row_has_cant_split(row) is True
    assert _row_cant_split_value(row) == "true"


@pytest.mark.asyncio
async def test_add_blocks_tool_marks_caption_only_table_as_non_split_without_tbl_header(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-caption-only-table-cantsplit"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="只有标题的表",
            output_name="caption-only-table.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "caption": "仅标题（无表头行）的表",
                "rows": [["数据 1"], ["数据 2"]],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert table.rows[0].cells[0].text == "仅标题（无表头行）的表"
    assert _row_has_cant_split(table.rows[0]) is True
    assert _row_is_repeated_header(table.rows[0]) is False
    assert _row_has_cant_split(table.rows[1]) is True
    assert _row_is_repeated_header(table.rows[1]) is False
    assert _row_has_cant_split(table.rows[2]) is True
    assert _row_is_repeated_header(table.rows[2]) is False


@pytest.mark.asyncio
async def test_add_blocks_tool_applies_custom_table_style_overrides(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.table import WD_TABLE_ALIGNMENT

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-custom-table-style")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="样式覆盖",
            output_name="custom-table-style.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "caption": "季度经营指标",
                "caption_emphasis": "strong",
                "header_groups": [
                    {"title": "经营数据", "span": 2},
                    {"title": "结果", "span": 2},
                ],
                "headers": ["区域", "目标", "完成值", "完成率"],
                "rows": [
                    ["华东", "120", "118", "98%"],
                    ["华南", "88", "91", "103%"],
                ],
                "header_fill": "1F4E79",
                "header_text_color": "FFFFFF",
                "banded_rows": True,
                "banded_row_fill": "EEF4FA",
                "first_column_bold": True,
                "table_align": "left",
                "border_style": "strong",
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]

    assert table.alignment == WD_TABLE_ALIGNMENT.LEFT
    assert _cell_fill(table.rows[0].cells[0]) == "1F4E79"
    assert _run_rgb(table.rows[0].cells[0]) == "FFFFFF"
    assert _cell_fill(table.rows[1].cells[0]) == "1F4E79"
    assert _run_rgb(table.rows[1].cells[0]) == "FFFFFF"
    assert _cell_fill(table.rows[3].cells[0]) == "EEF4FA"
    assert _cell_fill(table.rows[4].cells[0]) is None
    assert _run_bold(table.rows[3].cells[0]) is True
    assert _run_bold(table.rows[3].cells[1]) is False
    assert _table_border_size(table, "top") == "16"


@pytest.mark.asyncio
async def test_add_blocks_tool_treats_table_title_as_caption_alias(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-title-alias")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="表格标题别名",
            output_name="table-title-alias.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "table",
                "title": "季度经营指标总览",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"]],
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]
    assert table.rows[0].cells[0].text == "季度经营指标总览"
    assert table.rows[1].cells[0].text == "区域"


@pytest.mark.asyncio
async def test_add_blocks_tool_absorbs_heading_before_table_into_table_title(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-heading-merge")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="标题吸收",
            output_name="table-heading-merge.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "季度经营指标总览", "level": 2},
            {
                "type": "table",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    assert len(loaded_doc.tables) == 1
    table = loaded_doc.tables[0]
    assert table.rows[0].cells[0].text == "季度经营指标总览"
    assert table.rows[1].cells[0].text == "区域"
    assert "季度经营指标总览" not in [
        paragraph.text for paragraph in loaded_doc.paragraphs[1:]
    ]


@pytest.mark.asyncio
async def test_add_blocks_tool_drops_heading_that_duplicates_document_title(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-duplicate-title")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="项目阶段汇报",
            output_name="duplicate-title.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "项目阶段汇报", "level": 1},
            {"type": "paragraph", "text": "正文内容。"},
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph_texts = [paragraph.text for paragraph in loaded_doc.paragraphs]
    assert paragraph_texts.count("项目阶段汇报") == 1
    assert "正文内容。" in paragraph_texts


@pytest.mark.asyncio
async def test_add_blocks_tool_drops_duplicate_document_title_before_table_promotion(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-duplicate-title-before-table"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="项目阶段汇报",
            output_name="duplicate-title-before-table.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "项目阶段汇报", "level": 1},
            {
                "type": "table",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph_texts = [paragraph.text for paragraph in loaded_doc.paragraphs]
    table = loaded_doc.tables[0]

    assert paragraph_texts.count("项目阶段汇报") == 1
    assert table.rows[0].cells[0].text == "区域"
    assert "项目阶段汇报" not in table.rows[0].cells[0].text


@pytest.mark.asyncio
async def test_add_blocks_tool_does_not_absorb_long_heading_before_table_into_table_title(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-table-long-heading-no-merge"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}
    long_heading = "季度经营指标总览" * 20

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="长标题保留",
            output_name="table-long-heading-no-merge.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": long_heading, "level": 2},
            {
                "type": "table",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph_texts = [paragraph.text for paragraph in loaded_doc.paragraphs]
    assert long_heading in paragraph_texts[1:]

    table = loaded_doc.tables[0]
    assert table.rows[0].cells[0].text == "区域"
    assert long_heading not in table.rows[0].cells[0].text
    assert len(table.rows[0].cells) == len(table.rows[1].cells)


@pytest.mark.asyncio
async def test_add_blocks_tool_does_not_absorb_numbered_heading_before_table(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-numbered-heading-before-table"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="编号标题保留",
            output_name="numbered-heading-before-table.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "III. Training Schedule", "level": 1},
            {
                "type": "table",
                "headers": ["Date", "Time", "Session Title", "Trainer"],
                "rows": [["2026-04-10", "09:00 - 12:00", "Intro", "Alice"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    paragraph_texts = [paragraph.text for paragraph in loaded_doc.paragraphs]
    table = loaded_doc.tables[0]

    assert "III. Training Schedule" in paragraph_texts[1:]
    assert table.rows[0].cells[0].text == "Date"


@pytest.mark.asyncio
async def test_add_blocks_tool_does_not_absorb_bottom_border_heading_before_table(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-bordered-heading-before-table"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="带边线标题保留",
            output_name="bordered-heading-before-table.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {
                "type": "heading",
                "text": "Training Schedule",
                "level": 1,
                "bottom_border": True,
                "bottom_border_color": "D0D7DE",
            },
            {
                "type": "table",
                "headers": ["Date", "Time", "Session Title", "Trainer"],
                "rows": [["2026-04-10", "09:00 - 12:00", "Intro", "Alice"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    heading = _find_paragraph(loaded_doc, "Training Schedule")
    heading_divider = _paragraph_after(loaded_doc, heading)
    table = loaded_doc.tables[0]

    assert _paragraph_bottom_border_color(heading) is None
    assert _paragraph_bottom_border_color(heading_divider) == "D0D7DE"
    assert table.rows[0].cells[0].text == "Date"


@pytest.mark.asyncio
async def test_document_toolset_export_callback_runs(workspace_root: Path):
    pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-callback")
    callback_calls: list[str] = []

    async def after_export(_context, output_path: str) -> str:
        callback_calls.append(output_path)
        return "callback sent"

    toolset = build_document_toolset(
        workspace_dir=workspace_dir,
        after_export=after_export,
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Pytest Callback",
            output_name="pytest-callback.docx",
        )
    )

    exported = await tool_by_name["export_document"].call(
        object(),
        document_id=created["document"]["document_id"],
    )

    assert exported is None
    assert len(callback_calls) == 1
    assert Path(callback_calls[0]).exists()


@pytest.mark.asyncio
async def test_add_blocks_tool_rejects_updates_after_finalize():
    toolset = build_document_toolset()
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Finalized Guard",
            output_name="finalized-guard.docx",
        )
    )
    document_id = created["document"]["document_id"]

    finalized = json.loads(
        await tool_by_name["finalize_document"].call(
            None,
            document_id=document_id,
        )
    )
    assert finalized["success"] is True

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=[{"type": "paragraph", "text": "不应再追加"}],
        )
    )

    assert add_blocks_result["success"] is False
    assert (
        "add_blocks is only allowed while the document status is draft"
        in add_blocks_result["message"]
    )


@pytest.mark.asyncio
async def test_document_toolset_preserves_positional_after_export_callback(
    workspace_root: Path,
):
    pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-tools-positional-callback"
    )
    callback_calls: list[str] = []

    async def after_export(_context, output_path: str) -> str:
        callback_calls.append(output_path)
        return "callback sent"

    toolset = build_document_toolset(workspace_dir, after_export)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Positional Callback",
            output_name="positional-callback.docx",
        )
    )

    exported = await tool_by_name["export_document"].call(
        object(),
        document_id=created["document"]["document_id"],
    )

    assert exported is None
    assert len(callback_calls) == 1
    assert Path(callback_calls[0]).exists()
    assert Path(callback_calls[0]).name == "positional-callback.docx"


@pytest.mark.asyncio
async def test_export_pipeline_falls_back_to_python_backend(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-render-backend-fallback")
    store = DocumentSessionStore(workspace_dir=workspace_dir)
    document = store.create_document(
        CreateDocumentRequest(
            title="Backend Fallback",
            output_name="backend-fallback.docx",
        )
    )
    request = ExportDocumentRequest(document_id=document.document_id)
    called_backends: list[str] = []

    class _FailingBackend:
        name = "node"

        def render(self, _document, _output_path):
            called_backends.append(self.name)
            raise RuntimeError("node renderer unavailable")

    class _PythonBackend:
        name = "python"

        def render(self, _document, output_path):
            called_backends.append(self.name)
            output_path.write_bytes(b"fallback-ok")
            return RenderResult(backend_name=self.name, output_path=output_path)

    exported_document, output_path = await export_document_via_pipeline(
        store=store,
        render_backends=[_FailingBackend(), _PythonBackend()],
        request=request,
        source="pytest",
    )

    assert called_backends == ["node", "python"]
    assert output_path.read_bytes() == b"fallback-ok"
    assert exported_document.status.value == "exported"


@pytest.mark.asyncio
async def test_document_toolset_runs_after_export_hooks_before_delivery_callback(
    workspace_root: Path,
):
    pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-after-export")
    hook_calls: list[tuple[str, str]] = []
    callback_calls: list[str] = []

    async def after_export_hook(context):
        hook_calls.append(
            (context.document.status.value, Path(context.output_path).name)
        )
        return context

    async def after_export(_context, output_path: str) -> str:
        assert hook_calls == [("exported", "after-export-hook.docx")]
        callback_calls.append(output_path)
        return "callback sent"

    toolset = build_document_toolset(
        workspace_dir=workspace_dir,
        after_export_hooks=[after_export_hook],
        after_export=after_export,
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="After Export Hook",
            output_name="after-export-hook.docx",
        )
    )

    exported = await tool_by_name["export_document"].call(
        object(),
        document_id=created["document"]["document_id"],
    )

    assert exported is None
    assert hook_calls == [("exported", "after-export-hook.docx")]
    assert len(callback_calls) == 1
    assert Path(callback_calls[0]).exists()


@pytest.mark.asyncio
async def test_export_document_tool_keeps_success_when_callback_fails(
    workspace_root: Path,
):
    pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-callback-fail")

    async def after_export(_context, _output_path: str) -> str:
        raise RuntimeError("send failed")

    toolset = build_document_toolset(
        workspace_dir=workspace_dir,
        after_export=after_export,
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            session_id="pytest-session",
            title="Pytest Callback Failure",
            output_name="pytest-callback-failure.docx",
        )
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            object(),
            document_id=created["document"]["document_id"],
        )
    )

    assert exported["success"] is True
    assert "post-export delivery failed" in exported["message"]
    assert Path(exported["file_path"]).exists()


@pytest.mark.asyncio
async def test_document_toolset_runs_before_export_hooks(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-before-export")

    async def before_export(context):
        context.document.blocks.append(ParagraphBlock(text="Export hook note"))
        return context

    toolset = build_document_toolset(
        workspace_dir=workspace_dir,
        before_export_hooks=[before_export],
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Before Export Hook",
            output_name="before-export-hook.docx",
        )
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    assert "Export hook note" in [paragraph.text for paragraph in loaded_doc.paragraphs]


@pytest.mark.asyncio
async def test_mcp_registers_document_and_workbook_tools():
    server = create_server()
    tools = await server.list_tools()
    tool_names = [tool.name for tool in tools]

    assert tool_names == [
        "create_document",
        "add_blocks",
        "finalize_document",
        "export_document",
        "create_workbook",
        "write_rows",
        "export_workbook",
    ]

    add_blocks_tool = next(tool for tool in tools if tool.name == "add_blocks")
    write_rows_tool = next(tool for tool in tools if tool.name == "write_rows")

    assert add_blocks_tool.inputSchema["required"] == ["document_id", "blocks"]
    add_blocks_items = add_blocks_tool.inputSchema["properties"]["blocks"]["items"]
    assert add_blocks_items["type"] == "object"
    assert add_blocks_items["additionalProperties"] is True
    assert write_rows_tool.inputSchema["required"] == ["workbook_id", "sheet", "rows"]


@pytest.mark.asyncio
async def test_mcp_write_rows_returns_structured_failure_for_unknown_workbook():
    server = create_server()

    _, payload = await server.call_tool(
        "write_rows",
        {
            "workbook_id": "wb-missing",
            "sheet": "Data",
            "rows": [["value"]],
        },
    )

    assert payload["success"] is False
    assert payload["message"].startswith("write_rows failed:")


@pytest.mark.asyncio
async def test_mcp_write_rows_returns_structured_failure_for_validation_errors():
    server = create_server()
    _, created_payload = await server.call_tool(
        "create_workbook",
        {"filename": "validation.xlsx"},
    )

    _, payload = await server.call_tool(
        "write_rows",
        {
            "workbook_id": created_payload["workbook"]["workbook_id"],
            "sheet": "Data",
            "rows": [["=SUM(A1:A2)"]],
        },
    )

    assert payload["success"] is False
    assert "only fix invalid fields" in payload["message"]


@pytest.mark.asyncio
async def test_mcp_write_rows_rejects_oversized_row_window():
    server = create_server()
    _, created_payload = await server.call_tool(
        "create_workbook",
        {"filename": "validation.xlsx"},
    )

    _, payload = await server.call_tool(
        "write_rows",
        {
            "workbook_id": created_payload["workbook"]["workbook_id"],
            "sheet": "Data",
            "rows": [["value"], ["overflow"]],
            "start_row": MAX_WORKBOOK_ROW_INDEX,
        },
    )

    assert payload["success"] is False
    assert "only fix invalid fields" in payload["message"]
    assert "final written row must not exceed" in payload["message"]


@pytest.mark.asyncio
async def test_mcp_registers_only_document_tools_when_workbook_store_is_unavailable(
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(mcp_server_module, "_load_workbook_session_store", lambda: None)

    server = create_server()
    tools = await server.list_tools()

    assert [tool.name for tool in tools] == [
        "create_document",
        "add_blocks",
        "finalize_document",
        "export_document",
    ]
    assert "Excel" not in server.instructions
    assert "create_workbook" not in server.instructions


def test_mcp_workbook_loader_logs_import_error_and_disables_support(
    monkeypatch: pytest.MonkeyPatch,
):
    def _raise_import_error(_module_name: str):
        raise ModuleNotFoundError("missing workbook deps")

    monkeypatch.setattr(
        mcp_server_module.importlib,
        "import_module",
        _raise_import_error,
    )

    with patch.object(mcp_server_module.logger, "warning") as warning_mock:
        result = mcp_server_module._load_workbook_session_store()

    assert result is None
    warning_mock.assert_called_once()


def test_mcp_workbook_loader_reraises_non_import_errors(
    monkeypatch: pytest.MonkeyPatch,
):
    def _raise_runtime_error(_module_name: str):
        raise NameError("boom")

    monkeypatch.setattr(
        mcp_server_module.importlib,
        "import_module",
        _raise_runtime_error,
    )

    with patch.object(mcp_server_module.logger, "warning") as warning_mock:
        with pytest.raises(NameError, match="boom"):
            mcp_server_module._load_workbook_session_store()

    warning_mock.assert_not_called()


def test_mcp_create_server_constructs_workbook_store_once(
    monkeypatch: pytest.MonkeyPatch,
):
    constructed_workspaces: list[Path | None] = []

    class StubWorkbookSessionStore:
        def __init__(self, workspace_dir=None):
            constructed_workspaces.append(workspace_dir)

    register_mock = MagicMock()
    monkeypatch.setattr(
        mcp_server_module,
        "_load_workbook_session_store",
        lambda: StubWorkbookSessionStore,
    )
    monkeypatch.setattr(mcp_server_module, "register_workbook_tools", register_mock)

    create_server()

    assert constructed_workspaces == [None]
    register_mock.assert_called_once()


@pytest.mark.asyncio
async def test_mcp_export_document_runs_before_export_hooks(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-mcp-before-export")

    async def before_export(context):
        context.document.blocks.append(ParagraphBlock(text="MCP export hook note"))
        return context

    server = create_server(
        workspace_dir=workspace_dir,
        before_export_hooks=[before_export],
    )
    _, created_payload = await server.call_tool(
        "create_document",
        {"title": "MCP Hook", "output_name": "mcp-hook.docx"},
    )
    _, exported_payload = await server.call_tool(
        "export_document",
        {"document_id": created_payload["document"]["document_id"]},
    )

    loaded_doc = docx.Document(exported_payload["file_path"])
    assert "MCP export hook note" in [
        paragraph.text for paragraph in loaded_doc.paragraphs
    ]


@pytest.mark.asyncio
async def test_mcp_export_document_runs_after_export_hooks(workspace_root: Path):
    pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-mcp-after-export")
    hook_calls: list[tuple[str, str]] = []

    async def after_export(context):
        hook_calls.append(
            (context.document.status.value, Path(context.output_path).name)
        )
        return context

    server = create_server(
        workspace_dir=workspace_dir,
        after_export_hooks=[after_export],
    )
    _, created_payload = await server.call_tool(
        "create_document",
        {"title": "MCP After Hook", "output_name": "mcp-after-hook.docx"},
    )
    _, exported_payload = await server.call_tool(
        "export_document",
        {"document_id": created_payload["document"]["document_id"]},
    )

    assert exported_payload["success"] is True
    assert hook_calls == [("exported", "mcp-after-hook.docx")]


@pytest.mark.asyncio
async def test_mcp_server_exports_word_via_node_backend(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-mcp-node-word-export")
    server = create_server(
        workspace_dir=workspace_dir,
        render_backend_config=_node_render_backend_config_for_tests(),
    )

    _, created_payload = await server.call_tool(
        "create_document",
        {
            "title": "MCP Node Export",
            "output_name": "mcp-node-export.docx",
            "document_style": {"heading_color": "000000"},
            "header_footer": {
                "header_left": "季度经营复盘",
                "footer_right": "第 {PAGE} 页",
                "show_page_number": True,
            },
        },
    )
    document_id = created_payload["document"]["document_id"]

    _, add_blocks_payload = await server.call_tool(
        "add_blocks",
        {
            "document_id": document_id,
            "blocks": [
                {
                    "type": "heading",
                    "text": "一、经营总览",
                    "level": 1,
                    "bottom_border": True,
                },
                {
                    "type": "table",
                    "headers": ["日期", "时间", "内容"],
                    "rows": [
                        ["第一天", "09:00", "课程 A"],
                        ["第二天", "13:00", "课程 B"],
                    ],
                    "header_fill_enabled": False,
                    "header_bold": False,
                },
            ],
        },
    )
    assert add_blocks_payload["success"] is True

    _, finalized_payload = await server.call_tool(
        "finalize_document",
        {"document_id": document_id},
    )
    assert finalized_payload["success"] is True

    _, exported_payload = await server.call_tool(
        "export_document",
        {"document_id": document_id},
    )
    assert exported_payload["success"] is True

    loaded_doc = docx.Document(exported_payload["file_path"])
    table = loaded_doc.tables[0]

    assert table.rows[0].cells[0].text == "一、经营总览"
    assert _find_paragraph(loaded_doc, "MCP Node Export").text == "MCP Node Export"
    assert len(table.rows) >= 4
    assert table.rows[2].cells[0].text == "第一天"
    assert table.rows[3].cells[0].text == "第二天"


@pytest.mark.asyncio
async def test_document_toolset_exports_toc_and_section_break(workspace_root: Path):
    docx = pytest.importorskip("docx")
    from docx.enum.section import WD_ORIENT

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-toc-section")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="目录与分节",
            output_name="toc-section.docx",
            header_footer={
                "header_text": "默认页眉",
                "different_odd_even": True,
                "even_page_header_text": "默认偶数页页眉",
                "show_page_number": True,
            },
        )
    )

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks=[
                {
                    "type": "toc",
                    "title": "目录",
                    "levels": 2,
                    "start_on_new_page": True,
                },
                {"type": "heading", "text": "经营总览", "level": 2},
                {"type": "paragraph", "text": "第一节正文"},
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "inherit_header_footer": False,
                    "page_orientation": "landscape",
                    "margins": {
                        "top_cm": 1.6,
                        "bottom_cm": 1.7,
                        "left_cm": 1.8,
                        "right_cm": 1.9,
                    },
                    "restart_page_numbering": True,
                    "page_number_start": 5,
                    "header_footer": {
                        "header_text": "第二节页眉",
                        "footer_text": "第二节页脚",
                        "different_first_page": True,
                        "first_page_header_text": "第二节首页页眉",
                        "show_page_number": True,
                    },
                },
                {"type": "heading", "text": "行动计划", "level": 2},
            ],
        )
    )
    assert add_blocks_result["success"] is True

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    toc_index = next(
        index
        for index, paragraph in enumerate(loaded_doc.paragraphs)
        if paragraph.text == "目录"
    )
    assert len(loaded_doc.sections) == 2
    assert _document_updates_fields_on_open(loaded_doc) is True
    assert _document_uses_odd_even_headers(loaded_doc) is True
    assert "默认偶数页页眉" in _story_texts(loaded_doc.sections[0].even_page_header)
    assert "第二节页眉" in _story_texts(loaded_doc.sections[1].header)
    assert "第二节首页页眉" in _story_texts(loaded_doc.sections[1].first_page_header)
    assert _story_has_field_code(loaded_doc.sections[1].footer, "PAGE") is True
    assert _paragraph_field_nodes_use_runs(loaded_doc.paragraphs[toc_index + 1]) is True
    assert _section_page_number_start(loaded_doc.sections[1]) == 5
    assert loaded_doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert loaded_doc.sections[1].left_margin.cm == pytest.approx(1.8, abs=0.01)
    with zipfile.ZipFile(exported["file_path"]) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "w:fldSimple" in document_xml
    assert "\\o &quot;1-2&quot;" in document_xml
    assert 'w:type w:val="nextPage"' in document_xml


@pytest.mark.asyncio
async def test_document_toolset_exports_regular_blocks_in_portrait(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.section import WD_ORIENT

    workspace_dir = _make_workspace(
        workspace_root, "pytest-agent-tools-portrait-export"
    )
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="哲学知识回顾",
            output_name="philosophy-review.docx",
            theme_name="business_report",
        )
    )

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks=[
                {"type": "toc", "title": "目录", "levels": 2},
                {"type": "heading", "text": "一、哲学的含义", "level": 1},
                {
                    "type": "paragraph",
                    "text": "哲学是系统化、理论化的世界观，也是世界观和方法论的统一。",
                },
                {"type": "heading", "text": "二、哲学的基本问题", "level": 1},
                {
                    "type": "table",
                    "caption": "哲学基本问题分类",
                    "headers": ["方面", "核心内容", "主要派别"],
                    "rows": [
                        [
                            "第一方面（第一性）",
                            "思维和存在何者为第一性",
                            "唯物主义、唯心主义",
                        ],
                        [
                            "第二方面（同一性）",
                            "思维能否正确认识存在",
                            "可知论、不可知论",
                        ],
                    ],
                },
                {"type": "heading", "text": "三、其他相关知识", "level": 1},
                {"type": "paragraph", "text": "马克思主义哲学是科学的世界观和方法论。"},
            ],
        )
    )
    assert add_blocks_result["success"] is True

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    assert len(loaded_doc.sections) == 1
    assert all(
        section.orientation == WD_ORIENT.PORTRAIT for section in loaded_doc.sections
    )


@pytest.mark.asyncio
async def test_add_blocks_tool_normalizes_landscape_section_payload_aliases(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")
    from docx.enum.section import WD_ORIENT

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-raw-aliases")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="原始别名修复",
            output_name="raw-aliases.docx",
        )
    )

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks=[
                {"type": "toc", "text": "目录"},
                {
                    "type": "heading",
                    "text": "三、运营数据分析",
                    "level": 1,
                    "page_orientation": "landscape",
                    "start_on_new_page": True,
                    "restart_page_numbering": True,
                    "header_footer": {
                        "header_text": "横向页眉",
                        "footer_text": "横向页脚",
                        "show_page_number": True,
                    },
                },
                {
                    "type": "table",
                    "caption": "第一季度核心运营指标汇总",
                    "items": [
                        "用户增长数|10000|10500|105%",
                        "营收总额(万元)|5000|4800|96%",
                    ],
                    "columns": [
                        {"blocks": [{"type": "paragraph", "text": "指标名称"}]},
                        {"blocks": [{"type": "paragraph", "text": "Q1目标值"}]},
                        {"blocks": [{"type": "paragraph", "text": "Q1实际值"}]},
                        {"blocks": [{"type": "paragraph", "text": "达成率"}]},
                    ],
                },
                {
                    "type": "heading",
                    "text": "四、问题与挑战",
                    "level": 1,
                    "page_orientation": "portrait",
                    "start_on_new_page": True,
                    "header_footer": {
                        "header_text": "竖向页眉",
                        "footer_text": "竖向页脚",
                        "show_page_number": True,
                    },
                },
            ],
        )
    )

    assert add_blocks_result["success"] is True

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    assert len(loaded_doc.sections) == 3
    assert loaded_doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert loaded_doc.sections[2].orientation == WD_ORIENT.PORTRAIT
    assert "横向页眉" in _story_texts(loaded_doc.sections[1].header)
    assert any(
        text.startswith("横向页脚")
        for text in _story_texts(loaded_doc.sections[1].footer)
    )
    assert "竖向页眉" in _story_texts(loaded_doc.sections[2].header)
    assert any(
        text.startswith("竖向页脚")
        for text in _story_texts(loaded_doc.sections[2].footer)
    )
    toc_index = next(
        index
        for index, paragraph in enumerate(loaded_doc.paragraphs)
        if paragraph.text == "目录"
    )
    assert any(
        'TOC \\o "1-3"' in field_code
        for field_code in _paragraph_field_codes(loaded_doc.paragraphs[toc_index + 1])
    )
    assert loaded_doc.tables[0].rows[1].cells[0].text == "指标名称"
    assert loaded_doc.tables[0].rows[1].cells[3].text == "达成率"


@pytest.mark.asyncio
async def test_add_blocks_tool_clamps_block_ranges_and_drops_heading_color(
    workspace_root: Path,
):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-block-clamp")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="块级兜底",
            output_name="block-clamp.docx",
        )
    )

    add_blocks_result = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=created["document"]["document_id"],
            blocks=[
                {
                    "type": "paragraph",
                    "text": "第一季度经营复盘",
                    "style": {"font_scale": 3, "align": "center"},
                    "layout": {"spacing_before": 200, "spacing_after": -10},
                },
                {
                    "type": "heading",
                    "text": "一、第一季度整体经营概况",
                    "level": 1,
                    "heading_color": "1F4E79",
                },
            ],
        )
    )

    assert add_blocks_result["success"] is True

    document = tool_by_name["add_blocks"].store.require_document(
        created["document"]["document_id"]
    )
    paragraph_block = document.blocks[0]
    heading_block = document.blocks[1]

    assert paragraph_block.style.font_scale == pytest.approx(2.0)
    assert paragraph_block.layout.spacing_before == pytest.approx(72.0)
    assert paragraph_block.layout.spacing_after == pytest.approx(0.0)
    assert not hasattr(heading_block, "heading_color")


@pytest.mark.asyncio
async def test_document_toolset_falls_back_when_metrics_table_style_is_missing(
    workspace_root: Path, monkeypatch: pytest.MonkeyPatch
):
    docx = pytest.importorskip("docx")

    monkeypatch.setitem(DOCX_TABLE_STYLES, "metrics_compact", "Missing Docx Style")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-missing-style")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="Missing Style Fallback",
            output_name="missing-style-fallback.docx",
            table_template="metrics_compact",
        )
    )

    await tool_by_name["add_blocks"].call(
        None,
        document_id=created["document"]["document_id"],
        blocks=[
            {
                "type": "table",
                "headers": ["Metric", "Value"],
                "rows": [["Users", "42"]],
                "table_style": "metrics_compact",
            }
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=created["document"]["document_id"],
        )
    )

    assert exported["success"] is True
    loaded_doc = docx.Document(exported["file_path"])
    assert len(loaded_doc.tables) == 1
    assert loaded_doc.tables[0].style.name == "Table Grid"


@pytest.mark.asyncio
async def test_add_blocks_tool_ignores_blank_table_caption_when_absorbing_heading(
    workspace_root: Path,
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, "pytest-agent-table-blank-caption")
    toolset = build_document_toolset(workspace_dir=workspace_dir)
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(
        await tool_by_name["create_document"].call(
            None,
            title="空白表标题",
            output_name="blank-table-caption.docx",
        )
    )
    document_id = created["document"]["document_id"]

    await tool_by_name["add_blocks"].call(
        None,
        document_id=document_id,
        blocks=[
            {"type": "heading", "text": "季度经营指标总览", "level": 2},
            {
                "type": "table",
                "caption": "   ",
                "headers": ["区域", "营收（万元）"],
                "rows": [["华东", "1280"]],
            },
        ],
    )

    exported = json.loads(
        await tool_by_name["export_document"].call(
            None,
            document_id=document_id,
        )
    )

    loaded_doc = docx.Document(exported["file_path"])
    table = loaded_doc.tables[0]
    assert table.rows[0].cells[0].text == "季度经营指标总览"
