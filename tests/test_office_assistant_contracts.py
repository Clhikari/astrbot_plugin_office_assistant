import builtins
import importlib
import json
import shutil
import subprocess
import struct
import sys
import zipfile
import zlib
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch
from uuid import uuid4

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from astrbot_plugin_office_assistant.agent_tools import (
    build_document_toolset,
)
from astrbot_plugin_office_assistant.agent_tools.document_tools import (
    CreateDocumentTool,
)
from astrbot_plugin_office_assistant.document_core.builders.table_renderer import (
    DOCX_TABLE_STYLES,
    TableRenderer,
)
from astrbot_plugin_office_assistant.document_core.macros.summary_card import (
    build_summary_card_group,
)
from astrbot_plugin_office_assistant.domain.document.session_store import (
    DocumentSessionStore,
)
from astrbot_plugin_office_assistant.domain.document.export_pipeline import (
    export_document_via_pipeline,
)
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendConfig,
    DocumentRenderBackendError,
    NodeDocumentRenderBackend,
    build_document_render_backends,
    RenderResult,
    build_document_render_payload,
)
from astrbot_plugin_office_assistant.document_core.models.blocks import (
    BlockStyle,
    BusinessReviewCoverData,
    ColumnBlock,
    ColumnsBlock,
    GroupBlock,
    HeaderFooterConfig,
    HeadingBlock,
    HeroBannerBlock,
    ListItem,
    PageTemplateBlock,
    ParagraphBlock,
    ParagraphRun,
    ResumeSectionEntry,
    ResumeSection,
    SectionBreakBlock,
    SectionMarginsConfig,
    SummaryCardBlock,
    TableBlock,
    TechnicalResumeData,
    TocBlock,
)
from astrbot_plugin_office_assistant.document_core.models.document import (
    DocumentMetadata,
    DocumentModel,
    DocumentStyleConfig,
    DocumentSummaryCardDefaults,
)
from astrbot_plugin_office_assistant.domain.document.contracts import (
    AddBlocksRequest,
    AddHeadingRequest,
    AddListRequest,
    AddParagraphRequest,
    AddTableRequest,
    BlockHeadingInput,
    CreateDocumentRequest,
    ExportDocumentRequest,
    FinalizeDocumentRequest,
    SectionListInput,
    SectionParagraphInput,
    SectionTableInput,
    normalize_create_document_kwargs,
    normalize_raw_block_payloads,
)
from astrbot_plugin_office_assistant.mcp_server.server import (
    create_server,
)
from astrbot_plugin_office_assistant.tools.mcp_adapter import (
    register_document_tools_from_registry,
)
from astrbot_plugin_office_assistant.tools.astrbot_adapter import (
    build_document_toolset_from_registry,
)
from astrbot_plugin_office_assistant.tools.registry import (
    DocumentToolSpec,
    get_document_tool_specs,
)
from pydantic import ValidationError

from astrbot.core.utils.astrbot_path import get_astrbot_plugin_data_path


from tests._docx_test_helpers import *  # noqa: F401,F403
from tests._docx_test_helpers import _technical_resume_block



def test_create_document_tool_schema_exposes_document_style():
    toolset = build_document_toolset()
    create_document_tool = next(
        tool for tool in toolset.tools if tool.name == "create_document"
    )

    properties = create_document_tool.parameters["properties"]
    header_footer = properties["header_footer"]["properties"]
    document_style = properties["document_style"]["properties"]
    table_defaults = document_style["table_defaults"]["properties"]

    assert header_footer["header_text"]["type"] == "string"
    assert header_footer["footer_text"]["type"] == "string"
    assert header_footer["header_left"]["type"] == "string"
    assert header_footer["header_right"]["type"] == "string"
    assert header_footer["footer_left"]["type"] == "string"
    assert header_footer["footer_right"]["type"] == "string"
    assert header_footer["header_border_bottom"]["type"] == "boolean"
    assert header_footer["footer_border_top"]["type"] == "boolean"
    assert header_footer["header_border_color"]["type"] == "string"
    assert header_footer["footer_border_color"]["type"] == "string"
    assert header_footer["different_first_page"]["type"] == "boolean"
    assert header_footer["first_page_header_text"]["type"] == "string"
    assert header_footer["first_page_footer_text"]["type"] == "string"
    assert header_footer["first_page_show_page_number"]["type"] == "boolean"
    assert header_footer["different_odd_even"]["type"] == "boolean"
    assert header_footer["even_page_header_text"]["type"] == "string"
    assert header_footer["even_page_footer_text"]["type"] == "string"
    assert header_footer["even_page_show_page_number"]["type"] == "boolean"
    assert header_footer["show_page_number"]["type"] == "boolean"
    assert header_footer["page_number_align"]["enum"] == ["left", "center", "right"]
    assert header_footer["page_number_format"]["enum"] == [
        "decimal",
        "upperRoman",
        "lowerRoman",
        "upperLetter",
        "lowerLetter",
    ]
    assert document_style["brief"]["type"] == "string"
    assert document_style["heading_color"]["type"] == "string"
    assert document_style["heading_level_1_color"]["type"] == "string"
    assert document_style["heading_level_2_color"]["type"] == "string"
    assert document_style["heading_bottom_border_color"]["type"] == "string"
    assert document_style["heading_bottom_border_size_pt"]["type"] == "number"
    assert document_style["title_align"]["enum"] == [
        "left",
        "center",
        "right",
        "justify",
    ]
    assert document_style["body_font_size"]["type"] == "number"
    assert document_style["body_line_spacing"]["type"] == "number"
    assert document_style["font_name"]["type"] == "string"
    assert document_style["heading_font_name"]["type"] == "string"
    assert document_style["table_font_name"]["type"] == "string"
    assert document_style["code_font_name"]["type"] == "string"
    assert document_style["paragraph_space_after"]["type"] == "number"
    assert document_style["list_space_after"]["type"] == "number"
    assert document_style["summary_card_defaults"]["type"] == "object"
    assert document_style["summary_card_defaults"]["properties"]["title_align"][
        "enum"
    ] == ["left", "center", "right", "justify"]
    assert document_style["summary_card_defaults"]["properties"]["title_emphasis"][
        "enum"
    ] == ["normal", "strong", "subtle"]
    assert (
        "header_fill_enabled and header_bold belong on each table block"
        in document_style["table_defaults"]["description"]
    )
    assert table_defaults["preset"]["enum"] == [
        "report_grid",
        "metrics_compact",
        "minimal",
    ]
    assert table_defaults["body_fill"]["type"] == "string"
    assert table_defaults["table_align"]["enum"] == ["left", "center"]
    assert table_defaults["border_style"]["enum"] == [
        "minimal",
        "standard",
        "strong",
    ]
    assert table_defaults["cell_align"]["enum"] == ["left", "center", "right"]

def test_add_blocks_tool_schema_keeps_nested_array_items_for_gemini():
    toolset = build_document_toolset()
    add_blocks_tool = next(tool for tool in toolset.tools if tool.name == "add_blocks")

    block_properties = add_blocks_tool.parameters["properties"]["blocks"]["items"][
        "properties"
    ]
    assert block_properties["blocks"]["type"] == "array"
    assert block_properties["blocks"]["items"]["type"] == "object"
    assert block_properties["blocks"]["items"]["additionalProperties"] is True
    assert block_properties["columns"]["type"] == "array"
    assert block_properties["columns"]["items"]["type"] == "object"
    assert (
        block_properties["columns"]["items"]["properties"]["blocks"]["items"]["type"]
        == "object"
    )
    assert block_properties["numeric_columns"]["items"]["type"] == "integer"
    assert (
        block_properties["runs"]["items"]["properties"]["italic"]["type"] == "boolean"
    )
    run_properties = block_properties["runs"]["items"]["properties"]
    assert run_properties["bold"]["type"] == "boolean"
    assert run_properties["underline"]["type"] == "boolean"
    assert run_properties["code"]["type"] == "boolean"
    assert run_properties["color"]["type"] == "string"
    assert block_properties["template"]["type"] == "string"
    assert block_properties["data"]["type"] == "object"
    assert block_properties["data"]["properties"]["title"]["type"] == "string"
    assert block_properties["data"]["properties"]["summary_text"]["type"] == "string"
    assert block_properties["data"]["properties"]["metrics"]["type"] == "array"
    assert (
        block_properties["data"]["properties"]["metrics"]["items"]["properties"][
            "delta_color"
        ]["type"]
        == "string"
    )
    assert (
        block_properties["data"]["properties"]["auto_page_break"]["type"]
        == "boolean"
    )
    assert block_properties["text"]["type"] == "string"
    assert block_properties["subtitle"]["type"] == "string"
    assert block_properties["runs"]["type"] == "array"
    assert block_properties["runs"]["items"]["type"] == "object"
    assert block_properties["bottom_border"]["type"] == "boolean"
    assert block_properties["bottom_border_color"]["type"] == "string"
    assert block_properties["bottom_border_size_pt"]["type"] == "number"
    assert block_properties["header_groups"]["type"] == "array"
    assert block_properties["header_groups"]["items"]["type"] == "object"
    assert (
        block_properties["header_groups"]["items"]["properties"]["title"]["type"]
        == "string"
    )
    assert (
        block_properties["header_groups"]["items"]["properties"]["title"]["minLength"]
        == 1
    )
    assert (
        block_properties["header_groups"]["items"]["properties"]["span"]["type"]
        == "integer"
    )
    assert (
        block_properties["header_groups"]["items"]["properties"]["span"]["minimum"] == 1
    )
    assert block_properties["header_groups"]["items"]["required"] == [
        "title",
        "span",
    ]
    assert block_properties["header_fill"]["type"] == "string"
    assert block_properties["header_fill_enabled"]["type"] == "boolean"
    assert block_properties["header_text_color"]["type"] == "string"
    assert block_properties["header_bold"]["type"] == "boolean"
    assert block_properties["banded_rows"]["type"] == "boolean"
    assert block_properties["banded_row_fill"]["type"] == "string"
    assert block_properties["first_column_bold"]["type"] == "boolean"
    assert block_properties["table_align"]["enum"] == ["left", "center"]
    assert block_properties["border_style"]["enum"] == [
        "minimal",
        "standard",
        "strong",
    ]
    assert block_properties["caption_emphasis"]["enum"] == ["normal", "strong"]
    assert block_properties["start_type"]["enum"] == [
        "new_page",
        "continuous",
        "odd_page",
        "even_page",
        "new_column",
    ]
    assert block_properties["inherit_header_footer"]["type"] == "boolean"
    assert block_properties["page_orientation"]["enum"] == ["portrait", "landscape"]
    assert block_properties["margins"]["type"] == "object"
    assert block_properties["restart_page_numbering"]["type"] == "boolean"
    assert block_properties["page_number_start"]["type"] == "integer"
    assert block_properties["header_footer"]["type"] == "object"
    assert block_properties["items"]["items"]["anyOf"][0]["type"] == "string"
    assert (
        block_properties["items"]["items"]["anyOf"][1]["properties"]["runs"]["items"][
            "properties"
        ]["color"]["type"]
        == "string"
    )
    assert block_properties["rows"]["items"]["items"]["anyOf"][0]["type"] == "string"
    row_cell_properties = block_properties["rows"]["items"]["items"]["anyOf"][1][
        "properties"
    ]
    assert row_cell_properties["text"]["type"] == "string"
    assert row_cell_properties["row_span"]["minimum"] == 1
    assert row_cell_properties["col_span"]["minimum"] == 1
    assert row_cell_properties["fill"]["type"] == "string"
    assert row_cell_properties["text_color"]["type"] == "string"
    assert row_cell_properties["bold"]["type"] == "boolean"
    assert row_cell_properties["align"]["enum"] == ["left", "center", "right"]
    assert row_cell_properties["font_scale"]["type"] == "number"
    assert block_properties["theme_color"]["type"] == "string"
    assert block_properties["text_color"]["type"] == "string"
    assert block_properties["subtitle_color"]["type"] == "string"
    assert block_properties["min_height_pt"]["type"] == "number"
    assert block_properties["full_width"]["type"] == "boolean"
    assert block_properties["accent_color"]["type"] == "string"
    assert block_properties["fill_color"]["type"] == "string"
    assert block_properties["title_color"]["type"] == "string"
    assert block_properties["border_color"]["type"] == "string"
    assert block_properties["border_width_pt"]["type"] == "number"
    assert block_properties["accent_border_width_pt"]["type"] == "number"
    assert block_properties["divider_color"]["type"] == "string"
    assert block_properties["divider_width_pt"]["type"] == "number"
    assert block_properties["padding_pt"]["type"] == "number"
    assert block_properties["title_font_scale"]["type"] == "number"
    assert block_properties["body_font_scale"]["type"] == "number"
    assert block_properties["metrics"]["items"]["required"] == ["label", "value"]
    assert block_properties["metrics"]["items"]["properties"]["label_color"]["type"] == "string"
    assert block_properties["metrics"]["items"]["properties"]["note_color"]["type"] == "string"
    assert (
        block_properties["metrics"]["items"]["properties"]["value_font_scale"]["type"]
        == "number"
    )
    assert (
        block_properties["metrics"]["items"]["properties"]["delta_font_scale"]["type"]
        == "number"
    )
    assert block_properties["label_color"]["type"] == "string"
    assert block_properties["label_font_scale"]["type"] == "number"
    assert block_properties["value_font_scale"]["type"] == "number"
    assert block_properties["delta_font_scale"]["type"] == "number"
    assert block_properties["note_font_scale"]["type"] == "number"
    assert block_properties["cell_padding_horizontal_pt"]["type"] == "number"
    assert block_properties["cell_padding_vertical_pt"]["type"] == "number"
    assert block_properties["header_font_scale"]["type"] == "number"
    assert (
        block_properties["header_footer"]["properties"]["different_odd_even"]["type"]
        == "boolean"
    )
    create_header_footer_properties = next(
        tool for tool in toolset.tools if tool.name == "create_document"
    ).parameters["properties"]["header_footer"]["properties"]
    assert (
        block_properties["header_footer"]["properties"]
        == create_header_footer_properties
    )

    assert (
        block_properties["header_footer"]["properties"]["show_page_number"]["type"]
        == "boolean"
    )
    assert block_properties["levels"]["type"] == "integer"
    assert block_properties["start_on_new_page"]["type"] == "boolean"

def test_create_document_request_accepts_header_footer_defaults():
    request = CreateDocumentRequest(
        title="Paged Report",
        header_footer={
            "header_text": "季度经营复盘",
            "footer_text": "内部使用",
            "header_left": "Q3 经营复盘报告",
            "header_right": "机密 · 2024 年 10 月",
            "footer_left": "集团战略部 · 内部机密文件",
            "footer_right": "第 {PAGE} 页",
            "header_border_bottom": True,
            "footer_border_top": True,
            "header_border_color": "D0D7DE",
            "footer_border_color": "D0D7DE",
            "different_first_page": True,
            "first_page_header_text": "封面页眉",
            "first_page_footer_text": "封面页脚",
            "first_page_show_page_number": True,
            "different_odd_even": True,
            "even_page_header_text": "双面页眉",
            "even_page_footer_text": "双面页脚",
            "even_page_show_page_number": False,
            "show_page_number": True,
            "page_number_align": "center",
            "page_number_format": "upperRoman",
        },
    )

    assert request.header_footer.header_text == "季度经营复盘"
    assert request.header_footer.footer_text == "内部使用"
    assert request.header_footer.header_left == "Q3 经营复盘报告"
    assert request.header_footer.header_right == "机密 · 2024 年 10 月"
    assert request.header_footer.footer_left == "集团战略部 · 内部机密文件"
    assert request.header_footer.footer_right == "第 {PAGE} 页"
    assert request.header_footer.header_border_bottom is True
    assert request.header_footer.footer_border_top is True
    assert request.header_footer.header_border_color == "D0D7DE"
    assert request.header_footer.footer_border_color == "D0D7DE"
    assert request.header_footer.different_first_page is True
    assert request.header_footer.first_page_header_text == "封面页眉"
    assert request.header_footer.first_page_footer_text == "封面页脚"
    assert request.header_footer.first_page_show_page_number is True
    assert request.header_footer.different_odd_even is True
    assert request.header_footer.even_page_header_text == "双面页眉"
    assert request.header_footer.even_page_footer_text == "双面页脚"
    assert request.header_footer.even_page_show_page_number is False
    assert request.header_footer.show_page_number is True
    assert request.header_footer.page_number_align == "center"
    assert request.header_footer.page_number_format == "upperRoman"

def test_section_break_block_rejects_page_number_start_without_restart():
    with pytest.raises(
        ValidationError,
        match="page_number_start requires restart_page_numbering=True",
    ):
        SectionBreakBlock(
            page_number_start=2,
        )

def test_add_blocks_request_rejects_section_page_number_start_without_restart():
    with pytest.raises(
        ValidationError,
        match="page_number_start requires restart_page_numbering=True",
    ):
        AddBlocksRequest(
            document_id="doc-1",
            blocks=[
                {
                    "type": "section_break",
                    "page_number_start": 2,
                }
            ],
        )

def test_section_margins_config_rejects_zero_margin():
    with pytest.raises(ValidationError):
        SectionMarginsConfig(top_cm=0, bottom_cm=2, left_cm=2, right_cm=2)

def test_section_break_block_rejects_margin_greater_than_ten():
    with pytest.raises(ValidationError):
        SectionBreakBlock(
            margins={
                "top_cm": 2,
                "bottom_cm": 2,
                "left_cm": 20,
                "right_cm": 2,
            }
        )

def test_add_blocks_request_rejects_invalid_section_margins():
    with pytest.raises(ValidationError):
        AddBlocksRequest(
            document_id="doc-1",
            blocks=[
                {
                    "type": "section_break",
                    "margins": {
                        "top_cm": 0,
                        "bottom_cm": 2,
                        "left_cm": 2,
                        "right_cm": 2,
                    },
                }
            ],
        )

def test_normalize_raw_block_payloads_repairs_section_toc_and_table_aliases():
    normalized = normalize_raw_block_payloads(
        [
            {"type": "toc", "text": "目录"},
            {
                "type": "heading",
                "text": "三、运营数据分析",
                "level": 1,
                "page_orientation": "landscape",
                "start_on_new_page": True,
                "restart_page_numbering": True,
                "heading_color": "1F4E79",
            },
            {
                "type": "table",
                "title": "第一季度核心运营指标汇总",
                "items": ["用户增长数|10000|10500|105%"],
                "columns": [
                    {"blocks": [{"type": "paragraph", "text": "指标名称"}]},
                    {"blocks": [{"type": "paragraph", "text": "Q1目标值"}]},
                    {"blocks": [{"type": "paragraph", "text": "Q1实际值"}]},
                    {"blocks": [{"type": "paragraph", "text": "达成率"}]},
                ],
            },
            {
                "type": "heading",
                "text": "四、问题与挑战",
                "level": 1,
                "page_orientation": "portrait",
                "start_on_new_page": True,
            },
            {
                "type": "paragraph",
                "text": "封面标题",
                "style": {"font_scale": 3},
                "layout": {"spacing_before": 200, "spacing_after": -5},
            },
        ]
    )

    assert normalized[0] == {"type": "toc", "title": "目录"}
    assert normalized[1] == {
        "type": "section_break",
        "start_type": "new_page",
        "page_orientation": "landscape",
        "restart_page_numbering": True,
    }
    assert normalized[2]["type"] == "heading"
    assert normalized[2]["text"] == "三、运营数据分析"
    assert "heading_color" not in normalized[2]
    assert normalized[3]["type"] == "table"
    assert normalized[3]["headers"] == ["指标名称", "Q1目标值", "Q1实际值", "达成率"]
    assert normalized[3]["rows"] == [["用户增长数", "10000", "10500", "105%"]]
    assert "columns" not in normalized[3]
    assert normalized[4] == {
        "type": "section_break",
        "start_type": "new_page",
        "page_orientation": "portrait",
    }
    assert normalized[5]["type"] == "heading"
    assert normalized[5]["text"] == "四、问题与挑战"
    assert normalized[6]["style"]["font_scale"] == pytest.approx(2.0)
    assert normalized[6]["layout"]["spacing_before"] == pytest.approx(72.0)
    assert normalized[6]["layout"]["spacing_after"] == pytest.approx(0.0)

def test_normalize_raw_block_payloads_strips_markdown_table_edge_pipes():
    normalized = normalize_raw_block_payloads(
        [
            {
                "type": "table",
                "items": ["| 单元格1 | 单元格2 | 单元格3 |"],
            }
        ]
    )

    assert normalized[0]["rows"] == [["单元格1", "单元格2", "单元格3"]]

def test_normalize_raw_block_payloads_repairs_paragraph_items_alias():
    normalized = normalize_raw_block_payloads(
        [
            {
                "type": "paragraph",
                "items": [
                    {
                        "runs": [
                            {"text": "Training Title: ", "bold": True},
                            {"text": "Advanced Skills Workshop"},
                        ]
                    }
                ],
            }
        ]
    )

    assert normalized[0]["type"] == "paragraph"
    assert "items" not in normalized[0]
    assert normalized[0]["runs"][0]["text"] == "Training Title: "
    assert normalized[0]["runs"][1]["text"] == "Advanced Skills Workshop"

def test_normalize_create_document_kwargs_moves_top_level_style_fields():
    normalized = normalize_create_document_kwargs(
        {
            "title": "Sample Training Summary Report",
            "title_align": "center",
            "document_style": {"heading_color": "000000"},
        }
    )

    assert normalized["document_style"]["heading_color"] == "000000"
    assert normalized["document_style"]["title_align"] == "center"

def test_normalize_raw_block_payloads_rejects_excessive_nesting():
    nested_block: dict[str, object] = {"type": "paragraph", "text": "leaf"}
    for _ in range(34):
        nested_block = {"type": "group", "blocks": [nested_block]}

    with pytest.raises(ValueError, match="nesting exceeds limit"):
        normalize_raw_block_payloads([nested_block])

def test_create_document_request_normalizes_separator_only_output_name():
    request = CreateDocumentRequest(output_name="//")

    assert request.output_name == "document.docx"

def test_export_document_request_normalizes_dot_only_output_name():
    request = ExportDocumentRequest(document_id="doc-1", output_name=".")

    assert request.output_name == "document.docx"

def test_paragraph_schema_requires_text_or_runs():
    with pytest.raises(ValidationError, match="paragraph requires text or runs"):
        SectionParagraphInput.model_validate(
            {
                "type": "paragraph",
                "text": "",
                "runs": [],
            }
        )

def test_build_document_render_payload_preserves_runs_when_text_and_runs_exist():
    block = ParagraphBlock(
        text="plain text",
        runs=[
            ParagraphRun(text="rich"),
            ParagraphRun(text=" content"),
        ],
    )
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        metadata=DocumentMetadata(title="Rich Paragraph"),
        blocks=[block],
    )
    payload = build_document_render_payload(document)

    assert payload["blocks"][0]["text"] == "plain text"
    assert payload["blocks"][0]["runs"][0]["text"] == "rich"
    assert payload["blocks"][0]["runs"][1]["text"] == " content"


def test_build_document_render_payload_preserves_hyperlink_run_urls():
    block = ParagraphBlock(
        runs=[
            ParagraphRun(text="文档地址", url="https://example.com/docs"),
        ],
    )
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        metadata=DocumentMetadata(title="Hyperlink Paragraph"),
        blocks=[block],
    )

    payload = build_document_render_payload(document)

    assert payload["blocks"][0]["runs"][0]["url"] == "https://example.com/docs"


def test_paragraph_run_accepts_scheme_only_https_url():
    run = ParagraphRun(text="文档地址", url="https:example.com")

    assert run.url == "https:example.com"

def test_build_document_render_payload_keeps_default_metadata_fields():
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        metadata=DocumentMetadata(title="Default Metadata"),
    )

    payload = build_document_render_payload(document)

    assert payload["metadata"]["preferred_filename"] == "document.docx"
    assert payload["metadata"]["theme_name"] == "business_report"
    assert payload["metadata"]["table_template"] == "report_grid"
    assert payload["metadata"]["density"] == "comfortable"

def test_build_document_render_payload_keeps_page_template_required_defaults():
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        blocks=[
            PageTemplateBlock(
                template="technical_resume",
                data=TechnicalResumeData(
                    name="张明远",
                    contact_line="zhang@example.com",
                    sections=[
                        ResumeSection(
                            title="技术栈",
                            lines=["Go", "Python"],
                        )
                    ],
                ),
            )
        ],
    )

    payload = build_document_render_payload(document)

    assert payload["blocks"][0]["type"] == "page_template"
    assert payload["blocks"][0]["data"]["headline"] == ""


def test_build_document_render_payload_omits_none_in_page_template_runs():
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        blocks=[
            PageTemplateBlock(
                template="technical_resume",
                data=TechnicalResumeData(
                    name="张明远",
                    contact_line="zhang@example.com",
                    sections=[
                        ResumeSection(
                            title="实习经历",
                            entries=[
                                ResumeSectionEntry(
                                    heading="字节跳动",
                                    details=[
                                        ListItem(
                                            runs=[
                                                ParagraphRun(text="主导优化推荐引擎召回模块，", bold=True),
                                                ParagraphRun(text="将离线索引构建耗时压缩到 1/4。"),
                                            ]
                                        )
                                    ],
                                )
                            ],
                        )
                    ],
                ),
            )
        ],
    )

    payload = build_document_render_payload(document)
    runs = payload["blocks"][0]["data"]["sections"][0]["entries"][0]["details"][0]["runs"]

    assert all("color" not in run for run in runs)
    assert all("url" not in run for run in runs)


@pytest.mark.parametrize(
    "url",
    [
        "javascript:alert(1)",
        "https://exa mple.com",
        "ftp://example.com",
        "mailto:",
    ],
)
def test_paragraph_run_rejects_invalid_hyperlink_url(url: str):
    with pytest.raises(ValidationError, match="http, https, or mailto"):
        ParagraphRun(text="错误链接", url=url)

def test_build_document_render_payload_omits_unset_heading_bottom_border():
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        metadata=DocumentMetadata(title="Heading Defaults"),
        blocks=[
            HeadingBlock(text="默认分割线标题", level=1),
            HeadingBlock(text="显式关闭分割线标题", level=1, bottom_border=False),
        ],
    )

    payload = build_document_render_payload(document)

    assert "bottom_border" not in payload["blocks"][0]
    assert payload["blocks"][1]["bottom_border"] is False


def test_build_document_render_payload_omits_none_hero_banner_fields():
    document = DocumentModel(
        document_id="doc-1",
        session_id="",
        format="word",
        metadata=DocumentMetadata(title="Hero Banner Defaults"),
        blocks=[
            HeroBannerBlock(
                title="Q3 经营复盘报告",
                subtitle="围绕收入质量、区域表现与下季度动作的管理层复盘",
                theme_color="1F4E79",
                full_width=True,
            )
        ],
    )

    payload = build_document_render_payload(document)
    hero_banner = payload["blocks"][0]

    assert hero_banner["type"] == "hero_banner"
    assert hero_banner["title"] == "Q3 经营复盘报告"
    assert hero_banner["subtitle"] == "围绕收入质量、区域表现与下季度动作的管理层复盘"
    assert hero_banner["theme_color"] == "1F4E79"
    assert "text_color" not in hero_banner
    assert "subtitle_color" not in hero_banner
    assert "min_height_pt" not in hero_banner

def test_normalize_raw_block_payloads_moves_legacy_layout_alignment_into_style():
    normalized = normalize_raw_block_payloads(
        [
            {
                "paragraph": {
                    "text": "右对齐段落",
                    "layout": {"alignment": "right", "spacing_after": 6},
                }
            }
        ]
    )

    assert normalized[0]["type"] == "paragraph"
    assert normalized[0]["style"]["align"] == "right"
    assert normalized[0]["layout"] == {"spacing_after": 6}

def test_create_document_request_normalizes_document_style():
    request = CreateDocumentRequest(
        title="Styled",
        document_style={
            "brief": "deep blue business report",
            "heading_color": "#1f4e79",
            "heading_level_1_color": "0f4c81",
            "heading_level_2_color": "4b5563",
            "heading_bottom_border_color": "d0d7de",
            "heading_bottom_border_size_pt": 1.25,
            "title_align": "left",
            "body_font_size": 12,
            "body_line_spacing": 1.25,
            "font_name": "Microsoft YaHei",
            "heading_font_name": "Source Han Sans SC",
            "table_font_name": "SimSun",
            "code_font_name": "Consolas",
            "paragraph_space_after": 14,
            "list_space_after": 11,
            "summary_card_defaults": _summary_card_defaults(),
            "table_defaults": {
                "preset": "minimal",
                "header_fill": "#dce6f1",
                "body_fill": "f8fafc",
                "header_text_color": "ffffff",
                "banded_rows": True,
                "banded_row_fill": "eef4fa",
                "first_column_bold": True,
                "table_align": "left",
                "border_style": "strong",
                "caption_emphasis": "strong",
                "cell_align": "center",
            },
        },
    )

    assert request.document_style.brief == "deep blue business report"
    assert request.document_style.heading_color == "1F4E79"
    assert request.document_style.heading_level_1_color == "0F4C81"
    assert request.document_style.heading_level_2_color == "4B5563"
    assert request.document_style.heading_bottom_border_color == "D0D7DE"
    assert request.document_style.heading_bottom_border_size_pt == 1.25
    assert request.document_style.title_align == "left"
    assert request.document_style.body_font_size == 12
    assert request.document_style.body_line_spacing == 1.25
    assert request.document_style.font_name == "Microsoft YaHei"
    assert request.document_style.heading_font_name == "Source Han Sans SC"
    assert request.document_style.table_font_name == "SimSun"
    assert request.document_style.code_font_name == "Consolas"
    assert request.document_style.paragraph_space_after == 14
    assert request.document_style.list_space_after == 11
    assert request.document_style.summary_card_defaults.title_align == "center"
    assert request.document_style.summary_card_defaults.title_emphasis == "strong"
    assert request.document_style.summary_card_defaults.title_font_scale == 1.2
    assert request.document_style.summary_card_defaults.title_space_before == 12
    assert request.document_style.summary_card_defaults.title_space_after == 4
    assert request.document_style.summary_card_defaults.list_space_after == 8
    assert request.document_style.table_defaults.preset == "minimal"
    assert request.document_style.table_defaults.header_fill == "DCE6F1"
    assert request.document_style.table_defaults.body_fill == "F8FAFC"
    assert request.document_style.table_defaults.header_text_color == "FFFFFF"
    assert request.document_style.table_defaults.banded_rows is True
    assert request.document_style.table_defaults.banded_row_fill == "EEF4FA"
    assert request.document_style.table_defaults.first_column_bold is True
    assert request.document_style.table_defaults.table_align == "left"
    assert request.document_style.table_defaults.border_style == "strong"
    assert request.document_style.table_defaults.caption_emphasis == "strong"
    assert request.document_style.table_defaults.cell_align == "center"

def test_create_document_request_rejects_invalid_document_style_heading_color():
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Invalid heading color",
            document_style={
                "brief": "invalid color",
                "heading_color": "blue",
            },
        )

def test_create_document_request_rejects_invalid_table_defaults_header_fill():
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Invalid table header fill",
            document_style={
                "brief": "invalid table defaults",
                "table_defaults": {
                    "header_fill": "123",
                },
            },
        )

def test_add_blocks_request_rejects_invalid_hero_banner_colors():
    with pytest.raises(ValidationError, match="6-digit hex color"):
        AddBlocksRequest(
            document_id="doc-1",
            blocks=[
                {
                    "type": "hero_banner",
                    "title": "Q3 经营复盘报告",
                    "theme_color": "blue",
                }
            ],
        )

def test_create_document_request_defaults_nested_document_style_sections():
    request = CreateDocumentRequest(
        title="Nested Defaults",
        document_style={
            "brief": "defaults only",
        },
    )

    assert request.document_style.summary_card_defaults is not None
    assert request.document_style.table_defaults is not None
    assert request.document_style.summary_card_defaults.title_align is None
    assert request.document_style.table_defaults.header_fill is None

def test_create_document_request_normalizes_blank_document_style_colors_to_none():
    request = CreateDocumentRequest(
        title="Blank Colors",
        document_style={
            "brief": "blank colors",
            "heading_color": "   ",
            "table_defaults": {
                "header_fill": " ",
                "header_text_color": "",
                "banded_row_fill": "\t",
            },
        },
    )

    assert request.document_style.heading_color is None
    assert request.document_style.table_defaults.header_fill is None
    assert request.document_style.table_defaults.header_text_color is None
    assert request.document_style.table_defaults.banded_row_fill is None

def test_create_document_request_rejects_extra_document_style_keys():
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Unexpected document style key",
            document_style={
                "brief": "has extra key",
                "unknown_field": "nope",
            },
        )

    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Unexpected table default key",
            document_style={
                "brief": "has nested extra key",
                "table_defaults": {
                    "header_fill": "#FFFFFF",
                    "bogus": "value",
                },
            },
        )

@pytest.mark.parametrize("body_font_size", [8, 17])
def test_create_document_request_rejects_out_of_range_body_font_size(body_font_size):
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Styled",
            document_style={
                "brief": "deep blue business report",
                "body_font_size": body_font_size,
            },
        )

@pytest.mark.parametrize("body_line_spacing", [0.9, 2.6])
def test_create_document_request_rejects_out_of_range_body_line_spacing(
    body_line_spacing,
):
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Styled",
            document_style={
                "brief": "deep blue business report",
                "body_line_spacing": body_line_spacing,
            },
        )

@pytest.mark.parametrize("title_font_scale", [0.5, 2.5])
def test_create_document_request_rejects_out_of_range_title_font_scale(
    title_font_scale,
):
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Styled",
            document_style={
                "brief": "deep blue business report",
                "summary_card_defaults": {
                    "title_font_scale": title_font_scale,
                },
            },
        )

@pytest.mark.parametrize(
    "field_name,value",
    [
        ("paragraph_space_after", -1),
        ("paragraph_space_after", 73),
        ("list_space_after", -5),
        ("list_space_after", 100),
    ],
)
def test_create_document_request_rejects_out_of_range_document_spacing_fields(
    field_name, value
):
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Styled",
            document_style={
                "brief": "deep blue business report",
                field_name: value,
            },
        )

@pytest.mark.parametrize(
    "field_name,value",
    [
        ("title_space_before", -10),
        ("title_space_before", 100),
        ("title_space_after", -2),
        ("title_space_after", 80),
        ("list_space_after", -5),
        ("list_space_after", 100),
    ],
)
def test_create_document_request_rejects_out_of_range_summary_card_spacing_fields(
    field_name, value
):
    with pytest.raises(ValidationError):
        CreateDocumentRequest(
            title="Styled",
            document_style={
                "brief": "deep blue business report",
                "summary_card_defaults": {
                    field_name: value,
                },
            },
        )

def test_table_schema_normalizers_are_shared():
    request = AddTableRequest(
        document_id="doc-1",
        headers=["区域", "目标"],
        rows=[["华东", "120"]],
        header_groups=[{"title": "经营数据", "span": 2}],
        table_style="invalid-style",
        column_widths=[4.2, 0, -1.0, 3.0],
        numeric_columns=[2, -1, 1, 2],
        header_fill="#1f4e79",
        header_text_color="ffffff",
        banded_rows=True,
        banded_row_fill="eef4fa",
        first_column_bold=True,
        table_align="center",
        border_style="strong",
        caption_emphasis="strong",
    )
    section = SectionTableInput(
        headers=["区域"],
        rows=[["华东"]],
        header_groups=[{"title": "经营概览", "span": 1}],
        table_style="invalid-style",
        column_widths=[4.2, 0, -1.0, 3.0],
        numeric_columns=[2, -1, 1, 2],
        header_fill="1f4e79",
        header_text_color="#ffffff",
        banded_rows=False,
        banded_row_fill="#eef4fa",
        first_column_bold=False,
        table_align="left",
        border_style="minimal",
        caption_emphasis="normal",
    )

    assert request.table_style == ""
    assert request.column_widths == [4.2, 0, 0, 3.0]
    assert request.numeric_columns == [1, 2]
    assert request.header_groups[0].span == 2
    assert request.header_fill == "1F4E79"
    assert request.header_text_color == "FFFFFF"
    assert request.banded_rows is True
    assert request.banded_row_fill == "EEF4FA"
    assert request.first_column_bold is True
    assert request.table_align == "center"
    assert request.border_style == "strong"
    assert request.caption_emphasis == "strong"
    assert section.table_style == ""
    assert section.column_widths == [4.2, 0, 0, 3.0]
    assert section.numeric_columns == [1, 2]
    assert section.header_groups[0].title == "经营概览"
    assert section.header_fill == "1F4E79"
    assert section.header_text_color == "FFFFFF"
    assert section.banded_rows is False
    assert section.banded_row_fill == "EEF4FA"
    assert section.first_column_bold is False
    assert section.table_align == "left"
    assert section.border_style == "minimal"
    assert section.caption_emphasis == "normal"

def test_add_table_request_rejects_grouped_header_span_total_mismatch():
    with pytest.raises(
        ValidationError,
        match=r"header_groups span total \(1\) must equal column count \(2\)",
    ):
        AddTableRequest(
            document_id="doc-1",
            headers=["区域", "目标"],
            rows=[["华东", "120"]],
            header_groups=[{"title": "经营数据", "span": 1}],
        )

def test_section_table_input_rejects_grouped_header_span_below_minimum():
    with pytest.raises(ValidationError, match="greater than or equal to 1"):
        SectionTableInput(
            headers=["区域", "目标"],
            rows=[["华东", "120"]],
            header_groups=[{"title": "经营数据", "span": 0}],
        )

@pytest.mark.parametrize(
    "field_name", ["header_fill", "header_text_color", "banded_row_fill"]
)
@pytest.mark.parametrize(
    "invalid_color",
    [
        "blue",
        "#123",
        "12345g",
    ],
)
def test_add_table_request_rejects_invalid_color_fields(
    field_name: str,
    invalid_color: str,
):
    kwargs = {
        "document_id": "doc-1",
        "headers": ["区域"],
        "rows": [["华东"]],
        field_name: invalid_color,
    }

    with pytest.raises(ValidationError, match="6-digit hex color"):
        AddTableRequest(**kwargs)

@pytest.mark.parametrize(
    "field_name", ["header_fill", "header_text_color", "banded_row_fill"]
)
@pytest.mark.parametrize(
    "invalid_color",
    [
        "blue",
        "#123",
        "12345g",
    ],
)
def test_section_table_input_rejects_invalid_color_fields(
    field_name: str,
    invalid_color: str,
):
    kwargs = {
        "headers": ["区域"],
        "rows": [["华东"]],
        field_name: invalid_color,
    }

    with pytest.raises(ValidationError, match="6-digit hex color"):
        SectionTableInput(**kwargs)

def test_section_table_input_rejects_invalid_border_style():
    with pytest.raises(ValidationError):
        SectionTableInput(
            headers=["区域"],
            rows=[["华东"]],
            border_style="heavy",
        )

def test_table_block_rejects_header_groups_without_columns():
    with pytest.raises(
        ValidationError,
        match=r"header_groups require at least one column from headers or rows \(column_count=0\)",
    ):
        TableBlock(header_groups=[{"title": "经营数据", "span": 1}])

def test_table_schema_allows_empty_placeholder_tables():
    request = AddTableRequest(document_id="doc-1", headers=[], rows=[])
    section = SectionTableInput(headers=[], rows=[])
    block = TableBlock()

    assert request.headers == []
    assert request.rows == []
    assert section.headers == []
    assert section.rows == []
    assert block.headers == []
    assert block.rows == []

def test_add_table_request_supports_vertical_merge_cells():
    request = AddTableRequest(
        document_id="doc-1",
        headers=["日期", "时间", "课程"],
        rows=[
            [{"text": "第一天", "row_span": 2}, "09:00", "课程 A"],
            ["13:00", "课程 B"],
        ],
        header_fill_enabled=False,
        header_bold=False,
    )

    assert request.rows[0][0].row_span == 2
    assert request.header_fill_enabled is False
    assert request.header_bold is False


def test_add_table_request_supports_horizontal_merge_cells():
    request = AddTableRequest(
        document_id="doc-1",
        headers=["季度", "营收", "利润", "备注"],
        rows=[
            [
                {"text": "Q1 汇总", "col_span": 2},
                "18%",
                "达成",
            ],
        ],
    )

    assert request.rows[0][0].col_span == 2


def test_add_table_request_rejects_non_positive_col_span():
    with pytest.raises(ValidationError, match="greater than or equal to 1"):
        AddTableRequest(
            document_id="doc-1",
            headers=["季度", "营收", "利润", "备注"],
            rows=[
                [
                    {"text": "Q1 汇总", "col_span": 0},
                    "18%",
                    "达成",
                ],
            ],
        )


def test_add_table_request_rejects_combined_row_and_column_spans():
    with pytest.raises(
        ValidationError,
        match="table cell cannot combine row_span and col_span",
    ):
        AddTableRequest(
            document_id="doc-1",
            headers=["季度", "营收", "备注"],
            rows=[
                [
                    {"text": "Q1", "row_span": 2, "col_span": 2},
                    "达成",
                ],
            ],
        )


def test_add_table_request_rejects_horizontal_merge_overlapping_vertical_merge():
    with pytest.raises(
        ValidationError,
        match="table row 2 overlaps active row spans",
    ):
        AddTableRequest(
            document_id="doc-1",
            headers=["分类", "区域", "完成率"],
            rows=[
                ["单体", {"text": "上海", "row_span": 2}, "108%"],
                [{"text": "汇总", "col_span": 2}, "达成"],
            ],
        )


def test_add_table_request_rejects_horizontal_merge_exceeding_column_count():
    with pytest.raises(
        ValidationError,
        match="table row 1 exceeds column count",
    ):
        AddTableRequest(
            document_id="doc-1",
            headers=["季度", "营收", "利润", "备注"],
            rows=[
                [
                    {"text": "Q1 汇总", "col_span": 3},
                    "18%",
                    "达成",
                ],
            ],
        )

def test_add_table_request_accepts_empty_placeholder_cells_for_vertical_merge_rows():
    request = AddTableRequest(
        document_id="doc-1",
        headers=["日期", "时间", "课程"],
        rows=[
            [{"text": "第一天", "row_span": 2}, "09:00", "课程 A"],
            ["", "13:00", "课程 B"],
        ],
    )

    assert request.rows[1][0] == ""
    assert request.rows[1][1] == "13:00"

def test_add_table_request_rejects_underfilled_rows():
    with pytest.raises(ValidationError, match="table row 1 is missing cells"):
        AddTableRequest(
            document_id="doc-1",
            headers=["日期", "时间", "课程"],
            rows=[["第一天", "09:00"]],
        )

def test_add_blocks_request_accepts_accent_box_metric_cards_and_table_cell_styles():
    request = AddBlocksRequest(
        document_id="doc-1",
        blocks=[
            {
                "type": "accent_box",
                "title": "核心摘要",
                "items": [
                    {
                        "runs": [
                            {"text": "营业收入：", "bold": True},
                            {"text": "保持增长"},
                        ]
                    }
                ],
                "accent_color": "1F4E79",
                "fill_color": "F8FAFC",
            },
            {
                "type": "metric_cards",
                "metrics": [
                    {
                        "label": "营业收入",
                        "value": "¥4.82 亿",
                        "delta": "↑ 18.4% YoY",
                        "delta_color": "15803D",
                    }
                ],
                "accent_color": "1F4E79",
            },
            {
                "type": "table",
                "headers": ["区域", "预算完成率"],
                "rows": [
                    [
                        "华东",
                        {
                            "text": "112%",
                            "fill": "DCFCE7",
                            "text_color": "166534",
                            "bold": True,
                            "align": "right",
                        },
                    ]
                ],
            },
        ],
    )

    accent_box = request.blocks[0]
    metric_cards = request.blocks[1]
    table = request.blocks[2]

    assert accent_box.accent_color == "1F4E79"
    assert accent_box.items[0].runs[0].bold is True
    assert metric_cards.metrics[0].delta_color == "15803D"
    assert table.rows[0][1].fill == "DCFCE7"
    assert table.rows[0][1].text_color == "166534"
    assert table.rows[0][1].bold is True
    assert table.rows[0][1].align == "right"

def test_normalize_raw_block_payloads_flattens_nested_block_aliases():
    normalized = normalize_raw_block_payloads(
        [
            {
                "type": "hero_banner",
                "hero_banner": {
                    "text": "Q3 经营复盘报告",
                    "subtitle": "战略与增长委员会",
                    "title_color": "FFFFFF",
                },
            },
            {
                "type": "accent_box",
                "accent_box": {
                    "title": "核心摘要",
                    "content": "经营质量继续改善。",
                },
            },
            {
                "type": "metric_cards",
                "metric_cards": {
                    "cards": [
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "+12.4% YoY",
                        }
                    ]
                },
            },
            {
                "type": "table",
                "table": {
                    "headers": [
                        {"text": "区域"},
                        {"text": "营收（万）"},
                        {"text": "备注"},
                    ],
                    "rows": [
                        {
                            "cells": [
                                {"text": "华东大区"},
                                {"text": "18,450"},
                                {"text": "重点项目提前交付"},
                            ]
                        }
                    ],
                },
            },
            {
                "type": "paragraph",
                "paragraph": {
                    "runs": [
                        {"text": "供应链风险：", "bold": True},
                        {"text": "需建立多元供应商体系。"},
                    ],
                    "layout": {"alignment": "right"},
                },
            },
        ]
    )

    hero_banner = normalized[0]
    accent_box = normalized[1]
    metric_cards = normalized[2]
    table = normalized[3]
    paragraph = normalized[4]

    assert hero_banner["title"] == "Q3 经营复盘报告"
    assert hero_banner["subtitle"] == "战略与增长委员会"
    assert hero_banner["text_color"] == "FFFFFF"
    assert "hero_banner" not in hero_banner
    assert accent_box["text"] == "经营质量继续改善。"
    assert "content" not in accent_box
    assert metric_cards["metrics"][0]["value"] == "¥4.82 亿"
    assert "cards" not in metric_cards
    assert table["headers"] == ["区域", "营收（万）", "备注"]
    assert table["rows"][0][0]["text"] == "华东大区"
    assert table["rows"][0][1]["text"] == "18,450"
    assert table["rows"][0][2]["text"] == "重点项目提前交付"
    assert paragraph["runs"][0]["text"] == "供应链风险："
    assert paragraph["style"]["align"] == "right"

def test_add_blocks_request_accepts_hero_banner_and_report_style_fields():
    request = AddBlocksRequest(
        document_id="doc-1",
        blocks=[
            {
                "type": "hero_banner",
                "title": "Q3 经营复盘报告",
                "subtitle": "战略与增长委员会",
                "theme_color": "1F4E79",
                "text_color": "FFFFFF",
                "subtitle_color": "DCE6F1",
                "min_height_pt": 96,
                "full_width": True,
                "layout": {
                    "spacing_after": 12,
                    "padding_top_pt": 18,
                    "padding_left_pt": 20,
                },
            },
            {
                "type": "accent_box",
                "title": "核心摘要",
                "text": "经营质量继续改善。",
                "accent_color": "1F4E79",
                "border_color": "CBD5E1",
                "border_width_pt": 0.75,
                "accent_border_width_pt": 3.0,
                "padding_pt": 16,
                "title_font_scale": 1.2,
                "body_font_scale": 1.05,
            },
            {
                "type": "metric_cards",
                "accent_color": "1F4E79",
                "border_color": "D9E1E8",
                "divider_color": "CBD5E1",
                "border_width_pt": 0.75,
                "divider_width_pt": 0.75,
                "padding_pt": 14,
                "label_font_scale": 0.9,
                "value_font_scale": 1.8,
                "delta_font_scale": 0.9,
                "note_font_scale": 0.82,
                "metrics": [
                    {
                        "label": "营业收入",
                        "value": "¥4.82 亿",
                        "label_color": "475569",
                        "value_font_scale": 1.95,
                        "delta": "↑ 18.4% YoY",
                        "delta_font_scale": 1.1,
                        "note": "核心业务保持增长",
                        "note_color": "64748B",
                    }
                ],
            },
            {
                "type": "table",
                "headers": ["区域", "完成率"],
                "header_font_scale": 1.1,
                "body_font_scale": 0.95,
                "cell_padding_horizontal_pt": 8,
                "cell_padding_vertical_pt": 6,
                "rows": [
                    [
                        "华东",
                        {
                            "text": "112%",
                            "font_scale": 1.2,
                            "fill": "DCFCE7",
                            "text_color": "166534",
                        },
                    ]
                ],
            },
        ],
    )

    hero_banner = request.blocks[0]
    accent_box = request.blocks[1]
    metric_cards = request.blocks[2]
    table = request.blocks[3]

    assert hero_banner.title == "Q3 经营复盘报告"
    assert hero_banner.theme_color == "1F4E79"
    assert hero_banner.layout.padding_top_pt == pytest.approx(18)
    assert accent_box.border_color == "CBD5E1"
    assert accent_box.accent_border_width_pt == pytest.approx(3.0)
    assert accent_box.title_font_scale == pytest.approx(1.2)
    assert metric_cards.divider_color == "CBD5E1"
    assert metric_cards.note_font_scale == pytest.approx(0.82)
    assert metric_cards.metrics[0].label_color == "475569"
    assert metric_cards.metrics[0].value_font_scale == pytest.approx(1.95)
    assert metric_cards.metrics[0].delta_font_scale == pytest.approx(1.1)
    assert metric_cards.metrics[0].note_color == "64748B"
    assert table.header_font_scale == pytest.approx(1.1)
    assert table.body_font_scale == pytest.approx(0.95)
    assert table.cell_padding_horizontal_pt == pytest.approx(8)
    assert table.rows[0][1].font_scale == pytest.approx(1.2)

def test_add_blocks_request_accepts_page_template_business_review_cover():
    request = AddBlocksRequest(
        document_id="doc-1",
        blocks=[
            _business_review_cover_block(
                summary_text="Q3 营收同比增长 18.4%，整体毛利率保持稳定。",
                metrics=[
                    {
                        "label": "营业收入",
                        "value": "¥4.82 亿",
                        "delta": "↑ 18.4% YoY",
                        "delta_color": "15803D",
                    },
                    {
                        "label": "毛利率",
                        "value": "42.1%",
                        "delta": "↓ 0.3pp vs Q2",
                    },
                ],
                footer_note="编制：战略发展部 · 审核：CFO 办公室",
            )
        ],
    )

    block = request.blocks[0]
    assert block.template == "business_review_cover"
    assert block.data.summary_title == "核心摘要"
    assert block.data.metrics[0].delta_color == "15803D"
    assert block.data.auto_page_break is False


def test_add_blocks_request_accepts_page_template_technical_resume():
    request = AddBlocksRequest(
        document_id="doc-1",
        blocks=[
            _technical_resume_block(
                sections=[
                    {
                        "title": "教育背景",
                        "entries": [
                            {
                                "heading": "北京大学",
                                "date": "2019.09 – 2023.06",
                                "subtitle": "计算机科学与技术  |  工学学士",
                                "details": [
                                    "GPA 3.86/4.0，连续三年一等奖学金，排名前 5%"
                                ],
                            }
                        ],
                    },
                    {
                        "title": "技术栈",
                        "lines": ["语言：Go（熟练）、Java（熟练）、Python、SQL"],
                    },
                ]
            )
        ],
    )

    block = request.blocks[0]
    assert block.template == "technical_resume"
    assert block.data.name == "张明远"
    assert block.data.sections[0].entries[0].heading == "北京大学"
    assert block.data.sections[1].lines[0] == "语言：Go（熟练）、Java（熟练）、Python、SQL"

def test_add_table_request_rejects_vertical_merge_rows_that_exceed_header_columns():
    with pytest.raises(
        ValidationError,
        match=r"table row 2 exceeds column count \(2\)",
    ):
        AddTableRequest(
            document_id="doc-1",
            headers=["日期", "课程"],
            rows=[
                [{"text": "第一天", "row_span": 2}, "课程 A"],
                ["09:00", "课程 B"],
            ],
        )
