import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from tests._docx_test_helpers import (
    _business_report_metadata,
    _cell_border_color,
    _cell_border_size,
    _raw_row_cell_vertical_merge,
    _render_structured_payload_with_node,
    _table_border_color,
    _table_border_size,
    _table_grid_widths,
)


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _run_has_prop(run, prop_name: str) -> bool:
    from docx.oxml.ns import qn

    r_pr = run._r.rPr
    if r_pr is None:
        return False
    prop = r_pr.find(qn(f"w:{prop_name}"))
    if prop is None:
        return False
    value = prop.get(qn("w:val"))
    if value is None:
        return True
    return value not in {"0", "false", "off", "none"}


def _run_font_name(run) -> str | None:
    from docx.oxml.ns import qn

    r_pr = run._r.rPr
    if r_pr is None:
        return None
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        return None
    return fonts.get(qn("w:ascii"))


def _run_color(run) -> str | None:
    from docx.oxml.ns import qn

    r_pr = run._r.rPr
    if r_pr is None:
        return None
    color = r_pr.find(qn("w:color"))
    if color is None:
        return None
    return color.get(qn("w:val"))


def _run_size_pt(run) -> float | None:
    from docx.oxml.ns import qn

    r_pr = run._r.rPr
    if r_pr is None:
        return None
    size = r_pr.find(qn("w:sz"))
    if size is None:
        return None
    raw_value = size.get(qn("w:val"))
    if raw_value is None:
        return None
    return int(raw_value) / 2


def test_node_renderer_supports_table_cell_rich_style_controls(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-rich-style-controls",
        {
            "document_id": "table-rich-style-controls",
            "metadata": _business_report_metadata(title="表格富样式"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["项目", "富文本", "备注"],
                    "border_style": "standard",
                    "body_font_scale": 0.9,
                    "rows": [
                        [
                            "旧式单元格",
                            {
                                "text": "占位文本",
                                "font_name": "SimSun",
                                "font_scale": 1.2,
                                "bold": True,
                                "italic": True,
                                "underline": True,
                                "strikethrough": True,
                                "text_color": "666666",
                                "border": {
                                    "top": {
                                        "style": "double",
                                        "color": "1F4E79",
                                        "width_pt": 1.0,
                                    },
                                    "right": {
                                        "style": "dashed",
                                        "color": "0F766E",
                                        "width_pt": 0.75,
                                    },
                                },
                                "runs": [
                                    {"text": "默认"},
                                    {
                                        "text": "覆盖",
                                        "font_name": "Consolas",
                                        "font_scale": 1.5,
                                        "bold": False,
                                        "italic": False,
                                        "underline": False,
                                        "strikethrough": False,
                                        "color": "B91C1C",
                                    },
                                ],
                            },
                            "普通文本",
                        ],
                        [
                            "兼容",
                            "文本",
                            "表格",
                        ],
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    plain_cell = table.rows[1].cells[0]
    rich_cell = table.rows[1].cells[1]
    fallback_cell = table.rows[2].cells[2]

    assert plain_cell.text == "旧式单元格"
    assert rich_cell.text == "默认覆盖"
    assert fallback_cell.text == "表格"

    plain_run = plain_cell.paragraphs[0].runs[0]
    rich_first_run = rich_cell.paragraphs[0].runs[0]
    rich_second_run = rich_cell.paragraphs[0].runs[1]

    assert _run_font_name(rich_first_run) == "SimSun"
    assert _run_font_name(rich_second_run) == "Consolas"
    assert _run_color(rich_first_run) == "666666"
    assert _run_color(rich_second_run) == "B91C1C"
    assert _run_has_prop(rich_first_run, "b") is True
    assert _run_has_prop(rich_first_run, "i") is True
    assert _run_has_prop(rich_first_run, "u") is True
    assert _run_has_prop(rich_first_run, "strike") is True
    assert _run_has_prop(rich_second_run, "b") is False
    assert _run_has_prop(rich_second_run, "i") is False
    assert _run_has_prop(rich_second_run, "u") is False
    assert _run_has_prop(rich_second_run, "strike") is False
    assert _run_size_pt(rich_first_run) is not None
    assert _run_size_pt(rich_second_run) is not None
    assert _run_size_pt(rich_second_run) == _run_size_pt(rich_first_run)
    assert _run_size_pt(rich_first_run) > _run_size_pt(plain_run)

    assert _cell_border_color(rich_cell, "top") == "1F4E79"
    assert _cell_border_color(rich_cell, "right") == "0F766E"
    assert _cell_border_color(rich_cell, "left") == _table_border_color(table, "left")
    assert _cell_border_color(rich_cell, "bottom") == _table_border_color(table, "bottom")
    assert _cell_border_size(rich_cell, "top") == "8"
    assert _cell_border_size(rich_cell, "right") == "6"
    assert _cell_border_size(rich_cell, "left") == _table_border_size(table, "left")
    assert _cell_border_size(rich_cell, "bottom") == _table_border_size(table, "bottom")


def test_node_renderer_clamps_table_cell_run_font_scale_to_cell_default(
    workspace_root: Path,
):
    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-cell-font-scale-clamp",
        {
            "document_id": "table-cell-font-scale-clamp",
            "metadata": _business_report_metadata(title="表格字号收敛"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["项目", "备注"],
                    "rows": [
                        [
                            "毛利率",
                            {
                                "font_name": "SimSun",
                                "font_scale": 1.15,
                                "text_color": "666666",
                                "runs": [
                                    {"text": "基础盘稳定，"},
                                    {
                                        "text": "重点增长来自华东渠道，",
                                        "font_name": "Consolas",
                                        "font_scale": 1.35,
                                        "color": "B91C1C",
                                    },
                                    {"text": "查看明细", "url": "https://example.com/detail"},
                                ],
                            },
                        ]
                    ],
                }
            ],
        },
    )

    remark_runs = loaded_doc.tables[0].rows[1].cells[1].paragraphs[0].runs

    assert len(remark_runs) == 2
    assert _run_size_pt(remark_runs[0]) == _run_size_pt(remark_runs[1])

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    document_root = ET.fromstring(document_xml)
    remark_cell = document_root.findall(".//w:tbl", NS)[0].findall("./w:tr", NS)[1].findall(
        "./w:tc", NS
    )[1]
    hyperlink_run_size = (
        remark_cell.find(".//w:hyperlink//w:rPr/w:sz", NS)
        .attrib.get(f"{{{NS['w']}}}val")
    )
    first_run_size = remark_cell.findall(".//w:rPr/w:sz", NS)[0].attrib.get(
        f"{{{NS['w']}}}val"
    )

    assert hyperlink_run_size == first_run_size


def test_node_renderer_allows_table_cell_run_font_scale_to_shrink(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-cell-font-scale-shrink",
        {
            "document_id": "table-cell-font-scale-shrink",
            "metadata": _business_report_metadata(title="表格字号缩小"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["项目", "备注"],
                    "rows": [
                        [
                            "毛利率",
                            {
                                "font_name": "SimSun",
                                "font_scale": 1.15,
                                "text_color": "666666",
                                "runs": [
                                    {"text": "基础盘稳定，"},
                                    {
                                        "text": "附注",
                                        "font_scale": 0.95,
                                        "color": "B91C1C",
                                    },
                                ],
                            },
                        ]
                    ],
                }
            ],
        },
    )

    remark_runs = loaded_doc.tables[0].rows[1].cells[1].paragraphs[0].runs

    assert len(remark_runs) == 2
    assert _run_size_pt(remark_runs[1]) < _run_size_pt(remark_runs[0])


def test_node_renderer_renders_valid_hyperlink_in_table_cell(
    workspace_root: Path,
):
    url = "https://example.com/docs"
    link_text = "在线说明"

    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-cell-hyperlink",
        {
            "document_id": "table-cell-hyperlink",
            "metadata": _business_report_metadata(title="表格单元格超链接"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["项目", "说明"],
                    "rows": [
                        [
                            "链接",
                            {
                                "runs": [
                                    {"text": "查看："},
                                    {"text": link_text, "url": url},
                                ]
                            },
                        ]
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    assert table.rows[1].cells[1].text == f"查看：{link_text}"

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")

    document_root = ET.fromstring(document_xml)
    rels_root = ET.fromstring(rels_xml)
    hyperlink_targets = {
        rel.attrib.get("Id"): rel.attrib.get("Target")
        for rel in rels_root
        if rel.tag.endswith("Relationship")
    }

    found = False
    for hyperlink in document_root.findall(".//w:tbl//w:hyperlink", NS):
        text = "".join(
            node.text or ""
            for node in hyperlink.findall(".//w:t", NS)
            if node.text is not None
        )
        if text != link_text:
            continue
        relation_id = hyperlink.attrib.get(f"{{{NS['r']}}}id")
        assert relation_id is not None
        assert hyperlink_targets.get(relation_id) == url
        found = True
        break

    assert found is True


def test_node_renderer_flattens_rich_table_cell_line_breaks(
    workspace_root: Path,
):
    url = "https://example.com/detail"
    expected_text = "基础盘稳定，重点增长来自华东渠道，查看明细"

    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-cell-line-break-flattening",
        {
            "document_id": "table-cell-line-break-flattening",
            "metadata": _business_report_metadata(title="表格备注换行压平"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["指标", "Q1 目标", "Q1 实际", "同比变化", "备注"],
                    "rows": [
                        [
                            "营收（万元）",
                            "50,000",
                            "59,000",
                            "+18%",
                            {
                                "runs": [
                                    {"text": "基础盘稳定，\n重点增长来自华东渠道，\n"},
                                    {"text": "查看明细", "url": url},
                                ]
                            },
                        ]
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    assert table.rows[1].cells[4].text == expected_text

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")

    document_root = ET.fromstring(document_xml)
    rels_root = ET.fromstring(rels_xml)
    hyperlink_targets = {
        rel.attrib.get("Id"): rel.attrib.get("Target")
        for rel in rels_root
        if rel.tag.endswith("Relationship")
    }

    table_node = document_root.findall(".//w:tbl", NS)[0]
    body_row = table_node.findall("./w:tr", NS)[1]
    remark_cell = body_row.findall("./w:tc", NS)[4]

    assert remark_cell.findall(".//w:br", NS) == []

    hyperlinks = remark_cell.findall(".//w:hyperlink", NS)
    assert len(hyperlinks) == 1
    relation_id = hyperlinks[0].attrib.get(f"{{{NS['r']}}}id")
    assert relation_id is not None
    assert hyperlink_targets.get(relation_id) == url


def test_node_renderer_balances_business_review_remark_column_widths(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-business-review-width-balance",
        {
            "document_id": "business-review-width-balance",
            "metadata": _business_report_metadata(title="经营复盘列宽"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["指标", "Q1 目标", "Q1 实际", "同比变化", "备注"],
                    "rows": [
                        [
                            "营收（万元）",
                            "50,000",
                            "59,000",
                            "+18%",
                            {
                                "runs": [
                                    {"text": "基础盘稳定，\n重点增长来自华东渠道，\n"},
                                    {"text": "查看明细", "url": "https://example.com/detail"},
                                ]
                            },
                        ],
                        ["毛利率", "35.0%", "37.5%", "+2.5pts", "产品结构优化"],
                    ],
                }
            ],
        },
    )

    widths = _table_grid_widths(loaded_doc.tables[0])

    assert len(widths) == 5
    assert widths[4] > widths[0]
    assert widths[4] >= 3300
    assert widths[4] < 3400
    assert widths[0] <= 1650


def test_node_renderer_balances_preset_operating_review_column_widths(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-preset-width-balance",
        {
            "document_id": "preset-width-balance",
            "metadata": _business_report_metadata(title="经营类预设列宽"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "Q1 目标", "Q1 实际", "同比变化", "备注"],
                    "rows": [
                        [
                            "华东",
                            "50,000",
                            "59,000",
                            "+18%",
                            "重点增长来自直营网点和区域渠道联动",
                        ]
                    ],
                }
            ],
        },
    )

    widths = _table_grid_widths(loaded_doc.tables[0])

    assert len(widths) == 5
    assert widths[4] > widths[0]
    assert widths[4] >= 3100
    assert widths[4] < 3400
    assert widths[4] - widths[0] < 1200


def test_node_renderer_keeps_text_only_table_compatible(workspace_root: Path):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-text-only-compatibility",
        {
            "document_id": "table-text-only-compatibility",
            "metadata": _business_report_metadata(title="纯文本表格"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["A", "B"],
                    "border_style": "minimal",
                    "rows": [
                        ["第一行", "第二列"],
                        ["第二行", "第三列"],
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]

    assert table.rows[1].cells[0].text == "第一行"
    assert table.rows[1].cells[1].text == "第二列"
    assert table.rows[2].cells[0].text == "第二行"
    assert table.rows[2].cells[1].text == "第三列"


def test_node_renderer_supports_runs_only_cells_beneath_vertical_merge(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-table-runs-only-under-rowspan",
        {
            "document_id": "table-runs-only-under-rowspan",
            "metadata": _business_report_metadata(title="行合并下的富文本单元格"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["日期", "时间", "课程"],
                    "rows": [
                        [
                            {"text": "第一天", "row_span": 2},
                            {"runs": [{"text": "09:00"}]},
                            "课程 A",
                        ],
                        [
                            {"runs": [{"text": "13:00"}]},
                            "课程 B",
                        ],
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]

    assert table.rows[1].cells[1].text == "09:00"
    assert table.rows[2].cells[1].text == "13:00"
    assert _raw_row_cell_vertical_merge(table.rows[2], 0) == "continue"


def test_node_renderer_honors_empty_table_cell_border_side_defaults(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-default-table-cell-border-side",
        {
            "document_id": "default-table-cell-border-side",
            "metadata": _business_report_metadata(title="默认单元格边框"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "说明"],
                    "rows": [
                        [
                            "华东",
                            {
                                "text": "默认单元格边框",
                                "border": {
                                    "top": {},
                                },
                            },
                        ]
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    cell = table.rows[1].cells[1]
    assert _cell_border_size(cell, "top") == "4"
    assert _cell_border_color(cell, "top") is None


def test_node_renderer_does_not_create_cell_borders_without_explicit_border(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-object-cell-without-border",
        {
            "document_id": "object-cell-without-border",
            "metadata": _business_report_metadata(title="无显式边框单元格"),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "说明"],
                    "rows": [
                        [
                            "华东",
                            {
                                "text": "只有填充",
                                "fill": "EEF4FA",
                                "align": "right",
                            },
                        ]
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    cell = table.rows[1].cells[1]
    assert cell.text == "只有填充"
    assert _cell_border_size(cell, "top") is None
    assert _cell_border_size(cell, "left") is None
    assert _cell_border_size(cell, "bottom") is None
    assert _cell_border_size(cell, "right") is None
