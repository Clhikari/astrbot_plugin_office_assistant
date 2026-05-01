import json
import shutil
import subprocess
import struct
import zlib
from pathlib import Path
from uuid import uuid4

import pytest
from astrbot_plugin_office_assistant.agent_tools import build_document_toolset
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendConfig,
)
from astrbot_plugin_office_assistant.document_core.models.blocks import GroupBlock


def _cell_fill(cell) -> str | None:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _grid_span(cell) -> int:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1
    span = tc_pr.find(qn("w:gridSpan"))
    if span is None:
        return 1
    return int(span.get(qn("w:val"), "1"))


def _run_rgb(cell) -> str | None:
    runs = cell.paragraphs[0].runs
    if not runs:
        return None
    color = runs[0].font.color.rgb
    return str(color) if color is not None else None


def _run_bold(cell) -> bool | None:
    runs = cell.paragraphs[0].runs
    if not runs:
        return None
    return runs[0].bold


def _paragraph_run_rgb(paragraph) -> str | None:
    runs = paragraph.runs
    if not runs:
        return None
    color = runs[0].font.color.rgb
    return str(color) if color is not None else None


def _paragraph_run_size(paragraph) -> float | None:
    runs = paragraph.runs
    if not runs or runs[0].font.size is None:
        return None
    return runs[0].font.size.pt


def _find_paragraph(doc, text: str):
    return next(paragraph for paragraph in doc.paragraphs if paragraph.text == text)


def _schema_contains_key(schema: object, key: str) -> bool:
    if isinstance(schema, dict):
        return key in schema or any(
            _schema_contains_key(value, key) for value in schema.values()
        )
    if isinstance(schema, list):
        return any(_schema_contains_key(value, key) for value in schema)
    return False


def _schema_type_allows(schema: dict, expected_type: str) -> bool:
    schema_type = schema.get("type")
    if isinstance(schema_type, list):
        return expected_type in schema_type
    return schema_type == expected_type


def _schema_contains_type_list(schema: object) -> bool:
    if isinstance(schema, dict):
        if isinstance(schema.get("type"), list):
            return True
        return any(_schema_contains_type_list(value) for value in schema.values())
    if isinstance(schema, list):
        return any(_schema_contains_type_list(value) for value in schema)
    return False


def _paragraph_field_codes(paragraph) -> list[str]:
    from docx.oxml.ns import qn

    codes = [
        node.text or ""
        for node in paragraph._p.iter(qn("w:instrText"))
        if node.text is not None
    ]
    codes.extend(
        node.get(qn("w:instr")) or ""
        for node in paragraph._p.iter(qn("w:fldSimple"))
        if node.get(qn("w:instr")) is not None
    )
    return codes


def _paragraph_field_nodes_use_runs(paragraph) -> bool:
    from docx.oxml.ns import qn

    field_tags = {qn("w:fldChar"), qn("w:instrText")}
    return all(
        node.getparent() is not None and node.getparent().tag == qn("w:r")
        for node in paragraph._p.iter()
        if node.tag in field_tags
    )


def _story_texts(story) -> list[str]:
    return [paragraph.text for paragraph in story.paragraphs if paragraph.text]


def _story_has_field_code(story, token: str) -> bool:
    return any(
        token in field_code
        for paragraph in story.paragraphs
        for field_code in _paragraph_field_codes(paragraph)
    )


def _document_updates_fields_on_open(doc) -> bool:
    from docx.oxml.ns import qn

    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        return False
    return update_fields.get(qn("w:val")) in {None, "1", "true", "on"}


def _document_uses_odd_even_headers(doc) -> bool:
    from docx.oxml.ns import qn

    even_headers = doc.settings.element.find(qn("w:evenAndOddHeaders"))
    if even_headers is None:
        return False
    return even_headers.get(qn("w:val")) in {None, "1", "true", "on"}


def _section_page_number_start(section) -> int | None:
    from docx.oxml.ns import qn

    page_number = section._sectPr.find(qn("w:pgNumType"))
    if page_number is None:
        return None
    start = page_number.get(qn("w:start"))
    return int(start) if start is not None else None


def _section_page_number_format(section) -> str | None:
    from docx.oxml.ns import qn

    page_number = section._sectPr.find(qn("w:pgNumType"))
    if page_number is None:
        return None
    return page_number.get(qn("w:fmt"))


def _paragraph_has_page_break(paragraph) -> bool:
    from docx.oxml.ns import qn

    return any(
        node.get(qn("w:type")) == "page" for node in paragraph._p.iter(qn("w:br"))
    )


def _paragraph_has_keep_next(paragraph) -> bool:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        return False
    value = keep_next.get(qn("w:val"))
    return value in {None, "1", "true", "on"}


def _paragraph_after(doc, paragraph, offset: int = 1):
    for index, candidate in enumerate(doc.paragraphs):
        if candidate._p is paragraph._p:
            return doc.paragraphs[index + offset]
    raise ValueError("paragraph is not in document")


def _table_border_size(table, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None
    edge = tbl_borders.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:sz"))


def _table_border_color(table, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None
    edge = tbl_borders.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:color"))


def _table_cell_margin(table, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None
    tbl_cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_margins is None:
        return None
    edge = tbl_cell_margins.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:w"))


def _paragraph_bottom_border_color(paragraph) -> str | None:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    paragraph_border = p_pr.find(qn("w:pBdr"))
    if paragraph_border is None:
        return None
    bottom = paragraph_border.find(qn("w:bottom"))
    if bottom is None:
        return None
    return bottom.get(qn("w:color"))


def _paragraph_bottom_border_size(paragraph) -> str | None:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    paragraph_border = p_pr.find(qn("w:pBdr"))
    if paragraph_border is None:
        return None
    bottom = paragraph_border.find(qn("w:bottom"))
    if bottom is None:
        return None
    return bottom.get(qn("w:sz"))


def _paragraph_top_border_color(paragraph) -> str | None:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    paragraph_border = p_pr.find(qn("w:pBdr"))
    if paragraph_border is None:
        return None
    top = paragraph_border.find(qn("w:top"))
    if top is None:
        return None
    return top.get(qn("w:color"))


def _paragraph_tab_positions(paragraph) -> list[str]:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return []
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return []
    return [
        tab.get(qn("w:pos"), "")
        for tab in tabs.findall(qn("w:tab"))
        if tab.get(qn("w:pos")) is not None
    ]


def _cell_vertical_merge(cell) -> str | None:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    vertical_merge = tc_pr.find(qn("w:vMerge"))
    if vertical_merge is None:
        return None
    return vertical_merge.get(qn("w:val"))


def _cell_border_color(cell, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return None
    edge = tc_borders.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:color"))


def _cell_border_size(cell, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return None
    edge = tc_borders.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:sz"))


def _cell_margin(cell, edge_name: str) -> str | None:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    tc_margins = tc_pr.find(qn("w:tcMar"))
    if tc_margins is None:
        return None
    edge = tc_margins.find(qn(f"w:{edge_name}"))
    if edge is None:
        return None
    return edge.get(qn("w:w"))


def _cell_width(cell) -> tuple[str | None, str | None]:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None, None
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        return None, None
    return tc_width.get(qn("w:w")), tc_width.get(qn("w:type"))


def _table_width(table) -> tuple[str | None, str | None]:
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None, None
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        return None, None
    return tbl_width.get(qn("w:w")), tbl_width.get(qn("w:type"))


def _table_row_height(row) -> tuple[str | None, str | None]:
    from docx.oxml.ns import qn

    tr_pr = row._tr.trPr
    if tr_pr is None:
        return None, None
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        return None, None
    return tr_height.get(qn("w:val")), tr_height.get(qn("w:hRule"))


def _table_grid_widths(table) -> list[int]:
    from docx.oxml.ns import qn

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        return []
    return [
        int(grid_col.get(qn("w:w"), "0"))
        for grid_col in tbl_grid.iter(qn("w:gridCol"))
    ]


def _run_font_attr(run, attr_name: str) -> str | None:
    from docx.oxml.ns import qn

    r_pr = run._r.rPr
    if r_pr is None:
        return None
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        return None
    return r_fonts.get(qn(f"w:{attr_name}"))


def _raw_row_cell_vertical_merge(row, cell_index: int) -> str | None:
    from docx.oxml.ns import qn

    tc = row._tr.tc_lst[cell_index]
    tc_pr = tc.tcPr
    if tc_pr is None:
        return None
    vertical_merge = tc_pr.find(qn("w:vMerge"))
    if vertical_merge is None:
        return None
    return vertical_merge.get(qn("w:val"))


def _row_has_cant_split(row) -> bool:
    from docx.oxml.ns import qn

    tr_pr = row._tr.trPr
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:cantSplit")) is not None


def _row_cant_split_value(row) -> str | None:
    from docx.oxml.ns import qn

    tr_pr = row._tr.trPr
    if tr_pr is None:
        return None
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        return None
    return cant_split.get(qn("w:val"))


def _row_is_repeated_header(row) -> bool:
    from docx.oxml.ns import qn

    tr_pr = row._tr.trPr
    if tr_pr is None:
        return False
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        return False
    value = tbl_header.get(qn("w:val"))
    return value in {None, "1", "true", "on"}


def _make_workspace(workspace_root: Path, name: str) -> Path:
    workspace_dir = workspace_root / f"{name}-{uuid4().hex}"
    workspace_dir.mkdir(parents=True, exist_ok=True)
    return workspace_dir


def _write_png(path: Path, *, width: int, height: int) -> None:
    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + chunk_type
            + data
            + struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF)
        )

    scanline = b"\x00" + (b"\xff\x66\x33" * width)
    raw = scanline * height
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    path.write_bytes(png)


def _node_renderer_entry() -> Path:
    renderer_entry = (
        Path(__file__).resolve().parents[1] / "word_renderer_js" / "dist" / "cli.js"
    )
    if shutil.which("node") is None or not renderer_entry.exists():
        pytest.skip("node renderer build is not available")
    return renderer_entry


def _business_report_metadata(
    *,
    title: str = "Q3 经营复盘报告",
    theme_name: str = "business_report",
    table_template: str = "report_grid",
    density: str = "comfortable",
    accent_color: str | None = None,
    document_style: dict[str, object] | None = None,
    header_footer: dict[str, object] | None = None,
) -> dict[str, object]:
    metadata: dict[str, object] = {
        "title": title,
        "theme_name": theme_name,
        "table_template": table_template,
        "density": density,
        "document_style": document_style or {},
        "header_footer": header_footer or {},
    }
    if accent_color:
        metadata["accent_color"] = accent_color
    return metadata


def _business_review_cover_block(
    *,
    summary_text: str,
    metrics: list[dict[str, object]],
    footer_note: str = "",
    title: str = "Q3 经营复盘报告",
    subtitle: str = "战略与增长委员会 · 2024 年 10 月",
    summary_title: str = "核心摘要",
    auto_page_break: bool | None = None,
) -> dict[str, object]:
    data: dict[str, object] = {
        "title": title,
        "subtitle": subtitle,
        "summary_title": summary_title,
        "summary_text": summary_text,
        "metrics": metrics,
    }
    if footer_note:
        data["footer_note"] = footer_note
    if auto_page_break is not None:
        data["auto_page_break"] = auto_page_break
    return {
        "type": "page_template",
        "template": "business_review_cover",
        "data": data,
    }


def _technical_resume_block(
    *,
    name: str = "张明远",
    headline: str = "后端开发工程师  ·  分布式系统 / 高并发架构",
    contact_line: str = "zhangmingyuan@email.com  ·  138-0000-0000  ·  北京 | 可远程  ·  github.com/zhangmy",
    sections: list[dict[str, object]] | None = None,
) -> dict[str, object]:
    return {
        "type": "page_template",
        "template": "technical_resume",
        "data": {
            "name": name,
            "headline": headline,
            "contact_line": contact_line,
            "sections": sections
            or [
                {
                    "title": "教育背景",
                    "entries": [
                        {
                            "heading": "北京大学",
                            "date": "2019.09 – 2023.06",
                            "subtitle": "计算机科学与技术  |  工学学士",
                            "details": [
                                "GPA 3.86/4.0，连续三年一等奖学金，排名前 5%",
                                "荣誉：ACM-ICPC 亚洲区决赛银奖（2021）、校优秀毕业论文",
                            ],
                        }
                    ],
                },
                {
                    "title": "实习经历",
                    "entries": [
                        {
                            "heading": "字节跳动 · 基础架构部",
                            "date": "2022.07 – 2022.12",
                            "subtitle": "后端开发实习生 · 推荐系统组",
                            "details": [
                                {
                                    "runs": [
                                        {"text": "主导优化推荐引擎召回模块，", "bold": True},
                                        {
                                            "text": "将离线 Embedding 索引构建耗时从 4.2h 降至 1.1h，上线后 CTR 提升 3.2%。"
                                        },
                                    ]
                                },
                                "设计并落地 A/B 实验流量分桶系统，支持动态扩容，日均承载 12 亿次请求，P99 < 8ms。",
                            ],
                        }
                    ],
                },
                {
                    "title": "技术栈",
                    "lines": [
                        "语言：Go（熟练）、Java（熟练）、Python、SQL",
                        "框架/中间件：Spring Boot、Gin、gRPC、Kafka、Redis、MySQL、Elasticsearch",
                    ],
                },
            ],
        },
    }




def _summary_card_block(
    *,
    title: str = "Conclusion",
    items: list[str] | None = None,
    variant: str = "conclusion",
) -> dict[str, object]:
    return {
        "type": "summary_card",
        "title": title,
        "items": items or ["First takeaway"],
        "variant": variant,
    }


def _summary_card_defaults(
    *,
    title_align: str = "center",
    title_emphasis: str = "strong",
    title_font_scale: float = 1.2,
    title_space_before: float = 12,
    title_space_after: float = 4,
    list_space_after: float = 8,
) -> dict[str, object]:
    return {
        "title_align": title_align,
        "title_emphasis": title_emphasis,
        "title_font_scale": title_font_scale,
        "title_space_before": title_space_before,
        "title_space_after": title_space_after,
        "list_space_after": list_space_after,
    }


def _assert_summary_card_group(
    group: GroupBlock,
    *,
    title: str = "Conclusion",
    items: list[str] | None = None,
) -> None:
    assert isinstance(group, GroupBlock)
    assert group.blocks[0].text == title
    assert group.blocks[1].items == (items or ["First takeaway"])


def _render_structured_payload_with_node(
    workspace_root: Path,
    workspace_name: str,
    payload: dict[str, object],
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, workspace_name)
    renderer_entry = _node_renderer_entry()

    output_path = workspace_dir / "node-output.docx"
    payload_path = workspace_dir / "node-payload.json"
    normalized_payload = {
        "version": "v1",
        "format": "word",
        "render_mode": "structured",
        **payload,
    }
    payload_path.write_text(
        json.dumps(normalized_payload, ensure_ascii=False),
        encoding="utf-8",
    )
    subprocess.run(
        ["node", str(renderer_entry), str(payload_path), str(output_path)],
        cwd=str(renderer_entry.parents[1]),
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return docx.Document(output_path), output_path


def _node_render_backend_config_for_tests() -> DocumentRenderBackendConfig:
    renderer_entry = _node_renderer_entry()
    return DocumentRenderBackendConfig(
        preferred_backend="node",
        fallback_enabled=False,
        node_renderer_entry=str(renderer_entry),
    )


async def _export_docx_via_node_toolset(
    workspace_root: Path,
    workspace_name: str,
    *,
    create_kwargs: dict[str, object],
    blocks: list[dict[str, object]],
):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(workspace_root, workspace_name)
    toolset = build_document_toolset(
        workspace_dir=workspace_dir,
        render_backend_config=_node_render_backend_config_for_tests(),
    )
    tool_by_name = {tool.name: tool for tool in toolset.tools}

    created = json.loads(await tool_by_name["create_document"].call(None, **create_kwargs))
    document_id = created["document"]["document_id"]

    added = json.loads(
        await tool_by_name["add_blocks"].call(
            None,
            document_id=document_id,
            blocks=blocks,
        )
    )
    assert added["success"] is True, added["message"]
    finalized = json.loads(
        await tool_by_name["finalize_document"].call(None, document_id=document_id)
    )
    assert finalized["success"] is True, finalized["message"]
    exported = json.loads(
        await tool_by_name["export_document"].call(None, document_id=document_id)
    )
    assert exported["success"] is True, exported["message"]

    return docx.Document(exported["file_path"]), Path(exported["file_path"])


def _section_margin_twips(section, edge_name: str) -> int | None:
    from docx.oxml.ns import qn

    page_margin = section._sectPr.find(qn("w:pgMar"))
    if page_margin is None:
        return None
    value = page_margin.get(qn(f"w:{edge_name}"))
    return int(value) if value is not None else None


__all__ = [
    name
    for name in globals()
    if name.startswith("_") and not name.startswith("__")
]
