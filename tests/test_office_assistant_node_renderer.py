import json
import subprocess
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from astrbot_plugin_office_assistant.document_core.models.blocks import (
    HeaderFooterConfig,
    ParagraphBlock,
    SectionBreakBlock,
    SummaryCardBlock,
)
from astrbot_plugin_office_assistant.document_core.models.document import (
    DocumentMetadata,
    DocumentModel,
    DocumentStyleConfig,
    DocumentSummaryCardDefaults,
)
from astrbot_plugin_office_assistant.domain.document.contracts import (
    CreateDocumentRequest,
)
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendError,
    NodeDocumentRenderBackend,
    build_document_render_backends,
    build_document_render_payload,
)
from astrbot_plugin_office_assistant.domain.document.session_store import (
    DocumentSessionStore,
)
from tests._docx_test_helpers import *  # noqa: F401,F403
from tests._docx_test_helpers import (
    _export_docx_via_node_toolset,
    _render_structured_payload_with_node,
    _technical_resume_block,
)


def test_node_render_backend_serializes_payload_and_invokes_cli(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-node-render-backend")
    entry_path = workspace_dir / "dist" / "cli.js"
    entry_path.parent.mkdir(parents=True, exist_ok=True)
    entry_path.write_text("// fake cli", encoding="utf-8")

    document = DocumentSessionStore(workspace_dir=workspace_dir).create_document(
        CreateDocumentRequest(
            title="Node Backend",
            output_name="node-backend.docx",
        )
    )
    document.add_block(ParagraphBlock(text="Node payload paragraph"))
    output_path = workspace_dir / "node-output.docx"
    payloads: list[dict[str, object]] = []

    def _fake_run(command, cwd, check, capture_output, text, encoding):
        payload_path = Path(command[2])
        moved_payload_path = payload_path.with_suffix(".moved.json")
        payload_path.rename(moved_payload_path)
        payloads.append(json.loads(moved_payload_path.read_text(encoding="utf-8")))
        assert cwd == str(entry_path.parent)
        output_path.write_bytes(b"node-docx")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    with patch(
        "astrbot_plugin_office_assistant.domain.document.render_backends.subprocess.run",
        side_effect=_fake_run,
    ) as mocked_run:
        backend = NodeDocumentRenderBackend(entry_path=entry_path)
        result = backend.render(document, output_path)

    assert mocked_run.call_count == 1
    assert result.backend_name == "node"
    assert result.output_path == output_path
    assert payloads[0]["version"] == "v1"
    assert payloads[0]["render_mode"] == "structured"
    assert payloads[0]["document_id"] == document.document_id
    assert payloads[0]["blocks"][0]["type"] == "paragraph"


def test_node_render_backend_extracts_message_from_json_error(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-node-render-backend-error")
    entry_path = workspace_dir / "dist" / "cli.js"
    entry_path.parent.mkdir(parents=True, exist_ok=True)
    entry_path.write_text("// fake cli", encoding="utf-8")

    document = DocumentSessionStore(workspace_dir=workspace_dir).create_document(
        CreateDocumentRequest(
            title="Node Backend Error",
            output_name="node-backend-error.docx",
        )
    )
    document.add_block(ParagraphBlock(text="Node payload paragraph"))
    output_path = workspace_dir / "node-output.docx"

    with patch(
        "astrbot_plugin_office_assistant.domain.document.render_backends.subprocess.run",
        return_value=SimpleNamespace(
            returncode=1,
            stdout='{"success":false,"message":"stdout should be ignored"}',
            stderr='{"success":false,"code":"SCHEMA_VALIDATION_FAILED","message":"headline is required"}',
        ),
    ):
        backend = NodeDocumentRenderBackend(entry_path=entry_path)
        with pytest.raises(
            DocumentRenderBackendError,
            match=r"JS renderer failed: headline is required",
        ):
            backend.render(document, output_path)


def test_build_document_render_payload_keeps_only_explicit_header_footer_overrides():
    document = DocumentSessionStore().create_document(
        CreateDocumentRequest(
            title="Payload Header Footer",
            header_footer={"show_page_number": False},
        )
    )
    document.add_block(
        SectionBreakBlock(
            header_footer=HeaderFooterConfig(show_page_number=False),
        )
    )

    payload = build_document_render_payload(document)

    assert payload["metadata"]["header_footer"] == {"show_page_number": False}
    assert payload["blocks"][0]["header_footer"] == {"show_page_number": False}


def test_build_document_render_backends_reserves_excel_for_python():
    backends = build_document_render_backends("excel")

    assert [backend.name for backend in backends] == ["python-excel"]
    with pytest.raises(
        DocumentRenderBackendError,
        match="Excel render backend is reserved for Python implementation",
    ):
        backends[0].render(MagicMock(), Path("out.xlsx"))


def test_build_document_render_backends_defaults_to_node_for_word():
    backends = build_document_render_backends("word")

    assert [backend.name for backend in backends] == ["node"]


def test_node_render_backend_renders_sections_and_table_styles(workspace_root: Path):
    docx = pytest.importorskip("docx")

    workspace_dir = _make_workspace(
        workspace_root,
        "pytest-node-render-backend-structured-docx",
    )
    renderer_entry = _node_renderer_entry()

    output_path = workspace_dir / "node-structured.docx"
    payload_path = workspace_dir / "node-structured.json"
    payload_path.write_text(
        json.dumps(
            {
                "version": "v1",
                "render_mode": "structured",
                "document_id": "node-structured-doc",
                "metadata": _business_report_metadata(
                    document_style={
                        "heading_color": "000000",
                        "heading_level_1_color": "0F4C81",
                        "heading_bottom_border_color": "CBD5E1",
                        "heading_bottom_border_size_pt": 1.5,
                        "table_defaults": {
                            "body_fill": "F8FAFC",
                            "border_style": "minimal",
                        },
                    },
                    header_footer={
                        "header_left": "Q3 经营复盘报告 | 战略与增长委员会",
                        "header_right": "机密 · 2024 年 10 月",
                        "header_border_bottom": True,
                        "footer_left": "集团战略部 · 内部机密文件",
                        "footer_right": "第 {PAGE} 页",
                        "footer_border_top": True,
                    },
                ),
                "blocks": [
                    {
                        "type": "heading",
                        "text": "一、经营总览",
                        "level": 1,
                        "bottom_border": True,
                    },
                    {
                        "type": "paragraph",
                        "runs": [
                            {"text": "核心结论：", "bold": True},
                            {"text": "Q3 整体经营保持增长。"},
                        ],
                    },
                    {
                        "type": "table",
                        "headers": ["日期", "时间", "课程"],
                        "header_fill_enabled": False,
                        "header_text_color": "808080",
                        "header_bold": False,
                        "rows": [
                            [{"text": "第一天", "row_span": 2}, "09:00", "课程 A"],
                            ["13:00", "课程 B"],
                        ],
                        "border_style": "minimal",
                    },
                    {
                        "type": "section_break",
                        "start_type": "new_page",
                        "page_orientation": "landscape",
                        "restart_page_numbering": True,
                        "header_footer": {
                            "show_page_number": False,
                            "header_left": "第二节页眉",
                            "footer_left": "第二节页脚",
                        },
                    },
                    {
                        "type": "paragraph",
                        "text": "第二节内容",
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    subprocess.run(
        ["node", str(renderer_entry), str(payload_path), str(output_path)],
        cwd=str(renderer_entry.parents[1]),
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    loaded_doc = docx.Document(output_path)
    title_paragraph = _find_paragraph(loaded_doc, "Q3 经营复盘报告")
    header_paragraph = loaded_doc.sections[0].header.paragraphs[0]
    footer_texts = _story_texts(loaded_doc.sections[0].footer)
    table = loaded_doc.tables[0]

    assert _paragraph_run_rgb(title_paragraph) == "000000"
    overview_heading = _find_paragraph(loaded_doc, "一、经营总览")
    overview_divider = _paragraph_after(loaded_doc, overview_heading)
    assert _paragraph_run_rgb(overview_heading) == "0F4C81"
    assert _paragraph_bottom_border_color(overview_divider) == "CBD5E1"
    assert header_paragraph.text == "Q3 经营复盘报告 | 战略与增长委员会\t机密 · 2024 年 10 月"
    assert any(
        text.startswith("集团战略部 · 内部机密文件") for text in footer_texts
    )
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml
    assert _cell_vertical_merge(table.rows[1].cells[0]) == "restart"
    assert _raw_row_cell_vertical_merge(table.rows[2], 0) == "continue"
    assert _cell_fill(table.rows[1].cells[1]) == "F7FBFF"
    assert _section_page_number_start(loaded_doc.sections[1]) == 1
    assert (
        loaded_doc.sections[1].header.paragraphs[0].text
        == "第二节页眉\t机密 · 2024 年 10 月"
    )
    assert loaded_doc.sections[1].footer.paragraphs[0].text == "第二节页脚"
    assert "PAGE" not in loaded_doc.sections[1].footer._element.xml


def test_node_renderer_reports_ppt_placeholder_not_implemented(workspace_root: Path):
    workspace_dir = _make_workspace(
        workspace_root,
        "pytest-node-render-backend-ppt-placeholder",
    )
    renderer_entry = _node_renderer_entry()

    output_path = workspace_dir / "node-structured.pptx"
    payload_path = workspace_dir / "node-structured-ppt.json"
    payload_path.write_text(
        json.dumps(
            {
                "version": "v1",
                "render_mode": "structured",
                "document_id": "node-structured-ppt",
                "format": "ppt",
                "metadata": _business_report_metadata(title="PPT 占位"),
                "blocks": [],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        ["node", str(renderer_entry), str(payload_path), str(output_path)],
        cwd=str(renderer_entry.parents[1]),
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert completed.returncode != 0
    error_payload = json.loads(completed.stderr)
    assert error_payload["code"] == "FORMAT_NOT_IMPLEMENTED"
    assert "PPT structured renderer" in error_payload["message"]


def test_node_renderer_supports_toc_and_header_footer_variants(workspace_root: Path):
    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-toc-header-footer",
        {
            "document_id": "node-toc-header-footer",
            "metadata": _business_report_metadata(
                title="目录测试",
                header_footer={
                    "header_text": "季度经营复盘",
                    "footer_text": "内部使用",
                    "different_first_page": True,
                    "first_page_header_text": "封面页眉",
                    "first_page_footer_text": "封面页脚",
                    "first_page_show_page_number": True,
                    "different_odd_even": True,
                    "even_page_header_text": "偶数页页眉",
                    "even_page_footer_text": "偶数页页脚",
                    "even_page_show_page_number": False,
                    "show_page_number": True,
                    "page_number_align": "center",
                },
            ),
            "blocks": [
                {"type": "toc", "title": "目录", "levels": 2, "start_on_new_page": True},
                {"type": "heading", "text": "经营总览", "level": 2},
                {"type": "paragraph", "text": "正文"},
            ],
        },
    )

    assert _document_updates_fields_on_open(loaded_doc) is True
    assert _document_uses_odd_even_headers(loaded_doc) is True
    assert loaded_doc.sections[0].different_first_page_header_footer is True
    assert "季度经营复盘" in _story_texts(loaded_doc.sections[0].header)
    assert any(
        text.startswith("内部使用") for text in _story_texts(loaded_doc.sections[0].footer)
    )
    assert "封面页眉" in _story_texts(loaded_doc.sections[0].first_page_header)
    assert any(
        text.startswith("封面页脚")
        for text in _story_texts(loaded_doc.sections[0].first_page_footer)
    )
    assert "PAGE" in loaded_doc.sections[0].first_page_footer._element.xml
    assert "偶数页页眉" in _story_texts(loaded_doc.sections[0].even_page_header)
    assert "偶数页页脚" in _story_texts(loaded_doc.sections[0].even_page_footer)
    assert "PAGE" not in loaded_doc.sections[0].even_page_footer._element.xml
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml
    assert all(
        _paragraph_field_nodes_use_runs(paragraph)
        for paragraph in loaded_doc.sections[0].footer.paragraphs
    )
    assert any(
        _paragraph_has_page_break(paragraph)
        or paragraph.paragraph_format.page_break_before is True
        for paragraph in loaded_doc.paragraphs
    )
    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "w:fldSimple" in document_xml
    assert "TOC" in document_xml
    assert '\\o &quot;1-2&quot;' in document_xml


def test_node_renderer_schema_rejects_invalid_hero_banner_colors(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-node-schema-hero-banner")
    renderer_entry = _node_renderer_entry()

    output_path = workspace_dir / "invalid-hero-banner.docx"
    payload_path = workspace_dir / "invalid-hero-banner.json"
    payload_path.write_text(
        json.dumps(
            {
                "version": "v1",
                "render_mode": "structured",
                "document_id": "invalid-hero-banner",
                "metadata": _business_report_metadata(title=""),
                "blocks": [
                    {
                        "type": "hero_banner",
                        "title": "Q3 经营复盘报告",
                        "theme_color": "blue",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        ["node", str(renderer_entry), str(payload_path), str(output_path)],
        cwd=str(renderer_entry.parents[1]),
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert completed.returncode != 0
    error_payload = json.loads(completed.stderr)
    assert error_payload["code"] == "SCHEMA_VALIDATION_FAILED"
    assert "theme_color" in error_payload["message"]


def test_node_renderer_supports_business_review_cover_page_template(
    workspace_root: Path,
):
    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-page-template-business-review-cover",
        {
            "document_id": "page-template-business-review-cover",
            "metadata": _business_report_metadata(
                document_style={
                    "font_name": "Microsoft YaHei",
                    "heading_font_name": "Microsoft YaHei",
                }
            ),
            "blocks": [
                _business_review_cover_block(
                    summary_text="经营质量继续改善，现金流保持稳定，重点区域保持增长。",
                    metrics=[
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "↑ 18.4% YoY",
                            "delta_color": "15803D",
                        },
                        {
                            "label": "毛利率",
                            "value": "42.1%",
                            "delta": "↓ 0.3pp vs Q2",
                            "delta_color": "DC2626",
                        },
                    ],
                    footer_note="编制：战略发展部 · 审核：CFO 办公室",
                ),
                {
                    "type": "heading",
                    "text": "一、经营总览",
                    "level": 1,
                    "bottom_border": True,
                },
                {
                    "type": "paragraph",
                    "text": "第二页正文应该紧接在封面页之后。",
                },
            ],
        },
    )

    banner_table = loaded_doc.tables[0]
    summary_table = loaded_doc.tables[1]
    metrics_table = loaded_doc.tables[2]
    footer_table = loaded_doc.tables[3]
    heading = _find_paragraph(loaded_doc, "一、经营总览")

    assert len(loaded_doc.tables) == 4
    assert all(
        paragraph.text != "Q3 经营复盘报告" for paragraph in loaded_doc.paragraphs
    )
    assert banner_table.rows[0].cells[0].paragraphs[0].text == "Q3 经营复盘报告"
    assert banner_table.rows[0].cells[0].paragraphs[1].text == "战略与增长委员会 · 2024 年 10 月"
    assert _cell_fill(banner_table.rows[0].cells[0]) == "1F4E79"
    assert _table_width(banner_table) == ("9360", "dxa")
    assert _table_row_height(banner_table.rows[0]) == ("1600", "exact")
    assert _cell_margin(banner_table.rows[0].cells[0], "left") == "400"
    assert _cell_margin(banner_table.rows[0].cells[0], "right") == "400"
    assert _cell_margin(banner_table.rows[0].cells[0], "bottom") == "0"
    assert banner_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after.pt == 0
    assert _paragraph_run_size(banner_table.rows[0].cells[0].paragraphs[0]) == pytest.approx(
        26.0,
        abs=0.5,
    )
    assert summary_table.rows[0].cells[0].paragraphs[0].text == "核心摘要"
    assert "经营质量继续改善" in summary_table.rows[0].cells[0].text
    assert _table_width(summary_table) == ("9360", "dxa")
    assert _cell_width(summary_table.rows[0].cells[0]) == ("9360", "dxa")
    assert _cell_fill(summary_table.rows[0].cells[0]) == "EEF9F5"
    assert _cell_border_color(summary_table.rows[0].cells[0], "left") == "0F6E56"
    assert _cell_border_size(summary_table.rows[0].cells[0], "left") == "64"
    assert metrics_table.rows[0].cells[0].paragraphs[0].text == "营业收入"
    assert metrics_table.rows[0].cells[1].paragraphs[0].text == "毛利率"
    assert metrics_table.rows[0].cells[2].paragraphs[0].text == ""
    assert _table_width(metrics_table) == ("9360", "dxa")
    assert _cell_width(metrics_table.rows[0].cells[0]) == ("3120", "dxa")
    assert _cell_width(metrics_table.rows[0].cells[1]) == ("3120", "dxa")
    assert _cell_width(metrics_table.rows[0].cells[2]) == ("3120", "dxa")
    assert _cell_fill(metrics_table.rows[0].cells[0]) == "F2F7FC"
    assert _paragraph_run_size(metrics_table.rows[0].cells[0].paragraphs[1]) == pytest.approx(
        17.0,
        abs=0.2,
    )
    assert _paragraph_run_rgb(metrics_table.rows[0].cells[0].paragraphs[2]) == "15803D"
    assert _paragraph_run_rgb(metrics_table.rows[0].cells[1].paragraphs[2]) == "DC2626"
    assert _table_width(footer_table) == ("9360", "dxa")
    assert _cell_width(footer_table.rows[0].cells[0]) == ("160", "dxa")
    assert _cell_width(footer_table.rows[0].cells[1]) == ("9200", "dxa")
    assert _cell_fill(footer_table.rows[0].cells[0]) == "1F4E79"
    assert _cell_fill(footer_table.rows[0].cells[1]) == "EEF3FB"
    assert footer_table.rows[0].cells[1].paragraphs[0].text == "编制：战略发展部 · 审核：CFO 办公室"
    assert not any(
        _paragraph_has_page_break(paragraph) for paragraph in loaded_doc.paragraphs
    )
    assert heading.text == "一、经营总览"
    assert _find_paragraph(loaded_doc, "第二页正文应该紧接在封面页之后。").text == (
        "第二页正文应该紧接在封面页之后。"
    )

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert document_xml.count('w:pStyle w:val="Heading1"') == 1


def test_node_renderer_suppresses_page_template_auto_page_break_when_body_follows(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-page-template-auto-break-suppressed",
        {
            "document_id": "page-template-auto-break-suppressed",
            "metadata": _business_report_metadata(),
            "blocks": [
                _business_review_cover_block(
                    summary_text="经营质量继续改善，现金流保持稳定。",
                    metrics=[
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "↑ 18.4% YoY",
                        }
                    ],
                    auto_page_break=True,
                ),
                {"type": "heading", "text": "一、经营总览", "level": 1},
            ],
        },
    )

    assert not any(
        _paragraph_has_page_break(paragraph) for paragraph in loaded_doc.paragraphs
    )
    assert _find_paragraph(loaded_doc, "一、经营总览").text == "一、经营总览"


def test_node_renderer_defers_nested_page_template_footer_note_until_document_end(
    workspace_root: Path,
):
    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-nested-page-template-footer-note",
        {
            "document_id": "nested-page-template-footer-note",
            "metadata": _business_report_metadata(),
            "blocks": [
                {
                    "type": "group",
                    "blocks": [
                        _business_review_cover_block(
                            summary_text="封面页脚备注应该被放到全文最后。",
                            metrics=[
                                {
                                    "label": "营业收入",
                                    "value": "¥4.82 亿",
                                    "delta": "↑ 18.4% YoY",
                                }
                            ],
                            footer_note="编制：战略发展部 · 审核：CFO 办公室",
                        )
                    ],
                },
                {"type": "heading", "text": "一、经营总览", "level": 1},
                {"type": "paragraph", "text": "正文应该先于页脚备注出现。"},
            ],
        },
    )

    assert _find_paragraph(loaded_doc, "一、经营总览").text == "一、经营总览"

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert document_xml.index("一、经营总览") < document_xml.index(
        "编制：战略发展部 · 审核：CFO 办公室"
    )
    assert document_xml.count("编制：战略发展部 · 审核：CFO 办公室") == 1


def test_node_renderer_infers_compact_table_widths_for_short_business_tables(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-compact-table-widths",
        {
            "document_id": "compact-table-widths",
            "metadata": _business_report_metadata(
                title="经营复盘",
                document_style={
                    "font_name": "Microsoft YaHei",
                    "table_font_name": "Microsoft YaHei",
                },
            ),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "收入", "完成率", "备注"],
                    "rows": [
                        ["华东", "¥4.82 亿", "112%", "重点项目提前交付"],
                        ["华北", "¥3.14 亿", "95%", "渠道结构优化中"],
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]
    grid_widths = _table_grid_widths(table)

    assert _table_width(table)[1] == "dxa"
    assert len(grid_widths) == 4
    assert grid_widths[3] > grid_widths[0]
    assert grid_widths[1] < grid_widths[3]
    assert sum(grid_widths) < 9000


def test_node_renderer_uses_relaxed_report_grid_defaults_for_business_tables(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-relaxed-report-grid-defaults",
        {
            "document_id": "relaxed-report-grid-defaults",
            "metadata": _business_report_metadata(
                title="",
                document_style={
                    "table_font_name": "Microsoft YaHei",
                },
            ),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "营收 (万元)", "同比 (YoY)", "预算完成率", "备注"],
                    "rows": [
                        ["华东大区", "18,450", "+15.2%", "108.5%", "核心项目提前交付"],
                        ["华南大区", "14,200", "+10.8%", "105.2%", "新签客户放量明显"],
                    ],
                }
            ],
        },
    )

    table = loaded_doc.tables[0]

    assert _table_width(table) == ("9360", "dxa")
    assert _table_grid_widths(table) == [1720, 1280, 1280, 1280, 3800]
    assert _table_cell_margin(table, "left") == "108"
    assert _table_cell_margin(table, "top") == "136"
    assert _table_row_height(table.rows[0]) == ("520", "atLeast")
    assert _table_row_height(table.rows[1]) == ("480", "atLeast")


def test_node_renderer_uses_business_report_heading_dividers_when_theme_name_is_blank(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-blank-theme-name-heading-divider",
        {
            "document_id": "blank-theme-name-heading-divider",
            "metadata": _business_report_metadata(
                title="",
                theme_name="",
                document_style={
                    "heading_color": "1F4E79",
                    "table_font_name": "Microsoft YaHei",
                },
            ),
            "blocks": [
                {
                    "type": "heading",
                    "text": "一、分区业绩",
                    "level": 1,
                },
                {
                    "type": "heading",
                    "text": "1.1 各区营收完成情况",
                    "level": 2,
                },
            ],
        },
    )

    level_1_heading = _find_paragraph(loaded_doc, "一、分区业绩")
    level_1_divider = _paragraph_after(loaded_doc, level_1_heading)
    level_2_heading = _find_paragraph(loaded_doc, "1.1 各区营收完成情况")
    level_2_divider = _paragraph_after(loaded_doc, level_2_heading)

    assert _paragraph_bottom_border_color(level_1_divider) == "1F4E79"
    assert _paragraph_bottom_border_color(level_2_divider) == "1F4E79"


def test_node_renderer_business_report_metric_cards_use_fixed_cover_widths(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-business-report-metric-widths",
        {
            "document_id": "business-report-metric-widths",
            "metadata": _business_report_metadata(
                title="",
                density="compact",
                document_style={
                    "font_name": "Source Han Sans SC",
                    "heading_font_name": "Source Han Sans SC",
                    "table_font_name": "Arial",
                },
            ),
            "blocks": [
                {
                    "type": "metric_cards",
                    "metrics": [
                        {
                            "label": "营业收入",
                            "value": "4.28 亿",
                            "delta": "+12.4% YoY",
                        },
                        {
                            "label": "毛利率",
                            "value": "38.5%",
                            "delta": "+2.1 pct YoY",
                        },
                        {
                            "label": "净新增客户",
                            "value": "156",
                            "delta": "+12% vs Q2",
                        },
                    ],
                }
            ],
        },
    )

    metric_table = loaded_doc.tables[0]
    grid_widths = _table_grid_widths(metric_table)

    assert _table_width(metric_table) == ("9360", "dxa")
    assert grid_widths == [3120, 3120, 3120]
    assert _cell_width(metric_table.rows[0].cells[0]) == ("3120", "dxa")
    assert _cell_width(metric_table.rows[0].cells[1]) == ("3120", "dxa")
    assert _cell_width(metric_table.rows[0].cells[2]) == ("3120", "dxa")


@pytest.mark.asyncio
async def test_node_document_toolset_exports_business_review_cover_page_template(
    workspace_root: Path,
):
    loaded_doc, _ = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-page-template-business-review-cover",
        create_kwargs={
            "title": "Q3 经营复盘报告",
            "output_name": "business-review-cover-template.docx",
            "theme_name": "business_report",
            "header_footer": {
                "header_left": "Q3 经营复盘报告 | 战略与增长委员会",
                "header_right": "机密 · 2024 年 10 月",
                "footer_left": "集团战略部 · 内部机密文件",
                "footer_right": "第 {PAGE} 页",
            },
        },
        blocks=[
            _business_review_cover_block(
                summary_text="Q3 整体经营表现稳健，增长质量继续改善。",
                metrics=[
                    {
                        "label": "营业收入",
                        "value": "¥4.82 亿",
                        "delta": "↑ 18.4% YoY",
                        "delta_color": "15803D",
                    },
                    {
                        "label": "净新增客户",
                        "value": "184",
                        "delta": "↑ 12.0% QoQ",
                        "delta_color": "15803D",
                    },
                ],
                footer_note="编制：战略发展部 · 审核：CFO 办公室",
            ),
            {
                "type": "heading",
                "text": "一、分区业绩",
                "level": 1,
                "bottom_border": True,
            },
            {
                "type": "paragraph",
                "text": "华东和华南区域继续承担主要增长贡献。",
            },
        ],
    )

    assert len(loaded_doc.tables) == 4
    assert loaded_doc.tables[0].rows[0].cells[0].paragraphs[0].text == "Q3 经营复盘报告"
    assert _table_width(loaded_doc.tables[0]) == ("9360", "dxa")
    assert _table_row_height(loaded_doc.tables[0].rows[0]) == ("1600", "exact")
    assert loaded_doc.tables[1].rows[0].cells[0].paragraphs[0].text == "核心摘要"
    summary_width, summary_width_type = _cell_width(loaded_doc.tables[1].rows[0].cells[0])
    assert summary_width_type == "dxa"
    assert int(summary_width) > 9300
    assert loaded_doc.tables[2].rows[0].cells[0].paragraphs[0].text == "营业收入"
    assert loaded_doc.tables[2].rows[0].cells[2].paragraphs[0].text == ""
    assert _cell_width(loaded_doc.tables[2].rows[0].cells[0]) == ("3120", "dxa")
    assert _cell_width(loaded_doc.tables[2].rows[0].cells[1]) == ("3120", "dxa")
    assert _cell_width(loaded_doc.tables[2].rows[0].cells[2]) == ("3120", "dxa")
    assert _cell_fill(loaded_doc.tables[2].rows[0].cells[0]) == "F2F7FC"
    assert _cell_width(loaded_doc.tables[3].rows[0].cells[0]) == ("160", "dxa")
    assert _cell_width(loaded_doc.tables[3].rows[0].cells[1]) == ("9200", "dxa")
    assert loaded_doc.tables[3].rows[0].cells[1].paragraphs[0].text == "编制：战略发展部 · 审核：CFO 办公室"
    assert all(
        paragraph.text != "Q3 经营复盘报告" for paragraph in loaded_doc.paragraphs
    )
    assert not any(
        _paragraph_has_page_break(paragraph) for paragraph in loaded_doc.paragraphs
    )
    overview_heading = _find_paragraph(loaded_doc, "一、分区业绩")
    assert overview_heading.text == "一、分区业绩"
    overview_divider = _paragraph_after(loaded_doc, overview_heading)
    assert _paragraph_bottom_border_color(overview_divider) == "1F4E79"
    assert _paragraph_bottom_border_size(overview_divider) == "12"
    assert overview_divider.paragraph_format.space_before.pt == pytest.approx(2.5, abs=0.3)
    assert overview_heading.paragraph_format.space_before.pt == pytest.approx(8, abs=0.2)
    assert overview_heading.paragraph_format.space_after.pt == pytest.approx(0, abs=0.2)
    assert (
        _find_paragraph(loaded_doc, "华东和华南区域继续承担主要增长贡献。").text
        == "华东和华南区域继续承担主要增长贡献。"
    )


def test_node_renderer_supports_technical_resume_page_template(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-page-template-technical-resume",
        {
            "document_id": "page-template-technical-resume",
            "metadata": _business_report_metadata(title=""),
            "blocks": [_technical_resume_block()],
        },
    )

    name_paragraph = _find_paragraph(loaded_doc, "张明远")
    headline_paragraph = _find_paragraph(
        loaded_doc, "后端开发工程师  ·  分布式系统 / 高并发架构"
    )
    contact_paragraph = _find_paragraph(
        loaded_doc,
        "zhangmingyuan@email.com  ·  138-0000-0000  ·  北京 | 可远程  ·  github.com/zhangmy",
    )
    contact_divider = _paragraph_after(loaded_doc, contact_paragraph)
    education_heading = _find_paragraph(loaded_doc, "教育背景")
    education_entry = _find_paragraph(
        loaded_doc, "北京大学\t2019.09 – 2023.06\n计算机科学与技术  |  工学学士"
    )
    detail = _find_paragraph(
        loaded_doc,
        "GPA 3.86/4.0，连续三年一等奖学金，排名前 5%",
    )

    assert all(
        paragraph.text != "张明远 - 后端开发工程师简历"
        for paragraph in loaded_doc.paragraphs
    )
    assert _section_margin_twips(loaded_doc.sections[0], "top") == 1080
    assert _section_margin_twips(loaded_doc.sections[0], "right") == 1260
    assert _section_margin_twips(loaded_doc.sections[0], "bottom") == 1080
    assert _section_margin_twips(loaded_doc.sections[0], "left") == 1260
    assert name_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_run_size(name_paragraph) == pytest.approx(28.0, abs=0.2)
    assert _paragraph_run_rgb(name_paragraph) == "1A1A1A"
    assert headline_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_run_rgb(headline_paragraph) == "444444"
    assert _paragraph_run_size(headline_paragraph) == pytest.approx(10.0, abs=0.2)
    assert contact_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_bottom_border_color(contact_paragraph) is None
    assert _paragraph_bottom_border_color(contact_divider) == "000000"
    assert _paragraph_bottom_border_size(contact_divider) == "4"
    assert _paragraph_run_size(contact_paragraph) == pytest.approx(9.0, abs=0.2)
    assert _paragraph_bottom_border_color(education_heading) == "000000"
    assert _paragraph_bottom_border_size(education_heading) == "4"
    assert _paragraph_run_rgb(education_heading) == "1A1A1A"
    assert _paragraph_tab_positions(education_entry) == ["9026"]
    assert education_entry.runs[0].bold is True
    assert _paragraph_run_rgb(education_entry) == "1A1A1A"
    assert education_entry.runs[2].italic is True
    assert detail.runs[0].bold is False
    assert _paragraph_run_size(detail) == pytest.approx(10.0, abs=0.2)


def test_node_renderer_preserves_multiline_table_cell_text(workspace_root: Path):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-multiline-table-cell",
        {
            "document_id": "multiline-table-cell",
            "metadata": _business_report_metadata(title=""),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["阶段", "说明"],
                    "rows": [["Q1", {"text": "第一行\\n第二行"}]],
                }
            ],
        },
    )

    body_cell_paragraph = loaded_doc.tables[0].rows[1].cells[1].paragraphs[0]

    assert body_cell_paragraph.text == "第一行\n第二行"
    assert "<w:br" in body_cell_paragraph._p.xml


@pytest.mark.asyncio
async def test_node_document_toolset_exports_technical_resume_page_template(
    workspace_root: Path,
):
    loaded_doc, _ = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-page-template-technical-resume",
        create_kwargs={
            "title": "张明远_个人简历",
            "output_name": "technical-resume-template.docx",
        },
        blocks=[_technical_resume_block()],
    )

    contact_paragraph = _find_paragraph(
        loaded_doc,
        "zhangmingyuan@email.com  ·  138-0000-0000  ·  北京 | 可远程  ·  github.com/zhangmy",
    )
    experience_entry = _find_paragraph(
        loaded_doc, "字节跳动 · 基础架构部\t2022.07 – 2022.12\n后端开发实习生 · 推荐系统组"
    )
    skills_heading = _find_paragraph(loaded_doc, "技术栈")

    assert _find_paragraph(loaded_doc, "张明远").text == "张明远"
    assert _section_margin_twips(loaded_doc.sections[0], "left") == 1260
    assert _section_margin_twips(loaded_doc.sections[0], "right") == 1260
    assert _paragraph_bottom_border_color(_paragraph_after(loaded_doc, contact_paragraph)) == "000000"
    assert _paragraph_bottom_border_color(skills_heading) == "000000"
    assert _paragraph_tab_positions(experience_entry) == ["9026"]


def test_node_renderer_supports_table_cell_overrides_and_dashboard_blocks(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-dashboard-blocks",
        {
            "document_id": "node-dashboard-blocks",
            "metadata": _business_report_metadata(
                document_style={
                    "table_defaults": {
                        "body_fill": "F8FAFC",
                    }
                },
            ),
            "blocks": [
                {
                    "type": "table",
                    "headers": ["区域", "预算完成率"],
                    "rows": [
                        [
                            "华东",
                            {
                                "text": "112%",
                                "fill": "DCFCE7",
                                "text_color": "166534",
                                "bold": True,
                                "align": "right",
                            },
                        ]
                    ],
                    "border_style": "minimal",
                },
                {
                    "type": "accent_box",
                    "title": "核心摘要",
                    "text": "Q3 整体经营表现稳健，增长质量继续改善。",
                    "accent_color": "1F4E79",
                    "fill_color": "F8FAFC",
                },
                {
                    "type": "metric_cards",
                    "accent_color": "1F4E79",
                    "fill_color": "F8FAFC",
                    "metrics": [
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "↑ 18.4% YoY",
                            "delta_color": "15803D",
                        },
                        {
                            "label": "毛利率",
                            "value": "32.1%",
                            "delta": "↓ 0.3pp vs Q2",
                            "delta_color": "DC2626",
                        },
                    ],
                },
            ],
        },
    )

    table = loaded_doc.tables[0]
    accent_table = loaded_doc.tables[1]
    metric_table = loaded_doc.tables[2]

    assert _cell_fill(table.rows[1].cells[0]) == "F7FBFF"
    assert _cell_fill(table.rows[1].cells[1]) == "DCFCE7"
    assert _run_rgb(table.rows[1].cells[1]) == "166534"
    assert _run_bold(table.rows[1].cells[1]) is True
    assert _table_width(accent_table) == ("9360", "dxa")
    assert _table_grid_widths(accent_table) == [9360]
    assert accent_table.rows[0].cells[0].text.startswith("核心摘要")
    assert _cell_fill(accent_table.rows[0].cells[0]) == "F8FAFC"
    assert _cell_border_color(accent_table.rows[0].cells[0], "left") == "1F4E79"
    assert _cell_fill(metric_table.rows[0].cells[0]) == "F8FAFC"
    assert metric_table.rows[0].cells[0].paragraphs[0].text == "营业收入"
    assert metric_table.rows[0].cells[0].paragraphs[1].text == "¥4.82 亿"
    assert _paragraph_run_rgb(metric_table.rows[0].cells[0].paragraphs[2]) == "15803D"


def test_node_renderer_supports_hero_banner_fonts_and_report_box_styles(
    workspace_root: Path,
):
    loaded_doc, output_path = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-hero-banner-styles",
        {
            "document_id": "node-hero-banner-styles",
            "metadata": _business_report_metadata(
                title="",
                document_style={
                    "font_name": "Microsoft YaHei",
                    "heading_font_name": "Source Han Sans SC",
                    "table_font_name": "SimSun",
                    "code_font_name": "Consolas",
                },
            ),
            "blocks": [
                {
                    "type": "hero_banner",
                    "title": "Q3 经营复盘报告",
                    "subtitle": "战略与增长委员会",
                    "theme_color": "1F4E79",
                    "text_color": "FFFFFF",
                    "subtitle_color": "DCE6F1",
                    "min_height_pt": 92,
                },
                {"type": "toc", "title": "目录", "levels": 1},
                {"type": "heading", "text": "一、经营总览", "level": 1},
                {
                    "type": "paragraph",
                    "runs": [
                        {"text": "正文段落 "},
                        {"text": "const revenue = true;", "code": True},
                    ],
                },
                {
                    "type": "accent_box",
                    "title": "核心摘要",
                    "text": "经营质量稳中有升。",
                    "accent_color": "1F4E79",
                    "fill_color": "F8FAFC",
                    "border_color": "CBD5E1",
                    "border_width_pt": 0.75,
                    "accent_border_width_pt": 3.0,
                    "padding_pt": 16,
                    "title_font_scale": 1.2,
                    "body_font_scale": 1.05,
                },
                {
                    "type": "metric_cards",
                    "accent_color": "1F4E79",
                    "fill_color": "F8FAFC",
                    "border_color": "D9E1E8",
                    "divider_color": "CBD5E1",
                    "border_width_pt": 0.75,
                    "divider_width_pt": 0.75,
                    "padding_pt": 14,
                    "label_font_scale": 0.9,
                    "value_font_scale": 1.8,
                    "delta_font_scale": 0.9,
                    "note_font_scale": 0.82,
                    "metrics": [
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "↑ 18.4% YoY",
                            "note": "核心业务保持增长",
                            "delta_color": "15803D",
                            "note_color": "64748B",
                        }
                    ],
                },
                {
                    "type": "table",
                    "caption": "区域营收完成情况",
                    "header_groups": [{"title": "经营数据", "span": 3}],
                    "headers": ["区域", "完成率", "备注"],
                    "header_fill": "1F4E79",
                    "header_text_color": "FFFFFF",
                    "header_font_scale": 1.1,
                    "body_font_scale": 0.95,
                    "cell_padding_horizontal_pt": 8,
                    "cell_padding_vertical_pt": 6,
                    "rows": [
                        [
                            {"text": "华东", "font_scale": 1.15},
                            {
                                "text": "112%",
                                "fill": "DCFCE7",
                                "text_color": "166534",
                                "bold": True,
                                "align": "right",
                                "font_scale": 1.2,
                            },
                            "达成",
                        ]
                    ],
                },
            ],
        },
    )

    hero_table = loaded_doc.tables[0]
    accent_table = loaded_doc.tables[1]
    metric_table = loaded_doc.tables[2]
    data_table = loaded_doc.tables[3]
    heading = _find_paragraph(loaded_doc, "一、经营总览")
    body_paragraph = _find_paragraph(loaded_doc, "正文段落 const revenue = true;")

    assert hero_table.rows[0].cells[0].text.startswith("Q3 经营复盘报告")
    assert _cell_fill(hero_table.rows[0].cells[0]) == "1F4E79"
    assert _table_width(hero_table) == ("9360", "dxa")
    assert _run_font_attr(hero_table.rows[0].cells[0].paragraphs[0].runs[0], "ascii") == (
        "Source Han Sans SC"
    )
    assert _run_font_attr(
        hero_table.rows[0].cells[0].paragraphs[0].runs[0], "eastAsia"
    ) == "Source Han Sans SC"
    assert _run_font_attr(heading.runs[0], "ascii") == "Source Han Sans SC"
    assert _run_font_attr(heading.runs[0], "eastAsia") == "Source Han Sans SC"
    assert _run_font_attr(body_paragraph.runs[0], "ascii") == "Microsoft YaHei"
    assert _run_font_attr(body_paragraph.runs[0], "eastAsia") == "Microsoft YaHei"
    assert _run_font_attr(body_paragraph.runs[1], "ascii") == "Consolas"
    assert _run_font_attr(body_paragraph.runs[1], "eastAsia") == "Consolas"
    assert _table_width(accent_table) == ("9360", "dxa")
    assert _table_grid_widths(accent_table) == [9360]
    assert accent_table.rows[0].cells[0].text.startswith("核心摘要")
    accent_width, accent_width_type = _cell_width(accent_table.rows[0].cells[0])
    assert accent_width_type == "dxa"
    assert int(accent_width) > 9300
    assert _cell_fill(accent_table.rows[0].cells[0]) == "F8FAFC"
    assert _cell_border_color(accent_table.rows[0].cells[0], "left") == "1F4E79"
    assert _cell_border_color(accent_table.rows[0].cells[0], "top") == "CBD5E1"
    assert _cell_border_color(accent_table.rows[0].cells[0], "right") == "CBD5E1"
    assert _cell_border_color(accent_table.rows[0].cells[0], "bottom") == "CBD5E1"
    assert _cell_border_size(accent_table.rows[0].cells[0], "top") == "6"


def test_node_renderer_uses_font_name_for_heading_and_table_when_specific_fonts_are_unset(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-font-name-fallbacks",
        {
            "document_id": "node-font-name-fallbacks",
            "metadata": _business_report_metadata(
                title="",
                document_style={
                    "font_name": "Source Han Sans SC",
                },
            ),
            "blocks": [
                {"type": "heading", "text": "一、经营总览", "level": 1},
                {
                    "type": "table",
                    "headers": ["区域", "营收"],
                    "rows": [["华东", "¥4.82 亿"]],
                },
                {"type": "paragraph", "text": "正文段落"},
            ],
        },
    )

    heading = _find_paragraph(loaded_doc, "一、经营总览")
    body = _find_paragraph(loaded_doc, "正文段落")
    table = loaded_doc.tables[0]

    assert _run_font_attr(heading.runs[0], "ascii") == "Source Han Sans SC"
    assert _run_font_attr(heading.runs[0], "eastAsia") == "Source Han Sans SC"
    assert _run_font_attr(table.rows[1].cells[0].paragraphs[0].runs[0], "ascii") == (
        "Source Han Sans SC"
    )
    assert _run_font_attr(table.rows[1].cells[0].paragraphs[0].runs[0], "eastAsia") == (
        "Source Han Sans SC"
    )
    assert _run_font_attr(body.runs[0], "ascii") == "Source Han Sans SC"
    assert _run_font_attr(body.runs[0], "eastAsia") == "Source Han Sans SC"


def test_node_renderer_business_report_defaults_hero_divider_and_green_accent_box(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-business-report-default-strip-and-divider",
        {
            "document_id": "business-report-default-strip-and-divider",
            "metadata": _business_report_metadata(title=""),
            "blocks": [
                {
                    "type": "hero_banner",
                    "title": "Q3 经营复盘报告",
                    "subtitle": "战略与增长委员会",
                },
                {
                    "type": "accent_box",
                    "title": "核心摘要",
                    "text": "经营质量稳中有升。",
                },
            ],
        },
    )

    hero_divider = loaded_doc.paragraphs[0]
    accent_table = loaded_doc.tables[1]

    assert _paragraph_bottom_border_color(hero_divider) == "2E75B6"
    assert _paragraph_bottom_border_size(hero_divider) == "12"
    assert _cell_fill(accent_table.rows[0].cells[0]) == "EEF9F5"
    assert _cell_border_color(accent_table.rows[0].cells[0], "left") == "0F6E56"
    assert _cell_border_size(accent_table.rows[0].cells[0], "left") == "24"
    assert _cell_margin(accent_table.rows[0].cells[0], "left") == "320"
    assert metric_table.rows[0].cells[0].paragraphs[1].runs[0].font.size.pt > (
        metric_table.rows[0].cells[0].paragraphs[0].runs[0].font.size.pt
    )
    assert metric_table.rows[0].cells[0].paragraphs[2].runs[0].font.size.pt < (
        metric_table.rows[0].cells[0].paragraphs[1].runs[0].font.size.pt
    )
    assert _run_font_attr(data_table.rows[1].cells[0].paragraphs[0].runs[0], "ascii") == (
        "SimSun"
    )
    assert _run_font_attr(data_table.rows[2].cells[0].paragraphs[0].runs[0], "eastAsia") == (
        "SimSun"
    )
    assert _table_cell_margin(data_table, "left") == "160"
    assert _table_cell_margin(data_table, "top") == "120"
    assert _table_row_height(data_table.rows[2]) == ("520", "atLeast")
    assert _table_row_height(data_table.rows[3]) == ("480", "atLeast")
    assert _paragraph_run_size(data_table.rows[2].cells[0].paragraphs[0]) == pytest.approx(
        11.5, abs=0.2
    )
    assert _paragraph_run_size(data_table.rows[3].cells[1].paragraphs[0]) == pytest.approx(
        12.0, abs=0.2
    )
    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert document_xml.count('w:pStyle w:val="Heading1"') == 1


def test_node_renderer_suppresses_document_title_when_hero_banner_is_first(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-hero-banner-hides-document-title",
        {
            "document_id": "node-hero-banner-hide-title",
            "metadata": _business_report_metadata(),
            "blocks": [
                {
                    "type": "hero_banner",
                    "title": "Q3 经营复盘报告",
                    "subtitle": "战略与增长委员会 · 2024 年 10 月",
                    "theme_color": "1F4E79",
                },
                {"type": "paragraph", "text": "正文内容"},
            ],
        },
    )

    assert loaded_doc.tables[0].rows[0].cells[0].paragraphs[0].text == "Q3 经营复盘报告"
    assert sum(
        1 for paragraph in loaded_doc.paragraphs if paragraph.text == "Q3 经营复盘报告"
    ) == 0
    assert _find_paragraph(loaded_doc, "正文内容").text == "正文内容"


def test_node_renderer_supports_section_inheritance_and_cover_normalization(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-section-inheritance",
        {
            "document_id": "node-section-inheritance",
            "metadata": _business_report_metadata(
                title="封面页码测试",
                header_footer={
                    "header_text": "董事会季度经营汇报",
                    "different_first_page": True,
                    "first_page_header_text": "董事会季度经营汇报封面",
                    "first_page_show_page_number": False,
                    "different_odd_even": True,
                    "even_page_header_text": "董事会季度经营汇报（偶数页）",
                    "show_page_number": True,
                },
            ),
            "blocks": [
                {"type": "paragraph", "text": "封面内容"},
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "page_orientation": "landscape",
                    "restart_page_numbering": True,
                },
                {"type": "paragraph", "text": "横向节内容"},
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "header_footer": {"show_page_number": False},
                },
                {"type": "paragraph", "text": "纵向节内容"},
            ],
        },
    )

    assert len(loaded_doc.sections) == 3
    assert _document_uses_odd_even_headers(loaded_doc) is True
    assert loaded_doc.sections[0].different_first_page_header_footer is True
    assert loaded_doc.sections[1].different_first_page_header_footer is False
    assert loaded_doc.sections[2].different_first_page_header_footer is False
    assert "PAGE" in loaded_doc.sections[1].footer._element.xml
    assert _section_page_number_start(loaded_doc.sections[1]) == 1
    assert loaded_doc.sections[2].footer.is_linked_to_previous is False
    assert "董事会季度经营汇报" in _story_texts(loaded_doc.sections[2].header)
    assert loaded_doc.sections[2].footer.paragraphs[0].text == ""
    assert "PAGE" not in loaded_doc.sections[2].footer._element.xml


def test_node_renderer_supports_heading_styles_and_split_header_footer(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-heading-split-layout",
        {
            "document_id": "node-heading-split-layout",
            "metadata": _business_report_metadata(
                document_style={
                    "heading_color": "000000",
                    "heading_level_1_color": "0F4C81",
                    "heading_level_2_color": "4B5563",
                    "heading_bottom_border_color": "CBD5E1",
                    "heading_bottom_border_size_pt": 1.5,
                },
                header_footer={
                    "header_left": "Q3 经营复盘报告 | 战略与增长委员会",
                    "header_right": "机密 · 2024 年 10 月",
                    "header_border_bottom": True,
                    "header_border_color": "D0D7DE",
                    "footer_left": "集团战略部 · 内部机密文件",
                    "footer_right": "第 {PAGE} 页",
                    "footer_border_top": True,
                    "footer_border_color": "D0D7DE",
                },
            ),
            "blocks": [
                {
                    "type": "heading",
                    "text": "一、经营总览",
                    "level": 1,
                    "bottom_border": True,
                },
                {
                    "type": "heading",
                    "text": "1.1 各区营收完成情况",
                    "level": 2,
                    "bottom_border": True,
                },
                {
                    "type": "heading",
                    "text": "三级标题",
                    "level": 3,
                },
                {"type": "paragraph", "text": "正文"},
            ],
        },
    )

    title_paragraph = _find_paragraph(loaded_doc, "Q3 经营复盘报告")
    heading_one = _find_paragraph(loaded_doc, "一、经营总览")
    heading_two = _find_paragraph(loaded_doc, "1.1 各区营收完成情况")
    heading_three = _find_paragraph(loaded_doc, "三级标题")
    header_paragraph = loaded_doc.sections[0].header.paragraphs[0]
    footer_paragraph = loaded_doc.sections[0].footer.paragraphs[0]

    assert heading_one.style.style_id == "Heading1"
    assert heading_two.style.style_id == "Heading2"
    assert heading_three.style.style_id == "Heading3"
    assert _paragraph_run_rgb(title_paragraph) == "000000"
    assert _paragraph_run_rgb(heading_one) == "0F4C81"
    assert _paragraph_run_rgb(heading_two) == "4B5563"
    heading_one_divider = _paragraph_after(loaded_doc, heading_one)
    assert _paragraph_bottom_border_color(heading_one_divider) == "CBD5E1"
    assert _paragraph_bottom_border_size(heading_one_divider) == "12"
    assert heading_one_divider.paragraph_format.space_before.pt == pytest.approx(2.5, abs=0.3)
    assert "Q3 经营复盘报告 | 战略与增长委员会" in header_paragraph.text
    assert "机密 · 2024 年 10 月" in header_paragraph.text
    assert "集团战略部 · 内部机密文件" in footer_paragraph.text
    assert "第 " in footer_paragraph.text
    assert _paragraph_bottom_border_color(header_paragraph) == "D0D7DE"
    assert _paragraph_top_border_color(footer_paragraph) == "D0D7DE"
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml


def test_node_renderer_uses_accent_fallback_and_header_footer_fonts(
    workspace_root: Path,
):
    accent_color = "C2410C"
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-accent-fallback-header-footer-fonts",
        {
            "document_id": "node-accent-fallback-header-footer-fonts",
            "metadata": _business_report_metadata(
                title="经营复盘",
                accent_color=accent_color,
                document_style={
                    "font_name": "Microsoft YaHei",
                },
                header_footer={
                    "header_left": "经营复盘 | 战略组",
                    "header_right": "机密 · 2026 年 4 月",
                    "footer_left": "集团战略部",
                    "footer_right": "第 {PAGE} 页",
                },
            ),
            "blocks": [
                {
                    "type": "heading",
                    "text": "一、经营总览",
                    "level": 1,
                },
                {
                    "type": "accent_box",
                    "title": "核心摘要",
                    "text": "经营质量继续改善，现金流保持稳定。",
                },
            ],
        },
    )

    def _blend_hex(source: str, target: str, ratio: float) -> str:
        source_channels = [int(source[index : index + 2], 16) for index in (0, 2, 4)]
        target_channels = [int(target[index : index + 2], 16) for index in (0, 2, 4)]
        blended = [
            round(src * (1 - ratio) + dst * ratio)
            for src, dst in zip(source_channels, target_channels, strict=False)
        ]
        return "".join(f"{value:02X}" for value in blended)

    heading = _find_paragraph(loaded_doc, "一、经营总览")
    accent_table = loaded_doc.tables[0]
    header_paragraph = loaded_doc.sections[0].header.paragraphs[0]
    footer_paragraph = loaded_doc.sections[0].footer.paragraphs[0]

    assert _paragraph_run_rgb(heading) == accent_color
    assert _cell_fill(accent_table.rows[0].cells[0]) == accent_color
    assert _cell_fill(accent_table.rows[0].cells[1]) == _blend_hex(
        accent_color,
        "FFFFFF",
        0.92,
    )
    assert _run_font_attr(header_paragraph.runs[0], "ascii") == "Microsoft YaHei"
    assert _run_font_attr(header_paragraph.runs[0], "eastAsia") == "Microsoft YaHei"
    assert _run_font_attr(footer_paragraph.runs[0], "ascii") == "Microsoft YaHei"
    assert _run_font_attr(footer_paragraph.runs[0], "eastAsia") == "Microsoft YaHei"


def test_node_renderer_supports_default_header_footer_baseline(workspace_root: Path):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-default-header-footer",
        {
            "document_id": "node-default-header-footer",
            "metadata": _business_report_metadata(title="默认页眉页脚测试"),
            "blocks": [{"type": "paragraph", "text": "正文"}],
        },
    )

    assert _document_updates_fields_on_open(loaded_doc) is True
    assert _document_uses_odd_even_headers(loaded_doc) is False
    assert "PAGE" not in loaded_doc.sections[0].footer._element.xml
    assert "PAGE" not in loaded_doc.sections[0].first_page_footer._element.xml
    assert "PAGE" not in loaded_doc.sections[0].even_page_footer._element.xml


def test_node_renderer_supports_section_override_inheritance_and_nested_even_headers(
    workspace_root: Path,
):
    from docx.enum.section import WD_ORIENT

    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-section-variants",
        {
            "document_id": "node-section-variants",
            "metadata": _business_report_metadata(
                title="分节测试",
                header_footer={
                    "header_text": "默认页眉",
                    "footer_text": "默认页脚",
                    "show_page_number": True,
                },
            ),
            "blocks": [
                {"type": "paragraph", "text": "第一节"},
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "inherit_header_footer": False,
                    "page_orientation": "landscape",
                    "margins": {
                        "top_cm": 1.5,
                        "bottom_cm": 1.8,
                        "left_cm": 1.2,
                        "right_cm": 1.4,
                    },
                    "restart_page_numbering": True,
                    "page_number_start": 3,
                    "header_footer": {
                        "header_text": "第二节页眉",
                        "footer_text": "第二节页脚",
                        "show_page_number": True,
                        "different_odd_even": True,
                        "even_page_header_text": "第二节偶数页页眉",
                    },
                },
                {"type": "paragraph", "text": "第二节"},
                {"type": "section_break", "start_type": "new_page"},
                {"type": "paragraph", "text": "第三节"},
                {
                    "type": "group",
                    "blocks": [
                        {
                            "type": "columns",
                            "columns": [
                                {
                                    "blocks": [
                                        {
                                            "type": "section_break",
                                            "start_type": "new_page",
                                            "inherit_header_footer": False,
                                            "header_footer": {
                                                "different_odd_even": True,
                                                "even_page_header_text": "嵌套偶数页页眉",
                                            },
                                        },
                                        {"type": "paragraph", "text": "第四节"},
                                    ]
                                }
                            ],
                        }
                    ],
                },
            ],
        },
    )

    assert len(loaded_doc.sections) == 4
    assert "默认页眉" in _story_texts(loaded_doc.sections[0].header)
    assert "第二节页眉" in _story_texts(loaded_doc.sections[1].header)
    assert any(
        text.startswith("第二节页脚") for text in _story_texts(loaded_doc.sections[1].footer)
    )
    assert "第二节偶数页页眉" in _story_texts(loaded_doc.sections[1].even_page_header)
    assert "PAGE" in loaded_doc.sections[1].footer._element.xml
    assert _section_page_number_start(loaded_doc.sections[1]) == 3
    assert loaded_doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert loaded_doc.sections[1].top_margin.cm == pytest.approx(1.5, abs=0.01)
    assert loaded_doc.sections[1].bottom_margin.cm == pytest.approx(1.8, abs=0.01)
    assert loaded_doc.sections[1].left_margin.cm == pytest.approx(1.2, abs=0.01)
    assert loaded_doc.sections[1].right_margin.cm == pytest.approx(1.4, abs=0.01)
    assert loaded_doc.sections[2].header.is_linked_to_previous is True
    assert loaded_doc.sections[2].footer.is_linked_to_previous is True
    assert _section_page_number_start(loaded_doc.sections[2]) is None
    assert _document_uses_odd_even_headers(loaded_doc) is True
    assert "嵌套偶数页页眉" in _story_texts(loaded_doc.sections[3].even_page_header)


def test_node_renderer_clears_inherited_header_footer_when_disabled_without_override(
    workspace_root: Path,
):
    loaded_doc, _ = _render_structured_payload_with_node(
        workspace_root,
        "pytest-node-renderer-clear-inherited-header-footer",
        {
            "document_id": "node-clear-inherited-header-footer",
            "metadata": _business_report_metadata(
                title="分节清空测试",
                header_footer={
                    "header_text": "默认页眉",
                    "footer_text": "默认页脚",
                    "show_page_number": True,
                    "different_odd_even": True,
                    "even_page_header_text": "默认偶数页眉",
                },
            ),
            "blocks": [
                {"type": "paragraph", "text": "第一节"},
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "inherit_header_footer": False,
                },
                {"type": "paragraph", "text": "第二节"},
            ],
        },
    )

    assert len(loaded_doc.sections) == 2
    assert loaded_doc.sections[1].header.is_linked_to_previous is False
    assert loaded_doc.sections[1].footer.is_linked_to_previous is False
    assert loaded_doc.sections[1].header.paragraphs[0].text == ""
    assert loaded_doc.sections[1].footer.paragraphs[0].text == ""
    assert "PAGE" not in loaded_doc.sections[1].footer._element.xml
    assert loaded_doc.sections[1].even_page_header.paragraphs[0].text == ""


@pytest.mark.asyncio
async def test_node_document_toolset_exports_training_summary_golden_sample(
    workspace_root: Path,
):
    loaded_doc, _ = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-training-summary",
        create_kwargs={
            "title": "Sample Training Summary Report",
            "output_name": "sample-training-summary.docx",
            "theme_name": "business_report",
            "title_align": "center",
            "document_style": {"heading_color": "000000"},
            "header_footer": {
                "show_page_number": True,
                "page_number_align": "right",
            },
        },
        blocks=[
            {
                "type": "paragraph",
                "runs": [
                    {"text": "Training Title: ", "bold": True},
                    {"text": "Advanced Skills Workshop"},
                ],
            },
            {
                "type": "paragraph",
                "runs": [
                    {"text": "Date: ", "bold": True},
                    {"text": "April 10-11, 2026"},
                ],
            },
            {
                "type": "heading",
                "text": "I. Overview",
                "level": 1,
                "bottom_border": True,
                "bottom_border_color": "D0D7DE",
            },
            {
                "type": "paragraph",
                "text": "The workshop enhanced technical proficiency and collaboration across the team.",
            },
            {
                "type": "heading",
                "text": "II. Training Objectives",
                "level": 1,
                "bottom_border": True,
                "bottom_border_color": "D0D7DE",
            },
            {
                "type": "list",
                "items": [
                    {
                        "runs": [
                            {"text": "Enhance technical proficiency: ", "bold": True},
                            {"text": "Focus on the latest tools and methodologies."},
                        ]
                    },
                    {
                        "runs": [
                            {"text": "Improve collaboration: ", "bold": True},
                            {
                                "text": "Strengthen communication through shared exercises.",
                            },
                        ]
                    },
                ],
            },
            {
                "type": "heading",
                "text": "III. Training Schedule",
                "level": 1,
                "bottom_border": True,
                "bottom_border_color": "D0D7DE",
            },
            {
                "type": "table",
                "headers": ["Date", "Time", "Session Title", "Trainer"],
                "header_fill_enabled": False,
                "header_text_color": "666666",
                "header_bold": False,
                "border_style": "minimal",
                "rows": [
                    [
                        {"text": "2026-04-10", "row_span": 2},
                        "09:00 - 12:00",
                        "Introduction to Tools",
                        "Alice Smith",
                    ],
                    ["13:00 - 16:00", "Hands-on Practice", "Bob Johnson"],
                    [
                        {"text": "2026-04-11", "row_span": 2},
                        "09:00 - 12:00",
                        "Advanced Techniques",
                        "Alice Smith",
                    ],
                    ["13:00 - 16:00", "Group Project", "Bob Johnson"],
                ],
            },
            {
                "type": "heading",
                "text": "IV. Participant Feedback",
                "level": 1,
                "bottom_border": True,
                "bottom_border_color": "D0D7DE",
            },
            {
                "type": "table",
                "headers": ["Category", "Rating", "Comments"],
                "header_fill_enabled": False,
                "header_text_color": "666666",
                "header_bold": False,
                "border_style": "minimal",
                "rows": [
                    ["Content Quality", "4.8", "Very relevant and up-to-date"],
                    ["Trainer Expertise", "4.9", "Engaging and highly knowledgeable"],
                ],
            },
        ],
    )

    title_paragraph = _find_paragraph(loaded_doc, "Sample Training Summary Report")
    overview_heading = _find_paragraph(loaded_doc, "I. Overview")
    overview_divider = _paragraph_after(loaded_doc, overview_heading)
    schedule_table = loaded_doc.tables[0]
    feedback_table = loaded_doc.tables[1]

    assert _paragraph_run_rgb(title_paragraph) == "000000"
    assert _paragraph_bottom_border_color(overview_heading) is None
    assert _paragraph_bottom_border_color(overview_divider) == "D0D7DE"
    assert any(
        "Enhance technical proficiency:" in paragraph.text
        for paragraph in loaded_doc.paragraphs
    )
    assert schedule_table.rows[0].cells[0].text == "III. Training Schedule"
    assert _cell_vertical_merge(schedule_table.rows[2].cells[0]) == "restart"
    assert _raw_row_cell_vertical_merge(schedule_table.rows[3], 0) == "continue"
    assert _run_rgb(schedule_table.rows[1].cells[0]) == "666666"
    assert _run_bold(schedule_table.rows[1].cells[0]) is False
    assert feedback_table.rows[0].cells[0].text == "IV. Participant Feedback"
    assert feedback_table.rows[2].cells[0].text == "Content Quality"
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml


@pytest.mark.asyncio
async def test_node_document_toolset_exports_executive_brief_golden_sample(
    workspace_root: Path,
):
    loaded_doc, _ = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-executive-brief",
        create_kwargs={
            "title": "Executive Brief",
            "output_name": "executive-brief.docx",
            "theme_name": "executive_brief",
            "title_align": "center",
            "header_footer": {
                "header_left": "Executive Brief | Strategy Office",
                "header_right": "Confidential | October 2026",
                "header_border_bottom": True,
                "footer_left": "Prepared for Leadership Team",
                "footer_right": "Page {PAGE}",
                "footer_border_top": True,
            },
        },
        blocks=[
            {
                "type": "heading",
                "text": "Executive Summary",
                "level": 1,
                "bottom_border": True,
            },
            {
                "type": "accent_box",
                "title": "Key Message",
                "text": "Revenue momentum remained ahead of plan while cost pressure stayed manageable.",
                "accent_color": "1F4E79",
                "fill_color": "F8FAFC",
            },
            {
                "type": "columns",
                "columns": [
                    {
                        "blocks": [
                            {"type": "heading", "text": "What Changed", "level": 2},
                            {
                                "type": "paragraph",
                                "text": "Market momentum remains strong in enterprise accounts.",
                            },
                        ]
                    },
                    {
                        "blocks": [
                            {"type": "heading", "text": "Decision Requested", "level": 2},
                            {
                                "type": "paragraph",
                                "text": "Approve Q4 hiring for sales enablement and analytics.",
                            },
                        ]
                    },
                ],
            },
            {
                "type": "heading",
                "text": "Priority Actions",
                "level": 1,
                "bottom_border": True,
            },
            {
                "type": "table",
                "headers": ["Priority", "Owner", "Timing"],
                "border_style": "minimal",
                "rows": [
                    ["Pipeline quality review", "Revenue Ops", "Week 1"],
                    ["Pricing guardrail refresh", "Finance", "Week 2"],
                    ["Retention playbook launch", "Customer Success", "Week 3"],
                ],
            },
        ],
    )

    summary_heading = _find_paragraph(loaded_doc, "Executive Summary")
    header_paragraph = loaded_doc.sections[0].header.paragraphs[0]
    footer_paragraph = loaded_doc.sections[0].footer.paragraphs[0]
    accent_table = loaded_doc.tables[0]
    priorities_table = loaded_doc.tables[1]

    assert summary_heading.style.style_id == "Heading1"
    assert _paragraph_bottom_border_color(summary_heading) is not None
    assert "Executive Brief | Strategy Office" in header_paragraph.text
    assert "Confidential | October 2026" in header_paragraph.text
    assert "Prepared for Leadership Team" in footer_paragraph.text
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml
    assert accent_table.rows[0].cells[0].text.startswith("Key Message")
    assert _cell_border_color(accent_table.rows[0].cells[0], "left") == "1F4E79"
    assert any(
        "Market momentum remains strong in enterprise accounts." in paragraph.text
        for paragraph in loaded_doc.paragraphs
    )
    assert any(
        "Approve Q4 hiring for sales enablement and analytics." in paragraph.text
        for paragraph in loaded_doc.paragraphs
    )
    assert priorities_table.rows[0].cells[0].text == "Priority Actions"
    assert priorities_table.rows[2].cells[0].text == "Pipeline quality review"


@pytest.mark.asyncio
async def test_node_document_toolset_exports_q3_business_review_golden_sample(
    workspace_root: Path,
):
    loaded_doc, _ = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-q3-business-review",
        create_kwargs={
            "title": "Q3 经营复盘报告",
            "output_name": "q3-business-review.docx",
            "theme_name": "business_report",
            "title_align": "center",
            "document_style": {
                "heading_color": "000000",
                "heading_level_1_color": "0F4C81",
            },
            "header_footer": {
                "header_left": "Q3 经营复盘报告 | 战略与增长委员会",
                "header_right": "机密 · 2024 年 10 月",
                "header_border_bottom": True,
                "footer_left": "集团战略部 · 内部机密文件",
                "footer_right": "第 {PAGE} 页",
                "footer_border_top": True,
            },
        },
        blocks=[
            {
                "type": "accent_box",
                "title": "核心摘要",
                "text": "Q3 整体经营表现稳健，增长质量继续改善。",
                "accent_color": "1F4E79",
                "fill_color": "F8FAFC",
            },
            {
                "type": "metric_cards",
                "accent_color": "1F4E79",
                "fill_color": "F8FAFC",
                "metrics": [
                    {
                        "label": "营业收入",
                        "value": "¥4.82 亿",
                        "delta": "↑ 18.4% YoY",
                        "delta_color": "15803D",
                    },
                    {
                        "label": "毛利率",
                        "value": "32.1%",
                        "delta": "↓ 0.3pp vs Q2",
                        "delta_color": "DC2626",
                    },
                    {
                        "label": "净新增客户",
                        "value": "184",
                        "delta": "↑ 12.0% QoQ",
                        "delta_color": "15803D",
                    },
                ],
            },
            {
                "type": "heading",
                "text": "一、分区业绩",
                "level": 1,
                "bottom_border": True,
            },
            {
                "type": "table",
                "headers": ["区域", "收入", "预算完成率", "备注"],
                "header_fill": "1F4E79",
                "header_text_color": "FFFFFF",
                "border_style": "minimal",
                "rows": [
                    [
                        "华东",
                        "¥1.62 亿",
                        {
                            "text": "112%",
                            "fill": "DCFCE7",
                            "text_color": "166534",
                            "bold": True,
                        },
                        "超额完成",
                    ],
                    [
                        "华南",
                        "¥1.08 亿",
                        {
                            "text": "95%",
                            "fill": "FEF3C7",
                            "text_color": "92400E",
                            "bold": True,
                        },
                        "接近目标",
                    ],
                    [
                        "北区",
                        "¥0.94 亿",
                        {
                            "text": "89%",
                            "fill": "FEE2E2",
                            "text_color": "991B1B",
                            "bold": True,
                        },
                        "恢复中",
                    ],
                ],
            },
            {
                "type": "heading",
                "text": "二、风险与应对",
                "level": 1,
                "bottom_border": True,
            },
            {
                "type": "list",
                "items": [
                    {
                        "runs": [
                            {"text": "供应链风险：", "bold": True},
                            {"text": "部分关键器件交付周期仍有波动。"},
                        ]
                    },
                    {
                        "runs": [
                            {"text": "价格压力：", "bold": True},
                            {"text": "重点客户在续约谈判中要求更强折扣。"},
                        ]
                    },
                ],
            },
        ],
    )

    title_paragraph = _find_paragraph(loaded_doc, "Q3 经营复盘报告")
    header_paragraph = loaded_doc.sections[0].header.paragraphs[0]
    footer_paragraph = loaded_doc.sections[0].footer.paragraphs[0]
    accent_table = loaded_doc.tables[0]
    metric_table = loaded_doc.tables[1]
    data_table = loaded_doc.tables[2]

    assert _paragraph_run_rgb(title_paragraph) == "000000"
    assert "Q3 经营复盘报告 | 战略与增长委员会" in header_paragraph.text
    assert "集团战略部 · 内部机密文件" in footer_paragraph.text
    assert "PAGE" in loaded_doc.sections[0].footer._element.xml
    assert accent_table.rows[0].cells[0].text.startswith("核心摘要")
    assert _cell_fill(accent_table.rows[0].cells[0]) == "F8FAFC"
    assert _cell_border_color(accent_table.rows[0].cells[0], "left") == "1F4E79"
    assert metric_table.rows[0].cells[0].paragraphs[0].text == "营业收入"
    assert _paragraph_run_rgb(metric_table.rows[0].cells[0].paragraphs[2]) == "15803D"
    assert _paragraph_run_rgb(metric_table.rows[0].cells[1].paragraphs[2]) == "DC2626"
    assert data_table.rows[0].cells[0].text == "一、分区业绩"
    assert _cell_fill(data_table.rows[1].cells[0]) == "1F4E79"
    assert _run_rgb(data_table.rows[1].cells[0]) == "FFFFFF"
    assert _cell_fill(data_table.rows[2].cells[2]) == "DCFCE7"
    assert _cell_fill(data_table.rows[3].cells[2]) == "FEF3C7"
    assert _cell_fill(data_table.rows[4].cells[2]) == "FEE2E2"
    assert any("供应链风险：" in paragraph.text for paragraph in loaded_doc.paragraphs)


@pytest.mark.asyncio
async def test_node_document_toolset_exports_low_frequency_parity_sample(
    workspace_root: Path,
):
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    loaded_doc, exported_path = await _export_docx_via_node_toolset(
        workspace_root,
        "pytest-node-toolset-low-frequency-parity",
        create_kwargs={
            "title": "低频结构验收样例",
            "output_name": "low-frequency-parity.docx",
            "theme_name": "business_report",
            "document_style": {
                "summary_card_defaults": _summary_card_defaults(
                    title_font_scale=1.15,
                    title_space_before=10,
                    title_space_after=3,
                    list_space_after=7,
                ),
                "table_defaults": {
                    "header_fill": "1F4E79",
                    "header_text_color": "FFFFFF",
                    "banded_rows": True,
                    "banded_row_fill": "EEF4FA",
                    "table_align": "center",
                    "border_style": "minimal",
                    "cell_align": "center",
                },
            },
            "header_footer": {
                "header_text": "默认页眉",
                "footer_text": "默认页脚",
                "show_page_number": True,
            },
        },
        blocks=[
            _summary_card_block(
                title="结论与行动计划",
                items=["统一节奏", "聚焦续约"],
                variant="summary",
            ),
            {
                "type": "columns",
                "columns": [
                    {
                        "blocks": [
                            {
                                "type": "paragraph",
                                "text": "左栏提示：以下为横向节的详细看板。",
                            }
                        ]
                    },
                    {
                        "blocks": [
                            {
                                "type": "section_break",
                                "start_type": "new_page",
                                "inherit_header_footer": False,
                                "page_orientation": "landscape",
                                "restart_page_numbering": True,
                                "page_number_start": 2,
                                "header_footer": {
                                    "header_text": "横向明细页眉",
                                    "footer_text": "横向明细页脚",
                                    "show_page_number": True,
                                },
                            },
                            {
                                "type": "heading",
                                "text": "明细运营看板",
                                "level": 1,
                                "bottom_border": True,
                            },
                            {
                                "type": "table",
                                "header_groups": [
                                    {"title": "经营数据", "span": 3},
                                    {"title": "结果", "span": 1},
                                ],
                                "headers": ["区域", "Q3", "Q4", "完成率"],
                                "numeric_columns": [1, 2, 3],
                                "rows": [
                                    [
                                        "华东",
                                        "120",
                                        "135",
                                        {
                                            "text": "112%",
                                            "fill": "DCFCE7",
                                            "text_color": "166534",
                                            "bold": True,
                                        },
                                    ],
                                    [
                                        "华南",
                                        "108",
                                        "102",
                                        {
                                            "text": "95%",
                                            "fill": "FEF3C7",
                                            "text_color": "92400E",
                                            "bold": True,
                                        },
                                    ],
                                ],
                            },
                        ]
                    },
                ],
            },
        ],
    )

    title_paragraph = _find_paragraph(loaded_doc, "结论与行动计划")
    summary_item_paragraph = _find_paragraph(loaded_doc, "• 统一节奏")
    detail_heading = _find_paragraph(loaded_doc, "明细运营看板")
    detail_table = loaded_doc.tables[0]

    assert len(loaded_doc.sections) == 2
    assert title_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title_paragraph.runs[0].bold is True
    assert _paragraph_run_size(title_paragraph) == pytest.approx(12.5, abs=0.1)
    assert (_paragraph_run_size(title_paragraph) or 0) > (
        _paragraph_run_size(summary_item_paragraph) or 0
    )
    assert title_paragraph.paragraph_format.space_before.pt == pytest.approx(10, abs=0.5)
    assert title_paragraph.paragraph_format.space_after.pt == pytest.approx(3, abs=0.5)
    assert summary_item_paragraph.paragraph_format.space_after.pt == pytest.approx(
        7, abs=0.5
    )
    assert loaded_doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert _section_page_number_start(loaded_doc.sections[1]) == 2
    assert "横向明细页眉" in _story_texts(loaded_doc.sections[1].header)
    assert any(
        text.startswith("横向明细页脚")
        for text in _story_texts(loaded_doc.sections[1].footer)
    )
    assert "PAGE" in loaded_doc.sections[1].footer._element.xml
    assert _paragraph_bottom_border_color(_paragraph_after(loaded_doc, detail_heading)) is not None
    assert detail_table.rows[0].cells[0].text == "经营数据"
    assert detail_table.rows[1].cells[0].text == "区域"
    assert (
        detail_table.rows[2].cells[1].paragraphs[0].alignment
        == WD_ALIGN_PARAGRAPH.CENTER
    )
    assert detail_table.rows[2].cells[3].text == "112%"
    assert _cell_fill(detail_table.rows[2].cells[3]) == "DCFCE7"
    assert _run_rgb(detail_table.rows[2].cells[3]) == "166534"
    assert detail_table.rows[3].cells[3].text == "95%"
    assert _cell_fill(detail_table.rows[3].cells[3]) == "FEF3C7"
    with zipfile.ZipFile(exported_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:type w:val="nextPage"' in document_xml


def test_summary_card_defaults_apply_in_node_renderer(workspace_root: Path):
    docx = pytest.importorskip("docx")

    renderer_entry = _node_renderer_entry()

    workspace_dir = _make_workspace(workspace_root, "pytest-summary-card-node")
    document = DocumentModel(
        document_id="summary-card-parity",
        session_id="pytest-session",
        metadata=DocumentMetadata(
            title="Summary Card Parity",
            theme_name="business_report",
            document_style=DocumentStyleConfig(
                summary_card_defaults=DocumentSummaryCardDefaults(
                    **_summary_card_defaults(
                        title_font_scale=1.15,
                        title_space_before=10,
                        title_space_after=3,
                        list_space_after=7,
                    )
                )
            ),
        ),
        blocks=[
            SummaryCardBlock(**_summary_card_block(title="结论与行动计划", items=["统一节奏", "聚焦续约"], variant="summary"))
        ],
    )

    node_output = workspace_dir / "summary-card-node.docx"

    NodeDocumentRenderBackend(renderer_entry).render(document, node_output)

    node_doc = docx.Document(node_output)

    node_title = _find_paragraph(node_doc, "结论与行动计划")
    node_item = _find_paragraph(node_doc, "• 统一节奏")

    assert node_title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert node_title.runs[0].bold is True
    assert _paragraph_run_size(node_title) == pytest.approx(12.5, abs=0.1)
    assert node_title.paragraph_format.space_before.pt == pytest.approx(10, abs=0.1)
    assert node_title.paragraph_format.space_after.pt == pytest.approx(3, abs=0.1)
    assert node_item.paragraph_format.space_after.pt == pytest.approx(7, abs=0.1)
