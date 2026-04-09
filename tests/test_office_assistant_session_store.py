import builtins
import importlib
import json
import shutil
import subprocess
import struct
import sys
import zipfile
import zlib
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch
from uuid import uuid4

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from astrbot_plugin_office_assistant.agent_tools import (
    build_document_toolset,
)
from astrbot_plugin_office_assistant.agent_tools.document_tools import (
    CreateDocumentTool,
)
from astrbot_plugin_office_assistant.document_core.builders.table_renderer import (
    DOCX_TABLE_STYLES,
    TableRenderer,
)
from astrbot_plugin_office_assistant.document_core.macros.summary_card import (
    build_summary_card_group,
)
from astrbot_plugin_office_assistant.domain.document.session_store import (
    DocumentSessionStore,
    merge_document_style_defaults,
)
from astrbot_plugin_office_assistant.domain.document.export_pipeline import (
    export_document_via_pipeline,
)
from astrbot_plugin_office_assistant.domain.document.render_backends import (
    DocumentRenderBackendConfig,
    DocumentRenderBackendError,
    NodeDocumentRenderBackend,
    build_document_render_backends,
    RenderResult,
    build_document_render_payload,
)
from astrbot_plugin_office_assistant.document_core.models.blocks import (
    BlockStyle,
    BusinessReviewCoverData,
    ColumnBlock,
    ColumnsBlock,
    GroupBlock,
    HeaderFooterConfig,
    HeadingBlock,
    HeroBannerBlock,
    PageTemplateBlock,
    ParagraphBlock,
    ParagraphRun,
    SectionBreakBlock,
    SectionMarginsConfig,
    SummaryCardBlock,
    TableBlock,
    TocBlock,
)
from astrbot_plugin_office_assistant.document_core.models.document import (
    DocumentMetadata,
    DocumentModel,
    DocumentStyleConfig,
    DocumentSummaryCardDefaults,
)
from astrbot_plugin_office_assistant.domain.document.contracts import (
    AddBlocksRequest,
    AddHeadingRequest,
    AddListRequest,
    AddParagraphRequest,
    AddTableRequest,
    BlockHeadingInput,
    CreateDocumentRequest,
    ExportDocumentRequest,
    FinalizeDocumentRequest,
    SectionListInput,
    SectionParagraphInput,
    SectionTableInput,
    normalize_create_document_kwargs,
    normalize_raw_block_payloads,
)
from astrbot_plugin_office_assistant.mcp_server.server import (
    create_server,
)
from astrbot_plugin_office_assistant.tools.mcp_adapter import (
    register_document_tools_from_registry,
)
from astrbot_plugin_office_assistant.tools.astrbot_adapter import (
    build_document_toolset_from_registry,
)
from astrbot_plugin_office_assistant.tools.registry import (
    DocumentToolSpec,
    get_document_tool_specs,
)
from pydantic import ValidationError

from astrbot.core.utils.astrbot_path import get_astrbot_plugin_data_path


from tests._docx_test_helpers import *  # noqa: F401,F403
from tests._docx_test_helpers import _technical_resume_block



def test_document_session_store_keeps_exports_inside_workspace(workspace_root: Path):
    workspace_dir = _make_workspace(workspace_root, "pytest-agent-tools-paths")
    store = DocumentSessionStore(workspace_dir=workspace_dir)
    document = store.create_document(
        CreateDocumentRequest(
            session_id="pytest-session",
            output_name="../unsafe-name.docx",
        )
    )

    _, output_path = store.prepare_export_path(
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir="reports/q1",
            output_name="../final-report.docx",
        )
    )

    assert (
        output_path
        == (workspace_dir / "reports" / "q1" / "final-report.docx").resolve()
    )
    assert document.metadata.preferred_filename == "unsafe-name.docx"

    with pytest.raises(ValueError, match="relative to the document workspace"):
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir=str(workspace_dir.parent),
        )

    with pytest.raises(ValueError, match="cannot escape the document workspace"):
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir="../outside",
        )

    _, windows_style_output_path = store.prepare_export_path(
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir=r"reports\windows",
            output_name=r"nested\windows-report",
        )
    )

    assert (
        windows_style_output_path
        == (workspace_dir / "reports" / "windows" / "windows-report.docx").resolve()
    )

    with pytest.raises(ValueError, match="relative to the document workspace"):
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir=r"C:\temp",
        )

    with pytest.raises(ValueError, match="cannot escape the document workspace"):
        ExportDocumentRequest(
            document_id=document.document_id,
            output_dir=r"..\outside",
        )

def test_document_session_store_evicts_oldest_documents_when_capped():
    store = DocumentSessionStore(max_documents=2)
    first = store.create_document(CreateDocumentRequest(title="first"))
    second = store.create_document(CreateDocumentRequest(title="second"))
    third = store.create_document(CreateDocumentRequest(title="third"))

    assert store.get_document(first.document_id) is None
    assert store.get_document(second.document_id) is not None
    assert store.get_document(third.document_id) is not None

def test_document_session_store_evicts_expired_documents_by_ttl():
    store = DocumentSessionStore(ttl=timedelta(seconds=1))
    expired = store.create_document(CreateDocumentRequest(title="expired"))
    fresh = store.create_document(CreateDocumentRequest(title="fresh"))

    expired.metadata.updated_at = datetime.now(timezone.utc) - timedelta(seconds=5)
    fresh.metadata.updated_at = datetime.now(timezone.utc)

    assert store.get_document(expired.document_id) is None
    assert store.get_document(fresh.document_id) is not None

def test_word_document_builder_resolves_logical_table_styles():
    assert TableRenderer.resolve_docx_table_style("report_grid") == "Table Grid"
    assert (
        TableRenderer.resolve_docx_table_style("metrics_compact")
        == "Light List Accent 1"
    )
    assert TableRenderer.resolve_docx_table_style("minimal") == "Table Grid"
    assert TableRenderer.resolve_docx_table_style("") == "Table Grid"
    assert TableRenderer.resolve_docx_table_style("custom_style") == "custom_style"

def test_document_session_store_expands_summary_card_blocks():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="Summary Test"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_summary_card_block()],
        )
    )

    assert len(updated.blocks) == 1
    _assert_summary_card_group(updated.blocks[0])

def test_document_session_store_moves_landscape_intro_before_section_break():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="Landscape Flow"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[
                {
                    "type": "section_break",
                    "start_type": "new_page",
                    "page_orientation": "landscape",
                    "restart_page_numbering": True,
                    "page_number_start": 1,
                },
                {
                    "type": "heading",
                    "text": "二、各业务线详细数据盘点（宽表展示）",
                    "level": 1,
                },
                {
                    "type": "paragraph",
                    "text": "本章节详细列出了公司五大核心业务线在第一季度的关键经营指标。",
                },
                {
                    "type": "table",
                    "headers": ["业务线", "营收"],
                    "rows": [["核心电商业务", "21500"]],
                },
            ],
        )
    )

    assert [block.type for block in updated.blocks] == [
        "paragraph",
        "section_break",
        "heading",
        "table",
    ]
    assert (
        updated.blocks[0].text
        == "本章节详细列出了公司五大核心业务线在第一季度的关键经营指标。"
    )
    assert updated.blocks[1].page_orientation == "landscape"
    assert updated.blocks[2].text == "二、各业务线详细数据盘点（宽表展示）"

def test_document_session_store_applies_summary_card_defaults():
    store = DocumentSessionStore()
    document = store.create_document(
        CreateDocumentRequest(
            title="Summary Defaults",
            document_style={
                "summary_card_defaults": _summary_card_defaults()
            },
        )
    )

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_summary_card_block()],
        )
    )

    assert len(updated.blocks) == 1
    assert isinstance(updated.blocks[0], GroupBlock)
    title_block = updated.blocks[0].blocks[0]
    list_block = updated.blocks[0].blocks[1]
    assert title_block.style.align == "center"
    assert title_block.style.emphasis == "strong"
    assert title_block.style.font_scale == pytest.approx(1.2)
    assert title_block.layout.spacing_before == pytest.approx(12)
    assert title_block.layout.spacing_after == pytest.approx(4)
    assert list_block.layout.spacing_after == pytest.approx(8)

def test_document_session_store_tolerates_summary_card_default_resolution_errors():
    def _raise_defaults_error(_config):
        raise RuntimeError("bad defaults")

    store = DocumentSessionStore(summary_card_defaults_resolver=_raise_defaults_error)
    document = store.create_document(CreateDocumentRequest(title="Summary Fallback"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_summary_card_block()],
        )
    )

    assert len(updated.blocks) == 1
    _assert_summary_card_group(updated.blocks[0])

def test_domain_document_session_store_accepts_summary_card_defaults_resolver():
    from astrbot_plugin_office_assistant.domain.document.session_store import (
        DocumentSessionStore as DomainDocumentSessionStore,
    )

    def _raise_defaults_error(_config):
        raise RuntimeError("bad defaults")

    store = DomainDocumentSessionStore(
        summary_card_defaults_resolver=_raise_defaults_error
    )
    document = store.create_document(CreateDocumentRequest(title="Domain Summary"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_summary_card_block()],
        )
    )

    assert len(updated.blocks) == 1
    _assert_summary_card_group(updated.blocks[0])

def test_domain_document_session_store_uses_legacy_summary_card_patch_point(
    monkeypatch: pytest.MonkeyPatch,
):
    from astrbot_plugin_office_assistant.domain.document.session_store import (
        DocumentSessionStore as DomainDocumentSessionStore,
    )

    store = DomainDocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="Domain Summary"))

    def _raise_defaults_error(_config):
        raise RuntimeError("bad defaults")

    monkeypatch.setattr(
        "astrbot_plugin_office_assistant.mcp_server.session_store.summary_card_defaults_from_config",
        _raise_defaults_error,
    )

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_summary_card_block()],
        )
    )

    assert len(updated.blocks) == 1
    _assert_summary_card_group(updated.blocks[0])

def test_document_session_store_runs_internal_normalize_hooks():
    def _inject_heading(_context):
        return [
            BlockHeadingInput(text="Hook Title", level=2),
            SectionParagraphInput(text="正文"),
        ]

    store = DocumentSessionStore(normalize_block_hooks=[_inject_heading])
    document = store.create_document(CreateDocumentRequest(title="Hook Test"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[{"type": "paragraph", "text": "原始正文"}],
        )
    )

    assert len(updated.blocks) == 2
    assert updated.blocks[0].text == "Hook Title"
    assert updated.blocks[1].text == "正文"

def test_document_session_store_preserves_grouped_headers_for_table_blocks():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="Grouped Table"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[
                {
                    "type": "table",
                    "header_groups": [
                        {"title": "经营数据", "span": 2},
                        {"title": "结果", "span": 1},
                    ],
                    "headers": ["区域", "目标", "完成率"],
                    "rows": [["华东", "120", "98%"]],
                    "header_fill": "1F4E79",
                    "header_text_color": "FFFFFF",
                    "banded_rows": True,
                    "banded_row_fill": "EEF4FA",
                    "first_column_bold": True,
                    "table_align": "center",
                    "border_style": "strong",
                    "caption_emphasis": "strong",
                }
            ],
        )
    )

    table = updated.blocks[0]
    assert table.header_groups[0].title == "经营数据"
    assert table.header_groups[0].span == 2
    assert table.header_groups[1].title == "结果"
    assert table.header_fill == "1F4E79"
    assert table.header_text_color == "FFFFFF"
    assert table.banded_rows is True
    assert table.banded_row_fill == "EEF4FA"
    assert table.first_column_bold is True
    assert table.table_align == "center"
    assert table.border_style == "strong"
    assert table.caption_emphasis == "strong"

def test_document_session_store_preserves_hero_banner_and_report_style_fields():
    store = DocumentSessionStore()
    document = store.create_document(
        CreateDocumentRequest(
            title="经营复盘",
            document_style=DocumentStyleConfig(
                font_name="Microsoft YaHei",
                heading_font_name="Source Han Sans SC",
                table_font_name="SimSun",
                code_font_name="Consolas",
            ),
        )
    )

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[
                {
                    "type": "hero_banner",
                    "title": "Q3 经营复盘报告",
                    "subtitle": "战略与增长委员会",
                    "theme_color": "1F4E79",
                    "text_color": "FFFFFF",
                },
                {
                    "type": "accent_box",
                    "title": "核心摘要",
                    "text": "经营质量继续改善。",
                    "border_color": "CBD5E1",
                    "border_width_pt": 0.75,
                    "accent_border_width_pt": 3.0,
                    "padding_pt": 16,
                },
                {
                    "type": "metric_cards",
                    "metrics": [
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "value_font_scale": 1.9,
                            "note": "核心业务保持增长",
                            "note_color": "64748B",
                        }
                    ],
                    "border_color": "D9E1E8",
                    "divider_color": "CBD5E1",
                    "padding_pt": 14,
                },
                {
                    "type": "table",
                    "headers": ["区域", "完成率"],
                    "header_font_scale": 1.1,
                    "body_font_scale": 0.95,
                    "cell_padding_horizontal_pt": 8,
                    "cell_padding_vertical_pt": 6,
                    "rows": [
                        [
                            {
                                "text": "华东",
                                "font_scale": 1.15,
                            },
                            "112%",
                        ]
                    ],
                },
            ],
        )
    )

    hero_banner = updated.blocks[0]
    accent_box = updated.blocks[1]
    metric_cards = updated.blocks[2]
    table = updated.blocks[3]

    assert isinstance(hero_banner, HeroBannerBlock)
    assert hero_banner.subtitle == "战略与增长委员会"
    assert accent_box.border_color == "CBD5E1"
    assert accent_box.padding_pt == pytest.approx(16)
    assert metric_cards.border_color == "D9E1E8"
    assert metric_cards.metrics[0].value_font_scale == pytest.approx(1.9)
    assert metric_cards.metrics[0].note_color == "64748B"
    assert table.header_font_scale == pytest.approx(1.1)
    assert table.body_font_scale == pytest.approx(0.95)
    assert table.rows[0][0].font_scale == pytest.approx(1.15)


def test_merge_document_style_defaults_keeps_unset_font_fields_following_font_name():
    merged = merge_document_style_defaults(
        DocumentStyleConfig(),
        {
            "font_name": "Arial",
        },
    )

    assert merged.font_name == "Arial"
    assert "heading_font_name" not in merged.model_fields_set
    assert "table_font_name" not in merged.model_fields_set
    assert merged.model_dump(mode="python", exclude_unset=True) == {"font_name": "Arial"}


def test_document_session_store_preserves_page_template_block():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="经营复盘"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[
                _business_review_cover_block(
                    summary_text="Q3 营收同比增长 18.4%，整体毛利率保持稳定。",
                    metrics=[
                        {
                            "label": "营业收入",
                            "value": "¥4.82 亿",
                            "delta": "↑ 18.4% YoY",
                            "delta_color": "15803D",
                        }
                    ],
                    footer_note="编制：战略发展部 · 审核：CFO 办公室",
                )
            ],
        )
    )

    block = updated.blocks[0]

    assert isinstance(block, PageTemplateBlock)
    assert block.template == "business_review_cover"
    assert isinstance(block.data, BusinessReviewCoverData)
    assert block.data.title == "Q3 经营复盘报告"
    assert block.data.metrics[0].label == "营业收入"
    assert block.data.metrics[0].delta_color == "15803D"
    assert block.data.auto_page_break is False


def test_document_session_store_preserves_technical_resume_page_template_block():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="简历"))

    updated = store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[_technical_resume_block()],
        )
    )

    block = updated.blocks[0]

    assert isinstance(block, PageTemplateBlock)
    assert block.template == "technical_resume"
    assert block.data.name == "张明远"
    assert block.data.sections[0].title == "教育背景"

def test_document_session_store_add_table_preserves_grouped_headers():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="Legacy Table"))

    updated = store.add_table(
        AddTableRequest(
            document_id=document.document_id,
            headers=["区域", "目标", "完成率"],
            rows=[["华东", "120", "98%"]],
            header_groups=[
                {"title": "经营数据", "span": 2},
                {"title": "结果", "span": 1},
            ],
            header_fill="1F4E79",
            header_text_color="FFFFFF",
            banded_rows=True,
            banded_row_fill="EEF4FA",
            first_column_bold=True,
            table_align="center",
            border_style="strong",
            caption_emphasis="strong",
        )
    )

    table = updated.blocks[0]
    assert table.header_groups[0].title == "经营数据"
    assert table.header_groups[1].span == 1
    assert table.header_fill == "1F4E79"
    assert table.border_style == "strong"

def test_document_session_store_add_helpers_build_typed_blocks(
    monkeypatch: pytest.MonkeyPatch,
):
    store = DocumentSessionStore()
    captured: dict[str, AddBlocksRequest] = {}

    def _capture_add_blocks(self, request: AddBlocksRequest):
        captured["request"] = request
        return MagicMock()

    monkeypatch.setattr(DocumentSessionStore, "add_blocks", _capture_add_blocks)

    store.add_heading(
        AddHeadingRequest(
            document_id="doc-1",
            text="标题",
            level=2,
            bottom_border=True,
            bottom_border_color="D0D7DE",
            bottom_border_size_pt=1.25,
        )
    )
    assert isinstance(captured["request"].blocks[0], BlockHeadingInput)
    assert captured["request"].blocks[0].bottom_border is True
    assert captured["request"].blocks[0].bottom_border_color == "D0D7DE"
    assert captured["request"].blocks[0].bottom_border_size_pt == pytest.approx(1.25)

    store.add_paragraph(
        AddParagraphRequest(document_id="doc-1", text="正文", title="摘要")
    )
    assert isinstance(captured["request"].blocks[0], SectionParagraphInput)

    store.add_list(AddListRequest(document_id="doc-1", items=["要点 1"]))
    assert isinstance(captured["request"].blocks[0], SectionListInput)

    store.add_table(
        AddTableRequest(
            document_id="doc-1",
            headers=["区域"],
            rows=[["华东"]],
            title="表格标题",
        )
    )
    assert isinstance(captured["request"].blocks[0], SectionTableInput)

def test_document_session_store_builds_prompt_summary_for_draft_documents():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="季度复盘"))

    store.add_blocks(
        AddBlocksRequest(
            document_id=document.document_id,
            blocks=[
                {"type": "heading", "text": "概览", "level": 1},
                {"type": "paragraph", "text": "营收稳定增长。"},
                {
                    "type": "table",
                    "headers": ["区域", "完成率"],
                    "rows": [["华东", "98%"]],
                },
            ],
        )
    )

    summary = store.build_prompt_summary(document.document_id)

    assert summary == {
        "document_id": document.document_id,
        "title": "季度复盘",
        "status": "draft",
        "block_count": 3,
        "latest_block_types": ["heading", "paragraph", "table"],
        "next_allowed_actions": ["add_blocks", "finalize_document"],
    }

def test_document_session_store_builds_prompt_summary_for_later_states():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="经营复盘"))

    draft_summary = store.build_prompt_summary(document.document_id)
    assert draft_summary["latest_block_types"] == []
    assert draft_summary["next_allowed_actions"] == [
        "add_blocks",
        "finalize_document",
    ]

    store.finalize_document(FinalizeDocumentRequest(document_id=document.document_id))
    finalized_summary = store.build_prompt_summary(document.document_id)
    assert finalized_summary["status"] == "finalized"
    assert finalized_summary["next_allowed_actions"] == ["export_document"]

    store.complete_export(document.document_id)
    exported_summary = store.build_prompt_summary(document.document_id)
    assert exported_summary["status"] == "exported"
    assert exported_summary["next_allowed_actions"] == []


def test_document_session_store_rejects_add_blocks_after_finalize():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="经营复盘"))
    store.finalize_document(FinalizeDocumentRequest(document_id=document.document_id))

    with pytest.raises(
        ValueError,
        match="add_blocks is only allowed while the document status is draft",
    ):
        store.add_blocks(
            AddBlocksRequest(
                document_id=document.document_id,
                blocks=[BlockHeadingInput(text="一、经营总览", level=1)],
            )
        )

def test_document_session_store_builds_prompt_summary_with_unknown_block_type():
    store = DocumentSessionStore()
    document = store.create_document(CreateDocumentRequest(title="兼容块测试"))

    class _UnknownBlock:
        pass

    document.blocks.extend([_UnknownBlock()])

    summary = DocumentSessionStore._build_prompt_summary_locked(document)

    assert summary["latest_block_types"] == ["unknown"]
